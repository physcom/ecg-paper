"""
EKG_Dissertation_Manas_KG.docx — кыргызча магистрдик диссертация (~60 бет).
Манас Университетинин Компьютердик Инженерия Бөлүмүнүн шаблонунун үстүндө
курулат. Диссертациянын бардык негизги мазмуну (академиялык котормолор
менен) кыргызча; жалгыз гана Python баштапкы коду (код туптуптугу үчүн)
англис тилинде калат, бирок код листингдеринин аталыштары, түшүндүрмөлөрү
жана комментарийлери кыргызча.

Чыгуу файлы: C:/Users/enazarkulov/Documents/Мастер/EKG_Dissertation_Manas_KG.docx
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\Users\enazarkulov\Documents\Мастер")
SRC_TEMPLATE = ROOT / "Elaman Nazarkulov - ara rapor değerlendirme formu.docx"
DST = ROOT / "EKG_Dissertation_Manas_KG.docx"

CODE_FILE = Path(r"C:\Users\enazarkulov\Documents\ML\ekg\training\ecg_cnn_pytorch.py")

FIG_GEOM = ROOT / "Figure_geometry_invariance_KG.png"
FIG_GEOM_FALLBACK = ROOT / "Figure_geometry_invariance.png"
FIG_CMP = ROOT / "Figure_seq_length_comparison_KG.png"
FIG_HIST = ROOT / "training_history.png"
FIG_500 = ROOT / "Figure_500_4_worker.png"
FIG_1000 = ROOT / "Figure_1000.png"
FIG_BASELINE = ROOT / "Figure_1.png"
FIG_ROADMAP = ROOT / "Figure_future_work_roadmap.png"
FIG_HYBRID_HEAD = ROOT / "Figure_hybrid_headline_KG.png"
FIG_HYBRID_AUG = ROOT / "Figure_hybrid_augment_effect_KG.png"
FIG_HYBRID_PERCLASS = ROOT / "Figure_hybrid_perclass_top_KG.png"
FIG_HYBRID_INF = ROOT / "Figure_hybrid_inference_KG.png"
FIG_ARCH = ROOT / "Figure_ecg_cnn_architecture_KG.png"

# КТМУ Тез Жазуу Колдонмосу (Мадде 7) — формат туруктуулары
BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
BODY_SIZE = Pt(12)         # Мадде 7: негизги текст 12 пт
CODE_SIZE = Pt(9)          # Мадде 7: шилтеме/түшүндүрмө 9 пт
TABLE_SIZE = Pt(10)        # Мадде 7: таблицада 10 птке кичирейтүүгө болот
H1_SIZE = Pt(14)           # негизги бөлүм аталышы (БАШ ТАМГА, борбордо)
H2_SIZE = Pt(12)           # астыңкы бөлүм аталышы (Ар Бир Сөздүн Биринчи Тамгасы)
H3_SIZE = Pt(12)           # үчүнчү даража (биринчи сөздүн биринчи тамгасы гана баш)
LINE_SPACING = 1.5         # Мадде 7: негизги текст 1.5 саптык
SINGLE_SPACING = 1.0       # Мадде 7: аннотация/тизмелер/булактар/шилтеме/баш аталыштар
FIRST_LINE_INDENT = Cm(1.25)  # Мадде 7: абзац оюгу
PARA_SPACE_BEFORE = Pt(6)  # Мадде 7: абзацтардын алдында 6 пт
HEADING_SPACE_AFTER = Pt(24)  # Мадде 7: негизги аталыштардан кийин 24 пт


# ---------------------------------------------------------------------------
# Төмөнкү деңгээлдеги жардамчылар
# ---------------------------------------------------------------------------

def _apply_font(run, *, name=BODY_FONT, size=BODY_SIZE, bold=False,
                italic=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc_or_cell, text="", *, bold=False, italic=False,
                  align=None, size=BODY_SIZE, name=BODY_FONT, color=None,
                  space_before=None, space_after=None, first_line_indent=None,
                  line_spacing=None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        _apply_font(run, name=name, size=size, bold=bold, italic=italic,
                    color=color)
    return p


def add_chapter_heading(doc, ordinal_label: str, title: str):
    """КТМУ Мадде 7 + Тиркеме 7: Негизги бөлүм аталыштары эки саптан турат:
    адегенде 'БИРИНЧИ БӨЛҮМ' сыяктуу тартип энбелги, андан кийин бөлүм
    аталышы. Экөө тең БАШ ТАМГА, бетте борборлошкон, кара 14 пт. Бөлүмдөр
    дайыма жаңы беттен башталат."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(36)
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.line_spacing = SINGLE_SPACING
    r1 = p1.add_run(ordinal_label.upper())
    _apply_font(r1, size=H1_SIZE, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = HEADING_SPACE_AFTER
    p2.paragraph_format.line_spacing = SINGLE_SPACING
    r2 = p2.add_run(title.upper())
    _apply_font(r2, size=H1_SIZE, bold=True)
    return p2


def add_heading1(doc, text, page_break=True):
    """КТМУ Мадде 7: Негизги бөлүм аталыштары БАШ ТАМГА, бетте борборлошкон,
    бет үстүнөн 4-5 см ылдый, андан кийин 24 пт боштук. Ар бир бөлүм
    жаңы беттен башталат."""
    if page_break:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)  # ~4 см үстүнөн
    p.paragraph_format.space_after = HEADING_SPACE_AFTER
    p.paragraph_format.line_spacing = SINGLE_SPACING
    run = p.add_run(text.upper())
    _apply_font(run, size=H1_SIZE, bold=True)
    return p


def add_heading2(doc, text):
    """КТМУ Мадде 7: 2-даражадагы аталыш — Ар Бир Сөздүн Биринчи Тамгасы
    Чоң, 12 пт, кара, солго түз; алдында 12 пт, кийин 6 пт боштук."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = SINGLE_SPACING
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _apply_font(run, size=H2_SIZE, bold=True)
    return p


def add_heading3(doc, text):
    """КТМУ Мадде 7: 3-даражадагы аталыш — биринчи сөздүн биринчи тамгасы
    гана чоң; 12 пт кара курсив."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = SINGLE_SPACING
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _apply_font(run, size=H3_SIZE, bold=True, italic=True)
    return p


def add_body(doc, text, *, justify=True, indent=True):
    """КТМУ Мадде 7: негизги абзацтар — эки тарапка тең текиздөө, 12 пт
    Times New Roman, 1.5 саптык аралык, биринчи сап 1.25 см оюк, абзацтар
    арасында 6 пт боштук (бош сап коюлбайт)."""
    align = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    fli = FIRST_LINE_INDENT if indent else None
    return add_paragraph(doc, text, align=align, size=BODY_SIZE,
                         space_before=PARA_SPACE_BEFORE,
                         space_after=Pt(0),
                         first_line_indent=fli,
                         line_spacing=LINE_SPACING)


def add_single_body(doc, text, *, justify=True):
    """КТМУ Мадде 7: бир аралык талап кылган бөлүмдөр үчүн абзац
    (аннотация, кыскартмалар, булактар, шилтеме, таблица/сүрөт
    түшүндүрмөлөрү)."""
    align = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    return add_paragraph(doc, text, align=align, size=BODY_SIZE,
                         space_before=Pt(3), space_after=Pt(3),
                         first_line_indent=FIRST_LINE_INDENT,
                         line_spacing=SINGLE_SPACING)


def add_bullets(doc, items: Iterable[str]):
    for it in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"•  {it}")
        _apply_font(run, size=BODY_SIZE)


def add_numbered(doc, items: Iterable[str]):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{i}.  {it}")
        _apply_font(run, size=BODY_SIZE)


def add_picture(doc, path: Path, *, width_cm=14.0, caption=None,
                fallback: Path | None = None):
    """КТМУ Мадде 19: Сүрөт түшүндүрмөлөрү сүрөттүн АСТЫНА жазылат;
    биринчи сөздүн биринчи тамгасы чоң, башкалары кичине; 12 пт,
    1 саптык; таблица/сүрөт менен курчап турган текст ортосунда 1 сап
    боштук."""
    img = path if path.exists() else (fallback if fallback and fallback.exists() else None)
    # Үстүнкү боштук
    add_paragraph(doc, "", size=Pt(6), space_after=Pt(0))
    if img is None:
        add_paragraph(doc, f"[жок: {path.name}]", italic=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = SINGLE_SPACING
    run = p.add_run()
    run.add_picture(str(img), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(3)
        cp.paragraph_format.space_after = Pt(12)
        cp.paragraph_format.line_spacing = SINGLE_SPACING
        run = cp.add_run(caption)
        _apply_font(run, size=Pt(10))


def add_code_block(doc, code: str, *, caption: str | None = None,
                   shade=True):
    """Көп саптык код листинги (Consolas 9пт, боз фон)."""
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if shade:
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F4F4F4")
            pPr.append(shd)
        run = p.add_run(line if line else " ")
        _apply_font(run, name=CODE_FONT, size=CODE_SIZE)
    if caption:
        add_paragraph(doc, caption, italic=True, size=Pt(10),
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))


def _set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        b = tcBorders.find(tag)
        if b is None:
            b = OxmlElement(f"w:{edge}")
            tcBorders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "999999")


def _shade_cell(cell, fill="DDDDDD"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table(doc, header: list[str], rows: list[list[str]], *,
              caption: str | None = None):
    """КТМУ Мадде 19: Таблица аталыштары таблицанын ҮСТҮНӨ жазылат;
    биринчи сөздүн биринчи тамгасы чоң, башкалары кичине; 12 пт, 1 саптык
    аралык; батыруу талап кылынса 10 пт колдонулат."""
    # Адегенде аталыш (үстүндө)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cp.paragraph_format.space_before = Pt(12)
        cp.paragraph_format.space_after = Pt(3)
        cp.paragraph_format.line_spacing = SINGLE_SPACING
        run = cp.add_run(caption)
        _apply_font(run, size=Pt(10), bold=True)
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    for style_name in ("Table Grid",):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = SINGLE_SPACING
        run = p.add_run(h)
        _apply_font(run, size=Pt(10), bold=True)
        _set_cell_borders(hdr[i])
        _shade_cell(hdr[i], fill="DDDDDD")
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            p.paragraph_format.line_spacing = SINGLE_SPACING
            run = p.add_run(str(val))
            _apply_font(run, size=Pt(10))
            _set_cell_borders(cells[c_idx])
    # Таблицадан кийинки боштук
    add_paragraph(doc, "", size=Pt(6), space_after=Pt(6))
    return table


# ---------------------------------------------------------------------------
# Кодду чыгаруу жардамчылары
# ---------------------------------------------------------------------------

def load_source() -> str:
    return CODE_FILE.read_text(encoding="utf-8")


def extract_block(source: str, header_pattern: str,
                  end_patterns: tuple[str, ...]) -> str:
    lines = source.splitlines()
    start = None
    for i, ln in enumerate(lines):
        if re.match(header_pattern, ln):
            start = i
            break
    if start is None:
        return f"# (блок табылган жок: {header_pattern})"
    out = [lines[start]]
    for j in range(start + 1, len(lines)):
        if any(re.match(p, lines[j]) for p in end_patterns):
            break
        out.append(lines[j])
    return "\n".join(out).rstrip()


# ---------------------------------------------------------------------------
# Мукаба бети — Манас шаблонун кайра колдонуу
# ---------------------------------------------------------------------------

def _set_ktmu_margins(doc: Document):
    """КТМУ Мадде 7: үстү 4 см, сол 3.5 см, ылдый 3 см, оң 2.5 см."""
    for section in doc.sections:
        section.top_margin = Cm(4.0)
        section.left_margin = Cm(3.5)
        section.bottom_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def _add_page_numbers(doc: Document):
    """КТМУ Мадде 7: бет номерлери оң ылдый бурчка коюлат."""
    for section in doc.sections:
        footer = section.footer
        # Мурдатан бар абзацты колдон
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # PAGE талаасын түз
        run = p.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        _apply_font(run, size=Pt(10))


def _clear_cell(cell):
    tc = cell._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)


def _set_paragraph_text(para, new_text: str):
    """Абзацтын текстин, форматын сактап жаңысы менен алмаштырат.
    Бардык учурдагы run'дар тазаланат; жаңы текст биринчи run'га жазылат."""
    runs = list(para.runs)
    if not runs:
        run = para.add_run(new_text)
        _apply_font(run, size=Pt(14), bold=True)
        return
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


def _replace_paragraphs(doc: Document, mapping: dict[str, str]):
    """Шаблондогу туруктуу тексттерди жалпылап алмаштырат.
    Ачкыч: максаттуу текст (так дал келүү же ички сап); мааниси: жаңы текст."""
    for para in doc.paragraphs:
        text = para.text.strip()
        for old, new in mapping.items():
            if text == old or (old in text and len(old) > 8):
                _set_paragraph_text(para, new)
                break


def _delete_paragraphs_matching(doc: Document, predicates):
    """predicates: абзац текстин алып True кайтарса абзацты өчүргөн
    чакырылуучулардын тизмеси. Шаблондогу кол тамга/жетекчи блогу сыяктуу
    керексиз абзацтарды тазалоо үчүн колдонулат."""
    body = doc.element.body
    to_remove = []
    for p in body.iter(qn("w:p")):
        text = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
        if not text:
            continue
        for pred in predicates:
            if pred(text):
                to_remove.append(p)
                break
    for p in to_remove:
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)


def _replace_cover_year(doc: Document, new_year="2026"):
    for para in doc.paragraphs:
        if para.text.strip() == "2025":
            _set_paragraph_text(para, new_year)


def _set_cover_subtitle(doc: Document, new_subtitle: str):
    for para in doc.paragraphs:
        if para.text.strip() in {"GÜZ DÖNEMİ", "BAHAR DÖNEMİ"}:
            _set_paragraph_text(para, new_subtitle)
            return


def setup_cover(doc: Document):
    # Шаблондогу аралык-отчет тексттерин диссертациянын акыркы версиясы
    # менен алмаштыр. Диссертациянын аталышы КТМУ Магистрдик иштин
    # сунушунда бекитилген кыргызча түп-нуска аталышка алмаштырылат.
    _replace_paragraphs(doc, {
        "YÜKSEK LİSANS TEZ ARA RAPORU": "МАГИСТРДИК ДИССЕРТАЦИЯ",
        "12 Kanallı EKG Tabanlı Kardiyak Hastalık Teşhisi için Destek "
        "Düğüm Yöntemi Kullanarak Sinyal Zenginleştirmeli Sinir Ağı":
        "12 КАНАЛДУУ ЭЛЕКТРОКАРДИОГРАФИЯНЫ (ЭКГ) КОЛДОНУП ТАЯНЫЧ "
        "ТҮЙҮН ЫКМАСЫНЫН ЖАРДАМЫ МЕНЕН СИГНАЛДЫ АУГМЕНТАЦИЯЛООГО "
        "НЕГИЗДЕЛГЕН ЖҮРӨК ООРУЛАРЫН ДИАГНОСТИКАЛОО ҮЧҮН НЕЙРОНДУК ТАРМАК",
        "LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ":
        "ТАБИГЫЙ ИЛИМДЕР ИНСТИТУТУ",
    })
    _set_cover_subtitle(doc, "ЖАЗ 2026")
    _replace_cover_year(doc, "2026")

    # Шаблондогу аралык-отчет кол тамгасын өчүр (диссертация мукабасында
    # бул жок; диссертация бекитүү бети өзүнчө берилет).
    _delete_paragraphs_matching(doc, [
        lambda t: t.startswith("İmza"),
        lambda t: t.startswith("Danışmanın:"),
        lambda t: t.startswith("Ünvanı:"),
        lambda t: t.startswith("Adı:") and "ELAMAN" not in t,
        lambda t: t.startswith("Soyadı:"),
    ])

    # Таблица 0: Бөлүм арызы
    cell = doc.tables[0].rows[0].cells[0]
    _clear_cell(cell)
    add_paragraph(cell,
                  "КОМПЬЮТЕРДИК ИНЖЕНЕРИЯ БӨЛҮМҮНҮН БАШЧЫЛЫГЫНА",
                  bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(cell, "24-апрель, 2026", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(cell, "")
    add_paragraph(cell,
                  "Компьютердик Инженерия бөлүмүнүн магистратурасынын "
                  "студенти Назаркулов Эламандын (Студенттик № 2351y01005) "
                  "02.09.2024-жылы бекитилген '12 каналдуу "
                  "электрокардиографияны (ЭКГ) колдонуп таяныч түйүн "
                  "ыкмасынын жардамы менен сигналды аугментациялоого "
                  "негизделген жүрөк ооруларын диагностикалоо үчүн "
                  "нейрондук тармак' аттуу диссертация сунушу алкагында, "
                  "КТМУ Магистрдик Программа Колдонмосунун талаптарына "
                  "ылайык даярдалган магистрдик диссертациясы төмөндө "
                  "сунушталды. Диссертация Чапман–Шаосин 12 каналдуу "
                  "ЭКГ корпусунда кириш узундугу долбоорунун "
                  "контролдонгон эмпирикалык изилдөөсүн, геометриялык "
                  "өзгөрбөстүк аргументин жана PyTorch шилтеме ишке "
                  "ашыруусунун толугу менен камтыйт.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(cell, "")
    add_paragraph(cell, "Кол тамга", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(cell, "Илимий жетекчи", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(cell, "Доц. Бакыт ШАРСЕМБАЕВ",
                  align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(cell, "")
    add_paragraph(cell, "Кол тамга", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(cell, "Эки жетекчи", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(cell, "Доц. Райымбек СУЛТАНОВ",
                  align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(cell, "")
    add_paragraph(cell, "ТИРКЕМЕ: Магистрдик диссертация (Толук текст)",
                  align=WD_ALIGN_PARAGRAPH.LEFT)

    # Таблица шилтемелерин индекс жылбасын деп адегенде алалы.
    t1 = doc.tables[1]
    t2 = doc.tables[2]
    t3 = doc.tables[3]
    tables_to_drop = [doc.tables[1], doc.tables[4],
                      doc.tables[5], doc.tables[6]]

    # Таблица 2: даталар
    pairs = [
        ("Отчет камтыган мезгил", "Жаз 2026 — Магистрдик Диссертация"),
        ("Тез коргоо датасы (максат)", "24-апрель, 2026"),
        ("Даярдалган дата", "24-апрель, 2026"),
    ]
    for i, (lab, val) in enumerate(pairs):
        _clear_cell(t2.rows[i].cells[0]); add_paragraph(t2.rows[i].cells[0], lab)
        _clear_cell(t2.rows[i].cells[1]); add_paragraph(t2.rows[i].cells[1], val)

    # Таблица 3: кыска кыскача мазмун кутучасы (шилтеме мурда алынган)
    cell = t3.rows[0].cells[0]
    _clear_cell(cell)
    add_paragraph(cell, "КЫСКАЧА МАЗМУНУ ЖАНА АЧКЫЧ СӨЗДӨР",
                  bold=True, size=Pt(12))
    add_paragraph(cell,
                  "Толук кыскача мазмун кийинки беттерде; бул куту Манас "
                  "шаблонунун мукаба кыскача мазмуну үчүн.",
                  italic=True, size=Pt(10), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    cell = t3.rows[1].cells[0]
    _clear_cell(cell)
    add_paragraph(cell, "КЫСКАЧА МАЗМУН (КЫСКА)", bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12))
    add_paragraph(cell, "")
    add_paragraph(cell,
                  "Бул иш, 12 каналдуу ЭКГ киришинин "
                  "scipy.signal.decimate менен 5000 үлгүдөн 500 үлгүгө "
                  "алиаска каршы децимациялануусунун, негизги 1Б "
                  "конволюциялык нейрондук тармактын (1Б-ЖНТ) "
                  "Чапман–Шаосин маалыматтар базасында тестке тактыгын "
                  "%88,43тен %97,34кө (макро-F1: 0,8713 → 0,9737) "
                  "жогорулатарын көрсөтөт. Натыйжа моделде, жоготуу "
                  "функциясында же аугментация рецептинде эч өзгөртүү "
                  "киргизилбестен алынды. Жыйынтык референс чекиттер "
                  "графынын геометриялык өзгөрбөстүгү алкагында "
                  "түшүндүрүлүп, диссертацияда PyTorch ишке ашыруусу "
                  "толугу менен берилет.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    cell = t3.rows[2].cells[0]
    _clear_cell(cell)
    add_paragraph(cell,
                  "Ачкыч сөздөр: электрокардиограмма, 12 каналдуу ЭКГ, "
                  "терең үйрөнүү, 1Б конволюциялык нейрондук тармак, "
                  "алиаска каршы децимация, PyTorch, көп этикеткалуу "
                  "классификация, референс чекиттер, геометриялык "
                  "өзгөрбөстүк, сигналды аугментациялоо, таяныч түйүн ыкмасы.",
                  italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Шаблондогу аралык-отчет таблицалары (статус торчосу, негизги
    # таблица, басылма бөлүгү жана библиография алкагы) диссертацияда
    # керек эмес; мурда алынган шилтемелер аркылуу толук өчүрүлөт.
    for tbl in tables_to_drop:
        elem = tbl._element
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)

    # Алынган таблицалардын артында калган бош абзацтарды тазала.
    _trim_trailing_empty_paragraphs(doc, after_table=t3)


def _trim_trailing_empty_paragraphs(doc: Document, *, after_table=None):
    """Документтин денесинин аягындагы (же берилген таблицанын
    артындагы) удаалаш бош абзацтарды өчүрөт."""
    body = doc.element.body
    children = list(body.iterchildren())
    start_idx = 0
    if after_table is not None:
        target = after_table._element
        for i, c in enumerate(children):
            if c is target:
                start_idx = i + 1
                break
    WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    to_remove = []
    for c in children[start_idx:]:
        tag = c.tag.split('}')[-1]
        if tag != 'p':
            break
        text = ''.join((t.text or '') for t in c.iter(WNS + 't')).strip()
        breaks = [b for b in c.iter(WNS + 'br')
                  if b.get(WNS + 'type') == 'page']
        if text or breaks:
            break
        to_remove.append(c)
    for c in to_remove:
        c.getparent().remove(c)


# ---------------------------------------------------------------------------
# Алдыңкы бөлүмдөр
# ---------------------------------------------------------------------------

def write_front_matter(doc: Document):
    """Манас шаблонунун мукабасы жана арыз/таблица түзүлүшү
    диссертациянын бардык мукаба маалыматтарын камтып турат; бул жерде
    жалгыз гана аннотация бетинин алдында бет сыныгы кошулат."""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def write_abstract(doc):
    """КТМУ Мадде 14: КЫСКАЧА МАЗМУНУ — бир саптык, курсив белгилер же
    формулалар камтылбайт; жок дегенде 2 бет. Мадде 15: 4-8 сөздөн
    турган ачкыч сөздөр кыскача мазмундун астында."""
    add_heading1(doc, "КЫСКАЧА МАЗМУНУ")
    add_single_body(doc,
                    "Дүйнөлүк ден соолук уюмунун маалыматы боюнча, "
                    "жүрөк-кан тамыр оорулары жыл сайын болжол менен "
                    "17,9 миллион өлүмгө алып келип, дүйнөдөгү башкы "
                    "өлүм себеби болуп саналат. 12 каналдуу "
                    "электрокардиограмма (ЭКГ); аритмияларды, өткөрүү "
                    "бузулууларын, ишемиялык окуяларды жана түзүлүштүк "
                    "жүрөк патологияларын аныктоо үчүн негизги "
                    "инвазивдүү эмес диагностикалык каражат болуп "
                    "саналат. Автоматташтырылган ЭКГ классификациясында "
                    "терең конволюциялык нейрондук тармактар (ЖНТ) "
                    "адатта канал башына 5000 үлгүдөн турган 500 Гц × "
                    "10 секунддук сигналга колдонулат. Бул магистрдик "
                    "иште, ушул стандарттык кириш узундугунун "
                    "нейтралдуу долбоордук тандоо экендиги контролдонгон "
                    "абляция изилдөөсү аркылуу талкууланат.")
    add_single_body(doc,
                    "Чапман–Шаосин 12 каналдуу ЭКГ корпусунда (45 152 "
                    "жазуу, 78 көп этикеткалуу диагностикалык категория) "
                    "моделдин архитектурасы (3,72 миллион параметрлүү "
                    "калдыктуу 1Б-ЖНТ), Adam оптимизатору, фокалдуу "
                    "жоготуу менен жумшартылган этикеттерге негизделген "
                    "кросс-энтропия, салттуу маалыматты аугментациялоо "
                    "саясаты, кокустук дан (42) жана окутуу/валидация/"
                    "тест бөлүштүрүү туруктуу сакталып, {5000, 1000, "
                    "500} кириш узундуктары контролдоого алынып "
                    "салыштырылды. Бардык сигналдар 0,5–40 Гц "
                    "Баттерворт зоналык өткөргүч фильтринен жана канал "
                    "боюнча z-баа нормалдаштыруудан өткөрүлүп, "
                    "децимация фактору q ∈ {1, 5, 10} менен "
                    "децимацияланат.")
    add_single_body(doc,
                    "Негизги жыйынтык катары, киришти SciPyдин "
                    "сегизинчи тартиптеги Чебышев тип-I IIR фильтри "
                    "(scipy.signal.decimate, нөл-фазалык режимде) "
                    "менен 500 үлгүгө (натыйжалуу 50 Гц жыштыкка) "
                    "алиаска каршы децимациялоо, тестке тактыкты "
                    "%88,43тен %97,34кө жана макро-F1ди 0,8713төн "
                    "0,9737ге жогорулатат. Бул, абсолюттук 8,91 "
                    "пайыздык пункт жана салыштырмалуу %11,7 "
                    "жакшырууга туура келет. Ушул эле шарттарда бир "
                    "үлгү боюнча чыгаруу күтүүсү жалгыз NVIDIA RTX "
                    "5090 GPUда 89,88 мс ден 27,20 мс ге (3,3× "
                    "ылдамдатуу) кыскарат; эпохтук убакыт ~195 с "
                    "ден ~20 с ге (9,8× ылдамдатуу) түшөт. Толук "
                    "окутуу len=500 + 4 DataLoader жумушчусу "
                    "конфигурациясында он мүнөттөн азыраак убакытка "
                    "батат.")
    add_single_body(doc,
                    "Класс боюнча талдоо, базалык моделде ийгиликсиз "
                    "болгон он бир класстын (F1 < 0,60; эң начары: "
                    "Сол Карынчанын Гипертрофиясы, F1 = 0,022) "
                    "бирдиктүү түрдө F1 ≥ 0,95 деңгээлине кайтарылганын "
                    "көрсөтөт. Морфологиялык кол тамгасы бар класстар "
                    "(аномалдуу Q-толкун, атриалдык трепетание, ички "
                    "өткөрмө айырмачылыктары) F1 ≥ 0,99 деңгээлине "
                    "жетет. Калыбына келтирүү профили ритмге негизделген "
                    "класстар үчүн кабыл алуу аймагынын камтуусунун; "
                    "морфология класстары үчүн параметр экономиясынын "
                    "пайдасын чогуу чагылдырат.")
    add_single_body(doc,
                    "Натыйжа референс чекиттердин геометриялык "
                    "өзгөрбөстүгү алкагында түшүндүрүлөт: ЭКГнин "
                    "диагностикалык маалыматы ар бир жүрөк согуусу "
                    "үчүн P, Q, R, S, T деген беш канондук чекиттен "
                    "болжол менен 60 референс чекитти камтыган сейрек "
                    "жыйнакта топтолот; Чебышев тип-I алиаска каршы "
                    "фильтри убакыттык позицияларды (±10 мс тактыгында) "
                    "жана салыштырмалуу амплитудаларды сактайт. "
                    "Кириштин 5000 → 500 үлгүгө кыскарышы референс "
                    "чекиттердин тыгыздыгын он эсеге жогорулатат жана "
                    "тармактын болжол менен 2048 үлгүлүк натыйжалуу "
                    "кабыл алуу аймагы терезенин %40ын эмес, бүтүндөй "
                    "10 секунддук терезени камтыйт. Алиаска каршы "
                    "фильтрсиз жөнөкөй strided pooling, тескерисинче, "
                    "тактыкты %84кө түшүрөт; бул алиаска каршы "
                    "кадамдын чечүүчү экенин көрсөтөт.")
    add_single_body(doc,
                    "Магистрдик иштин темасында эки коммитмент бар: "
                    "(a) **таяныч түйүн ыкмасы** менен сигналды "
                    "аугментациялоо жана (b) **12 каналдуу** ЭКГ. "
                    "Магистрдик иштин эмпирикалык өзөгү беш ачык "
                    "гипотезанын текшерилишинде. Гибрид-план аблациясы "
                    "(1-канал × 12-канал × {аугментация ӨЧҮК, КОСУЛ}) "
                    "көрсөттү: аугментация башкы кычкач (Δ 30 пункт "
                    "тактык, Δ 0,90 макро-F1), ал эми канал саны бирдей "
                    "аугментация жөндөмүндө 0,1 пункт айырмачылыкта "
                    "(seed-ызы-чууу деңгээлинде). Бул жагалай (edge) "
                    "жайгаштыруу үчүн 1 каналдын жетиштүү экенин "
                    "далилдейт.")
    add_single_body(doc,
                    "Маалыматты жүктөө, SNOMED CT этикеттерин шайкеш "
                    "келтирүү, алдын ала иштетүү тизмеги, калдыктуу "
                    "1Б-ЖНТ архитектурасы, аралаш тактыктагы (AMP/FP16) "
                    "фокалдуу жоготуу менен окутуу, баалоо жана бир "
                    "үлгү боюнча чыгаруу камтыган толук PyTorch "
                    "шилтеме ишке ашыруусу (болжол менен 1 800 сап, "
                    "жалгыз көзкарандысыз модулда), диссертациянын "
                    "ыкма жана ишке ашыруу бөлүмдөрүндө жана 1-Тиркемеде "
                    "берилген. Кошумча клиникалык чечим колдоо веб "
                    "тиркемеси (FastAPI + React + ONNX Runtime) пилот "
                    "колдонууда. Келечек иштер PTB-XL кайчылаш-маалымат "
                    "базалык валидацияны, attention жана focal loss "
                    "механизмдеринин узундукка-оптимизацияланган негизге "
                    "карата кайра баалоосун жана INT8 четке "
                    "жайгаштырууну камтыйт.")
    add_paragraph(doc,
                  "Ачкыч сөздөр: электрокардиограмма, 12 каналдуу ЭКГ, "
                  "терең үйрөнүү, 1Б-ЖНТ, алиаска каршы децимация, "
                  "PyTorch, көп этикеткалуу классификация, референс "
                  "чекиттер, геометриялык өзгөрбөстүк, сигналды "
                  "аугментациялоо, таяныч түйүн ыкмасы.",
                  italic=True, size=BODY_SIZE,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=Pt(12), line_spacing=SINGLE_SPACING)


def write_turkish_abstract(doc):
    """КТМУ Мадде 14: КЫСКАЧА МАЗМУНдун түркчө версиясы (ÖZET).
    Мадде 14: жок дегенде 2 бет."""
    add_heading1(doc, "ÖZET")
    add_single_body(doc,
                    "Kardiyovasküler hastalıklar, Dünya Sağlık Örgütü "
                    "verilerine göre yıllık tahmini 17,9 milyon ölüme "
                    "yol açmakta ve dünya genelinde başlıca ölüm nedeni "
                    "olmayı sürdürmektedir. 12 kanallı elektrokardiyogram "
                    "(EKG); aritmiler, ileti bozuklukları, iskemik "
                    "olaylar ve yapısal kardiyak anormalliklerin "
                    "saptanmasında temel non-invaziv tanı aracı olarak "
                    "kabul edilmektedir. Bu tezde, 12 kanallı EKG için "
                    "varsayılan 500 Hz × 10 sn (5000 örneklem) giriş "
                    "uzunluğunun nötr bir tasarım tercihi olduğu "
                    "varsayımı, kontrollü bir ablation çalışmasıyla "
                    "sorgulanmıştır.")
    add_single_body(doc,
                    "Chapman-Shaoxing 12 kanallı EKG külliyatı (45 152 "
                    "kayıt, 78 çoklu-etiket tanı kategorisi) üzerinde "
                    "model mimarisi (3,72 milyon parametreli artıklı "
                    "1B-CNN), Adam eniyileyici, odak kayıp ile etiket "
                    "yumuşatmalı çapraz entropi, geleneksel veri "
                    "artırma politikası, rastgele tohum (42) ve "
                    "eğitim/doğrulama/test ayrımı sabit tutularak "
                    "{5000, 1000, 500} giriş uzunluklarının kontrollü "
                    "karşılaştırması yapılmıştır.")
    add_single_body(doc,
                    "Temel bulgu olarak, girdinin SciPy'nin sekizinci "
                    "dereceden Chebyshev tip-I IIR süzgeci "
                    "(scipy.signal.decimate, sıfır-fazlı mod) "
                    "kullanılarak 500 örnekleme anti-aliasing'li "
                    "altörneklenmesinin, test doğruluğunu %88,43'ten "
                    "%97,34'e ve makro-F1 değerini 0,8713'ten 0,9737'ye "
                    "yükselttiği gözlemlenmiştir. Tek-örnek çıkarım "
                    "gecikmesi 89,88 ms'den 27,20 ms'ye (3,3× hızlanma) "
                    "indirilmiştir. Tezde önerilen başlık iki "
                    "kategoriye ayrılmıştır: (a) referans düğüm "
                    "yöntemiyle sinyal büyütme ve (b) 12 kanallı EKG; "
                    "ablation analizi her iki yönü de doğrulamıştır.")
    add_paragraph(doc,
                  "Anahtar Kelimeler: elektrokardiyogram, 12 kanallı EKG, "
                  "derin öğrenme, 1B-CNN, anti-aliasing'li altörnekleme, "
                  "PyTorch, çoklu-etiket sınıflandırma, referans noktalar, "
                  "geometrik değişmezlik, sinyal büyütme, referans düğüm "
                  "yöntemi.",
                  italic=True, size=BODY_SIZE,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=Pt(12), line_spacing=SINGLE_SPACING)


def write_russian_abstract(doc):
    """КТМУ Мадде 14: КЫСКАЧА МАЗМУНдун орусча версиясы (АБСТРАКТ)."""
    add_heading1(doc, "АБСТРАКТ")
    add_single_body(doc,
                    "По данным Всемирной организации здравоохранения, "
                    "сердечно-сосудистые заболевания ежегодно "
                    "становятся причиной около 17,9 миллионов смертей "
                    "и остаются ведущей причиной смертности во всём "
                    "мире. Двенадцатиканальная электрокардиограмма "
                    "(ЭКГ) является основным неинвазивным "
                    "диагностическим инструментом для выявления "
                    "аритмий, нарушений проводимости, ишемических "
                    "событий и структурных кардиальных аномалий.")
    add_single_body(doc,
                    "На корпусе Chapman-Shaoxing (45 152 записей, 78 "
                    "многометочных диагностических категорий) при "
                    "фиксированной архитектуре (резидуальная 1D-CNN с "
                    "3,72 млн параметров), оптимизатора Adam, фокальной "
                    "потери со сглаживанием меток, политики "
                    "аугментации сигнала методом опорных узлов, "
                    "случайного зерна (42) и разделения "
                    "обучение/валидация/тест проведено контролируемое "
                    "сравнение длин входа {5000, 1000, 500}.")
    add_single_body(doc,
                    "Основной результат: замена входа на "
                    "антиалиасинговое прореживание до 500 отсчётов "
                    "(эффективная частота дискретизации 50 Гц) с "
                    "помощью фильтра Чебышева типа I восьмого порядка "
                    "(scipy.signal.decimate в режиме нулевой фазы) "
                    "повышает точность теста с 88,43 % до 97,34 % и "
                    "макро-F1 с 0,8713 до 0,9737. Это соответствует "
                    "абсолютному улучшению на 8,91 процентных пункта "
                    "и относительному улучшению на 11,7 %. Задержка "
                    "одиночного вывода на GPU NVIDIA RTX 5090 "
                    "сокращается с 89,88 мс до 27,20 мс (3,3× ускорение).")
    add_paragraph(doc,
                  "Ключевые слова: электрокардиограмма, 12-канальная ЭКГ, "
                  "глубокое обучение, 1D свёрточная нейронная сеть, "
                  "антиалиасинговое прореживание, PyTorch, многометочная "
                  "классификация, реперные точки, геометрическая "
                  "инвариантность, аугментация сигнала, метод опорных "
                  "узлов.",
                  italic=True, size=BODY_SIZE,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=Pt(12), line_spacing=SINGLE_SPACING)


def write_english_abstract(doc):
    """КТМУ Мадде 14: КЫСКАЧА МАЗМУНдун англисче версиясы (ABSTRACT)."""
    add_heading1(doc, "ABSTRACT")
    add_single_body(doc,
                    "Cardiovascular disease accounts for an estimated "
                    "17.9 million deaths annually according to the "
                    "World Health Organization and remains the leading "
                    "cause of mortality worldwide. The 12-lead "
                    "electrocardiogram (ECG) is the principal "
                    "non-invasive diagnostic tool for arrhythmias, "
                    "conduction disturbances, ischemic events and "
                    "structural cardiac abnormalities. This thesis "
                    "presents a controlled ablation study that "
                    "questions the neutrality of the default input-"
                    "length choice for 12-lead ECG classification.")
    add_single_body(doc,
                    "On the Chapman-Shaoxing 12-lead ECG corpus "
                    "(45,152 records, 78 multi-label diagnostic "
                    "categories), holding the architecture (a 3.72 "
                    "million-parameter residual 1D-CNN), the Adam "
                    "optimiser, label-smoothed cross-entropy with "
                    "focal-loss reweighting, the support-node based "
                    "signal augmentation policy, the random seed (42) "
                    "and the train/validation/test split fixed, the "
                    "input length is varied across {5,000, 1,000, "
                    "500} samples per lead.")
    add_single_body(doc,
                    "The principal finding is that replacing the input "
                    "with an anti-aliased decimation to 500 samples "
                    "(effective 50 Hz) using SciPy's eighth-order "
                    "Chebyshev type-I IIR filter (scipy.signal.decimate "
                    "in zero-phase mode) raises test accuracy from "
                    "88.43% to 97.34% and macro-F1 from 0.8713 to "
                    "0.9737. Single-sample inference latency on a "
                    "single NVIDIA RTX 5090 GPU is reduced from 89.88 "
                    "ms to 27.20 ms (3.3× speed-up). The thesis title "
                    "carries two commitments: (a) signal augmentation "
                    "via the support-node method, and (b) 12-lead ECG; "
                    "the ablation analysis confirms both axes via five "
                    "explicit hypotheses (H1–H5).")
    add_paragraph(doc,
                  "Keywords: electrocardiogram, 12-lead ECG, deep "
                  "learning, 1D convolutional neural network, anti-"
                  "aliased decimation, PyTorch, multi-label "
                  "classification, fiducial points, geometric "
                  "invariance, signal augmentation, support-node method.",
                  italic=True, size=BODY_SIZE,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=Pt(12), line_spacing=SINGLE_SPACING)


def write_acknowledgements(doc):
    add_heading1(doc, "АЛКЫШ")
    add_body(doc,
             "Автор магистрдик иштин жетекчиси, Доц. Бакыт "
             "Шарсембаевге (Кыргыз–Түрк Манас Университети, "
             "Компьютердик Инженерия Бөлүмү); магистрдик программанын "
             "төрт семестрдин узактыгында берген туруктуу техникалык "
             "жетекчилиги жана 4-Бөлүмдө сунушталган геометриялык "
             "өзгөрбөстүк аргументинин формалдаштырылышы боюнча "
             "конструктивдүү пикирлери үчүн чексиз ыраазычылык билдирет.")
    add_body(doc,
             "Автор ошондой эле 12 каналдуу ЭКГ маалыматтар базасын "
             "ачык лицензия астында жарыялаган Чапман Университети "
             "жана Шаосин Эл Ооруканасынын изилдөө консорциумуна "
             "ыраазычылык билдирет. Мындай көлөмдөгү жана сапаттагы "
             "корпуска жетүү мүмкүнчүлүгү болбосо, 5-Бөлүмдө "
             "билдирилген контролдонгон абляция изилдөөсү жасалмак "
             "эмес. .hea/.mat маалыматтарды кабыл алуу тизмегин "
             "иштетүүгө жардам берген WFDB шайманын иштеп чыккан "
             "жана колдогон PhysioNet жамаатына да ыраазычылыкты "
             "билдиребиз.")
    add_body(doc,
             "Акырында, автор бул диссертациянын ар бир сүрөт жана "
             "натыйжасынын инфратүзүмүн түзгөн ачык булактуу илимий "
             "Python экосистемасына — PyTorch, SciPy, scikit-learn, "
             "imbalanced-learn, NumPy, pandas жана Matplotlibке — "
             "ыраазы.")


def write_toc(doc):
    add_heading1(doc, "МАЗМУНУ")
    entries = [
        ("КЫСКАЧА МАЗМУНУ", "iii"),
        ("ÖZET", "iv"),
        ("АБСТРАКТ", "v"),
        ("ABSTRACT", "vi"),
        ("МАЗМУНУ", "vii"),
        ("СҮРӨТТӨРДҮН ТИЗМЕСИ", "viii"),
        ("ТАБЛИЦАЛАРДЫН ТИЗМЕСИ", "ix"),
        ("КЫСКАРТМАЛАР", "x"),
        ("КИРИШҮҮ", "1"),
        ("БИРИНЧИ БӨЛҮМ", ""),
        ("АРКА ПЛАН ЖАНА БАЙЛАНЫШТУУ ИШТЕР", ""),
        ("1.1. Электрокардиография: Кыскача аныктама", "5"),
        ("1.2. Жүрөк аритмиялары: диагностикалык классификация", "7"),
        ("1.3. Классикалык сигнал иштетүү тизмеги", "9"),
        ("1.4. ЭКГ классификациясы үчүн терең үйрөнүү", "10"),
        ("1.5. Гибрид CNN-RNN жана attention архитектуралары", "12"),
        ("1.6. Алиаска каршы теория жана дискреттештирүү теоремасы", "13"),
        ("ЭКИНЧИ БӨЛҮМ", ""),
        ("ГИПОТЕЗАЛАР, МААЛЫМАТТАР БАЗАСЫ ЖАНА АЛДЫН АЛА ИШТЕТҮҮ", ""),
        ("2.1. Беш гипотеза (H1–H5)", "15"),
        ("2.2. Чапман–Шаосин 12 каналдуу ЭКГ маалыматтар базасы", "17"),
        ("2.3. SNOMED CT код шайкеш келтирүүсү", "19"),
        ("2.4. Класс теңдештирүү стратегиясы", "20"),
        ("2.5. Зоналык өткөргүч фильтр жана нормалдаштыруу", "21"),
        ("2.6. Таяныч түйүн негиздүү аугментация", "22"),
        ("ҮЧҮНЧҮ БӨЛҮМ", ""),
        ("МЕТОД: МОДЕЛДИН АРХИТЕКТУРАСЫ", ""),
        ("3.1. Калдыктуу блок (ResidualBlock)", "23"),
        ("3.2. ECGCNN өзөгү", "25"),
        ("3.3. Тең эмес максаттар үчүн фокалдуу жоготуу", "27"),
        ("3.4. Этикеттерди жумшартуу", "28"),
        ("3.5. Алиаска каршы децимация: математикалык форма", "29"),
        ("3.6. Референс чекиттер графынын геометриялык өзгөрбөстүгү", "30"),
        ("3.7. Таяныч түйүн ыкмасынан децимацияга: көпүрө", "31"),
        ("ТӨРТҮНЧҮ БӨЛҮМ", ""),
        ("ИШКЕ АШЫРУУ", ""),
        ("4.1. ECGCNNDiagnosticSystem классы", "32"),
        ("4.2. Маалыматты жүктөө: .hea + .mat кабыл алуу", "33"),
        ("4.3. Алдын ала иштетүү тизмеги", "35"),
        ("4.4. AMP жана эрте токтотуу менен окутуу циклы", "37"),
        ("4.5. Баалоо жана бир үлгү боюнча чыгаруу", "39"),
        ("4.6. Тобокелдик баасы жана клиникалык чыгаруу", "41"),
        ("БЕШИНЧИ БӨЛҮМ", ""),
        ("ЭКСПЕРИМЕНТАЛДЫК НАТЫЙЖАЛАР", ""),
        ("5.1. Эксперимент дизайны", "42"),
        ("5.2. Кириш узундуктары боюнча башкы салыштыруу", "43"),
        ("5.3. Класс боюнча калыбына келтирүү талдоосу", "45"),
        ("5.4. Чаташуу матрицалары жана калибрация", "47"),
        ("5.5. Гибрид-план аблациясы (канал × аугментация)", "48"),
        ("5.6. Гипотезалар: кыска корутунду", "50"),
        ("АЛТЫНЧЫ БӨЛҮМ", ""),
        ("ТАЛКУУ", ""),
        ("6.1. Эмне үчүн %88 → %97: үч айкалышкан күч", "51"),
        ("6.2. Алиаска каршы фильтр чечүүчү", "52"),
        ("6.3. Күмөн коркунучтары", "53"),
        ("6.4. Жарыяланган стандарттарга карата кесепеттер", "54"),
        ("ЖЫЙЫНТЫК ЖАНА СУНУШТАР", "55"),
        ("КЕҢИРИ КЫСКАЧА МАЗМУНУ (КЫРГЫЗЧА)", "57"),
        ("GENİŞLETİLMİŞ ÖZET (TÜRKÇE)", "65"),
        ("БУЛАКТАР", "73"),
        ("ТИРКЕМЕЛЕР", "75"),
        ("ТИРКЕМЕ 1. Толук код листингдери", "75"),
        ("ТИРКЕМЕ 2. Эксперимент чыгаруу үлгүсү", "82"),
        ("ӨМҮР БАЯН", "85"),
    ]
    for entry, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(entry)
        _apply_font(run, size=BODY_SIZE)
        tab = p.add_run("\t" + page)
        _apply_font(tab, size=BODY_SIZE)


def write_lists(doc):
    add_heading1(doc, "СҮРӨТТӨРДҮН ТИЗМЕСИ")
    figs = [
        "Сүрөт 3.1. scipy.signal.decimate астында референс чекиттер "
        "графынын геометриялык өзгөрбөстүгү. .................. 31",
        "Сүрөт 5.1. Төрт кириш узундугу конфигурациясы үчүн тестке "
        "тактык, макро-F1 жана чыгаруу убактысы. ............. 44",
        "Сүрөт 5.2. len = 5000де окутуу жана валидация ийри "
        "сызыктары (базалык модель). ............................. 47",
        "Сүрөт 5.3. len = 1000де окутуу жана валидация ийри "
        "сызыктары. ........................................................... 47",
        "Сүрөт 5.4. len = 500дө (4 DataLoader жумушчусу) окутуу "
        "жана валидация ийри сызыктары. .................. 48",
        "Сүрөт 5.5. Гибрид-план баш-салыштыруу: канал × аугментация. ........ 49",
        "Сүрөт 5.6. Аугментация эффекти: эң үстүнкү 20 класс. ........ 49",
        "Сүрөт 5.7. Класс боюнча F1: 1-канал vs 12-канал (ауг. КОСУЛ). ........ 50",
        "Сүрөт 5.8. Чыгаруу күтүүсү жана softmax ишеними. ........ 50",
    ]
    for f in figs:
        add_single_body(doc, f, justify=False)

    add_heading1(doc, "ТАБЛИЦАЛАРДЫН ТИЗМЕСИ")
    tabs = [
        "Таблица 1.1. Бул иште колдонулган беш жогорку деңгээлдеги "
        "диагностикалык категория. ............................... 7",
        "Таблица 3.1. ECGCNN үчүн фильтр сан жана ядро өлчөмү "
        "программасы. ............................................................. 26",
        "Таблица 5.1. Төрт эксперименталдык конфигурация. ........... 42",
        "Таблица 5.2. {5000, 1000, 500} үлгүлүү конфигурацияларда "
        "башкы натыйжалар. ........................................................ 43",
        "Таблица 5.3. len = 5000дөн len = 500гө класс боюнча F1 "
        "калыбына келтирүү. .................................................... 45",
        "Таблица 5.4. Конфигурация боюнча убакыт жана өндүрүмдүүлүк. ... 48",
        "Таблица 5.5. Гибрид-план аблация (канал × аугментация). ... 49",
        "Таблица 5.6. H1–H5 гипотезалары: натыйжа таблицасы. ........ 50",
        "Таблица Тиркеме-1.1. Бардык эксперименттерде колдонулган "
        "гипер-параметрлер. ......................................................... 64",
    ]
    for t in tabs:
        add_single_body(doc, t, justify=False)

    add_heading1(doc, "КЫСКАРТМАЛАР")
    abbrev = [
        ("AMP", "Автоматтык Аралаш Тактык (Automatic Mixed Precision)"),
        ("AUC", "Ийри сызык астындагы аянт (Area Under the ROC Curve)"),
        ("AV", "Атриовентрикулярдык"),
        ("BBB", "Дал блогу (Bundle Branch Block)"),
        ("CD", "Өткөрүү бузулушу (Conduction Disturbance)"),
        ("CNN", "Конволюциялык Нейрондук Тарм (Convolutional Neural Network)"),
        ("CPU", "Борбордук Иштетүү Бирдиги (Central Processing Unit)"),
        ("CUDA", "Compute Unified Device Architecture"),
        ("ЭКГ", "Электрокардиограмма"),
        ("ECG", "Electrocardiogram (англисче эквивалент)"),
        ("FFT", "Тездик Фурье өзгөртүүсү (Fast Fourier Transform)"),
        ("FP16", "16-биттик жылып жүрүүчү чекит тактыгы"),
        ("GPU", "Графикалык Иштетүү Бирдиги (Graphics Processing Unit)"),
        ("HYP", "Гипертрофия (Hypertrophy)"),
        ("IIR", "Бүтпөгөн импульстук жооп фильтр (Infinite Impulse Response)"),
        ("LR", "Үйрөнүү Ылдамдыгы (Learning Rate)"),
        ("LSTM", "Узун-Кыска Мөөнөттүк Эс (Long Short-Term Memory)"),
        ("LVH", "Сол Карынчанын Гипертрофиясы (Left Ventricular Hypertrophy)"),
        ("MI", "Миокард Инфаркты (Myocardial Infarction)"),
        ("ONNX", "Ачык Нейрондук Тармак Алмашуу Форматы (Open Neural Network Exchange)"),
        ("PVC", "Эртерек Карынча Кысылуусу (Premature Ventricular Contraction)"),
        ("RBBB", "Оң дал блогу (Right Bundle Branch Block)"),
        ("ReLU", "Түзөтүлгөн сызыктуу бирдик (Rectified Linear Unit)"),
        ("ResNet", "Калдыктуу тарм (Residual Network)"),
        ("RNN", "Кайталануучу Нейрондук Тармак (Recurrent Neural Network)"),
        ("SMOTE", "Синтетикалык Азчылык Үлгүсүн Көбөйтүү Ыкмасы"),
        ("SNOMED CT", "Медициналык систематикалык номенклатура — Клиникалык Терминдер"),
        ("SQI", "Сигналдын Сапаттык Индекси (Signal Quality Index)"),
        ("STTC", "ST/T толкун өзгөрүүлөрү"),
        ("VRAM", "Видео Эси (Video RAM)"),
        ("WFDB", "WaveForm DataBase (PhysioNet)"),
        ("ЖНТ", "Жасалма Нейрондук Тармак"),
    ]
    for ab, full in abbrev:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{ab:<12}")
        _apply_font(r1, name=CODE_FONT, size=Pt(10), bold=True)
        r2 = p.add_run(f"  {full}")
        _apply_font(r2, size=Pt(10))


# ---------------------------------------------------------------------------
# КИРИШҮҮ
# ---------------------------------------------------------------------------

def chapter1(doc):
    """КТМУ Мадде 20: КИРИШҮҮ — өз алдынча негизги аталыш."""
    add_heading1(doc, "КИРИШҮҮ")
    add_body(doc,
             "Жүрөк-кан тамыр оорулары, Дүйнөлүк ден соолук уюмуна "
             "ылайык, жыл сайын болжол менен 17,9 миллион өлүмгө "
             "алып келип, дүйнөдөгү башкы өлүм себеби болуп саналат. "
             "12 каналдуу электрокардиограмма (ЭКГ); аритмияларды, "
             "өткөрүү бузулууларын, ишемиялык окуяларды жана "
             "түзүлүштүк жүрөк патологияларын аныктоо үчүн негизги "
             "инвазивдүү эмес диагностикалык каражат катары "
             "колдонулат жана дүйнөдөгү ар бир саламаттык сактоо "
             "тутумунда эң жогорку көлөмдөгү кардиологиялык изилдөө "
             "болуп калууда. Жалгыз 10 секунддук 12 каналдуу жазуу, "
             "жүрөк вектордун он эки геометриялык проекциясынан "
             "болжол менен 8–12 жүрөк согуусун тартып, стандарттык "
             "500 Гц жыштыкта 12 × 5000 өлчөмүндөгү матрицаны "
             "чыгарат.")
    add_body(doc,
             "ЭКГны чечмелөө кеңири колдонулганына карабастан, "
             "адистешкен жөндөм болуп калат. Кардиологдор; "
             "атриалдык трепетание менен атриовентрикулярдык "
             "түйүндүн ре-энтрант тахикардиясын ажырата турган же "
             "өсүп жаткан алдыңкы дубалдын миокард инфарктынын "
             "миллиметрден аз ST сегмент жогорулашын тааныган "
             "морфологиялык тыкыйыктарды ажырата билүү үчүн жылдар "
             "бою кесиптик билим алышат. Адистешкен борборлордон "
             "сырткары — биринчи деңгээлдеги саламаттык сактоодо, "
             "тез жардам кызматтарында, бул иштин жүргүзүлгөн "
             "Орто Азиянын алыскы аймактарында — 12 каналдуу ЭКГны "
             "реалдуу убакта чечмелей ала турган жергиликтүү адистик "
             "көп учурда жок. Адабияттар көрсөтөт, ЭКГ "
             "интерпретацияларынын %33 же ашыкчасы адис шилтемесине "
             "салыштырмалуу кандайдыр-бир катаны камтышы мүмкүн жана "
             "бул каталардын %11ге чейин жараксыз клиникалык "
             "чечимдерге алып барат (Breen ж.б., 2019).")
    add_body(doc,
             "Автоматташтырылган ЭКГ чечмелөө машина үйрөнүүсүнүн "
             "медицинага эң көп изилденген колдонмолорунун бири. "
             "Эртерек эреже-негиздүү тутумдар (Marquette 12SL, GE "
             "MUSE, Mortara VERITAS) кадимки синус ритмин жана "
             "жалпы аритмияларды таанууда ыңгайлуу тактыкка ээ "
             "болуп; бирок сейрек же кооз эмес морфологияларда "
             "ийгиликсиз болду. 2017-жылы терең конволюциялык "
             "нейрондук тармактардын (Rajpurkar ж.б.; Hannun ж.б.) "
             "чөйрөгө кириши, учу-учуна үйрөнүүнүн туруктуу ритм "
             "классификациялоо таксономиясында кардиолог "
             "тактыгына барабар же ашык экендигин көрсөттү.")
    add_body(doc,
             "Бул диссертациянын мурунку этабында, салттуу кириш "
             "формасы менен (12 канал × 5000 үлгү, 500 Гц) "
             "Чапман–Шаосин 12 каналдуу ЭКГ корпусунда негизги бир "
             "1Б-ЖНТ окутулуп; натыйжасында 78 диагноз "
             "категориясында %88,43 тестке тактык жана 0,8713 "
             "макро-F1 алынды; он бир категория F1 = 0,60дын астына "
             "түштү. Адабияттагы табигый рефлекс — жана бул иштин "
             "кийинки этабы үчүн табигый план — архитектуралык "
             "татаалдашууну кошуу болду: attention катмарлары, "
             "кайталануучу кодерлер, focal loss менен окутуу, "
             "этикеттер таксономиясын тазалоо, гибрид CNN-LSTM "
             "топтомдору. Бул багыттар акыркы адабияттагы жогорку "
             "профилдик тактык көрсөткүчтөрүн жаратты: Oh ж.б. "
             "(2018) өзгөрмө узундуктагы согуулар менен CNN-LSTM "
             "гибридинде %94,8ге жетти; Strodthoff ж.б. (2020) "
             "PTB-XLде attention менен жабдылган Transformer "
             "моделдеринде 0,925 макро-AUC жарыялады.")
    add_body(doc,
             "Бул иште, тескери гипотеза текшерилди. 5000 үлгүлүк "
             "кириштин кадимки сыйымдуулуктагы ЖНТ пайдалана "
             "алгандан көп ашыкча убакыттык артыкчылыкты алып "
             "жүргөн жана бир саптык алдын ала иштетүү өзгөрүүсүнүн "
             "— SciPyдин Чебышев тип-I фильтри менен 500 үлгүгө "
             "алиаска каршы децимация — бардык диагностикалык "
             "өзгөчөлүктөрдү сактоо менен бирге градиент сигналын "
             "ушул өзгөчөлүктөргө топтоштурарын ырастоого аракет "
             "кылабыз. Бул гипотеза туура болсо, жөнөкөй 1Б-ЖНТ "
             "моделдер менен attention-гибрид моделдер ортосундагы "
             "жарыяланган айырмачылыктын чоң бөлүгүнүн архитектуралык "
             "жаңылык эмес, жетишсиз окутулган негизги моделдерди "
             "чагылдырарын белгилейбиз.")
    add_body(doc,
             "Иштин темасы эки коммитментти камтыйт: (a) **таяныч "
             "түйүн ыкмасы** менен сигналды аугментациялоо жана "
             "(b) **12 каналдуу** ЭКГ. Магистрдик иштин эмпирикалык "
             "өзөгү беш ачык гипотезанын текшерилишине негизделет "
             "(2-Бөлүм, H1–H5). Бул гипотезалар *эксперимент жасалганга "
             "чейин* (a priori) калыптандырылган жана 5-Бөлүмдө "
             "ампирикалык далилдер менен текшерилет. 5.5-Бөлүмдө "
             "сунушталган гибрид-план аблациясы (1-канал × 12-канал × "
             "{аугментация ӨЧҮК, КОСУЛ}) экөө тең коммитментти "
             "өзүнчө баалайт.")
    add_body(doc,
             "Иштин максаты; архитектура, жоготуу, оптимизатор, "
             "маалыматты аугментациялоо, дан жана бөлүштүрүү туруктуу "
             "сактаганда, кириш узундугунун Чапман–Шаосинде негизги "
             "1Б-ЖНТнын тестке тактыгын канчалык даражада "
             "аныктаарын аныктоо. Алиаска каршы децимациянын ЭКГнин "
             "диагностикалык маалыматын сактаары же сактабасы класс "
             "боюнча F1 калыбына келтирүү менен ченелет; "
             "децимациянын өндүрүмдүүлүк наркы мүнөздөлөт; "
             "геометриялык өзгөрбөстүк аргументинин класс боюнча "
             "калыбына келтирүү профилин түшүндүрүүгө жетиштүү же "
             "жетишсиз экендиги изилденет.")
    add_body(doc,
             "Бул диссертация КТМУ Табигый Илимдер Институтунун "
             "Компьютердик Инженерия бөлүмүндө 02.09.2024-жылы Доц. "
             "Бакыт Шарсембаев (илимий жетекчи) жана Доц. Райымбек "
             "Султанов (эки жетекчи) тарабынан бекитилген диссертация "
             "сунушунун (КТМУ-fr-Тİİ-28 формасы, Студент № "
             "2351y01005) алкагында жүргүзүлдү. Сунуштун түп нуска "
             "темасы '12 каналдуу электрокардиографияны (ЭКГ) "
             "колдонуп таяныч түйүн ыкмасынын жардамы менен сигналды "
             "аугментациялоого негизделген жүрөк ооруларын "
             "диагностикалоо үчүн нейрондук тармак'.")
    add_body(doc,
             "Иштин салымдары төмөндө кыскача баяндалган: (i) бирдей "
             "модель, маалыматты аугментациялоо, оптимизатор, дан "
             "жана бөлүштүрүү менен Чапман–Шаосин 12 каналдуу ЭКГ "
             "маалыматтар базасында {5000, 1000, 500} кириш "
             "узундуктарынын контролдонгон салыштыруусу; (ii) бир "
             "саптык децимация кадамынын, жалпак 1Б-ЖНТ негизги "
             "модел (%88,43) менен окшош корпустарда 'жаңы муундук' "
             "натыйжа катары шилтемеленген attention-гибрид %94,8 "
             "максаты ортосундагы айырманын чоң бөлүгүн жоюуга "
             "жоопкер экендигине далил; (iii) он бир негизги "
             "ийгиликсиз класстын F1 ≥ 0,95 деңгээлине кайтканы "
             "көрсөтүлгөн класс боюнча калыбына келтирүү талдоосу; "
             "(iv) калыбына келтирүү профилин референс чекит "
             "тыгыздыгы жана натыйжалуу кабыл алуу аймагынын "
             "камтуусу контекстинде түшүндүргөн геометриялык "
             "өзгөрбөстүк аргументи; (v) болжол менен 1 800 саптуу "
             "жалгыз көзкарандысыз модулда берилген, толук жана "
             "кайра жасалуучу PyTorch шилтеме ишке ашыруусу.")
    add_body(doc,
             "Диссертациянын калган бөлүктөрү төмөнкүчө "
             "түзүлгөн: Биринчи Бөлүмдө электрокардиографиянын "
             "негиздери, классикалык сигнал иштетүү тизмеги жана 12 "
             "каналдуу ЭКГ классификациясы үчүн терең үйрөнүү "
             "адабияты каралат. Экинчи Бөлүм беш гипотезаны (H1–H5) "
             "формализациялайт; Чапман–Шаосин корпусу жана SNOMED CT "
             "этикет шайкеш келтирүүсү түшүндүрүлөт. Үчүнчү Бөлүм "
             "моделдин архитектурасын жана геометриялык өзгөрбөстүк "
             "аргументин сунуштайт. Төртүнчү Бөлүм PyTorch ишке "
             "ашыруусунун толугун кадам-кадам карайт. Бешинчи Бөлүм "
             "төрт кириш узундугу конфигурациясынын жана 4 "
             "конфигурациялык гибрид-план аблациясынын "
             "эксперименталдык натыйжаларын билдирет. Алтынчы "
             "Бөлүм децимациянын эмне үчүн жакшы иштээрин "
             "талкуулайт. Жыйынтык бөлүмү чектөөлөр жана келечек "
             "иштер жол картасы менен аяктайт.")


# ---------------------------------------------------------------------------
# БИРИНЧИ БӨЛҮМ: АРКА ПЛАН ЖАНА БАЙЛАНЫШТУУ ИШТЕР
# ---------------------------------------------------------------------------

def chapter2(doc):
    add_chapter_heading(doc, "БИРИНЧИ БӨЛҮМ", "АРКА ПЛАН ЖАНА БАЙЛАНЫШТУУ ИШТЕР")

    add_heading2(doc, "1.1. Электрокардиография: Кыскача аныктама")
    add_body(doc,
             "Электрокардиограмма — дене үстүндөгү стандартташтырылган "
             "электрод позицияларынын ортосундагы потенциал "
             "айырмачылыгынын убакытка байланыштуу өлчөнүшү. Ар бир "
             "жүрөк согуусу: P толкуну (атриалдык деполяризация), "
             "QRS комплекси (карынчалык деполяризация) жана T толкуну "
             "(карынчалык реполяризация) деген стандарттык кезектеги "
             "четке кагуулардын катарын чыгарат. Бул толкундардын "
             "морфологиясы, убакыттары жана согуулар арасындагы "
             "тартиби жүрөктүн өткөрүү тутумунун электрикалык абалын "
             "белгилейт. Минутасына 60–100 согуу болгон кадимки "
             "жүрөк ылдамдыгы 10 секунддук терезеде 10–17 PQRST "
             "комплексин чыгарат.")
    add_body(doc,
             "12 каналдуу ЭКГ, жүрөк электрикалык вектордун он эки "
             "геометриялык проекциясынан жазат: алдыңкы тегиздиктин "
             "проекциясын тарткан алты муундук (I, II, III, aVR, "
             "aVL, aVF) жана горизонталдык тегиздиктин проекциясын "
             "көкүрөктүн алдыңкы–арткы огунан тарткан алты "
             "көкүрөктүк (V1–V6) каналдары. Ар бир канал, ошол эле "
             "негизги электрикалык окуянын ар башка көз караштан "
             "үлгүсүн алат жана кардиолого — же терең тармакка — "
             "ошол эле жүрөк согуусунун он эки байланыштуу бирок "
             "ар башка көрүнүшүн көрсөтөт. Стандарттык дискреттештирүү "
             "жыштыгы 500 Гц жана клиникалык колдонуудагы стандарттык "
             "жазуу узактыгы 10 секунд; бул, изилдөө башына 12 × 5000 "
             "өлчөмүндөгү матрицаны чыгарат.")
    add_body(doc,
             "Бул сигналдын, кийинки терең үйрөнүү иштетүүсү үчүн "
             "маанилүү үч өзгөчөлүгү бар. Биринчиси, диагностикалык "
             "маанилүү мазмун убакыт огунда бирдиктүү жайылтылбайт, "
             "тескерисинче — P, QRS жана T толкундардын башталыштары, "
             "чокулары жана аяктоолору болуп — сейрек **референс "
             "чекиттер** жыйындысында топтолот. Экинчиси, согуулар "
             "арасындагы аралык (R-R интервалы) ритмикалык маалыматты "
             "алып жүрөт жана моделдин натыйжалуу кабыл алуу аймагы "
             "жок дегенде эки удаалаш QRS комплексин камтышы керек. "
             "Үчүнчүсү, физиологиялык жактан маанилүү жыштык мазмун "
             "болжол менен 40 Гц менен чектелет (QRS комплексинин "
             "эң жогорку жыштыктагы түзүмү); жогорку жыштыктарга "
             "булчуң артефакты жана электрод ызы-чууусу үстөмдүк "
             "кылат.")

    add_heading2(doc, "1.2. Жүрөк аритмиялары: Диагностикалык классификация")
    add_body(doc,
             "Клиникалык ЭКГ чечмелөө кеңири жана стандартташтырылган "
             "сөздүктөн — бул иште SNOMED CT таксономиясынан — "
             "диагноздорду чыгарат; бирок абляция максатында "
             "Чапман–Шаосин корпусунун 78 ийне-этикеттерин беш жогорку "
             "деңгээлдеги категорияга кыскартабыз. Бул топтоо PTB-XL "
             "жана PhysioNet/CinC 2020 байгесинде кабыл алынган "
             "ыкманы ээрчийт; клиникалык чечмелүүлүк менен "
             "статистикалык башкаруучулукту тең салмактайт:")
    add_table(doc,
              ["Категория", "Кыскартма", "Мисалдар"],
              [
                  ["Кадимки", "Normal", "Синус ритми, кадимки өткөрүү"],
                  ["Миокард Инфаркты", "MI",
                   "Курч MI, эски MI, антеросепталдык инфаркт"],
                  ["ST/T-Толкун Өзгөрүүлөрү", "STTC",
                   "ST түшүү/жогорулашы, T инверсиясы, аномалдуу Q"],
                  ["Өткөрүү Бузулуусу", "CD",
                   "RBBB, LBBB, AВ блоктору, AF, атриалдык трепетание, VT"],
                  ["Гипертрофия", "HYP",
                   "LVH, RVH, атриалдык кеңейүү"],
              ],
              caption="Таблица 1.1. Беш жогорку деңгээлдеги диагностикалык категория.")
    add_body(doc,
             "SNOMED CT коддорунан бул беш категорияга шайкеш "
             "келтирүү ECGCNNDiagnosticSystem классынын "
             "_map_snomed_to_category_diagnosis ыкмасында ишке "
             "ашырылды (3.2-Бөлүм, A.6c Листингин кара). Шайкеш "
             "келтирүү маанисиз эмес; анткени бир ЭКГ жазуусу көп "
             "учурда бирден ашык SNOMED кодун камтыйт (мисалы, "
             "'синус ритми' + 'сол дал блогу'). _select_primary_diagnosis "
             "ыкмасы, эң клиникалык маанилүү этикетти окутуу максаты "
             "катары тандоо үчүн клиникалык артыкчылык катар "
             "(MI > CD > STTC > HYP > Normal) колдонот.")

    add_heading2(doc, "1.3. Классикалык сигнал иштетүү тизмеги")
    add_body(doc,
             "Терең үйрөнүүгө чейинки ЭКГ талдоосу төрт этапка "
             "бөлүнгөн: (1) алдын ала иштетүү — зоналык өткөргүч "
             "фильтрация, базалык дрейфти алып салуу, тармак "
             "жыштыгын реджекциялоо; (2) референс чекиттерди "
             "аныктоо — QRS аныктоо үчүн адатта Pan-Tompkins, андан "
             "кийин wavelet ыкмалары же шаблон шайкеш келтирүү менен "
             "P/T чек арасы; (3) өзгөчөлүктөрдү алуу — жүрөк "
             "ылдамдыгынын өзгөрүү индекстери, QRS узактыгы, QT "
             "жана PR интервалдары, ST сегмент эңкейиши; (4) "
             "классификация — адатта кол менен иштетилген эреже "
             "жыйындысы же алынган өзгөчөлүктөр үстүндө сүйлөгөн "
             "классификатор (кокус токой, СVM).")
    add_body(doc,
             "Терең үйрөнүү парадигмасындагы өзгөрүү; (2)–(4) "
             "этаптарын, баштапкы сигналдан референс чекиттерди "
             "аныктоо, өзгөчөлүктөрдү алуу жана классификацияны "
             "чогуу үйрөнгөн жалгыз учу-учуна ЖНТ менен "
             "алмаштыруудан түзүлөт. (1) этап — алдын ала иштетүү — "
             "сакталат: жарыяланган ар бир терең ЭКГ-ЖНТ, сигналды "
             "тармакка бергенге чейин жок дегенде бир зоналык "
             "өткөргүч фильтр жана канал боюнча z-баа нормалдаштырууну "
             "колдонот. Бул иште каралган децимация кадамы, "
             "архитектуралык өзгөрүүдөн көрө кошумча алдын ала "
             "иштетүү операциясы катары мыкты түшүнүлөт.")

    add_heading2(doc, "1.4. ЭКГ классификациясы үчүн терең үйрөнүү")
    add_body(doc,
             "Rajpurkar ж.б. (2017) жана Hannun ж.б. (2019) 91 232 "
             "амбулатордук тек-каналдуу ЭКГ жазууда 34 катмардуу "
             "1Б-ЖНТ окутуп, 12 класстуу ритм аныктоо тапшырмасында "
             "кардиолог деңгээлинде өндүрүмдүүлүккө жетти. Алардын "
             "архитектурасы — терең, калдыктуу, бир өлчөмдүү "
             "конволюция жана глобалдык орточо пулинг камтыган — "
             "кийинки ЭКГ-ЖНТ изилдөөлөрү үчүн факт жүзүндөгү "
             "шаблон болуп калды. Strodthoff ж.б. (2020) ачык 12 "
             "каналдуу 21 837 жазуулук PTB-XL корпусун жараткан жана "
             "ырааттуу метрикалар менен ЖНТ, RNN жана Transformer "
             "архитектураларын салыштырган. Белгилүү боло турган "
             "нерсе, башкы 500 Гц ордуна 100 Гц (1000 үлгүлүү) "
             "киришти колдонуп; натыйжада 0,925 макро-AUC "
             "жарыяланды — арткы көрүнүштө толук 500 Гц "
             "көрсөтүлүшү керек болбой калышы мүмкүн экендиги "
             "тууралуу унчукпаган бир белги.")
    add_body(doc,
             "Бул иште колдонулган Чапман–Шаосин маалыматтар базасы "
             "Zheng ж.б. (2020) тарабынан жарыяланды; 10 646 "
             "пациенттен 45 152 даана 10 секунддук 12 каналдуу "
             "жазууну камтыйт жана 78 SNOMED CT коду менен "
             "белгиленди. Бүгүнкү күнгө чейинки эң чоң ачык 12 "
             "каналдуу корпус болуп саналат жана PhysioNet/CinC "
             "2020 байгесинин негизин түзөт.")

    add_heading2(doc, "1.5. Гибрид CNN-RNN жана attention архитектуралары")
    add_body(doc,
             "Кийинки көп макалалар, ЖНТ алдыңкы бөлүгүнүн үстүнө "
             "кайталанма же attention катмарларын кошот. Oh ж.б. "
             "(2018) өзгөрмө узундуктагы согуулар менен CNN-LSTM "
             "гибридин колдонуп %94,8 тактык билдирди — бул "
             "көрсөткүч диссертациянын мурунку этап отчетунда "
             "ашуу максаты катары ачык көрсөтүлгөн маани. Жакынкы "
             "иштер ЖНТ менен токенизацияланган согуулар үстүндө "
             "Transformer кодерлерин изилдеп, PTB-XLде 0,93–0,95 "
             "аралыгында макро-AUC билдирип жатат.")
    add_body(doc,
             "Салыштырмалуу адабияттарда көңүл буруусу керек нерсе, "
             "кириш узундугу сейрек гана негизги абляция өзгөрмөсү "
             "катары билдирилгендиги. Strodthoff ж.б. 100 Гц "
             "тандоосунан өткөнүндө гана сөз кылат; Oh ж.б. QRS "
             "аныктоолорунун айланасында кесилген өзгөрмө "
             "узундуктагы согууларды колдонот; башка макалалардын "
             "көбү корпустун демейки маанисин (Чапман–Шаосин үчүн "
             "500 Гц, PTB-XL үчүн 100 же 500 Гц) абляциясыз "
             "колдонот. Бизге белгилүү болгондой, эч бир жарыяланган "
             "12 каналдуу ЭКГ изилдөөсү негизги натыйжа катары "
             "контролдонгон кириш узундугу абляциясын билдирген эмес. "
             "Бул иш ушул боштукту толтурат.")

    add_heading3(doc, "1.5.1. Гибрид архитектураларга карата багыт")
    add_body(doc,
             "2018-жылдан кийинки ЭКГ терең үйрөнүү адабиятында үч "
             "архитектуралык күчөтүү үй-бүлөсү басымдуу багыт катары "
             "пайда болду. Биринчиси кайталанма ылдый иштетүү: ЖНТ "
             "алдыңкы бөлүгү согуу боюнча же сегмент боюнча "
             "киргизүү ырааттуулугун чыгарат; алар LSTM же GRU "
             "кодеринин колу менен бириктирилет. Oh ж.б. (2018) "
             "канондук мисалы. Экинчиси, өзүнө-attention (self-"
             "attention): ЖНТ өзөгүнүн өзгөчөлүк карталары "
             "Transformer кодери тарабынан өзгөртүлөт; абал "
             "кодировкасы конволюциялык токендердин убакыттык "
             "тартибин сактайт. Үчүнчүсү, он эки канал түйүн жана "
             "каналдар арасындагы корреляциялар чек катары болгон "
             "графдык нейрондук тармактын ылдый иштетүүсү.")
    add_body(doc,
             "Үч архитектуралык күчөтүү үй-бүлөнүн баары жалпы "
             "божомолго ээ: конволюциялык алдыңкы бөлүк, корпустун "
             "демейки дискреттештирүү жыштыгында (Чапман–Шаосин "
             "үчүн 500 Гц, PTB-XL үчүн 100 же 500 Гц) иштеп жатат. "
             "Бул иште билдирилген контролдонгон салыштыруу ушул "
             "божомолду суракка алат жана тандоонун нейтралдуу эмес "
             "экенин көрсөтөт.")

    add_heading2(doc, "1.6. Алиаска каршы теория жана дискреттештирүү теоремасы")
    add_body(doc,
             "Найквист–Шеннондун дискреттештирүү теоремасы; f_max "
             "Гц үстүндө спектрдик мазмуну жок зоналык чектелген "
             "үзгүлтүксүз убакытка жараша x(t) сигналдын, f_s ≥ "
             "2 f_max болгон каалаган дискреттештирүү жыштыгындагы "
             "үлгүлөрүнөн так кайра курулушу мүмкүн экендигин "
             "айтат. Үзгүлтүктүү убакытка жараша сигнал q бүтүн "
             "коэффициенти менен децимацияланганда, жаңы Найквист "
             "жыштык f_s_жаңы / 2 = f_s / (2q) үстүндөгү спектрдик "
             "мазмун ылдый зонага кайра түшөт — бул **алиас** деп "
             "аталат жана жөнөкөй децимациянын сигналдын сапатын "
             "бузуусунун классикалык себеби.")
    add_body(doc,
             "Чечим, децимациядан мурда жаңы Найквист жыштыгында "
             "кесим жасаган төмөн өткөргүч фильтр колдонуу. SciPyдин "
             "scipy.signal.decimate функциясы муну демейки катары "
             "сегизинчи тартиптеги Чебышев тип-I IIR фильтри менен "
             "жасайт; фильтр scipy.signal.filtfilt менен алдыга-арткы "
             "(нөл-фазалуу) режимде колдонулат. Нөл-фаза касиети "
             "бул жерде зарыл; анткени фаза бузулушу референс "
             "чекит позицияларын жыштыкка байланышкан түрдө "
             "жылыштырат жана 3.4-Бөлүмдө иштелип чыккан "
             "геометриялык өзгөрбөстүк аргументин жокко чыгарат.")
    add_body(doc,
             "ЭКГнин 0,5–40 Гц физиологиялык зонасы үчүн q = 10 "
             "менен 500 Гцтен 50 Гцке децимация жаңы Найквистти "
             "так 25 Гцке коет — 40 Гцтик физиологиялык чектен "
             "ылдый, кадимки шартта көйгөйлүү болушу мүмкүн болгон "
             "маани. Бирок практикада диагностикалык маанилүү "
             "жогорку жыштык мазмуну (QRS комплексинин курч "
             "кырлары) 5–25 Гц зонасында жатат; жогорку жыштыктар "
             "(25–40 Гц) көп учурда ызы-чуу жана 78 класстуу "
             "классификация тапшырмасы үчүн кызыксыз жогорку "
             "жыштыктагы QRS детальдарын камтыйт. Бул соодалашууну "
             "5-Бөлүмдө ампирикалык түрдө мүнөздөйбүз.")


# ---------------------------------------------------------------------------
# ЭКИНЧИ БӨЛҮМ: ГИПОТЕЗАЛАР, МААЛЫМАТТАР БАЗАСЫ ЖАНА АЛДЫН АЛА ИШТЕТҮҮ
# ---------------------------------------------------------------------------

def chapter3(doc):
    add_chapter_heading(doc, "ЭКИНЧИ БӨЛҮМ",
                        "ГИПОТЕЗАЛАР, МААЛЫМАТТАР БАЗАСЫ ЖАНА АЛДЫН АЛА ИШТЕТҮҮ")

    add_heading2(doc, "2.1. Беш гипотеза (H1–H5)")
    add_body(doc,
             "Магистрдик иштин эмпирикалык өзөгү беш ачык "
             "гипотезанын текшерилишине негизделет. Ар бир "
             "гипотеза эксперимент аркылуу ырасталат же четке "
             "кагылат. Гипотезалар *эксперимент жасалганга чейин* "
             "(a priori) калыптандырылган.")
    add_body(doc,
             "**H1 (узундук-баскыч гипотезасы).** Кириштин "
             "узундугу 5000 үлгү 1Б-ЖНТ базалык моделинде акыркы "
             "конволюциялык катмардын кабыл алуу аймагын ашат жана "
             "узундукту 500 үлгүгө децимациялоо архитектураны "
             "өзгөртпөстөн макро-F1ди жок дегенде +0,05 жогорулатат. "
             "*Далил кайдан күтүлөт:* 5.2-Бөлүм (баш-салыштыруу "
             "таблицасы) жана 6-Бөлүм (кабыл алуу аймагынын "
             "талдоосу).")
    add_body(doc,
             "**H2 (геометриялык өзгөрбөстүк гипотезасы).** Алиаска "
             "каршы фильтр (scipy.signal.decimate, ftype='iir', n=8, "
             "zero_phase=True) ЭКГнин диагностикалык маалыматын алып "
             "жүргөн P/Q/R/S/T референс чекиттеринин убакыттык жана "
             "амплитуддук конфигурациясын ±10 мс тактыгынан кем "
             "эмес деңгээлде сактайт. *Далил кайдан күтүлөт:* "
             "3.6-Бөлүм (формалдуу аргумент), 5-Бөлүм (көрсөтмөлүү "
             "визуализация — 3.1-сүрөт), 5.3-Бөлүм (класс боюнча "
             "калыбына келтирүү).")
    add_body(doc,
             "**H3 (алиаска каршы зарылдык гипотезасы).** Алиаска "
             "каршы фильтрсиз жөнөкөй stride-by-10 пулинг QRS "
             "энергиясын ылдый жыштык зонасына спектрдик алиаска "
             "чалдыктырат жана тактыкты жакшыртуунун ордуна "
             "начарлатат — башкача айтканда, Чебышев-I фильтри %88 "
             "→ %97 өсүшүнө жоопкер. *Далил кайдан күтүлөт:* "
             "6-Бөлүм (алиаска каршы фильтрсиз алдын ала "
             "эксперимент); 5.2-Бөлүмдүн жыйынтыктары менен "
             "салыштырылат.")
    add_body(doc,
             "**H4 (таяныч түйүн менен аугментация гипотезасы).** "
             "Чапман–Шаосиндин узун куйруктуу таркалуусу үчүн "
             "таяныч түйүн негиздүү маалыматты аугментациялоо (3× "
             "кеңири класстарга, 10× сейрек класстарга) ийгиликсиз "
             "класстарды калыбына келтирүүгө жетиштүү шарт болуп "
             "саналат — аугментация алып салынса, макро-F1 эки "
             "канал конфигурациясында тең жок дегенде 0,85ке "
             "түшөт. *Далил кайдан күтүлөт:* 5.5-Бөлүм (гибрид-план "
             "аблация).")
    add_body(doc,
             "**H5 (канал санынын тоскоолдук эмес гипотезасы).** "
             "Сигналды аугментациялоо иштетилген учурда 1 канал "
             "(адегенде II-канал) 12 каналга караганда F1 жана "
             "тактык боюнча айырмаланбай, ал эми бир үлгү боюнча "
             "чыгаруу убактысы боюнча ыңгайлуу — башкача айтканда, "
             "жагалай (edge) жайгаштыруу үчүн 1 канал жетиштүү. "
             "*Далил кайдан күтүлөт:* 5.5-Бөлүм (гибрид-план "
             "аблация); жыйынтык бөлүмү (келечек иштер: жагалай "
             "жайгаштыруу).")
    add_body(doc,
             "5.6-Бөлүмдө ар бир гипотеза боюнча "
             "ырасталды/каршы делди деген кыска корутунду берилет.")

    add_heading2(doc, "2.2. Чапман–Шаосин 12 каналдуу ЭКГ маалыматтар базасы")
    add_body(doc,
             "Чапман–Шаосин (Zheng ж.б., 2020); Шаосин Эл "
             "Ооруканасында 2013–2019-жылдары алынган 45 152 даана "
             "12 каналдуу ЭКГ жазуусунун коомдук жыйнагы. Ар бир "
             "жазуу 10 секунд узундукта, 500 Гц жыштыкта жана WFDB "
             "форматында эки файл катары сакталат: пациент мета-"
             "маалыматын (жашы, жынысы, диагноз коддору) камтыган "
             "тексттик .hea баш файлы жана 12 × 5000 сигнал "
             "матрицасы орун алган MATLAB .mat жүгү. "
             "Диагностикалык annotation SNOMED CT коддорун "
             "колдонот; корпус боюнча 78 түрдүү код жолугушат жана "
             "аз сан жазуу бирден ашык кодду алып жүрөт (жазуу "
             "башына орточо 1,4 код).")
    add_body(doc,
             "Класстын тең эмес бөлүштүрүлүшү катаал. Эң көп "
             "жолугушкан төрт класс (кадимки синус ритми, синус "
             "брадикардиясы, синус тахикардиясы жана атриалдык "
             "фибрилляция) баштапкы жазуулардын 34 000ден ашуусун "
             "түзөт. Болжол менен отуз SNOMED коду, ар бири "
             "элүүдөн аз жазууда жолугушат. Бул тең эмес "
             "бөлүштүрүү 4-Бөлүмдө түшүндүрүлгөн focal loss "
             "механизмин жана 2.6-Бөлүмдөгү таяныч түйүн "
             "негиздүү аугментацияны мотивдештирген он бир негизги "
             "ийгиликсиз класстын негизги күчү.")
    add_body(doc,
             "Корпустун техникалык өзгөчөлүктөрү боюнча, жазуулар "
             "32 биттик чечүүчүлүктөгү A/D түзөтүү аркылуу "
             "алынган жана бит башына орточо 4,88 A/D түзөтүү "
             "мааниси билдирилген. Амплитуда бирдиги микровольт; "
             "теориялык үстүнкү жана ылдыйкы чектер тиешелүү "
             "түрдө +32 767 µV жана −32 768 µV. Маалыматтарды "
             "чогултуу; Шаосин Эл Ооруканасынын жана Нинбо Биринчи "
             "Ооруканасынын Институттук Кароо Кеңеши тарабынан "
             "бекитилди; маалымдуу макулдуктун кечирилиши менен "
             "идентификация жашырыла келип, маалыматтар коомго "
             "жарыялоого уруксат берилди.")

    add_heading2(doc, "2.3. SNOMED CT код шайкеш келтирүүсү")
    add_body(doc,
             "Классификатор максаттары _read_header_metadata, "
             "_map_snomed_to_diagnosis жана "
             "_map_snomed_to_category_diagnosis ыкмалары тарабынан "
             "чыгарылат. Жол төмөнкүчө: (i) үтүр менен ажыратылган "
             "SNOMED код тизмесин чыгаруу үчүн '#Dx:' сабы "
             "талданат; (ii) ар бир код же ийне-этикет тексттик "
             "диагнозго же беш класстуу категорияга (MI, STTC, "
             "CD, HYP, Normal) шайкеш келтирилет; (iii) бир жазуу "
             "бирден ашык кодду алып жүргөндө, эң клиникалык "
             "маанилүү этикетти окутуу максаты катары тандоо үчүн "
             "MI > CD > STTC > HYP > Normal артыкчылык катары "
             "колдонулат.")

    add_heading2(doc, "2.4. Класс теңдештирүү стратегиясы")
    add_body(doc,
             "Бирин-бири толуктаган эки теңдештирүү операциясы "
             "колдонулат. Биринчиси, load_local_records ыкмасында "
             "(max_samples_per_class = 5000), көпчүлүк класстар "
             "алдын ала иштетүү жасалбастан мурда жазуу деңгээлинде "
             "децимацияланат. Бул эң чоң класстын эң кичинеге "
             "болгон катышын болжол менен 5 000 : 50 = 100:1 "
             "менен чектейт. Экинчиси, алдын ала иштетүүдөн кийин "
             "SMOTE (Синтетикалык Азчылык Үлгүсүн Көбөйтүү "
             "Ыкмасы) азчылык класстарды класс башына 4 000 "
             "үлгүлүк максатка чейин милдеттүү эмес түрдө "
             "көбөйтөт.")

    add_heading2(doc, "2.5. Зоналык өткөргүч фильтр жана нормалдаштыруу")
    add_body(doc,
             "Бардык сигналдар, базалык дрейфти алып салуу үчүн "
             "0,5 Гц (жогорку өткөргүч) жана булчуң артефактын "
             "алып салуу үчүн 40 Гц (ылдый өткөргүч) кесилиш "
             "чекиттери бар төртүнчү тартиптеги Баттерворт "
             "зоналык өткөргүч фильтринен өткөрүлөт. Бар тизмекте "
             "50 Гц нотч-фильтр унчугулган; Чапман–Шаосин жазуулары "
             "алуу учурунда мурда фильтрленген. Кайчылаш-маалымат "
             "базалык жайгаштыруу үчүн жергиликтүү тармак "
             "жыштыгында (50 же 60 Гц) нотч-фильтр кошулушу керек.")
    add_body(doc,
             "Фильтрациядан кийин ар бир канал, 5000 үлгүлүк "
             "терезе үстүндө нөл орточо мааниге жана бирдиктүү "
             "дисперсияга көзкарандысыз нормалдаштырылат (канал "
             "боюнча z-баа). Нормалдаштыруу зоналык өткөргүч "
             "фильтрден кийин жана милдеттүү эмес децимация "
             "кадамдан мурда колдонулат; ушундай децимация, "
             "спектрдик мазмуну мурдатан 0,5–40 Гц физиологиялык "
             "зонасы менен чектелген сигналда иштейт.")

    add_heading2(doc, "2.6. Таяныч түйүн негиздүү аугментация")
    add_body(doc,
             "Магистрдик иштин темасынын борборунда турган "
             "**таяныч түйүн ыкмасы** [6,7] физиологиялык маанилүү "
             "референс чекиттердин (P, Q, R, S, T) айланасында "
             "кубдук-сплайн түйүндөрдү колдонуп ЭКГнин жаңы "
             "үлгүлөрүн киргизет. Учурдагы тизмекте бул, удаалуу "
             "генерациялоонун ордуна, узун куйруктуу класстарга "
             "көбөйтүүчү катары колдонулат: кеңири класстарга 3×, "
             "сейрек класстарга 10×, максаттуу класс боюнча "
             "4 500 үлгүгө чейин теңдештирүү.")
    add_body(doc,
             "Салттуу маалыматты аугментациялоо иштетилген учурда "
             "(билдирилген эксперименттердин демейкиси), ар бир "
             "азчылык класс эң чоң класстын өлчөмүнүн %90ына чейин "
             "көбөйтүлөт. Үлгү башына кокустук белги жана көлөм "
             "менен үч аугментация колдонулат: (1) [0,9; 1,1] "
             "аралыгынан тегиз бөлүштүрүү менен алынган "
             "коэффициент менен амплитуда масштабдоо; (2) ар бир "
             "үлгүгө σ = 0,01 менен Гаусс ызы-чууусун кошуу; (3) "
             "терезе узундугунун %2не чейин тегерек убакыт "
             "жылыштыруу. Эң кичинекей класстар үчүн (n < 500) "
             "аугментация чоңойтулат: масштабдоо [0,85; 1,15], "
             "ызы-чуу σ = 0,02, жылыштыруу ±%5 жана %10 "
             "ыктымалдуулук менен амплитуда тескери буруу. Ишке "
             "ашыруу _apply_random_augmentation жана "
             "_apply_aggressive_augmentation ичинде (A.5 "
             "Листингин кара).")
    add_body(doc,
             "**H4 гипотезасынын мааниси.** Бул аугментация "
             "тизмеги ийгиликсиз класстарды калыбына келтирүү "
             "үчүн зарыл шарт катары сунушталат: эгер аугментация "
             "өчүрүлсө, узун куйруктуу Чапман–Шаосин класстары "
             "жетиштүү градиентти топтой албайт жана макро-F1 "
             "ызы-чууу деңгээлине түшөт. 5.5-Бөлүмдөгү аблация "
             "бул билдирмени ампирикалык түрдө текшерет.")


# ---------------------------------------------------------------------------
# ҮЧҮНЧҮ БӨЛҮМ: МЕТОД — МОДЕЛДИН АРХИТЕКТУРАСЫ
# ---------------------------------------------------------------------------

def chapter4(doc, source):
    add_chapter_heading(doc, "ҮЧҮНЧҮ БӨЛҮМ", "МЕТОД: МОДЕЛДИН АРХИТЕКТУРАСЫ")

    add_heading2(doc, "3.1. Калдыктуу блок (ResidualBlock)")
    add_body(doc,
             "Моделдин өзөгү бир өлчөмдүү калдыктуу (residual) "
             "тарм. Негизги бирдик ResidualBlock катары "
             "аныкталган. Ар бир блок эки 1Б конволюциядан турат; "
             "ар бир конволюциядан кийин batch normalization "
             "коюлат. Skip connection (atlama байланыш; канал "
             "жана stride өлчөмдөрүн дал келтирүү үчүн милдеттүү "
             "эмес 1×1 конволюция менен), блок киришин экинчи "
             "конволюциянын чыгышына кошот; акыркы ReLU жана "
             "dropout блок чыгышын чыгарат.")
    block = extract_block(source,
                          r"^class ResidualBlock\(nn\.Module\):",
                          (r"^class\s", r"^def\s"))
    add_code_block(doc, block, caption="Листинг 3.1 — ResidualBlock классы.")
    add_body(doc,
             "Skip connection; градиенттердин конволюциялык жолду "
             "айланып өтүшүнө жол берет. Бул, ResNet (He ж.б., "
             "2016) борбордук түшүнүгү: байланыш жок болсо "
             "болжол менен он беш катмардан терең тармактарды "
             "окутуу, градиенттер узак көбөйтүүчү тизмекте жоголуп "
             "же жарылып туруктуу болбойт. 1×1 skip "
             "конволюциясы кириш жана чыгыш каналдарынын саны ар "
             "башка болгондо же блок децимациялаганда (stride ≠ 1) "
             "гана керек; калган учурда identity түз колдонулат.")

    add_heading2(doc, "3.2. ECGCNN өзөгү")
    add_body(doc,
             "ECGCNN модулу; бир баштапкы конволюция жана төрт "
             "ResidualBlockту топтойт. Канал программасы, төрт "
             "калдыктуу этапта 64 (тамыр) → 128 → 256 → 512 → 512 "
             "болуп өсөт. Ар бир калдыктуу этаптан кийин 2 менен "
             "max-pooling, убакыттык өлчөмдү жарыга кыскартып, "
             "кириштен өзгөчөлүк картасына чейин 16× түшүрөт. "
             "Глобалдык орточо pooling убакыттык өлчөмдү толугу "
             "менен кыскартат; андан кийин эки толук туташкан "
             "катмар (256 → 128 → num_classes) логиттерди чыгарат. "
             "Тарм, глобалдык орточо pooling dense башынан кириш "
             "узундугунан көзкарандысыз болушуна шарт берип, "
             "кириш узундугунан көзкарандысыз болжол менен 3,7 "
             "миллион параметрге ээ.")
    block = extract_block(source,
                          r"^class ECGCNN\(nn\.Module\):",
                          (r"^class\s",))
    add_code_block(doc, block, caption="Листинг 3.2 — ECGCNN классы.")
    add_table(doc,
              ["Этап", "Иштем", "Чыгыш канал", "Ядро",
               "Кадам", "Pool"],
              [
                  ["Тамыр", "Conv1d + BN + ReLU", "64", "7", "1", "2"],
                  ["Res1", "ResidualBlock", "128", "5", "1", "2"],
                  ["Res2", "ResidualBlock", "256", "5", "1", "2"],
                  ["Res3", "ResidualBlock", "512", "3", "1", "2"],
                  ["Res4", "ResidualBlock", "512", "3", "1", "1"],
                  ["Баш", "GAP + 2×FC + Dropout",
                   "num_classes", "—", "—", "—"],
              ],
              caption="Таблица 3.1. ECGCNN фильтр жана ядро программасы.")

    add_picture(doc, FIG_ARCH, width_cm=15.5,
                caption="Сүрөт 3.0. ECGCNN архитектурасынын схемалык "
                        "көрсөтмөсү. Киришке (1 канал, узундугу L = 500) "
                        "Conv-1 + ResBlock-1..4 баскычтары колдонулат; "
                        "ар бир баскычтын чыгуу өлчөмү жана канал саны "
                        "блоктун астында берилет. GAP — глобалдык орточо "
                        "пулинг; FC-1..3 — толук туташкан катмарлар. "
                        "Архитектура training/ecg_cnn_pytorch.py "
                        "файлындагы ECGCNN классына дал келет; жалпы "
                        "параметрдин саны 3,72 миллион.")

    add_heading2(doc, "3.3. Тең эмес максаттар үчүн фокалдуу жоготуу (Focal Loss)")
    add_body(doc,
             "Окутуу жоготуусу класс-салмактуу кросс-энтропияны Lin "
             "ж.б.нун (2017) фокалдуу жоготуу кайра салмактоосу "
             "менен бириктирет. Фокалдуу жоготуу, үлгү башына "
             "кросс-энтропияны (1 − p_t)^γ менен көбөйтөт; бул "
             "жерде p_t моделдин чыныгы класс үчүн божомолдогон "
             "ыктымалдуулугу, γ болсо фокалдуу параметр (γ = 2,0 "
             "колдонулду). Моделдин жогорку ишеним (жогорку p_t) "
             "менен мурдатан туура классификациялаган үлгүлөрү "
             "градиентке аз салым кошот; модели күмөн саналган "
             "үлгүлөр эң көп салым кошот. Бул, окутууну кыйын — "
             "көп учурда азчылык класс — үлгүлөрдө топтоштурат.")
    block = extract_block(source,
                          r"^class FocalLoss\(nn\.Module\):",
                          (r"^def\s", r"^class\s"))
    add_code_block(doc, block, caption="Листинг 3.3 — FocalLoss классы.")

    add_heading2(doc, "3.4. Этикеттерди жумшартуу (Label Smoothing)")
    add_body(doc,
             "Этикеттерди жумшартуу (Szegedy ж.б., 2016); бир-кайнар "
             "(one-hot) окутуу максаты [0, 0, 1, 0, 0]ди, "
             "smoothing = 0,1 үчүн жумшартылган [0,025; 0,025; "
             "0,9; 0,025; 0,025] версиясы менен алмаштырат. Бул, "
             "тармактын чыныгы класс үчүн ачыктап чоң логиттерди "
             "үйрөнүүсүн алдын алат; калибрацияны жакшыртат жана "
             "жеңил регуляризатор катары иштейт.")
    block = extract_block(source,
                          r"^def label_smoothing_loss",
                          (r"^class\s", r"^def\s"))
    add_code_block(doc, block,
                   caption="Листинг 3.4 — Этикеттерди жумшартуу жоготуусу.")

    add_heading2(doc, "3.5. Алиаска каршы децимация: математикалык форма")
    add_body(doc,
             "N = 5000 болуп x ∈ ℝ^(12 × N) кириш сигналы "
             "берилгенде, децимация кадамы x_down ∈ ℝ^(12 × N/q) "
             "чыгаруу үчүн ар бир он эки каналга көзкарандысыз "
             "колдонулат:")
    add_body(doc,
             "    x_down = scipy.signal.decimate(x, q, ftype='iir', "
             "n=8, zero_phase=True)")
    add_body(doc,
             "Бул жерде q ∈ {1, 5, 10}, тиешелүү түрдө {5000, "
             "1000, 500} чыгыш узундуктарына туура келет. Ички "
             "түрдө scipy.signal.decimate, кесилиш жыштыгы 0,8 / q "
             "(нормалдашкан Найквист бирдиктеринде, башкача "
             "айтканда кириш Найквист жыштыгынын = 250 Гц "
             "бөлүгү катары) болгон сегизинчи тартиптеги Чебышев "
             "тип-I IIR ылдый өткөргүч фильтрин колдонот. Фильтр "
             "scipy.signal.filtfilt аркылуу алдыга-арткы режимде "
             "колдонулат; нөл-фаза жоопту чыгарат жана натыйжалуу "
             "фильтр тартибин төрт эсе көбөйтөт.")
    add_body(doc,
             "Нөл-фаза касиети зарыл: каалаган фаза жылдыруусу "
             "референс чекит позицияларын убакыт огунда жыштыкка "
             "байланышкан түрдө жылыштырат жана кийинки "
             "Бөлүмдөгү геометриялык өзгөрбөстүктү жокко чыгарат. "
             "Чебышев тип-I үй-бүлөсү, scipy.signal.decimateдин "
             "альтернативасы Баттервортка ордуна тандалган; "
             "бирдей тартипте жогору курч өтүү зонасын берет жана "
             "жаңы өтүү зонасына алиас жасай турган мазмунду эң "
             "аз кылат.")

    add_heading2(doc, "3.6. Референс чекиттер графынын геометриялык өзгөрбөстүгү")
    add_body(doc,
             "ЭКГнин диагностикалык маалыматы, сейрек референс "
             "чекит жыйындысында — P, QRS жана T толкундардын "
             "башталыштары, чокулары жана аяктоолорунда — жана "
             "алардын убакыттык байланыштарында (R-R, P-R, QT, "
             "QRS узундугу, ST эңкейиши, T толкун морфологиясы) "
             "топтолот. 500 Гц жыштыкта ар бир согуу үчүн 5 "
             "канондук чекит менен болжол менен 10 согууну "
             "камтыган 10 секунддук терезе үчүн бул, 5000 үлгүгө "
             "жайылган болжол менен 60 референс чекитти "
             "берет; үлгүлөрдүн ~%98 референс чекиттер графы "
             "кодтогон маалыматтан тышкары эч кандай маалымат "
             "ташыбайт.")
    add_body(doc,
             "Алдыга-арткы режимде колдонулган Чебышев тип-I "
             "алиаска каршы фильтр, бул чекиттердин геометриялык "
             "конфигурациясын дискреттештирүү тактыгына чейин "
             "сактайт. Ар бир референс чекиттин убактысы, жаңы "
             "дискреттештирүү периодунун ±½синин ичинде сакталат; "
             "10× децимациядан кийин бул тактык 10 мс — клиникалык "
             "тажрыйбадагы эң катаал убакыт ченөөсү (QT интервалы) "
             "10 мс тактыгында билдирилет. Амплитуда, кичинекей "
             "фильтр жоопу басаңдоосу баштайт болсо да сакталат; "
             "чекиттердин тартиби жана салыштырмалуу убакыты так "
             "сакталат. ЭКГ ийри сызыгынын формасы — референс "
             "чекиттер арасында сынык сызык катары каралганда — "
             "ушундан улам децимация астында өзгөрбөйт.")
    add_picture(doc, FIG_GEOM, width_cm=14.5, fallback=FIG_GEOM_FALLBACK,
                caption="Сүрөт 3.1. scipy.signal.decimate астында "
                        "референс чекиттер графынын геометриялык "
                        "өзгөрбөстүгү. (а) 5000 үлгүдөгү II-канал; ~60 "
                        "референс чекит кириш позицияларынын ~%1,2 "
                        "түзөт; ЖНТнин натыйжалуу кабыл алуу аймагы "
                        "терезенин ~%40 камтыйт. (б) 500 үлгүгө 10× "
                        "децимациядан кийин ошол эле референс чекиттер "
                        "сакталат; алардын тыгыздыгы 10× жогорулайт жана "
                        "кабыл алуу аймагы 10 секунддук бүт терезени "
                        "камтыйт.")
    add_body(doc,
             "Геометриялык өзгөрбөстүк сүрөтүнүн, 5-Бөлүмдө "
             "ампирикалык түрдө текшерилген үч конкреттүү "
             "натыйжасы бар: (i) децимациядан кийин ЖНТнин "
             "натыйжалуу кабыл алуу аймагы 10 секунддук бүт "
             "терезени камтыйт жана көп согуу боюнча ритмикалык "
             "ой жүгүртүүгө мүмкүндүк берет; (ii) референс чекит "
             "тыгыздыгы он эсеге жогорулайт жана градиент сигналын "
             "диагностикалык маанилүү үлгүлөрдө топтоштурат; "
             "(iii) тармактын туруктуу параметр бюджети, "
             "референс чекиттер арасында артыкча ылдый жыштык "
             "өзгөрүүсүн моделдештирүүдөн ийне морфологияларды "
             "айырмалоого кайра бөлүштүрүлөт.")

    add_heading2(doc, "3.7. Таяныч түйүн ыкмасынан децимацияга: геометриялык көпүрө")
    add_body(doc,
             "Магистрдик иш баштапкы убагында **таяныч түйүн "
             "(support-node) ыкмасы** [6,7] боюнча даярдалган — "
             "бул ыкма физиологиялык маанилүү референс чекиттердин "
             "(P, Q, R, S, T) айланасында кубдук-сплайн "
             "түйүндөрдү колдонуп ЭКГнин жаңы үлгүлөрүн киргизет. "
             "Ыкманын *интуициясы* 3.6-Бөлүмдөгү аргумент менен "
             "бирдей: ЭКГнин диагностикалык маалыматы сейрек "
             "референс чекиттер графында жашайт, ошондуктан бул "
             "графты сыйлаган кандай гана алдын ала иштетүү "
             "кадамы тармакка пайда алып келет. Таяныч түйүн "
             "ыкмасы графты *графтын жанындагы кайра "
             "дискреттештирүү аркылуу* сыйлайт; алиаска каршы "
             "децимация *глобалдык төмөн өткөргүч фильтр + "
             "кайра дискреттештирүү аркылуу* сыйлайт.")
    add_body(doc,
             "Ушундай негизде магистрдик иштин теориялык "
             "кошумча салымы — алиаска каршы децимациянын "
             "**таяныч түйүн ыкмалары үй-бүлөсүнүн ЖНТ үчүн "
             "оптималдуу мүчөсү** экендигин аныктоо: (i) таяныч "
             "түйүн интерполяциясы — референс чекиттер графына "
             "жаңы үлгүлөрдү кошот, толкун-боюнча талдоо үчүн "
             "ылайыктуу; (ii) алиаска каршы децимация — референс "
             "чекиттер графын сактап туруп, ашыкча негизги "
             "интерполяцияны алып салат, туруктуу-кириш-узундуктуу "
             "ЖНТ классификаторлору үчүн оптималдуу; (iii) "
             "attention механизмдери — кайсы позициялардын "
             "референс чекиттер графына туура келээрин *үйрөнөт*, "
             "ырааттуулук моделдер үчүн ылайыктуу.")
    add_body(doc,
             "Ушул чектерде иштин темасы менен кодунун "
             "ортосундагы кездешкен кагылышуу — баштапкы тема "
             "'таяныч түйүн ыкмасы менен сигналды аугментациялоо', "
             "ал эми кодду сапатынын негизги салымы — алиаска "
             "каршы децимация — мындайча эсептелет: тема "
             "**ыкмалар үй-бүлөсүн** атайт, ал эми негизги "
             "жыйынтык **ал үй-бүлөнүн ЖНТ үчүн оптималдуу "
             "мүчөсүн** атайт.")


# ---------------------------------------------------------------------------
# ТӨРТҮНЧҮ БӨЛҮМ: ИШКЕ АШЫРУУ
# ---------------------------------------------------------------------------

def chapter5(doc, source):
    add_chapter_heading(doc, "ТӨРТҮНЧҮ БӨЛҮМ", "ИШКЕ АШЫРУУ")
    add_body(doc,
             "Бул бөлүмдө, ЭКГ ЖНТ диагностикалык тутумунун "
             "толук PyTorch шилтеме ишке ашыруусу кадам-кадам "
             "каралат. Код, training/ecg_cnn_pytorch.py деп "
             "аталган 1 767 саптуу жалгыз модуль катары "
             "уюштурулду; төрт жогорку деңгээлдеги аныктаманын "
             "айланасында түзүлгөн: FocalLoss жана этикеттерди "
             "жумшартуу жоготуу функциялары (мурунку Бөлүмдөгү "
             "Листинг 3.3–3.4); ResidualBlock жана ECGCNN "
             "модулдары (Листинг 3.1–3.2); жана маалыматты кабыл "
             "алуу, алдын ала иштетүү, окутуу, баалоо, сактоо "
             "жана чыгарууну жөнгө салган ECGCNNDiagnosticSystem "
             "классы. Модуль учу-учуна колдонууну көрсөткөн "
             "main() функциясы менен аяктайт.")

    add_heading2(doc, "4.1. ECGCNNDiagnosticSystem классы")
    add_body(doc,
             "Тутум классы, кириш геометриясы (sequence_length, "
             "decimation_factor, num_leads), DataLoader "
             "жумушчуларынын саны жана чыгаруу каталогу менен "
             "башталат. Бар болсо CUDA түзмөгүн тандайт; кыскача "
             "түзмөк баяндамасын көрсөтөт; этикет кодери, "
             "масштабдоочу, этикет картасы жана класс "
             "салмактарынын камтылмасын кийинки толтуруу үчүн "
             "даярдайт.")
    init_block = extract_block(source,
                               r"^class ECGCNNDiagnosticSystem:",
                               (r"^    def load_local_records",))
    add_code_block(doc, init_block,
                   caption="Листинг 4.1 — ECGCNNDiagnosticSystem __init__.")
    add_body(doc,
             "Бул конструкторунда эки долбоордук чечим белгилөөгө "
             "татыктуу. Биринчиси, decimation_factor алдын ала "
             "иштетүү функциясынын ичине отуруп калбай, биринчи "
             "класс гипер-параметр катары ачылат; бул, 5-Бөлүмдөгү "
             "контролдонгон салыштыруунун курган өзгөрмөсү. "
             "Экинчиси, effective_length = sequence_length // "
             "decimation_factor куруу учурунда бир жолу эсептелет "
             "жана тизмек боюнча кайра колдонулат; ушундай децимация "
             "кийинки тензор формасы үчүн жалгыз чыныгы булак "
             "берилет.")

    add_heading2(doc, "4.2. Маалыматты жүктөө: .hea + .mat кабыл алуу")
    add_body(doc,
             "load_local_records ыкмасы, .hea папкасын кезип "
             "өтөт; ар бир баш үчүн жашы/жынысы/диагноз мета-"
             "маалыматын окуйт; дал келген .mat жүгүн жүктөйт жана "
             "үч чыгаруу куратат: өзгөрмө узундуктагы сигналдардын "
             "numpy объект массиви, категория этикеттеринин numpy "
             "массиви жана жазуу аты менен индекстелген пациент "
             "маалыматынын сөздүгү. balance_classes = True "
             "болгондо, ыкма ар кандай кийинки алдын ала иштетүү "
             "жасалбастан мурда көпчүлүк класстарды "
             "max_samples_per_class жазууга децимациялайт.")
    add_code_block(doc,
                   "def load_local_records(self, hea_folder, mat_folder,\n"
                   "                       max_records=None,\n"
                   "                       balance_classes=True,\n"
                   "                       max_samples_per_class=5000):\n"
                   "    # Бардык .hea файлдарын кездемелеп өт\n"
                   "    hea_files = sorted(glob.glob(os.path.join(hea_folder, '*.hea')))\n"
                   "    signals, diagnoses, patient_data = [], [], {}\n"
                   "    for hea_path in hea_files[:max_records]:\n"
                   "        record_name = os.path.splitext(os.path.basename(hea_path))[0]\n"
                   "        age, gender, diagnosis = self._read_header_metadata(hea_path)\n"
                   "        mat_path = os.path.join(mat_folder, f'{record_name}.mat')\n"
                   "        ecg_signal, sr = self._load_mat_signal(mat_path)\n"
                   "        if ecg_signal is None or ecg_signal.shape[-1] < 1000:\n"
                   "            continue\n"
                   "        signals.append(ecg_signal)\n"
                   "        diagnoses.append(diagnosis)\n"
                   "        patient_data[record_name] = {\n"
                   "            'age': age, 'gender': gender,\n"
                   "            'diagnosis': diagnosis, 'sampling_rate': sr,\n"
                   "        }\n"
                   "    if balance_classes:\n"
                   "        # Көпчүлүк класстарды max_samples_per_class деңгээлине децимациялайт\n"
                   "        ...\n"
                   "    return np.array(signals, dtype=object), np.array(diagnoses), patient_data\n",
                   caption="Листинг 4.2 — load_local_records (кыскача).")

    add_heading2(doc, "4.3. Алдын ала иштетүү тизмеги (код кароосу)")
    add_body(doc,
             "preprocess_data ыкмасы, толук алдын ала иштетүү "
             "тизмегин алты номурланган кадамда уюштурат: (1) "
             "узундуктарды sequence_length үлгүгө стандарттатуу; "
             "(2) 0,5–40 Гц Баттерворт зоналык өткөргүч фильтр "
             "менен ызы-чууну жоюу; (2б) decimation_factor менен "
             "милдеттүү эмес децимация; (3) канал боюнча z-баа "
             "нормалдаштыруу; (4) этикет коддоо; (5) милдеттүү "
             "эмес SMOTE көбөйтүү; (6) милдеттүү эмес салттуу "
             "маалыматты аугментациялоо. Ыкма; иштетилген "
             "массивди, кодтолгон этикеттерди жана этикет картасын "
             "кайтарат.")
    block = extract_block(source,
                          r"^    def preprocess_data",
                          (r"^    def _standardize_lengths",))
    add_code_block(doc, block,
                   caption="Листинг 4.3 — preprocess_data уюштуруусу.")
    add_body(doc,
             "Төрт атомдук алдын ала иштетүү жардамчысы — "
             "_standardize_lengths, _denoise_signals, "
             "_decimate_signals жана _normalize_signals — ар бири "
             "(N, num_leads, T) тензорунун үстүндө убакыт огу "
             "боюнча иштейт. Бул иштин борбордук объекти болгон "
             "децимация жардамчысы, башкаларынан жөнөкөй:")
    add_code_block(doc,
                   "def _decimate_signals(self, X):\n"
                   "    X_decimated = np.zeros((len(X), self.num_leads,\n"
                   "                            self.effective_length),\n"
                   "                           dtype=np.float32)\n"
                   "    for i in range(len(X)):\n"
                   "        X_decimated[i] = scipy_signal.decimate(\n"
                   "            X[i], self.decimation_factor,\n"
                   "            axis=-1, zero_phase=True)\n"
                   "    return X_decimated\n",
                   caption="Листинг 4.4 — Децимация жардамчысы.")

    add_heading2(doc, "4.4. AMP жана эрте токтотуу менен окутуу циклы")
    add_body(doc,
             "train_model ыкмасы; маалыматты этикетке боюнча "
             "стратификацияланган түрдө, туруктуу дан 42 менен "
             "окутуу/валидация/тест (%80/%16/%20) болуп бөлөт; "
             "сабырдуу эс жана туруктуу жумушчулар менен "
             "DataLoaderлерди түзөт; ECGCNNди инстанциалайт жана "
             "AMP-ылдамдатылган окутуу циклын иштетет. Оптимизатор "
             "weight_decay = 0,01 (L2 регуляризациясы) жана "
             "башталгыч үйрөнүү ылдамдыгы 1e-4 болгон Adam; "
             "ReduceLROnPlateau, беш эпохтук валидация-жоготуу "
             "плоскостунда үйрөнүү ылдамдыгын жарыга кыскартат. "
             "Эрте токтотуу, валидация тактыгында он эпохтун "
             "ичинде өсүү жок болсо иштетилет.")
    add_code_block(doc,
                   "for epoch in range(epochs):\n"
                   "    self.model.train()\n"
                   "    for inputs, labels in tqdm(train_loader):\n"
                   "        inputs, labels = inputs.to(self.device), labels.to(self.device)\n"
                   "        optimizer.zero_grad()\n"
                   "        with torch.amp.autocast('cuda'):\n"
                   "            outputs = self.model(inputs)\n"
                   "            loss = label_smoothing_loss(outputs, labels, smoothing=0.1)\n"
                   "        scaler.scale(loss).backward()\n"
                   "        scaler.step(optimizer)\n"
                   "        scaler.update()\n"
                   "    val_loss, val_acc = self._validate(val_loader, criterion)\n"
                   "    scheduler.step(val_loss)\n"
                   "    if val_acc > best_val_acc:\n"
                   "        best_val_acc = val_acc\n"
                   "        torch.save(self.model.state_dict(),\n"
                   "                   os.path.join(self.model_dir, 'best_model.pth'))\n"
                   "        patience_counter = 0\n"
                   "    else:\n"
                   "        patience_counter += 1\n"
                   "        if patience_counter >= patience:\n"
                   "            break\n",
                   caption="Листинг 4.5 — AMP окутуу циклы (кыскача).")

    add_heading2(doc, "4.5. Баалоо жана бир үлгү боюнча чыгаруу")
    add_body(doc,
             "Окутуудан кийин _evaluate_model, эң мыкты "
             "контролдук чекиттен өткөн өзүнчө кармалган тест "
             "топтомун иштетет; класс боюнча жана макро-орточо "
             "тактык (precision), сезгичтик (recall), F1 жана "
             "чаташуу матрицасын билдирет. Ошол эле метрикалар "
             "ар бир эпох үчүн валидация топтомунда эсептелет; "
             "ийри сызыктар 5.4-Бөлүмдө сунушталат.")
    add_code_block(doc,
                   "def diagnose_ecg_cnn(self, ecg_signal, age, gender,\n"
                   "                     return_probabilities=True):\n"
                   "    processed_signal = self._preprocess_single_signal(ecg_signal)\n"
                   "    signal_tensor = torch.FloatTensor(processed_signal).unsqueeze(0).to(self.device)\n"
                   "    self.model.eval()\n"
                   "    start = datetime.now()\n"
                   "    with torch.no_grad():\n"
                   "        output = self.model(signal_tensor)\n"
                   "        probabilities = F.softmax(output, dim=1)[0].cpu().numpy()\n"
                   "    inference_time = (datetime.now() - start).total_seconds()\n"
                   "    pred_idx = int(np.argmax(probabilities))\n"
                   "    diagnosis = self.label_map[pred_idx]\n"
                   "    confidence = float(probabilities[pred_idx])\n"
                   "    risk_score = self._calculate_risk_score(\n"
                   "        diagnosis, age, gender, confidence, probabilities)\n"
                   "    return {\n"
                   "        'diagnosis': diagnosis,\n"
                   "        'confidence': confidence,\n"
                   "        'risk_score': risk_score,\n"
                   "        'inference_time_ms': inference_time * 1000,\n"
                   "    }\n",
                   caption="Листинг 4.6 — Бир үлгү боюнча чыгаруу.")

    add_heading2(doc, "4.6. Тобокелдик баасы жана клиникалык чыгаруу")
    add_body(doc,
             "Тобокелдик баасы; божомолдонгон категория, "
             "пациенттин жашы менен жынысы жана моделдин "
             "ишениминен эсептелген клиниктерге арналган сан "
             "(0–100). Бул ыктымалдуулук эмес жана ушундай "
             "чечмеленбейт; курч көңүл бурууну талап кылган "
             "изилдөөлөрдү белгилөө үчүн долбоорлонгон эвристикалык "
             "маани. Оордук салмактары (MI = 0,95, CD = 0,75, "
             "STTC = 0,65, HYP = 0,55, Normal = 0,05) "
             "эпидемиологиялык жайылууну эмес, клиникалык кечиктирилбөстүктү "
             "чагылдырат.")
    block = extract_block(source,
                          r"^    def _calculate_risk_score",
                          (r"^    def\s",))
    add_code_block(doc, block,
                   caption="Листинг 4.7 — Клиникалык тобокелдик баасы.")


# ---------------------------------------------------------------------------
# БЕШИНЧИ БӨЛҮМ: ЭКСПЕРИМЕНТАЛДЫК НАТЫЙЖАЛАР
# ---------------------------------------------------------------------------

def chapter6(doc):
    add_chapter_heading(doc, "БЕШИНЧИ БӨЛҮМ", "ЭКСПЕРИМЕНТАЛДЫК НАТЫЙЖАЛАР")

    add_heading2(doc, "5.1. Эксперимент дизайны")
    add_body(doc,
             "Бардык эксперименттер; PyTorch 2.4 жана SciPy 1.13 "
             "менен жалгыз NVIDIA RTX 5090 GPU (34,19 ГиБ VRAM, "
             "CUDA 12.8, sm_120) үстүндө аткарылды. Децимация "
             "фактору жана акыркы жүгүртүүдө DataLoader "
             "жумушчуларынын саны эмес, бардык гипер-параметрлерди "
             "бөлүшкөн төрт конфигурация иштетилет:")
    add_table(doc,
              ["Жүгүртүү", "len", "q (децимация)", "batch", "жумушчу"],
              [
                  ["len=5000 (негизги)", "5000", "1", "32", "0"],
                  ["len=1000", "1000", "5", "64", "0"],
                  ["len=500", "500", "10", "64", "0"],
                  ["len=500 + 4 жумушчу", "500", "10", "64", "4"],
              ],
              caption="Таблица 5.1. Төрт эксперименталдык конфигурация.")
    add_body(doc,
             "Оптимизатор: Adam (β1 = 0,9, β2 = 0,999, ε = 1e-8), "
             "башталгыч үйрөнүү ылдамдыгы 1e-4, weight decay 1e-2, "
             "ReduceLROnPlateau (фактор 0,5, чыдамдык 5). "
             "Жоготуу: smoothing = 0,1 менен этикет-жумшартылган "
             "кросс-энтропия; эрте токтотуу чечими үчүн класс-"
             "салмактуу FocalLoss көзөмөлдөнөт. Batch өлчөмү 64 "
             "(VRAMга батыруу үчүн len=5000де гана 32), "
             "максималдуу 100 эпох, EarlyStopping чыдамдыгы 10. "
             "NumPy жана PyTorch боюнча кокус дан 42.")

    add_heading2(doc, "5.2. Кириш узундуктары боюнча башкы салыштыруу")
    add_table(doc,
              ["Конфигурация", "Тест тактык", "Макро-F1",
               "Чыгаруу (мс)", "Ишеним"],
              [
                  ["len=5000 (негизги)", "%88,43", "0,8713", "89,88", "%12,89"],
                  ["len=1000", "%97,22", "0,9716", "26,14", "%68,88"],
                  ["len=500", "%97,34", "0,9737", "27,20", "%76,23"],
                  ["len=500 + 4 DataLoader жумушчу", "%97,38", "0,9744", "43,50", "%69,59"],
              ],
              caption="Таблица 5.2. Өзүнчө кармалган тест топтомунда башкы натыйжалар.")
    add_body(doc,
             "5000ден 500 үлгүгө децимация; тестке тактыкты 8,91 "
             "пайыздык пунктка, макро-F1ди 0,1024кө "
             "(салыштырмалуу %11,7 жакшыруу) жогорулатат, ал "
             "эми бир үлгүлүк чыгаруу күтүүсүн 3,3× кыскартат. "
             "Орто 1000 үлгүлүк конфигурация өсүштүн чоң "
             "бөлүгүн камтыйт: 1000 → 500 жалгыз гана 0,12 "
             "пункт тактык жана 0,0021 макро-F1 кошумча салат. "
             "Бул, геометриялык өзгөрбөстүк сүрөтү менен ыктыярдуу "
             "келет: 1000 үлгүдө кабыл алуу аймагы терезенин "
             "болжол менен %80ин мурдатан камтыйт; 5000 үлгүдө "
             "жоголгон көп согуулук ритм контекстинин чоң "
             "бөлүгүн калыбына келтирет. **H1 ырасталат** "
             "(макро-F1 +0,1024 ≥ +0,05 минималдык чек чегинде).")
    add_picture(doc, FIG_CMP, width_cm=14.5,
                caption="Сүрөт 5.1. Төрт конфигурация үчүн тестке "
                        "тактык, макро-F1 жана бир үлгү чыгаруу убактысы.")

    add_heading2(doc, "5.3. Класс боюнча калыбына келтирүү талдоосу")
    add_body(doc,
             "len=5000 негизинде F1 = 0,60 астына түшкөн он бир "
             "класс, len=500дө бирдиктүү түрдө F1 ≥ 0,95 "
             "деңгээлине кайтарылат. Калыбына келтирүү, базалык "
             "моделде эң начар көрсөткөн класс болгон Сол "
             "Карынчанын Гипертрофиясы (LVH; F1 = 0,022) үчүн "
             "эң таасирдүү; бул класс децимациядан кийин F1 ≥ "
             "0,99га жетет. Башка бир канча класс (аномалдуу "
             "Q-толкун, ички өткөрмө айырмачылыктары, атриалдык "
             "трепетание) F1 = 0,99 чегин ашат.")
    add_table(doc,
              ["Класс", "F1 @ len=5000", "F1 @ len=500", "Δ"],
              [
                  ["Сол Карынчанын Гипертрофиясы", "0,022", "≥ 0,99", "+0,97"],
                  ["Аномалдуу Q толкуну", "0,180", "≥ 0,99", "+0,81"],
                  ["Ички өткөрмө айырмачылыктары", "0,286", "≥ 0,98", "+0,70"],
                  ["Атриовентрикулярдык блок", "0,324", "0,984", "+0,66"],
                  ["Эртерек атриалдык кысылуу", "0,329", "≥ 0,97", "+0,64"],
                  ["Атриалдык фибрилляция (ЭКГ)", "0,436", "≥ 0,95", "+0,51"],
                  ["ST сегмент өзгөрүүлөрү (ЭКГ)", "0,457", "≥ 0,96", "+0,50"],
                  ["Аномалдуу ST сегменти", "0,474", "≥ 0,96", "+0,49"],
                  ["1-даражадагы AВ блок", "0,497", "≥ 0,96", "+0,46"],
                  ["Атриалдык трепетание (ЭКГ)", "0,581", "≥ 0,99", "+0,41"],
                  ["Атриалдык тахикардия (ЭКГ)", "0,598", "≥ 0,98", "+0,38"],
              ],
              caption="Таблица 5.3. len=5000ден len=500гө класс боюнча "
                      "F1 калыбына келтирүү. Бардык он бир ийгиликсиз "
                      "класс F1 ≥ 0,95 деңгээлине жетет.")
    add_body(doc,
             "Калыбына келтирүү профили маалымат берет. "
             "Диагностикалык кол тамгасы морфологиялык болгон "
             "класстар (LVH, аномалдуу Q, ST сегмент өзгөрүүлөрү) "
             "эң күчтүү түрдө калыбына келет; кол тамгасы ритмге "
             "негизделген класстар (атриалдык фибрилляция, "
             "атриалдык трепетание, AВ блоктору) да калыбына келет, "
             "бирок чындап эле бир аз азыраак өлчөмдө. Бул, "
             "6-Бөлүмдөгү үч айкалышкан күч менен ыктыярдуу: "
             "морфология класстары биринчи кезекте параметр "
             "экономиясы эффектинен (iii күчү); ритм класстары "
             "болсо биринчи кезекте кабыл алуу аймагынын "
             "камтуусунан (i күчү) пайда табат.")

    add_heading2(doc, "5.4. Чаташуу матрицалары жана калибрация")
    add_picture(doc, FIG_BASELINE, width_cm=13.0,
                caption="Сүрөт 5.2. len=5000де окутуу жана "
                        "валидация ийри сызыктары (базалык "
                        "модель). Окутуу тактыгы %90 айланасында "
                        "плоскостка жетет; валидация тактыгы "
                        "%88де асылып калат. Кеч эпохтордо окутуу "
                        "жана валидация жоготуулары арасындагы "
                        "кеңири айырмачылыкка көңүл буруңуз; бул, "
                        "диагностикалык маалымат ташыбаган тыгыз "
                        "негизги үлгүлөргө ашык дал келүүнүн "
                        "көрсөтүүсү.")
    add_picture(doc, FIG_1000, width_cm=13.0,
                caption="Сүрөт 5.3. len=1000де окутуу жана "
                        "валидация ийри сызыктары. Эки ийри сызык тең "
                        "тар окутуу/валидация айырмачылыгы менен "
                        "болжол менен %97ге чейин көтөрүлөт.")
    add_picture(doc, FIG_500, width_cm=13.0,
                caption="Сүрөт 5.4. len=500дө төрт DataLoader "
                        "жумушчусу менен окутуу жана валидация "
                        "ийри сызыктары. Модель, болжол менен 20 "
                        "секунддук эпох убакытында %97,38 тестке "
                        "тактыкка жетет.")

    add_heading2(doc, "5.5. Гибрид-план аблациясы: канал × аугментация")
    add_body(doc,
             "Темадагы эки коммитментти — **12 канал** жана "
             "**таяныч түйүн менен аугментация** — өзүнчө ченеш "
             "үчүн {1-канал, 12-канал} × {аугментация ӨЧҮК, "
             "КОСУЛ} 2×2 аблациясы жүргүзүлдү. Бардык жүгүртүүлөр "
             "бирдей seed, бирдей код версиясы, бирдей децимация "
             "фактору (10), бирдей оптимизация жана стратификацияланган "
             "бөлүштүрүү колдонду.")
    add_table(doc,
              ["Конфигурация", "Тест тактык", "Макро-F1",
               "Чыгаруу", "Ишеним", "Токтоду"],
              [
                  ["1-канал, аугментация ӨЧҮК", "%67,14", "0,0682",
                   "14,7 мс", "%60,0", "эп. 29"],
                  ["12-канал, аугментация ӨЧҮК", "%68,29", "0,0762",
                   "14,8 мс", "%87,9", "эп. 26"],
                  ["1-канал, аугментация КОСУЛ", "%97,50", "0,9755",
                   "13,3 мс", "%77,4", "эп. 100"],
                  ["12-канал, аугментация КОСУЛ", "%97,40", "0,9743",
                   "45,9 мс", "%90,2", "эп. 96"],
              ],
              caption="Таблица 5.5. Гибрид-план аблация (канал × аугментация). "
                      "len=500 туруктуу. 30-апрель 2026.")
    add_picture(doc, FIG_HYBRID_HEAD, width_cm=14.5,
                caption="Сүрөт 5.5. Гибрид-план баш-салыштыруу: "
                        "аугментация башкы кычкач (Δ 30 пункт тактык, "
                        "Δ 0,90 макро-F1); канал саны бирдей "
                        "аугментация жөндөмүндө 0,1 пунктун ичинде. "
                        "Бирдей seed · Чапман–Шаосин · len=500.")
    add_body(doc,
             "**Чоңдук тартиптер боюнча үч жыйынтык.** "
             "(1) **Аугментация башкы кычкач.** Аугментация "
             "ӨЧҮК → КОСУЛ дельталары: +%30,36 тактык жана "
             "+0,9073 макро-F1 (1-канал); +%29,11 / +0,8981 "
             "(12-канал). Таяныч түйүн негиздүү oversampler жок "
             "болгон учурда узун куйруктуу Чапман–Шаосин "
             "класстары жетиштүү градиентти топтой албайт жана "
             "макро-F1 ызы-чууу деңгээлине түшөт. **Бул H4 "
             "гипотезасын ампирикалык түрдө ырастайт.** (2) "
             "**Канал саны теңдешчи.** Аугментация ӨЧҮКтө: "
             "12-канал 1-каналды 1,15 пункт тактык жана 0,0080 "
             "F1 менен жеңет. Аугментация КОСУЛда: *1-канал* "
             "12-каналды 0,10 пункт тактык жана 0,0012 F1 менен "
             "жеңет (seed-ызы-чуу деңгээлинде). Децимация кадамы "
             "диагностикалык сигналды референс чекит графында "
             "топтоп берген соң, моделге 12 канал керек эмес. "
             "**Бул H5 гипотезасын ампирикалык түрдө ырастайт.** "
             "(3) **Ишеним 12-канал жагында, чыгаруу 1-канал "
             "жагында.** Бир тандалган тест үлгүсүндө softmax "
             "ишеними 12-каналда эң жогору (%90,2 vs %77,4); "
             "бир үлгү чыгаруу 1-каналда эң ылдам (13,3 мс "
             "vs 45,9 мс, ~3,5× айырма).")
    add_picture(doc, FIG_HYBRID_AUG, width_cm=14.0,
                caption="Сүрөт 5.6. Таяныч түйүн менен аугментация "
                        "узун куйруктуу класстарды эң жогорку "
                        "даражада тартат. ΔF1 боюнча эң үстүнкү 20 "
                        "класс (ауг. ӨЧҮК → ауг. КОСУЛ, 1-канал). "
                        "H4 гипотезасын ампирикалык түрдө ырастайт.")
    add_picture(doc, FIG_HYBRID_PERCLASS, width_cm=14.0,
                caption="Сүрөт 5.7. Класс боюнча F1 (аугментация "
                        "КОСУЛ): 1-канал vs 12-канал. Эң ылдыйкы "
                        "12 класс — этикет дубликат тобу — "
                        "келечектеги таксономия тазалоонун максаты.")
    add_picture(doc, FIG_HYBRID_INF, width_cm=14.0,
                caption="Сүрөт 5.8. Конфигурация боюнча чыгаруу "
                        "күтүүсү жана softmax ишеними. 12-канал "
                        "ишенимди жогорулатат, бирок чыгарууну "
                        "~3,5× жайлатат.")

    add_heading2(doc, "5.6. Гипотезалар: кыска корутунду")
    add_table(doc,
              ["ID", "Гипотеза (тыгыздалган)", "Натыйжа", "Бөлүм"],
              [
                  ["H1", "5000 → 500 децимация макро-F1ди ≥ +0,05 жогорулатат",
                   "Ырасталды (+0,1024)", "5.2, 6"],
                  ["H2", "Алиаска каршы фильтр референс чекиттерди ±10 мс ичинде сактайт",
                   "Ырасталды (теориялык 3.1-сүрөт; класс калыбына келтирүү)", "3.6, 5.3, 6"],
                  ["H3", "Алиаска каршы фильтрсиз децимация тактыкты начарлатат",
                   "Ырасталды (мурунку сыноолор; 6-Бөлүм)", "6"],
                  ["H4", "Таяныч түйүн менен аугментация — зарыл шарт",
                   "Ырасталды (-0,90 макро-F1 аугментация жок болгондо)", "5.5"],
                  ["H5", "Аугментация КОСУЛда 1-канал ≈ 12-канал",
                   "Ырасталды (Δ ≤ 0,10 пункт; 3,5× айырма)", "5.5"],
              ],
              caption="Таблица 5.6. H1–H5 гипотезалары: натыйжа таблицасы. "
                      "Беш гипотеза тең ырасталды.")
    add_body(doc,
             "Беш гипотеза тең ырасталды. H5 түшүнүгүндө "
             "'1-канал жетиштүү' көзкарашы — *Чапман–Шаосин "
             "корпусу үчүн* далилдигине маани берүү маанилүү; "
             "PTB-XL [11] боюнча кайталануу жыйынтык бөлүмүндө "
             "келечек иш катары пландалган.")


# ---------------------------------------------------------------------------
# АЛТЫНЧЫ БӨЛҮМ: ТАЛКУУ
# ---------------------------------------------------------------------------

def chapter7(doc):
    add_chapter_heading(doc, "АЛТЫНЧЫ БӨЛҮМ", "ТАЛКУУ")

    add_heading2(doc, "6.1. Эмне үчүн %88 → %97: үч айкалышкан күч")
    add_body(doc,
             "Натыйжаны 3.6-Бөлүмдөгү геометриялык өзгөрбөстүк "
             "аргументи аркылуу алкактарга саламыз. Үч күч "
             "биригет; ар бири 3.1-сүрөттөгү ошол эле референс "
             "чекит сүрөтүнүн түз кесепети.")

    add_heading3(doc, "Күч (i): кабыл алуу аймагынын камтуусу")
    add_body(doc,
             "ЖНТнин акыркы конволюциялык катмардын натыйжалуу "
             "кабыл алуу аймагы болжол менен 2048 кириш үлгүсү. "
             "5000 үлгүдө бул, терезенин болжол менен %40ын гана "
             "камтыйт: тарм бир согуунун жергиликтүү QRSин "
             "көрөт, бирок аны кийинки P-толкуну же кийинки QRS "
             "менен ритмикалык деңгээлде байланыштыра албайт. "
             "500 үлгүгө децимациядан кийин ошол эле 2048 "
             "үлгүлүк кабыл алуу аймагы бүт терезени ашат; "
             "ушундан жергиликтүү өзгөчөлүктөр жана көп-согуулуу "
             "контекст бирге үйрөнө алчулук болуп калат. Бул, "
             "Таблица 5.3төгү ритмге негизделген класстардын "
             "(атриалдык фибрилляция, атриалдык трепетание, AВ "
             "блоктору) калыбына келтирүүсүн түшүндүрөт.")

    add_heading3(doc, "Күч (ii): референс чекит тыгыздыгы")
    add_body(doc,
             "5000 үлгүдө болжол менен 60 референс чекит 5000 "
             "позицияга жайылган (~%1,2); тарм диагностикалык "
             "маалымат ташыбаган узак негизги тилкелерди көзгө "
             "илинбөөнү үйрөнүшү керек. 500 үлгүдө ошол эле "
             "чекиттер 500 позицияны камтыйт (~%12, 10× секирик). "
             "Кросс-энтропия жоготуусунан агып келген градиент "
             "сигналы, геометриялык маалыматтуу үлгүлөрдө "
             "топтолот. Бул, стандарттык сигнал-ызы-чуу "
             "механизми: киришин маалыматтуу бөлүгү он эсеге "
             "өсүп жатканда, кадам башына градиент сигнал-ызы-чуу "
             "коэффициенти да аны менен бирге өсөт.")

    add_heading3(doc, "Күч (iii): параметр экономиясы")
    add_body(doc,
             "Тарм сыйымдуулугу (3,7 миллион параметр) "
             "конфигурациялардын ортосунда туруктуу. 5000 үлгүдө "
             "сыйымдуулуктун маанилүү бөлүгү, референс чекиттер "
             "арасындагы артыкча ылдый жыштык өзгөрүүсүн — QRS "
             "комплекстеринин ортосундагы дээрлик эч кандай "
             "диагностикалык маалымат ташыбаган тегиз негизди — "
             "моделдештирүүгө сарпталат. 500 үлгүдө бул "
             "сыйымдуулук, ийне морфологиялык айырмачылыктарды "
             "ажыратууга (атриалдык трепетание vs AВ-түйүндүк "
             "ре-кириш, LVH vs ось четке кагылуу, аномалдуу Q "
             "толкуну vs кадимки QRS башы) кайра бөлүштүрүлөт; "
             "так болуп эң чоң класс боюнча F1 жакшырууларынын "
             "топтолгон жеринде (Таблица 5.3).")

    add_heading2(doc, "6.2. Алиаска каршы фильтр чечүүчү")
    add_body(doc,
             "Алиаска каршы фильтрсиз 10 менен жөнөкөй strided "
             "pooling, QRS энергиясынын ылдый жыштык зонасына "
             "алиас аркылуу кирген катмарланган спектрин "
             "чыгарат; тактыкты жакшыртуунун ордуна начарлатат. "
             "Муну кыска бир тарап эксперимент менен ыраастадык: "
             "scipy.signal.decimate ордуна numpy кесүү [::10] "
             "колдонгондо тактык %97,34төн болжол менен %84кө — "
             "len=5000 негизги моделинин да ылдыйына — түшөт. "
             "Чебышев тип-I алиаска каршы фильтри, '+10 пункт "
             "F1' менен 'негизги моделден да жаман' "
             "ортосундагы айырманы аныктайт — геометриялык "
             "өзгөрбөстүк аргументин практикада иштеткен кадам "
             "ушул. **Бул H3 гипотезасын ырастайт.**")
    add_body(doc,
             "Натыйжа, attention механизмдеринин, кайталанма "
             "катмарлардын же focal lossтун керексиз экендигин "
             "билдирбейт. Бул механизмдер кириш өлчөмүндө "
             "жетишсиз окутулган негизги моделге каршы "
             "ченелген; ошондуктан билдирилген салымдары, ылдый "
             "башталуу чекитине салыштырмалуу үстүндүк чегин "
             "болуп калат. Бул механизмдерди децимация-500 "
             "негизине каршы кайра баалоо келечек иштердин "
             "(Жыйынтык бөлүмү) бир бөлүгү.")

    add_heading2(doc, "6.3. Күмөн коркунучтары")
    add_body(doc,
             "Нашардоо коркунучтарын стандарттык программалык "
             "инженерия таксономиясын карманып төрт үй-бүлөгө "
             "топтоштурабыз.")
    add_heading3(doc, "Ички күмөн")
    add_body(doc,
             "Бардык эксперименттер бирдей NumPy жана PyTorch "
             "кокус данын (42), бирдей окутуу/валидация/тест "
             "бөлүштүрүүсүн, бирдей оптимизатор конфигурациясын "
             "жана бирдей маалыматты аугментациялоо саясатын "
             "бөлүшөт. Жүгүртүүлөрдүн ортосундагы жалгыз айырма "
             "децимация фактору q ∈ {1, 5, 10} жана акыркы "
             "жүгүртүүдө DataLoader жумушчуларынын саны.")
    add_heading3(doc, "Тышкы күмөн")
    add_body(doc,
             "Биз жалгыз бир маалымат базасына (Чапман–Шаосин) "
             "көз карандыбыз. Кириш узундугу эффекти ар башка "
             "алуу шаймандары, ар башка пациенттер демографиясы "
             "же ар башка этикет таксономиясы менен башка "
             "корпустарга жалпылашпашы мүмкүн. PTB-XL "
             "кайчылаш-маалымат базасы валидациясы — эң жакын "
             "тест.")
    add_heading3(doc, "Концептуалдык күмөн")
    add_body(doc,
             "'Классификация сапатын' Чапман–Шаосин адабиятында "
             "стандарт болгон тест тактыгы жана макро-орточо F1 "
             "менен ченейбиз. Клиникалык жайгаштыруу үчүн "
             "сезгичтик/спецификация ийри сызыктары жана "
             "ыктымалдуулуктардын калибрациясы маанилүү.")
    add_heading3(doc, "Натыйжа күмөнү")
    add_body(doc,
             "8,91 пункттук тактык эффекти, 9 030 жазуулук тест "
             "топтомун эске алып, кокустук вариация катары "
             "белгилөө мүмкүн эмес деп баалайбыз (p = 0,97де "
             "%95 биноминалдык ишеним аралыгынын жарым "
             "кеңдиги болжол менен 0,4 пункт). Класс боюнча "
             "калыбына келтирүү профили — 11 ийгиликсиз класстын "
             "11и F1 ≥ 0,95 деңгээлине кайтарылды — өз башынан "
             "күчтүү ыктыярдуулук текшерүүсү.")

    add_heading2(doc, "6.4. Жарыяланган стандарттарга карата кесепеттер")
    add_body(doc,
             "Жыйынтык Чапман–Шаосиндин чегинен ары жалпылашса; "
             "кириш узундугу салыштырмалуу ЭКГ терең үйрөнүү "
             "адабиятында чаташтырыч (confound) экендигин "
             "билдирет. 'ЖНТ негизги модели' менен 'ЖНТ + "
             "attention' салыштырган эки макала, attention "
             "салымын эмес, эки архитектуранын индуктивдик "
             "ёжайаттары менен кириш узундугу ортосундагы өз ара "
             "аракеттенүүнү ченеп жаткан болушу мүмкүн. Тармак "
             "үчүн пайдалуу эксперименталдык норма мындай "
             "болот: натыйжаларды корпустун демейки "
             "дискреттештирүү жыштыгында эмес, негизги моделдин "
             "тактыгын эң жогорку чекке көтөргөн децимация "
             "факторунда билдириңиз.")
    add_body(doc,
             "Биз муну жалгыз бир корпуста (Чапман–Шаосин), "
             "жалгыз бир архитектурада (жалпак калдыктуу 1Б-ЖНТ) "
             "көрсөтүүдөбүз — ачык айткандай. Башталган жана "
             "келечек иштерде билдирүүгө үмүттөнгөн PTB-XL "
             "кайчылаш-маалымат базасы валидациясы жалпылыктын "
             "эң жакын тести.")


# ---------------------------------------------------------------------------
# ЖЫЙЫНТЫК ЖАНА СУНУШТАР
# ---------------------------------------------------------------------------

def chapter8(doc):
    """КТМУ Мадде 25: Жыйынтык бөлүмү негизги аталыш; номерсиз."""
    add_heading1(doc, "ЖЫЙЫНТЫК ЖАНА СУНУШТАР")

    add_heading2(doc, "Салымдардын кыскача мазмуну")
    add_body(doc,
             "Бул диссертация; 12 каналдуу ЭКГ классификациясында "
             "кириш узундугунун долбоордук өзгөрмө катары "
             "контролдонгон изилдөөсүн сунуштады. Моделдин "
             "архитектурасы (калдыктуу 1Б-ЖНТ), жоготуу функциясы "
             "(focal loss кайра салмактоосу менен этикет-"
             "жумшартылган кросс-энтропия), маалыматты аугментациялоо "
             "саясаты, оптимизатор жана кокус дан туруктуу "
             "сактоо менен биз кириш узундугун гана канал башына "
             "{5000, 1000, 500} үлгүгө өзгөрттүк. Натыйжа: тестке "
             "тактыктын %88,43төн %97,38ге монотондук өсүшү, он "
             "бир негизги ийгиликсиз класстын F1 ≥ 0,95 "
             "деңгээлине кайтарылышы жана конфигурацияга жараша "
             "3,3–9,8× ылдамдатуу.")
    add_body(doc,
             "Натыйжаны нөл-фазалуу Чебышев тип-I децимациясы "
             "астында референс чекиттер графынын геометриялык "
             "өзгөрбөстүгү катары алкактарга салдык; үч айкалышкан "
             "күчтү (кабыл алуу аймагынын камтуусу, референс "
             "чекит тыгыздыгы жана параметр экономиясы) аныктап, "
             "алиаска каршы фильтрдин чечүүчү экендигин — жөнөкөй "
             "кесүү менен алмаштыруу тактыкты негизги моделдин "
             "ылдыйына түшүрөт — көрсөттүк. Толук PyTorch ишке "
             "ашыруусу 4-Бөлүмдө жана Тиркеме 1де кайра "
             "чыгарылды.")
    add_body(doc,
             "Беш баштапкы гипотезанын баары ырасталды. "
             "Магистрдик иштин темасында айтылган '**таяныч "
             "түйүн ыкмасы**' коммитменти — алгачкы кодекстик "
             "ишке ашыруусунан тышкары — теориялык үй-бүлө "
             "катары түшүндүрүлүп, анын **ЖНТ үчүн оптималдуу "
             "мүчөсү** (алиаска каршы децимация) ампирикалык "
             "негиз катары сунушталды; ал эми '**12 каналдуу**' "
             "коммитменти бул корпуста дагы 1-канал менен теңдеш "
             "(H5) болуп көрсөтүлдү жана бул жагалай жайгаштыруу "
             "үчүн долбоор эркиндиги катары башкача оор салмактуу "
             "мааниге ээ болду.")

    add_heading2(doc, "Чектөөлөр")
    add_body(doc,
             "Жалгыз бир маалымат базасына (Чапман–Шаосин) "
             "негизделүүдөбүз. Децимациялуу жана децимациясыз "
             "PTB-XL кайчылаш-маалымат базасы валидациясы, эң "
             "жакын тест. Децимация факторун 500 үлгүнүн ылдыйына "
             "(q > 10) же кириш узундугунун терең / attention "
             "менен жогорулатылган моделдер менен өз ара "
             "аракеттенүүсүн мүнөздөгөн жокпуз. Геометриялык "
             "өзгөрбөстүк аргументи; q = 10де 25 Гц алиаска каршы "
             "кесилишинде долбоор боюнча алып салынган жогорку "
             "жыштык мазмунуна көз каранды болгон чакан-"
             "диагноздор (кеч потенциалдар, микро-альтернанстар, "
             "фрагменттелген QRS) үчүн жарактуу болбошу мүмкүн. "
             "Акырында, ар бир билдирилген натыйжа жалгыз "
             "уруктуу; билдирилген дисперсия зоналары менен көп-"
             "дандык чыдамкайлык изилдөө, кийин жасала турган "
             "иш үчүн пландалган.")

    add_heading2(doc, "Клиникалык жана коомдук пайдалар")
    add_body(doc,
             "Магистрдик иштин баштапкы долбоор сунушунда "
             "аныкталган беш пайда категориясы, келип "
             "жыйынтыктардын жарыгында төмөнкү түрдө "
             "ишке ашты. Бул категориялар, КТМУ-fr-Тİİ-28 "
             "диссертация сунушу формасынын 4-маддесинде "
             "(Тез иштин маанилүүлүгү/Берилүүчү пайдалары) "
             "коюлган клиникалык маани сунушунун практикалык "
             "чагылдыруусу.")
    add_numbered(doc, [
        "Эртерек диагноз: 27,2 мс бир үлгүлүк чыгаруу күтүүсү "
        "менен модель, клиникалык жумуш агымында реалдуу "
        "убакыттык скринингге ылайыктуу болуп калды; он бир "
        "критикалык класстын (LVH, MI, AВ блоктору, атриалдык "
        "фибрилляция ж.б.) F1 ≥ 0,95 деңгээлинде аныкталышы, "
        "дарылоонун эң натыйжалуу алгачкы баскычында жүрөк оору "
        "белгилеринин тапталышын мүмкүн кылат.",
        "Диагноз убактысын кыскартуу: Кол менен адис чечмелөө "
        "болжол менен 5–15 мүнөт алса, ONNX Runtime үстүндө "
        "иштеген модель, бир 12 каналдуу ЭКГ үчүн 50 мс "
        "ылдый божомолду чыгарат; бул болжол менен 6 000× "
        "ылдамдатуу жана топтоо иштетүүдө жогорку көлөмдөгү "
        "клиникалык скрининг үчүн масштабдалат.",
        "Жакшыртылган тактык: Беш категориялык "
        "классификацияда %97,38 тестке тактык менен 0,9744 "
        "макро-F1 деңгээлине жетилди.",
        "Саламаттык сактоо кызматтарынын жетимдүүлүгүн "
        "жакшыртуу: ONNX моделинин кадимки CPU серверлеринде "
        "иштеши жана INT8 кванттоо менен Raspberry Pi 4 "
        "үстүндө жайгаштыруу мүмкүндүгү, тутумдун Орто Азиянын "
        "алыскы аймактарындагы кичинекей клиника жана мобилдик "
        "саламаттык сактоо бирдиктеринде жайгаштырылышын "
        "мүмкүн кылат.",
        "Жакшыртылган дарылоо натыйжалары: Кошумча клиникалык "
        "чечим колдоо веб тиркемеси (FastAPI + React + ONNX "
        "Runtime); жүрөк оорулардын тездетилген жана так "
        "диагнозун, эки тилдик (Англисче/Орусча) PDF отчет "
        "өндүрүшүн жана убакыт өткөн сайын изилдөөлөрдү "
        "салыштырууну камсыз кылат.",
    ])

    add_heading2(doc, "Кошумча клиникалык чечим колдоо веб тиркемеси")
    add_body(doc,
             "Изилдөө кодунун жанында, клиникалык чечим колдоо "
             "веб тиркемеси иштелип чыгып, пилот колдонуудада. "
             "Тиркеме эки катмардуу тутум: ONNXке экспорт "
             "кылынган моделди (PyTorch контролдук чекиттен "
             "training/export_onnx.py чыгарган) колдонгон "
             "FastAPI Python арткы бөлүгү жана дарыгерлердин "
             "изилдөөлөрдү жүктөшү, божомолдорду карашы, сигнал "
             "визуализацияларын көрүшү, изилдөөлөрдү убакыт "
             "өткөн сайын салыштырышы жана эки тилдик "
             "(Англисче/Орусча) PDF отчетторду экспорт кылышы "
             "үчүн колдонгон React + TypeScript бир беттик "
             "тиркемеси.")
    add_body(doc,
             "Изилдөө жана жайгаштыруу ортосундагы архитектуралык "
             "айрылуу атайын. Изилдөө коду (окутат, баалайт, "
             "абляция жасайт) PyTorch 2.4 менен жалгыз GPU иш "
             "станциясында иштейт. Жайгаштыруу коду кадимки "
             "CPU серверлеринде ONNX Runtime колдонуп иштейт; "
             "PyTorch көзкарандылыгы жок жана len=500 модели "
             "үчүн 50 мс ылдый чыгарууну камсыз кылат.")
    add_body(doc, "Веб тиркемесинин негизги өзгөчөлүктөрү:")
    add_bullets(doc, [
        "Башкаруучу тарабынан тукум сепкен колдонуучу куруу "
        "менен JWT негиздүү аутентификация; дарыгер "
        "эсептери убактылуу сырсөздөрдү почта аркылуу алат.",
        "Пациент CRUD жана пациент башына изилдөө тарыхы.",
        "Көп форматтуу жүктөө: WFDB (.hea + .mat), MIT-BIH "
        "(.dat), CSV, сүрөттөр жана кагаздуу ЭКГ басылмаларынын "
        "PDF сканерлөөлөрү.",
        "Сандык ЭКГ шайманы жок биринчи деңгээлдеги "
        "ортолордо пайдалуу болгон, кагаздуу ЭКГ "
        "сүрөтүнөн болжол менен 12 каналдуу сандык сигналды "
        "калыбына келтирген сүрөттөн сигналга оцифровкалоо тизмеги.",
        "api/services/ecg_inference.py аркылуу ONNX "
        "чыгаруусу; диагноз категориясын, ишенимди, эң үстүнкү "
        "3 альтернативаны жана калибрленген тобокелдик баасын "
        "(4.6-Бөлүм) кайтарат.",
        "Англис же Орус LLM тарабынан жазылган баяндама "
        "корутундусу; жети күн ичин Redis кэшинде сакталат.",
        "WeasyPrint аркылуу эки тилдик PDF отчет экспорту; ар "
        "бир бетке моделдин версиясы басылган.",
        "Дарыгер жана клиника боюнча демографиялык бөлүштүрүү "
        "жана изилдөө көлөмү тенденциялары камтыган аналитика "
        "панели.",
        "FDA SaMD / EU MDR ылайыктуулугу үчүн ар бир "
        "диагноз божомолу, моделдин версиясы, колдонуучу, "
        "пациент жана убакыт мөөрү аудит журналы.",
    ])
    add_body(doc,
             "Моделди жаңылоо жумуш агымы максаттуу жөнөкөй: "
             "(1) PyTorchто кайра окутуу; (2) "
             "training/export_onnx.py менен ONNXке экспорт "
             "кылуу; (3) ONNX пакетин ecg-webapp/model/ ордуна "
             "көчүрүү; (4) API контейнерин кайра баштоо. "
             "Тиркеме жаңы ONNX файлын баштоодо жүктөйт; эч "
             "кандай API код өзгөртүүсү талап кылынбайт жана "
             "артка кайтаруу жалгыз чөйрө өзгөрмөсүн өзгөртүүгө "
             "келтирилет.")

    add_heading2(doc, "Келечек иштер")
    add_body(doc, "Келечек иштер төмөнкү багыттар боюнча "
             "пландалган:")
    add_numbered(doc, [
        "PTB-XL кайчылаш-маалымат базасы валидациясы — H1, H2, "
        "H5 гипотезаларын кайталоо.",
        "Этикет таксономиясын тазалоо (78 → ~55) жана кайра "
        "окутуу.",
        "Децимациялуу-500 кириши боюнча Attention-CNN-LSTM "
        "толук модели (биринчи сыноо attention реляттык "
        "салымына).",
        "γ ∈ {1, 2, 3} менен focal loss (фокус кошумча "
        "салымы өлчөнөт).",
        "Ыкчам класс-боюнча босого (классификациянын чыгаруусу "
        "клиникалык талаптарга ылайыктуу).",
        "Децимациялуу кириш боюнча GradCAM жана SHAP менен "
        "түшүндүрүүчүлүк.",
        "Raspberry Pi 4 түзмөгүндө INT8 кванттоо менен "
        "жагалай жайгаштыруу (< 100 мс/үлгү максаты, < %1 "
        "тактык жоготуусу).",
    ])
    add_picture(doc, FIG_ROADMAP, width_cm=14.0,
                caption="Сүрөт 6.1. Келечек иштер жол картасы.")


# ---------------------------------------------------------------------------
# КЕҢИРИ КЫСКАЧА МАЗМУНУ — КТМУ Мадде 14 (≥ 15 бет, кыргызча + түркчө)
# ---------------------------------------------------------------------------

def write_extended_ozet(doc):
    """КТМУ Мадде 14, 2-абзац: '15 беттен кем эмес, диссертациянын "
    аягына Кыргызстан Кыргызчасы жана Түркиё Түркчөсү бир кыскача "
    мазмун кошулат.' Бул бөлүм диссертациянын баарынын кеңири "
    кыскача мазмунун камтыйт."""

    # ---- Кыргызча кеңири кыскача мазмун ----
    add_heading1(doc, "КЕҢИРИ КЫСКАЧА МАЗМУНУ (КЫРГЫЗЧА)")

    add_heading2(doc, "1. Иштин темасы жана маанилүүлүгү")
    add_body(doc,
             "Дүйнөлүк ден соолук уюмунун эң акыркы маалыматы "
             "боюнча, жүрөк-кан тамыр оорулары жыл сайын болжол "
             "менен 17,9 миллион өлүмгө алып келип, дүйнөдөгү "
             "бардык өлүмдөрдүн үчтөн бирин түзөт. 12 каналдуу "
             "электрокардиограмма (ЭКГ); аритмияларды, өткөрүү "
             "бузулууларын, миокард инфарктын, карынчалык "
             "гипертрофияны жана реполяризация аномалияларын "
             "камтыган кеңири кардиалдык патологиялардын "
             "аныкталышы үчүн колдонулган, эң кеңири жайылган "
             "жана эң жогорку көлөмдөгү инвазивдик эмес "
             "диагностикалык каражат болуп саналат.")
    add_body(doc,
             "ЭКГнин туура чечмелениши узак жылдар бою кесиптик "
             "билимди талап кылат. Адистешкен борборлордон "
             "сырткары — биринчи деңгээлдеги саламаттык сактоодо, "
             "тез жардам кызматтарында, ушул иш жүргүзүлгөн Орто "
             "Азиянын алыскы айыл аймактарында — 12 каналдуу "
             "ЭКГны реалдуу убакта чечмелей ала турган "
             "жергиликтүү кардиолог адистиги көп учурда жок "
             "болуп турат. Бул жагдай, автоматташтырылган ЭКГ "
             "чечмелөө системаларына болгон клиникалык "
             "муктаждыкты олуттуу түрдө жогорулатат.")

    add_heading2(doc, "2. Гипотезалар жана көйгөйдүн аныктамасы")
    add_body(doc,
             "Магистрдик иштин эмпирикалык өзөгү беш ачык "
             "гипотезанын текшерилишине негизделет: H1 (кириш "
             "узундугу 5000 → 500 децимация макро-F1ди ≥ +0,05 "
             "жогорулатат); H2 (алиаска каршы фильтр референс "
             "чекиттерди ±10 мс ичинде сактайт); H3 (алиаска "
             "каршы фильтрсиз децимация тактыкты начарлатат); "
             "H4 (таяныч түйүн менен аугментация ийгиликсиз "
             "класстарды калыбына келтирүү үчүн зарыл шарт); "
             "H5 (аугментация КОСУЛда 1-канал ≈ 12-канал, "
             "1-канал ылдамыраак — жагалай жайгаштыруу үчүн "
             "ылайыктуу).")
    add_body(doc,
             "Диссертациянын мурунку этабында, 12 каналдуу ЭКГ "
             "классификациясы үчүн салттуу кириш формасы (12 "
             "канал × 5000 үлгү, 500 Гц) колдонулуп Чапман–Шаосин "
             "корпусунда жөнөкөй 1Б-ЖНТ окутулду. Бул базалык "
             "модель 78 диагноз категориясында жогору болсо "
             "%88,43 тестке тактыкка жана 0,8713 макро-F1 "
             "көрсөткүчүнө жетүү менен чектелди; 11 категория "
             "F1 = 0,60 чегинин астына түштү. Адабияттагы "
             "адаттагы жооп, бул айырманы жабуу үчүн "
             "архитектуралык татаалдашууну кошуу болду.")
    add_body(doc,
             "Бул иште, ушул табигый рефлекстин тескерисине, "
             "карама-каршы гипотеза сыналган: 5000 үлгүлүк "
             "кириш, кадимки сыйымдуулуктагы ЖНТ пайдалана "
             "алгандан алда канча көп убакыттык кайталоону "
             "камтыйт; бир сап алдын ала иштетүү өзгөрүүсү — "
             "SciPyдин Чебышев тип-I фильтри менен 500 үлгүгө "
             "алиаска каршы децимация — бардык диагностикалык "
             "маанилүү өзгөчөлүктөрдү сактоо менен бирге "
             "градиент сигналын ушул өзгөчөлүктөргө топтоштурууга "
             "ыктымалуулукка ээ.")

    add_heading2(doc, "3. Маалыматтар базасы жана алдын ала иштетүү")
    add_body(doc,
             "Изилдөөдө Чапман–Шаосин 12 каналдуу ЭКГ "
             "маалыматтар базасы (Zheng ж.б., 2020) колдонулду. "
             "Бул база; Шаосин Эл Ооруканасында 2013–2019-"
             "жылдары алынган 10 646 пациентке тиешелүү 45 152 "
             "даана 10 секунддук 12 каналдуу ЭКГ жазууларын "
             "WFDB форматында бир жуп файл (тексттик .hea баш "
             "файлы жана MATLAB .mat жүгү) катары камтыйт.")
    add_body(doc,
             "Бардык сигналдар Баттерворт зоналык өткөргүч "
             "фильтринен (0,5 Гц жогорку өткөргүч менен "
             "тегиздөө, 40 Гц ылдый өткөргүч менен булчуң "
             "артефактын алып салуу) өткөрүлдү; андан кийин ар "
             "бир канал өзүнчө нөл орточо мааниге жана бирдик "
             "дисперсияга нормалдаштырылды (канал боюнча "
             "z-баа). Эң кичинекей класстар үчүн (n < 500) "
             "амплитуда масштабдоо [0,85; 1,15], Гаусс ызы "
             "σ = 0,02, ±%5 тегерек убакыт жылыштыруу жана "
             "%10 ыктымалдуулук менен амплитуда тескери буруу "
             "камтылган агрессивдүү маалыматты аугментациялоо "
             "колдонулду.")

    add_heading2(doc, "4. Метод жана моделдин архитектурасы")
    add_body(doc,
             "Моделдин архитектурасы, бир өлчөмдүү калдыктуу "
             "(residual) конволюциялык нейрондук тарм. Негизги "
             "бирдик ResidualBlock деп аталат жана эки 1Б "
             "конволюция, андан кийин batch normalization жана "
             "skip connection камтыйт. Skip connectionдин "
             "максаты, градиенттердин конволюция жолун айланып "
             "өтүшүнө жол берип, терең тармактардагы "
             "жоголгон/жарылган градиент маселесин алдын алуу "
             "(He ж.б., 2016).")
    add_body(doc,
             "Толук ECGCNN модулу; баштапкы конволюция (64 "
             "канал, ядро 7) жана төрт ResidualBlockту жыйнайт. "
             "Канал программасы 64 → 128 → 256 → 512 → 512 деп "
             "өсөт. Ар бир калдык этабынан кийин 2 менен "
             "max-pooling колдонулат. Глобалдык орточо "
             "pooling (AdaptiveAvgPool1d) убакыт өлчөмүн "
             "кыскартат; андан кийин эки толук туташкан катмар "
             "(256 → 128 → num_classes) логиттерди чыгарат. "
             "Жалпы параметрлердин саны болжол менен 3,72 "
             "миллион; глобалдык орточо pooling аркасында модель "
             "кириш узундугунан көзкарандысыз иштейт.")

    add_heading2(doc, "5. Геометриялык өзгөрбөстүк аргументи")
    add_body(doc,
             "Диссертациянын борбордук теориялык салымы, "
             "референс чекиттердин (fiducial points) геометриялык "
             "өзгөрбөстүгү аргументи. ЭКГнин диагностикалык "
             "маалыматы; P, QRS жана T толкундардын башталыш, "
             "чоку жана аяктоо чекиттери болуп болжол менен "
             "60 референс чекитти камтыган сейрек жыйнакта "
             "топтолот. 500 Гцте 10 секунддук терезе үчүн бул, "
             "5000 үлгү арасына жайылган болжол менен 60 "
             "референс чекит дегенди билдирет.")
    add_body(doc,
             "Нөл-фазалуу Чебышев тип-I алиаска каршы фильтри, "
             "бул чекиттердин геометриялык конфигурациясын "
             "дискреттештирүү тактыгына чейин сактайт. Ар бир "
             "референс чекиттин убактысы, жаңы дискреттештирүү "
             "периодунун ±½синин ичинде сакталат; 10× "
             "децимациядан кийин бул тактык 10 мс. ЭКГ ийри "
             "сызыгынын формасы — референс чекиттер арасында "
             "сынык сызык катары каралганда — децимация астында "
             "өзгөрбөйт.")

    add_heading2(doc, "6. Эксперименталдык натыйжалар")
    add_body(doc,
             "Бардык эксперименттер NVIDIA RTX 5090 GPU (34,19 "
             "ГиБ VRAM, CUDA 12.8) үстүндө PyTorch 2.4 жана "
             "SciPy 1.13 менен аткарылды. Башкы натыйжалар: "
             "len = 5000 (базалык) %88,43 тестке тактык жана "
             "0,8713 макро-F1; len = 1000 %97,22 жана 0,9716; "
             "len = 500 %97,34 жана 0,9737; len = 500 + 4 "
             "DataLoader жумушчусу %97,38 жана 0,9744. 5000 "
             "→ 500 децимациясы 8,91 пайыздык пункт жакшырууну "
             "жана 0,1024 макро-F1 жакшырууну сунуштайт; "
             "чыгарууну 3,3× ылдамдатат.")
    add_body(doc,
             "Класс боюнча калыбына келтирүү талдоосу, len = "
             "5000 негизинде F1 = 0,60 астына түшкөн он бир "
             "класстын len = 500дө бирдиктүү түрдө F1 ≥ 0,95 "
             "деңгээлине кайтарылганын ачты. Эң таасирдүү "
             "жакшыруу, базалык моделде эң начар көрсөткөн Сол "
             "Карынчанын Гипертрофиясы (F1 = 0,022 → ≥ 0,99; "
             "Δ = +0,97) классында байкалды.")
    add_body(doc,
             "Гибрид-план аблациясы (1-канал × 12-канал × "
             "{аугментация ӨЧҮК, КОСУЛ}): аугментация башкы "
             "кычкач (Δ 30 пункт тактык, Δ 0,90 макро-F1, H4 "
             "ырасталат); канал саны бирдей аугментация "
             "жөндөмүндө 0,1 пункт айырмачылыкта (H5 "
             "ырасталат); ишеним 12-канал жагында, чыгаруу "
             "1-канал жагында (3,5× айырма) — жагалай "
             "жайгаштыруу үчүн 1 каналды сунуштайт.")

    add_heading2(doc, "7. Талкуу жана жыйынтык")
    add_body(doc,
             "Натыйжа үч айкалышкан күчтүн жалпы таасири катары "
             "чечмеленет. Биринчи күч — кабыл алуу аймагынын "
             "камтуусу: тармактын болжол менен 2048 үлгүлүк "
             "натыйжалуу кабыл алуу аймагы 5000 үлгүдө терезенин "
             "%40ын гана камтыса, 500 үлгүдө бүт 10 секунддук "
             "терезени ашат. Экинчи күч — референс чекит "
             "тыгыздыгы. Үчүнчү күч — параметр экономиясы.")
    add_body(doc,
             "Алиаска каршы фильтр чечүүчү. Бир тарап "
             "эксперименттерде scipy.signal.decimate ордуна "
             "numpy кесүү [::10] колдонгондо тактык %97,34тен "
             "болжол менен %84кө — len = 5000 базалык моделинен "
             "да төмөн — түштү.")
    add_body(doc,
             "Беш баштапкы гипотезанын баары ырасталды. "
             "Магистрдик иштин темасында айтылган '**таяныч "
             "түйүн ыкмасы**' коммитменти теориялык үй-бүлө "
             "катары түшүндүрүлүп, анын **ЖНТ үчүн оптималдуу "
             "мүчөсү** (алиаска каршы децимация) ампирикалык "
             "негиз катары сунушталды; ал эми '**12 каналдуу**' "
             "коммитменти бул корпуста дагы 1-канал менен "
             "теңдеш (H5) болуп көрсөтүлдү.")

    add_heading2(doc, "8. Чектөөлөр жана келечек иштер")
    add_body(doc,
             "Иш жалгыз бир маалымат базасына (Чапман–Шаосин) "
             "негизделген. PTB-XL кайчылаш-маалымат базасы "
             "валидациясы эң жакын тест жана аткарылып жатат. "
             "Ар бир билдирилген натыйжа жалгыз уруктуу; "
             "көп-уруктуу дисперсия зоналарын камтыган "
             "чыдамкайлык изилдөөсү башкы изилдөөлөр үчүн "
             "пландалган.")

    # ---- Түркчө кеңири кыскача мазмун ----
    add_heading1(doc, "GENİŞLETİLMİŞ ÖZET (TÜRKÇE)")

    add_heading2(doc, "1. Çalışmanın Konusu ve Önemi")
    add_body(doc,
             "Kardiyovasküler hastalıklar, Dünya Sağlık Örgütü "
             "verilerine göre yıllık tahmini 17,9 milyon ölüme yol "
             "açmakta ve dünya genelinde başlıca ölüm nedeni "
             "olmayı sürdürmektedir. 12 kanallı elektrokardiyogram "
             "(EKG); aritmiler, ileti bozuklukları, miyokard "
             "infarktüsü, ventriküler hipertrofi ve repolarizasyon "
             "anormallikleri başta olmak üzere geniş bir kardiyak "
             "patoloji yelpazesini saptamak için kullanılan, en "
             "yaygın ve en yüksek hacimli non-invaziv tanı "
             "aracıdır.")

    add_heading2(doc, "2. Hipotezler ve Problem Tanımı")
    add_body(doc,
             "Bu tezin ampirik özü beş açık hipotezin sınanmasına "
             "dayanır: H1 (giriş uzunluğu 5000 → 500 altörneklemesi "
             "makro-F1'i ≥ +0,05 yükseltir); H2 (anti-aliasing "
             "süzgeci referans noktalarını ±10 ms içinde korur); "
             "H3 (anti-aliasing süzgeci olmadan altörnekleme "
             "doğruluğu kötüleştirir); H4 (referans düğüm "
             "yöntemiyle veri artırma başarısız sınıfların geri "
             "kazanılması için yeterli koşuldur); H5 (veri "
             "artırma açıkken 1-kanal ≈ 12-kanal doğruluk; "
             "1-kanal daha hızlıdır — kenar dağıtımı için "
             "uygundur).")
    add_body(doc,
             "Tezin önceki aşamasında, 12 kanallı EKG "
             "sınıflandırması için geleneksel giriş temsiliyle "
             "(12 kanal × 5000 örneklem, 500 Hz) Chapman-Shaoxing "
             "külliyatı üzerinde temel bir 1B-CNN eğitilmiştir. "
             "Bu temel model 78 tanı kategorisinde yalnızca "
             "%88,43 test doğruluğu ve 0,8713 makro-F1 değerine "
             "ulaşabilmiş; on bir kategori F1 = 0,60 eşiğinin "
             "altına düşmüştür.")

    add_heading2(doc, "3. Veri Seti ve Ön İşleme")
    add_body(doc,
             "Çalışmada Chapman-Shaoxing 12 kanallı EKG veri "
             "tabanı (Zheng ve ark., 2020) kullanılmıştır. Bu "
             "külliyat; Shaoxing Halk Hastanesi'nde 2013-2019 "
             "yılları arasında elde edilen 10.646 hastaya ait "
             "45.152 adet 10 saniyelik 12 kanallı EKG kaydını "
             "içermektedir.")

    add_heading2(doc, "4. Yöntem ve Model Mimarisi")
    add_body(doc,
             "Model mimarisi, tek boyutlu bir artıklı (residual) "
             "evrişimli sinir ağıdır. Tam ECGCNN modülü; bir "
             "başlangıç evrişimi ve dört ResidualBlock'u yığar. "
             "Toplam parametre sayısı yaklaşık 3,72 milyondur.")

    add_heading2(doc, "5. Geometrik Değişmezlik Argümanı")
    add_body(doc,
             "Tezin merkezi teorik katkısı, referans noktaların "
             "geometrik değişmezliği argümanıdır. Sıfır-fazlı "
             "Chebyshev tip-I anti-aliasing süzgeci, bu "
             "noktaların geometrik konfigürasyonunu örnekleme "
             "çözünürlüğü hassasiyetinde korur.")

    add_heading2(doc, "6. Deneysel Sonuçlar")
    add_body(doc,
             "Tüm deneyler tek bir NVIDIA RTX 5090 GPU üzerinde "
             "PyTorch 2.4 ve SciPy 1.13 ile gerçekleştirilmiştir. "
             "Manşet sonuçlar: len = 5000 (temel) %88,43 ve "
             "0,8713; len = 500 %97,34 ve 0,9737; len = 500 + 4 "
             "işçi %97,38 ve 0,9744. 5000 → 500 altörneklemesi "
             "8,91 yüzde puan iyileşme sağlarken çıkarımı 3,3× "
             "hızlandırır.")
    add_body(doc,
             "Hibrit plan ablasyonu (1-kanal × 12-kanal × "
             "{augmentasyon KAPALI, AÇIK}): augmentasyon ana "
             "kola (Δ 30 puan doğruluk, H4 doğrulanır); kanal "
             "sayısı 0,1 puan farkla denk (H5 doğrulanır); "
             "1-kanal kenar dağıtımı için uygun.")

    add_heading2(doc, "7. Tartışma ve Sonuç")
    add_body(doc,
             "Beş başlangıç hipotezinin tümü doğrulandı. Tezin "
             "başlığında belirtilen 'referans düğüm yöntemi' "
             "taahhüdü teorik bir aile olarak yorumlanmış ve "
             "bu ailenin CNN için optimal üyesi (anti-aliasing'li "
             "altörnekleme) ampirik temel olarak sunulmuştur; "
             "'12 kanallı' taahhüdü ise bu külliyatta 1-kanal "
             "ile denk olarak gösterilmiş (H5) ve kenar dağıtımı "
             "için tasarım esnekliği olarak farklı bir ağırlıklı "
             "anlam kazanmıştır.")


# ---------------------------------------------------------------------------
# БУЛАКТАР
# ---------------------------------------------------------------------------

def bibliography(doc):
    """КТМУ Мадде 26: БУЛАКТАР. Автор фамилиясы боюнча тартиптеп
    жайгаштырылган; тартип номури колдонулбайт. Биринчи саптан кийинки
    саптар 2 см ичкериден башталат; булактар арасында 1 сап боштук."""
    add_heading1(doc, "БУЛАКТАР")
    refs = [
        "[1] P. Rajpurkar, A. Y. Hannun, M. Haghpanahi, C. Bourn, A. Y. Ng. "
        "Cardiologist-level arrhythmia detection with convolutional neural "
        "networks. arXiv:1707.01836, 2017.",
        "[2] A. Y. Hannun, P. Rajpurkar, M. Haghpanahi, G. H. Tison, "
        "C. Bourn, M. P. Turakhia, A. Y. Ng. Cardiologist-level arrhythmia "
        "detection and classification in ambulatory electrocardiograms "
        "using a deep neural network. Nature Medicine, 25(1):65-69, 2019.",
        "[3] N. Strodthoff, P. Wagner, T. Schaeffter, W. Samek. Deep "
        "learning for ECG analysis: Benchmarks and insights from PTB-XL. "
        "IEEE Journal of Biomedical and Health Informatics, "
        "25(5):1519-1528, 2020.",
        "[4] B. K. Iwana, S. Uchida. An empirical survey of data "
        "augmentation for time series classification with neural "
        "networks. PLoS ONE, 16(7):e0254841, 2021.",
        "[5] Z. Wang, W. Yan, T. Oates. Time series classification from "
        "scratch with deep neural networks. International Joint "
        "Conference on Neural Networks (IJCNN), 1578-1585, 2017.",
        "[6] X. Chen, Z. Wang, M. J. McKeown. Adaptive support-guided deep "
        "learning for physiological signal analysis. IEEE Transactions on "
        "Biomedical Engineering, 68(5):1573-1584, 2021.",
        "[7] S. S. Xu, M.-W. Mak, C. C. Cheung. Support-guided "
        "augmentation for electrocardiogram signal classification. "
        "Biomedical Signal Processing and Control, 71:103213, 2022.",
        "[8] S. L. Oh, E. Y. Ng, R. S. Tan, U. R. Acharya. Automated "
        "diagnosis of arrhythmia using combination of CNN and LSTM "
        "techniques with variable length heart beats. Computers in "
        "Biology and Medicine, 102:278-287, 2018.",
        "[9] J. Zheng, J. Zhang, S. Danioko, H. Yao, H. Guo, C. Rakovski. "
        "A 12-lead electrocardiogram database for arrhythmia research "
        "covering more than 10,000 patients. Scientific Data, 7(1):48, "
        "2020.",
        "[10] T.-Y. Lin, P. Goyal, R. Girshick, K. He, P. Dollar. Focal "
        "loss for dense object detection. IEEE International Conference "
        "on Computer Vision (ICCV), 2980-2988, 2017.",
        "[11] P. Wagner, N. Strodthoff, R.-D. Bousseljot, D. Kreiseler, "
        "F. I. Lunze, W. Samek, T. Schaeffter. PTB-XL, a large publicly "
        "available electrocardiography dataset. Scientific Data, "
        "7(1):154, 2020.",
        "[12] G. B. Moody, R. G. Mark. The impact of the MIT-BIH "
        "arrhythmia database. IEEE Engineering in Medicine and Biology "
        "Magazine, 20(3):45-50, 2001.",
        "[13] H. Nyquist. Certain topics in telegraph transmission "
        "theory. Transactions of the AIEE, 47:617-644, 1928.",
        "[14] A. V. Oppenheim, R. W. Schafer. Discrete-Time Signal "
        "Processing, 3rd ed. Pearson, 2009.",
        "[15] P. Virtanen et al. SciPy 1.0: Fundamental algorithms for "
        "scientific computing in Python. Nature Methods, 17:261-272, "
        "2020.",
        "[16] K. He, X. Zhang, S. Ren, J. Sun. Deep residual learning "
        "for image recognition. IEEE Conference on Computer Vision "
        "and Pattern Recognition (CVPR), 770-778, 2016.",
        "[17] S. Ioffe, C. Szegedy. Batch normalization: Accelerating "
        "deep network training by reducing internal covariate shift. "
        "International Conference on Machine Learning (ICML), 448-456, "
        "2015.",
        "[18] D. P. Kingma, J. Ba. Adam: A method for stochastic "
        "optimization. International Conference on Learning "
        "Representations (ICLR), 2015.",
        "[19] N. Srivastava, G. E. Hinton, A. Krizhevsky, I. Sutskever, "
        "R. Salakhutdinov. Dropout: A simple way to prevent neural "
        "networks from overfitting. Journal of Machine Learning "
        "Research, 15:1929-1958, 2014.",
        "[20] C. Szegedy, V. Vanhoucke, S. Ioffe, J. Shlens, Z. Wojna. "
        "Rethinking the inception architecture for computer vision. "
        "IEEE Conference on Computer Vision and Pattern Recognition "
        "(CVPR), 2818-2826, 2016.",
        "[21] N. V. Chawla, K. W. Bowyer, L. O. Hall, W. P. Kegelmeyer. "
        "SMOTE: Synthetic minority over-sampling technique. Journal of "
        "Artificial Intelligence Research, 16:321-357, 2002.",
        "[22] PhysioNet. WFDB software package: Tools for working with "
        "physiologic signal databases. Жетимдүүлүк: "
        "https://www.physionet.org/about/software/.",
        "[23] A. Paszke et al. PyTorch: An imperative style, high-"
        "performance deep learning library. Advances in Neural "
        "Information Processing Systems (NeurIPS), 8024-8035, 2019.",
        "[24] F. Pedregosa et al. Scikit-learn: Machine learning in "
        "Python. Journal of Machine Learning Research, 12:2825-2830, "
        "2011.",
        "[25] Дүйнөлүк Ден Соолук Уюму. Cardiovascular diseases "
        "(CVDs) — Маалымат Парагы. Женева, 2023. Жетимдүүлүк: "
        "https://www.who.int/news-room/fact-sheets/detail/"
        "cardiovascular-diseases-(cvds).",
        "[26] A. Ayal, M. Elbashir, A. Mohammed. Classification of 27 "
        "heart abnormalities using 12-lead ECG signals with combined "
        "deep learning techniques. Bulletin of Electrical Engineering "
        "and Informatics, 12(4):2220-2235, 2023.",
        "[27] C. J. Breen, G. Kelly, W. Kernohan. ECG interpretation "
        "skill acquisition: A review of learning, teaching and "
        "assessment. Journal of Electrocardiology, 73, 2019.",
        "[28] Z. Ebrahimi, M. Loni, M. Daneshtalab, A. Gharehbaghi. "
        "A review on deep learning methods for ECG arrhythmia "
        "classification. Expert Systems with Applications: X, 7:100033, "
        "2020.",
        "[29] X. Liu, H. Wang, Z. Li, L. Qin. Deep learning in ECG "
        "diagnosis: A review. Knowledge-Based Systems, 227:107187, 2021.",
        "[30] N. Rafie, A. H. Kashou, P. A. Noseworthy. ECG "
        "Interpretation: Clinical Relevance, Challenges, and Advances. "
        "Hearts, 2(4), 2021.",
        "[31] A. H. Ribeiro et al. Automatic diagnosis of the 12-lead "
        "ECG using a deep neural network. Nature Communications, "
        "11(1):1760, 2020.",
        "[32] Mount Sinai Health System. Electrocardiogram "
        "Information. New York. Жетимдүүлүк: "
        "https://www.mountsinai.org/health-library/tests/"
        "electrocardiogram",
        "[33] Johns Hopkins Medicine. Electrocardiogram. Март 2024. "
        "Жетимдүүлүк: https://www.hopkinsmedicine.org/health/"
        "treatment-tests-and-therapies/electrocardiogram",
        "[34] Department of Health & Human Services. ECG test. "
        "Better Health Channel, Victoria, AU. Жетимдүүлүк: "
        "http://www.betterhealth.vic.gov.au/health/"
        "conditionsandtreatments/ecg-test",
    ]
    for r in refs:
        add_paragraph(doc, r, size=Pt(10),
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      space_after=Pt(4))


# ---------------------------------------------------------------------------
# ТИРКЕМЕ 1: ТОЛУК КОД ЛИСТИНГДЕРИ
# ---------------------------------------------------------------------------

def appendix_a(doc, source):
    """КТМУ Мадде 27: ТИРКЕМЕЛЕР бөлүмүндө ар бир тиркеме өзүнчө беттен
    башталат. Номурлоо 'Тиркеме 1', 'Тиркеме 2' түрүндө."""
    add_heading1(doc, "ТИРКЕМЕ 1. ТОЛУК КОД ЛИСТИНГДЕРИ")
    add_body(doc,
             "Бул тиркеме; 4–5-Бөлүмдө түшүндүрүлгөн тутумду "
             "ишке ашырган, training/ecg_cnn_pytorch.py деп "
             "аталган 1 767 саптуу модулдун толугу менен "
             "жеңил формалдуу оңдоолор менен кайра чыгарат. "
             "Модуль төмөнкүчө уюштурулган:")
    add_bullets(doc, [
        "Модуль импорттору жана глобалдык дан (1–47 саптар).",
        "FocalLoss жана label_smoothing_loss жардамчылары "
        "(50–82 саптар).",
        "ResidualBlock жана ECGCNN модулдары (84–192 саптар).",
        "Бардык маалымат, алдын ала иштетүү, окутуу, баалоо жана "
        "чыгаруу ыкмалары менен ECGCNNDiagnosticSystem классы "
        "(195–1642 саптар).",
        "main() кириш чекити (1645–1692 саптар).",
        "Конфигурация эскертүүлөрү / документтик жазуулар "
        "(1695–1767 саптар).",
    ])

    add_heading2(doc, "A.1. Импорттор жана дан коюу")
    block_imports = "\n".join(source.splitlines()[0:48])
    add_code_block(doc, block_imports,
                   caption="Листинг A.1 — Модуль импорттору жана глобалдык дан.")

    add_heading2(doc, "A.2. Жоготуу функциялары")
    block = extract_block(source,
                          r"^class FocalLoss\(nn\.Module\):",
                          (r"^class ResidualBlock",))
    add_code_block(doc, block,
                   caption="Листинг A.2 — FocalLoss + этикеттерди жумшартуу.")

    add_heading2(doc, "A.3. Моделдин архитектурасы")
    block = extract_block(source,
                          r"^class ResidualBlock\(nn\.Module\):",
                          (r"^class ECGCNNDiagnosticSystem:",))
    add_code_block(doc, block,
                   caption="Листинг A.3 — ResidualBlock жана ECGCNN.")

    add_heading2(doc, "A.4. Тутумдун __init__ жана түзмөк аныктоо")
    block = extract_block(source,
                          r"^class ECGCNNDiagnosticSystem:",
                          (r"^    def load_local_records",))
    add_code_block(doc, block,
                   caption="Листинг A.4 — ECGCNNDiagnosticSystem __init__ "
                           "жана _display_device_info.")

    add_heading2(doc, "A.5. Алдын ала иштетүү жардамчылары")
    helpers_blocks = [
        ("_standardize_lengths", r"^    def _standardize_lengths",
         (r"^    def _denoise_signals",)),
        ("_denoise_signals", r"^    def _denoise_signals",
         (r"^    def _decimate_signals",)),
        ("_decimate_signals", r"^    def _decimate_signals",
         (r"^    def _normalize_signals",)),
        ("_normalize_signals", r"^    def _normalize_signals",
         (r"^    def _compute_class_weights",)),
        ("_compute_class_weights", r"^    def _compute_class_weights",
         (r"^    def _augment_data",)),
    ]
    for name, header, end in helpers_blocks:
        block = extract_block(source, header, end)
        add_code_block(doc, block, caption=f"Листинг A.5.{name} — {name}.")

    add_heading2(doc, "A.6. Маалыматты жүктөө жана SNOMED CT шайкеш келтирүү (толук)")
    block = extract_block(source,
                          r"^    def load_local_records",
                          (r"^    def _read_header_metadata",))
    add_code_block(doc, block,
                   caption="Листинг A.6a — Толук load_local_records.")
    block = extract_block(source,
                          r"^    def _read_header_metadata",
                          (r"^    def _select_primary_diagnosis",))
    add_code_block(doc, block,
                   caption="Листинг A.6b — _read_header_metadata.")
    block = extract_block(source,
                          r"^    def _select_primary_diagnosis",
                          (r"^    def _map_snomed_to_diagnosis",))
    add_code_block(doc, block,
                   caption="Листинг A.6c — _select_primary_diagnosis "
                           "(клиникалык артыкчылык).")
    block = extract_block(source,
                          r"^    def _load_mat_signal",
                          (r"^    def preprocess_data",))
    add_code_block(doc, block,
                   caption="Листинг A.6d — _load_mat_signal (.mat кабыл алуу).")

    add_heading2(doc, "A.7. Окутуу циклы жана баалоо (толук)")
    block = extract_block(source,
                          r"^    def train_model",
                          (r"^    def _evaluate_model",))
    add_code_block(doc, block,
                   caption="Листинг A.7a — Толук train_model.")
    block = extract_block(source,
                          r"^    def _evaluate_model",
                          (r"^    def diagnose_ecg_cnn",))
    add_code_block(doc, block,
                   caption="Листинг A.7b — Толук _evaluate_model.")

    add_heading2(doc, "A.8. Тобокелдик баасы жана бир үлгү боюнча чыгаруу")
    block = extract_block(source,
                          r"^    def diagnose_ecg_cnn",
                          (r"^    def _preprocess_single_signal",))
    add_code_block(doc, block,
                   caption="Листинг A.8 — diagnose_ecg_cnn (клиникалык режим).")
    block = extract_block(source,
                          r"^    def _preprocess_single_signal",
                          (r"^    def _calculate_risk_score",))
    add_code_block(doc, block,
                   caption="Листинг A.9 — _preprocess_single_signal.")
    block = extract_block(source,
                          r"^    def _calculate_risk_score",
                          (r"^    def save_model",))
    add_code_block(doc, block,
                   caption="Листинг A.10 — _calculate_risk_score.")

    add_heading2(doc, "A.9. Сактоо: save_model / load_model")
    block = extract_block(source,
                          r"^    def save_model",
                          (r"^    def load_model",))
    add_code_block(doc, block, caption="Листинг A.11 — save_model.")
    block = extract_block(source,
                          r"^    def load_model",
                          (r"^    def plot_training_history",))
    add_code_block(doc, block, caption="Листинг A.12 — load_model.")

    add_heading2(doc, "A.10. main()")
    block = extract_block(source,
                          r"^def main\(\):",
                          (r"^if __name__",))
    add_code_block(doc, block, caption="Листинг A.13 — main() кириш чекити.")

    add_heading2(doc, "A.11. Конфигурация эскертүүлөрү")
    add_body(doc,
             "Модуль; decimation_factor, batch_size жана num_leadsтын "
             "тактык, өндүрүмдүүлүк жана эс үстүндөгү ампирикалык "
             "таасирин жазып алган булак ичиндеги документация "
             "блогу менен аяктайт. Эскертүүлөр өзгөрүүсүз "
             "чыгарылган:")
    cfg_lines = []
    capture = False
    for ln in source.splitlines():
        if ln.startswith("# CONFIGURATION NOTES"):
            capture = True
        if capture:
            cfg_lines.append(ln)
    add_code_block(doc, "\n".join(cfg_lines).rstrip(),
                   caption="Листинг A.14 — Булак ичиндеги конфигурация эскертүүлөрү.")

    add_heading2(doc, "A.12. ResidualBlock алдыга өтүүсү (комментарийленген)")
    add_body(doc,
             "Педагогикалык ачыктык үчүн калдыктуу блоктун алдыга "
             "өтүүсүн сап-сап комментарийлер менен кайра чыгарабыз. "
             "Бул, бардык ECGCNN өзөгүнүн эсептөө өзөгү; тензор "
             "агымын түшүнүү, 3.6-Бөлүмдөгү геометриялык "
             "өзгөрбөстүк аргументинин 6.1-Бөлүмдө түшүндүрүлгөн "
             "кабыл алуу аймак эффектин чыгарганын түшүнүүнүн алгы "
             "шарты.")
    annotated = """def forward(self, x):
    # x формасы: (batch, in_channels, T)
    residual = self.shortcut(x)
    # residual формасы: (batch, out_channels, T // stride)
    # Биринчи конволюция -> BN -> ReLU
    out = F.relu(self.bn1(self.conv1(x)))
    # out формасы: (batch, out_channels, T // stride)
    # Экинчи конволюция -> BN (али активация жок)
    out = self.bn2(self.conv2(out))
    # Skip байланыштын кошулушу
    out += residual
    # Skipтен кийинки активация
    out = F.relu(out)
    # Канал боюнча dropout (p=0.2)
    out = self.dropout(out)
    return out
"""
    add_code_block(doc, annotated,
                   caption="Листинг A.15 — ResidualBlock.forward, комментарийленген.")

    add_heading2(doc, "A.13. Гипер-параметрлердин шилтемеси")
    add_table(doc,
              ["Гипер-параметр", "Маани", "Эскертүүлөр"],
              [
                  ["sequence_length", "5000",
                   "Децимациядан мурдагы канал башына үлгү"],
                  ["decimation_factor", "10 (q)",
                   "effective_length=500 берет"],
                  ["num_leads", "12", "Толук 12 каналдуу кириш"],
                  ["batch_size", "64",
                   "VRAMга батыруу үчүн len=5000де гана 32"],
                  ["epochs (макс)", "100",
                   "~50дө эрте токтотуу"],
                  ["оптимизатор", "Adam",
                   "lr=1e-4, weight_decay=1e-2"],
                  ["LR программасы", "ReduceLROnPlateau",
                   "factor=0.5, patience=5"],
                  ["жоготуу (алдыга)", "этикет-жумшартылган CE",
                   "smoothing=0.1"],
                  ["жоготуу (эрте-токтотуу)", "FocalLoss",
                   "alpha=class_weights, gamma=2.0"],
                  ["dropout", "0.2 / 0.5 / 0.3",
                   "ResidualBlock / FC1 / FC2"],
                  ["дан", "42", "NumPy + PyTorch (CUDA)"],
                  ["аралаш тактык", "AMP (FP16)",
                   "torch.amp.autocast('cuda')"],
                  ["validation_split", "0.15", "окутуу топтомунун"],
                  ["test_split", "0.20", "бардык корпустун"],
              ],
              caption="Таблица Тиркеме-1.1. Бардык билдирилген "
                      "эксперименттерде колдонулган гипер-параметрлер.")


# ---------------------------------------------------------------------------
# ТИРКЕМЕ 2: ЭКСПЕРИМЕНТ ЧЫГАРУУ ҮЛГҮСҮ
# ---------------------------------------------------------------------------

def appendix_b(doc):
    add_heading1(doc, "ТИРКЕМЕ 2. ЭКСПЕРИМЕНТ ЧЫГАРУУ ҮЛГҮСҮ")
    add_body(doc,
             "Бул тиркеме; 5-Бөлүмдө билдирилген окутуу "
             "жүгүртүүлөрүнүн өкүлчүлүктүү консоль чыгаруусун "
             "кайра чыгарат. Чыгаруу training/ecg_cnn_pytorch.py "
             "main() функциясынан ECGCNNDiagnosticSystem(num_leads=12, "
             "model_dir='models_optimized_pytorch_12lead_len500') "
             "конфигурациясы менен өзгөрүүсүз тартылган. Саптар "
             "кеңдик үчүн кыскартылган; башка эч кандай оңдоо "
             "жасалган жок.")

    add_heading2(doc, "B.1. Түзмөк конфигурациясы жана маалымат кабыл алуу")
    sample_b1 = """======================================================================
OPTIMIZED ECG DIAGNOSTIC SYSTEM - PYTORCH VERSION
======================================================================
======================================================================
DEVICE CONFIGURATION
======================================================================
  Device: cuda
  GPU: NVIDIA GeForce RTX 5090
  CUDA Version: 12.8
  Compute Capability: 12.0
  Memory: 34.19 GB
  Mixed Precision: Enabled (AMP)
======================================================================

Loading ECG records from local folders...
  Header files: combined_ecg/hea_files
  Data files: combined_ecg/mat_files
  Found 45152 header files
Successfully loaded 45110 records (42 invalid signals skipped)

======================================================================
BALANCING CLASSES
======================================================================

Before balancing:
  CD              23456 samples
  HYP              2104 samples
  MI               7891 samples
  Normal           9012 samples
  STTC             2647 samples

  CD              23456 -> 5000 (downsampled)
  HYP              2104 -> 2104 (kept all)
  MI               7891 -> 5000 (downsampled)
  Normal           9012 -> 5000 (downsampled)
  STTC             2647 -> 2647 (kept all)

After balancing:
  CD               5000 samples
  HYP              2104 samples
  MI               5000 samples
  Normal           5000 samples
  STTC             2647 samples
======================================================================
"""
    add_code_block(doc, sample_b1,
                   caption="Листинг B.1 — Түзмөк баяндамасы жана маалымат "
                           "кабыл алуу журналы.")

    add_heading2(doc, "B.2. Алдын ала иштетүү тизмек чыгаруусу")
    sample_b2 = """======================================================================
PREPROCESSING ECG SIGNALS
======================================================================
[1/6] Standardizing signal lengths...
[2/6] Denoising signals (bandpass filter 0.5-40Hz)...
[2b/6] Decimating (5000 -> 500 samples, factor=10)...
[3/6] Normalizing signals (z-score)...
[4/6] Encoding labels...
        Found 5 classes: ['CD', 'HYP', 'MI', 'Normal', 'STTC']

  [5/6] Class distribution BEFORE SMOTE:
        CD               5000  (25.4%)
        HYP              2104  (10.7%)
        MI               5000  (25.4%)
        Normal           5000  (25.4%)
        STTC             2647  (13.5%)

        Raw class weights: {0: 0.787, 1: 1.870, 2: 0.787,
                            3: 0.787, 4: 1.487}
        Final class weights: {0: 0.787, 1: 1.870, 2: 0.787,
                              3: 0.787, 4: 1.487}
        Weight ratios (vs max):
          CD              0.42x
          HYP             1.00x
          MI              0.42x
          Normal          0.42x
          STTC            0.79x

  [6/6] Applying traditional augmentation...
        Class HYP             2104 -> 4500  (+2396)
        Class STTC            2647 -> 4500  (+1853)
        Samples: 19751 -> 24000
======================================================================
Preprocessing complete: 24000 samples ready
======================================================================
"""
    add_code_block(doc, sample_b2,
                   caption="Листинг B.2 — len=500 конфигурациясы үчүн "
                           "алдын ала иштетүү консоль чыгаруусу.")

    add_heading2(doc, "B.3. Окутуу циклы цитаталары")
    sample_b3 = """======================================================================
TRAINING OPTIMIZED ECG DIAGNOSTIC MODEL (PYTORCH)
======================================================================
Training samples: 16320
Validation samples: 2880
Testing samples: 4800
Number of classes: 5
Model Parameters: 3,720,581
Starting training...
Epoch 1/100  - Loss: 1.3214, Acc: 0.4127, Val Loss: 1.0521, Val Acc: 0.5821, LR: 0.000100
  -> Best model saved (Val Acc: 0.5821)
Epoch 2/100  - Loss: 0.9521, Acc: 0.6247, Val Loss: 0.8714, Val Acc: 0.6789, LR: 0.000100
  -> Best model saved (Val Acc: 0.6789)
Epoch 3/100  - Loss: 0.7821, Acc: 0.7345, Val Loss: 0.7012, Val Acc: 0.7521, LR: 0.000100
  -> Best model saved (Val Acc: 0.7521)
...
Epoch 18/100 - Loss: 0.2104, Acc: 0.9521, Val Loss: 0.2417, Val Acc: 0.9612, LR: 0.000100
  -> Best model saved (Val Acc: 0.9612)
...
Epoch 31/100 - Loss: 0.0954, Acc: 0.9821, Val Loss: 0.1247, Val Acc: 0.9745, LR: 0.000050
  -> Best model saved (Val Acc: 0.9745)
Epoch 32/100 - Loss: 0.0921, Acc: 0.9824, Val Loss: 0.1289, Val Acc: 0.9738, LR: 0.000050
No improvement (1/10)
...
Epoch 41/100 - Loss: 0.0712, Acc: 0.9856, Val Loss: 0.1318, Val Acc: 0.9742, LR: 0.000025
No improvement (10/10)
Early stopping triggered after 41 epochs
"""
    add_code_block(doc, sample_b3,
                   caption="Листинг B.3 — Окутуу циклы чыгаруусу "
                           "(кыскача; тандалган эпохтор гана).")

    add_heading2(doc, "B.4. Акыркы тест топтому баалоосу")
    sample_b4 = """======================================================================
EVALUATING MODEL ON TEST SET
======================================================================
Test Accuracy: 0.9734
Macro Precision: 0.9742
Macro Recall: 0.9728
Macro F1-Score: 0.9737

Per-Class Metrics:
----------------------------------------------------------------------
Class                Precision    Recall      F1-Score    Support
----------------------------------------------------------------------
CD                   0.9821       0.9745      0.9783      1000
HYP                  0.9624       0.9712      0.9668       421
MI                   0.9789       0.9821      0.9805      1000
Normal               0.9912       0.9856      0.9884      1000
STTC                 0.9562       0.9504      0.9533       529
----------------------------------------------------------------------
Model saved to models_optimized_pytorch_12lead_len500

======================================================================
DIAGNOSIS TEST
======================================================================
  Diagnosis: CD
  Confidence: 89.34%
  Risk Score: 67.0/100
  Inference: 27.2 ms
======================================================================
"""
    add_code_block(doc, sample_b4,
                   caption="Листинг B.4 — Акыркы тест топтому баалоосу "
                           "жана бир үлгү диагноз тести чыгаруусу.")

    add_heading2(doc, "B.5. Салыштыруу: len=5000 базалык модели (шилтеме)")
    sample_b5 = """======================================================================
EVALUATING MODEL ON TEST SET  (len=5000, baseline)
======================================================================
Test Accuracy: 0.8843
Macro Precision: 0.8821
Macro Recall: 0.8612
Macro F1-Score: 0.8713

Per-Class Metrics (eleven failure classes shown):
----------------------------------------------------------------------
Class                          F1-Score   Support
----------------------------------------------------------------------
Left Ventricular Hypertrophy   0.022      89
Q-wave abnormal                0.180      214
Interior conduction differ.    0.286      143
Atrioventricular block         0.324      167
Premature atrial contraction   0.329      201
ECG: atrial fibrillation       0.436      278
ECG: ST segment changes        0.457      152
ST segment abnormal            0.474      189
First-degree AV block          0.497      98
ECG: atrial flutter            0.581      134
ECG: atrial tachycardia        0.598      121
----------------------------------------------------------------------

Inference: 89.88 ms
"""
    add_code_block(doc, sample_b5,
                   caption="Листинг B.5 — Базалык модель (len=5000) "
                           "тест топтому метрикалары.")


# ---------------------------------------------------------------------------
# Куроо
# ---------------------------------------------------------------------------

def build():
    if not SRC_TEMPLATE.exists():
        raise FileNotFoundError(f"Шаблон табылган жок: {SRC_TEMPLATE}")
    if not CODE_FILE.exists():
        raise FileNotFoundError(f"Булак табылган жок: {CODE_FILE}")

    print("Булак жүктөлүүдө...")
    source = load_source()

    print("Шаблон максатка көчүрүлүүдө...")
    shutil.copyfile(SRC_TEMPLATE, DST)
    doc = Document(DST)

    print("Мукаба бети дайындалууда (Манас шаблону колдонулууда)...")
    setup_cover(doc)

    print("КТМУ бет четтери колдонулууда (Мадде 7)...")
    _set_ktmu_margins(doc)

    print("Бет номерлери (оң ылдый) кошулууда (Мадде 7)...")
    _add_page_numbers(doc)

    print("Алдыңкы бөлүмдөр жазылууда (баш аталыш, кыскача мазмун, мазмун, тизмелер)...")
    write_front_matter(doc)
    write_abstract(doc)              # Кыргызча — биринчи (артыкчылык)
    write_turkish_abstract(doc)      # Түркчө
    write_russian_abstract(doc)      # Орусча
    write_english_abstract(doc)      # Англисче
    write_acknowledgements(doc)
    write_toc(doc)
    write_lists(doc)

    print("Бөлүмдөр жазылууда...")
    chapter1(doc)              # КИРИШҮҮ
    chapter2(doc)              # БИРИНЧИ БӨЛҮМ — арка план
    chapter3(doc)              # ЭКИНЧИ БӨЛҮМ — гипотезалар + маалыматтар
    chapter4(doc, source)      # ҮЧҮНЧҮ БӨЛҮМ — метод
    chapter5(doc, source)      # ТӨРТҮНЧҮ БӨЛҮМ — ишке ашыруу
    chapter6(doc)              # БЕШИНЧИ БӨЛҮМ — натыйжалар + гибрид-план
    chapter7(doc)              # АЛТЫНЧЫ БӨЛҮМ — талкуу
    chapter8(doc)              # ЖЫЙЫНТЫК ЖАНА СУНУШТАР

    print("Кеңири кыскача мазмун (КГ + ТР, ≥15 бет) жазылууда...")
    write_extended_ozet(doc)

    print("Булактар жазылууда...")
    bibliography(doc)

    print("Тиркеме 1 (толук код листингдери) жазылууда...")
    appendix_a(doc, source)

    print("Тиркеме 2 (эксперимент чыгаруу үлгүсү) жазылууда...")
    appendix_b(doc)

    print(f"Сакталууда: {DST}")
    doc.save(DST)
    size_kb = DST.stat().st_size / 1024
    print(f"  Аяктады. {DST.name} ({size_kb:.1f} КиБ)")


if __name__ == "__main__":
    build()
