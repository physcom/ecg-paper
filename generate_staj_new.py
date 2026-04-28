"""Generate a new 16-week staj plan docx by replacing weekly content in the original,
preserving format, headers, styles. Topic: Agentic Risk & Analytics Platform (tax-agent)."""

import copy
import re
from pathlib import Path

from docx import Document

SRC = Path("staj-Назаркулов Эламан.docx")
DST = Path("staj-Назаркулов Эламан - agentic platform.docx")

WEEKS = [
    # 1
    "Практиканын биринчи жумасында мен тармактык компаниянын алкагында иштелип жаткан "
    "\"Agentic Risk & Analytics Platform\" долбоорунун жалпы архитектурасы менен таанышып, "
    "иштеп чыгуу чөйрөсүн даярдадым. Python 3.11+ орнотуп, conda виртуалдык чөйрөсүн түзүп, "
    "pyproject.toml аркылуу көз карандылыктарды орноттум. FastAPI, Pydantic v2, SQLAlchemy 2.0 "
    "(async) жана Anthropic Claude SDK негизги стек катары кабылданды. Docker Desktop жана "
    "docker-compose аркылуу PostgreSQL, Redis, Qdrant, ClickHouse сервистерин жергиликтүү "
    "ишке киргизип, проекттин CLAUDE.md жана README.md документтерин деталдуу изилдедим.",
    # 2
    "Экинчи жумада долбоордун негизги абстракцияларын, башкача айтканда core катмарын "
    "үйрөндүм. BaseAgent жана BaseTool абстракт класстары, AgentConfig, ToolConfig жана "
    "ExecutionContext структуралары менен таанышып, агенттер жана куралдардын жашоо "
    "циклин, алардын натыйжаларын (AgentResult, ToolResult) карадым. ExecutionContext "
    "контекст объектиси агенттер жана куралдар ортосунда маалыматты айлантуу үчүн кантип "
    "колдонулганын практикада текшерип, тесттик агент түзүп көрдүм.",
    # 3
    "Үчүнчү жумада аутентификация модулун үйрөндүм. JWT (access 30 мин + refresh 7 күн) "
    "жана bcrypt негизиндеги сырсөз хэштөөсү - passlibден четтеп, түз bcrypt 5.x менен "
    "иштелгенин карадым. FastAPIнин Depends механизминин негизинде get_current_user жана "
    "require_admin көз карандылыктарын түшүнүп, /api/auth/register, /login, /refresh, /me "
    "эндпоинттеринин агымын Postman аркылуу текшерип, коргоого муктаж болгон маршруттарды "
    "байкадым.",
    # 4
    "Төртүнчү жумада long-lived LLM кардар катмарын, башкача айтканда llm_client.py "
    "файлын карадым. create_llm_client() фабрикасы Anthropic API жана OpenAI-шайкеш "
    "endpoint (Ollama, vLLM, LM Studio) ортосунда кантип жасалып алынганын изилдедим. "
    "ROUTER_LLM_BACKEND тууралоосу аркылуу ар түрдүү провайдерлерге которуу практикасын "
    "жасап, жергиликтүү Ollama менен Claude Sonnet ортосунда салыштыруу жүргүзүп көрдүм.",
    # 5
    "Бешинчи жумада IntentRouter (router.py) модулун терең үйрөндүм. Маршрутташтыруу "
    "кантип иштээрин - keyword-based эрежелер (30% салмак) + LLM Chain-of-Thought "
    "классификациялоо (70% салмак) - чечмелеп чыктым. SHA256(query|agent_ids) негизиндеги "
    "TTL-кэш (5 мүнөт) жана score blending формуласы, hard floor эрежелери blending "
    "операциясынан кийин кайра колдонулганын байкадым. Өзүмдүн тесттик суроолорум менен "
    "router.route() функциясын чакырып, классификация натыйжаларын карап чыктым.",
    # 6
    "Алтынчы жумада QueryNormalizer жана SemanticCache модулдарын үйрөндүм. "
    "QueryNormalizer ЭСФ, карыз, мүлк, ETTN сыяктуу ачкыч сөздөр үчүн fast-path "
    "эрежелерди аткарса, SemanticCache sentence-transformers аркылуу векторлук fuzzy "
    "match (cosine similarity ≥ 0.92) жасайт. Иш жүзүндө окшош суроолор кэштен кайтарылып "
    "LLM чыгымдары азайганын көрдүм. Sentence-transformers орнотулбаганда графалуу "
    "деградация (exact-match fallback) механизмин текшердим.",
    # 7
    "Жетинчи жумада MultiAgentPlanner жана MultiAgentExecutor компоненттеринин иш "
    "принцибин изилдедим. Суроодон объекттерди (ИНН, дата, аталыштар, суммалар) "
    "бөлүп чыгаруу, агенттер ортосундагы көз карандылык графын куруу жана "
    "SEQUENTIAL/PARALLEL режимдеринде аткаруу концепцияларын үйрөндүм. ParallelScheduler "
    "семафор, timeout, dependency order аркылуу бир нече агентти параллелдүү аткаргандыгын "
    "иш жүзүндө текшердим.",
    # 8
    "Сегизинчи жумада шарттуу (conditional, эгерде-анда) суроолор түтүгү менен иштедим. "
    "SemanticQueryParser regex gate аркылуу 'если/if-then/mismatch' түрдөгү суроолорду "
    "аныктап, бир LLM чакыруусу менен QueryPlan объектин кайтарат. ConditionalPlanBuilder "
    "ошол QueryPlanды 3-фазалуу MultiAgentPlanга айлантат: Phase 1 - PRIMARY (ЗАГС, ККМ "
    "реестри), Phase 2 - FILTER (параллелдүү филтрлөөчү агенттер), Phase 3 - "
    "ResultCorrelator.",
    # 9
    "Тогузунчу жумада ResultCorrelator жана ChainOfThoughtVerifier модулдарын көрдүм. "
    "ResultCorrelator агенттердин натыйжаларын ИНН боюнча бириктирип, LLM тарабынан "
    "түзүлгөн filter_expr чыпкасын AST sandbox (Compare, BoolOp, Subscript, Constant "
    "гана уруксат берилген) аркылуу коопсуз баалайт. CoT Verifier ар бир IF-THEN "
    "шартын өзүнчө чечмелеп, канааттандырылбаса retry_directive менен кайра аткарууну "
    "буйруйт. Ретрай чексиз циклге кирип кетпеши үчүн 1 раундга чектелгенин түшүндүм.",
    # 10
    "Онунчу жумада берилиштер базасы куралдары катмарын үйрөндүм. PostgresTool (asyncpg, "
    "read-only, parameterized query), MSSQLTool (pyodbc, Driver 17/18 SSL fallback), "
    "ClickHouseTool (TPPayment, ETTN аналитикасы), MongoDBESFTool (motor async) жана "
    "HealthAwareTool wrapper концепциялары менен иштеп, алардын жалпы BaseTool интерфейси "
    "аркылуу кантип бириктирилгенин изилдедим. Ар бир куралга өзүмдүн тест-сурамдарды "
    "жазып, max_rows жана таймаут параметрлерин текшердим.",
    # 11
    "Он биринчи жумада билимдер базасы (Knowledge Base) жана RAG түтүгү менен "
    "тааныштым. Qdrant коллекциясында гибриддик поиск - BAAI/bge-m3 dense (1024-dim) "
    "+ BM42 sparse векторлорунун Reciprocal Rank Fusion (RRF) бириктирүүсү, андан "
    "кийин cross-encoder reranker тарабынан топ-5 кайтаруу. Chunking стратегиясы "
    "(1500/300, similarity 0.65), HyDE query rewriting жана көп тилдүүлүк "
    "(орусча/кыргызча/англисче) колдоосу практикада карап чыгылды.",
    # 12
    "Он экинчи жумада адистештирилген агенттерди (tax_risk, fraud_detection, "
    "analytics, declaration, ETTN, payment, property, tax_debt, taxpayer_info, "
    "knowledge_base) жана алардын ар биринин system prompt'торун изилдедим. "
    "DATA_INTEGRITY_RULES бардык агенттердин prompt'торуна ыргытылып, аттардын же "
    "ИНН-дардын тразитерацияланбашын, кыскартылбашын жана \"оңдолбошун\" кантип "
    "камсыздагандыгын көрдүм. Tax Risk жана Fraud Detection агенттеринин "
    "карусел-алдамчылыкты жана shell-компанияларды табуу логикасы өзгөчө кызыктуу "
    "болду.",
    # 13
    "Он үчүнчү жумада WebSocket real-time жаңылоолору жана сессиялык эс-тутум (Session "
    "Memory) модулдарын үйрөндүм. ConversationSession объекти жеке адамдарды, "
    "компанияларды, автоунааларды жана мүлктү көзөмөлдөп, director_of, founder_of, "
    "employed_at сыяктуу мамилелерди куруп, D3.js форматына экспорттой алат. Redis-TTL "
    "(24 саат) аркылуу колдонуучунун акыркы N сүйлөшүү кезегин сактоо жана Redis жок "
    "учурда graceful no-op деградация логикасын иш жүзүндө текшердим.",
    # 14
    "Он төртүнчү жумада Evaluator модулу, OpenTelemetry жана мониторинг катмарын "
    "изилдедим. Evaluator жооптун completeness, grounding, precision, actionability "
    "өлчөмдөрү боюнча балл коюп, боштуктар табылса максимум N жолу gap-filling retry "
    "жасайт. telemetry.py OTEL_ENABLED=false учурунда нөл-чыгым noop катары иштеп, true "
    "учурда OTLP/gRPC аркылуу трейстерди жиберерин көрдүм. HealthMonitor пул "
    "метрикалары (HEALTHY/DEGRADED/UNHEALTHY) жана JobPoller 202 Accepted паттерни "
    "менен иштөө практикасы жасалды.",
    # 15
    "Он бешинчи жумада Airflow + Spark интеграциясы жана чоң маалыматтарды иштетүү "
    "(batch-processing) сценарийлерин карадым. jobs/ каталогундагы DAGдар "
    "(fraud_detection_dag, risk_scoring_dag, network_analysis_dag) жана ага "
    "тиешелүү PySpark жумуштары аркылуу миллиондогон ЭСФ жазууларын пакет режиминде "
    "эсептөө жана натыйжаларды кайра PostgreSQL/ClickHouseге жазуу түтүгүнүн "
    "архитектурасы үйрөнүлдү. REQUIRE_APPROVAL_FOR_HEAVY желе аркылуу адамдын "
    "бекитүүсүн талап кылуу механизмин карап чыктым.",
    # 16
    "Он алтынчы жумада өндүрүш деңгээлиндеги жайгаштырууну (deployment) даярдадым. "
    "Docker-compose жана k8s-api-deployment.yaml конфигурацияларын окуп чыгып, .env "
    "баалуулуктарын Secret менен алмаштыруу, JWT_SECRET_KEY үчүн коопсуз маанилерди "
    "түзүү, Qdrant collection миграциясы (vector size дал келбегенде авто-кайра "
    "түзүлүшү), бэкап стратегиясын жана өлчөө (load test) сценарийлерин иштеп чыктым. "
    "Практика жыйынтыгында долбоордун толук агымын - колдонуучу суроосунан тартып "
    "жооп кайтарылганга чейин - жекече отладкалап көрүп, документациясын жаңыладым.",
]

FINAL_SUMMARY = (
    "Бул 16 жумалык практиканын негизги жыйынтыгы - көп агенттик "
    "(multi-agent) системанын иштөө принцибин өздөштүрүү. Системанын "
    "ишмердүүлүгү төрт кадамдан турат: (1) IntentRouter колдонуучунун "
    "суроосун классификациялап, ага жооп бере ала турган адистештирилген "
    "агенттерди тандайт; (2) MultiAgentPlanner суроодон ИНН, дата, сумма "
    "сыяктуу объекттерди бөлүп чыгарып, агенттер ортосундагы көз "
    "карандылыктын планын түзөт; (3) MultiAgentExecutor ал агенттерди "
    "(tax risk, fraud detection, analytics, ETTN, property ж.б.) "
    "параллелдүү же ирет менен аткарат; (4) ResultCorrelator алардын "
    "натыйжаларын ИНН боюнча бир жоопко бириктирет. Ушул төрт кадамдын "
    "аркасында бир татаал суроо бир нече агенттин макулдашкан жообуна "
    "айланарын иш жүзүндө көрдүм."
)


def replace_in_paragraph(paragraph, new_text):
    """Replace text in paragraph, preserving the formatting of the first run."""
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return
    first_run = paragraph.runs[0]
    first_run.text = new_text
    for r in paragraph.runs[1:]:
        r.text = ""


def iter_cell_paragraphs(cell):
    for p in cell.paragraphs:
        yield p


def set_cell_text(cell, new_text):
    """Replace all text in a cell with new_text, preserving formatting of first run."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.add_paragraph(new_text)
        return
    first_p = paragraphs[0]
    replace_in_paragraph(first_p, new_text)
    for p in paragraphs[1:]:
        for r in p.runs:
            r.text = ""


def normalize_run_formatting(cell, font_name="Times New Roman", half_points=22, color_hex="000000"):
    """Force every run in `cell` to use a uniform font, size, and color.
    half_points=22 means 11pt (Word stores size in half-points).
    """
    for p in cell.paragraphs:
        for run in p.runs:
            rPr = run._element.find(f"{{{W_NS}}}rPr")
            if rPr is None:
                rPr = etree.SubElement(run._element, f"{{{W_NS}}}rPr")
                run._element.remove(rPr)
                run._element.insert(0, rPr)

            # Font family - set all ascii/hAnsi/cs/eastAsia
            rFonts = rPr.find(f"{{{W_NS}}}rFonts")
            if rFonts is None:
                rFonts = etree.SubElement(rPr, f"{{{W_NS}}}rFonts")
            rFonts.set(f"{{{W_NS}}}ascii", font_name)
            rFonts.set(f"{{{W_NS}}}hAnsi", font_name)
            rFonts.set(f"{{{W_NS}}}cs", font_name)
            rFonts.set(f"{{{W_NS}}}eastAsia", font_name)

            # Size - <w:sz w:val="22"/> and <w:szCs w:val="22"/>
            for tag in ("sz", "szCs"):
                elem = rPr.find(f"{{{W_NS}}}{tag}")
                if elem is None:
                    elem = etree.SubElement(rPr, f"{{{W_NS}}}{tag}")
                elem.set(f"{{{W_NS}}}val", str(half_points))

            # Color - <w:color w:val="000000"/>
            color = rPr.find(f"{{{W_NS}}}color")
            if color is None:
                color = etree.SubElement(rPr, f"{{{W_NS}}}color")
            color.set(f"{{{W_NS}}}val", color_hex)

            # Remove any highlight/shading that may be leftover
            for tag in ("highlight",):
                leftover = rPr.find(f"{{{W_NS}}}{tag}")
                if leftover is not None:
                    rPr.remove(leftover)


def find_week_cells(doc):
    """
    Find for each week (1..16) the corresponding 'task' cell (the middle column).
    Return dict {week_number: cell}.
    """
    week_cells = {}
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            first = cells[0].text.strip()
            m = re.match(r"^(\d+)\s*\.?\s*Hafta", first)
            if m:
                week = int(m.group(1))
                if 1 <= week <= 16:
                    week_cells[week] = cells[1]
    return week_cells


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

from lxml import etree

NSMAP = {"w": W_NS}


def _paragraph_text(p_elem):
    return "".join(t.text or "" for t in p_elem.iter(f"{{{W_NS}}}t"))


def _set_paragraph_text(p_elem, new_text):
    """Replace all text inside an OOXML paragraph element with new_text.
    Preserves formatting of first w:r; removes other runs' text."""
    t_elems = list(p_elem.iter(f"{{{W_NS}}}t"))
    if not t_elems:
        return False
    t_elems[0].text = new_text
    for t in t_elems[1:]:
        t.text = ""
    return True


def find_and_replace_summary_in_textbox(doc, new_text):
    """Locate the summary paragraph inside any text box (wps:txbx) and replace it."""
    body = doc.element.body
    for p_elem in body.iter(f"{{{W_NS}}}p"):
        txt = _paragraph_text(p_elem).strip()
        if txt.startswith("Бул 16 жумалык практикада"):
            _set_paragraph_text(p_elem, new_text)
            return True
    return False


def _first_child(elem, tag):
    found = elem.find(f"{{{W_NS}}}{tag}")
    return found


def _ensure_child_first(parent, tag):
    """Ensure a child element exists as the FIRST child of parent. Return it."""
    existing = parent.find(f"{{{W_NS}}}{tag}")
    if existing is not None:
        # move to first if not already
        if parent.index(existing) != 0:
            parent.remove(existing)
            parent.insert(0, existing)
        return existing
    new_elem = etree.SubElement(parent, f"{{{W_NS}}}{tag}")
    parent.remove(new_elem)
    parent.insert(0, new_elem)
    return new_elem


def apply_table_row_fixes(doc, page_break_before_weeks=(11, 15)):
    """
    For every weekly-task table row:
      - add <w:cantSplit/> to <w:trPr> so the row never splits across pages
    For each week N in page_break_before_weeks:
      - add <w:pageBreakBefore/> to the first paragraph's <w:pPr> in that row,
        so weeks N and N+1 start on a fresh page.
    """
    body = doc.element.body
    rows_patched = 0
    breaks_applied = set()
    target_weeks = set(page_break_before_weeks)

    for tr in body.iter(f"{{{W_NS}}}tr"):
        first_tc = tr.find(f"{{{W_NS}}}tc")
        if first_tc is None:
            continue
        first_text = "".join(t.text or "" for t in first_tc.iter(f"{{{W_NS}}}t")).strip()
        m = re.match(r"^(\d+)\s*\.?\s*Hafta", first_text)
        if not m:
            continue

        week = int(m.group(1))

        # 1) Ensure <w:trPr> exists and add <w:cantSplit/>
        trPr = tr.find(f"{{{W_NS}}}trPr")
        if trPr is None:
            trPr = etree.SubElement(tr, f"{{{W_NS}}}trPr")
            tr.remove(trPr)
            tr.insert(0, trPr)

        if trPr.find(f"{{{W_NS}}}cantSplit") is None:
            etree.SubElement(trPr, f"{{{W_NS}}}cantSplit")
        rows_patched += 1

        # 2) Page break before target weeks
        if week in target_weeks and week not in breaks_applied:
            first_p = first_tc.find(f"{{{W_NS}}}p")
            if first_p is not None:
                pPr = first_p.find(f"{{{W_NS}}}pPr")
                if pPr is None:
                    pPr = etree.Element(f"{{{W_NS}}}pPr")
                    first_p.insert(0, pPr)
                if pPr.find(f"{{{W_NS}}}pageBreakBefore") is None:
                    etree.SubElement(pPr, f"{{{W_NS}}}pageBreakBefore")
                breaks_applied.add(week)

    print(f"  cantSplit added to {rows_patched} weekly rows")
    print(f"  pageBreakBefore applied to weeks: {sorted(breaks_applied)}")


def update_semester(doc):
    """Change 'Küzgü' semester label to Jazgi 2026 (spring)."""
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = p.text
                    if "Küzgü" in txt or "Күзгү" in txt or "күзгү" in txt:
                        new_txt = (txt.replace("2025 Күзгү семестр", "2026 Жазгы семестр")
                                       .replace("2025 Güz yarıyılı", "2026 Bahar yarıyılı"))
                        if new_txt != txt:
                            replace_in_paragraph(p, new_txt)


def main():
    doc = Document(str(SRC))

    week_cells = find_week_cells(doc)
    missing = [w for w in range(1, 17) if w not in week_cells]
    if missing:
        raise SystemExit(f"Could not locate week cells for weeks: {missing}")

    for w in range(1, 17):
        set_cell_text(week_cells[w], WEEKS[w - 1])
        normalize_run_formatting(week_cells[w])

    if find_and_replace_summary_in_textbox(doc, FINAL_SUMMARY):
        print("Replaced summary paragraph in text box.")
    else:
        print("WARN: final summary paragraph not found; appending at end")
        doc.add_paragraph(FINAL_SUMMARY)

    update_semester(doc)

    apply_table_row_fixes(doc, page_break_before_weeks=(11, 15))

    doc.save(str(DST))
    print(f"Saved: {DST}")


if __name__ == "__main__":
    main()
