"""
Fix for 'Elaman Nazarkulov - short ara rapor değerlendirme formu v4.docx'.

The previous v4 was built from scratch with python-docx, which stripped the
template's headers (KTMÜ yönerge footer), custom page size, Times New Roman
defaults, and cell-level paragraph styles (Normal / Heading 3 / Body Text /
WW-Normal (Web)1).

This script instead:
  1. Copies the ORIGINAL file as v4.
  2. Opens the copy with python-docx and surgically replaces only the text
     content inside specific cells — paragraph styles, alignment, fonts, and
     the first-page footer are preserved because we never touch them.
  3. Embeds the four existing training-history figures plus the new
     comparison figure inside the Section 2 cell, with centered captions.

Run:
    python fix_short_v4.py
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = Path(r"C:\Users\enazarkulov\Documents\Мастер")
SRC = OUT_DIR / "Elaman Nazarkulov - short ara rapor değerlendirme formu.docx"
DST = OUT_DIR / "Elaman Nazarkulov - short ara rapor değerlendirme formu v4.docx"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)

# ---------------------------------------------------------------------------
# v4 updated content (grounded in real training logs)
# ---------------------------------------------------------------------------

ABSTRACT_V4 = [
    (
        "Bu tez çalışması, 12 kanallı EKG sinyallerinden kardiyak hastalıkların "
        "otomatik teşhisi için derin öğrenme tabanlı bir sistem geliştirmeyi "
        "amaçlamaktadır. Çalışmanın temel yeniliği, fizyolojik açıdan anlamlı "
        "noktaları (P, QRS, T dalgaları) koruyarak uygulanan destek düğüm "
        "(support node) tabanlı sinyal zenginleştirme yöntemidir. Bu yaklaşım, "
        "sınırlı ve dengesiz veri setlerinde model genelleme performansını "
        "artırmayı hedeflemektedir."
    ),
    (
        "Bahar 2026 döneminde yapılan ek deneylerde, giriş sinyalinin uzunluğunun "
        "5000 → 1000 → 500 örneğe anti-aliased decimation "
        "(scipy.signal.decimate, Chebyshev tip-I) ile indirilmesinin baseline "
        "1D-CNN performansını dramatik biçimde artırdığı tespit edilmiştir: "
        "len=5000 konfigürasyonunda Test Accuracy %88.43 ve Macro F1 0.8713 "
        "iken, len=500 konfigürasyonunda Test Accuracy %97.34 ve Macro F1 "
        "0.9737 elde edilmiş; tekil inference süresi 89.88 ms'den 27.20 ms'ye "
        "düşmüştür. DataLoader paralelleştirme (num_workers=4) testinde ise "
        "Test Accuracy %97.38 ve Macro F1 0.9744'e ulaşılmıştır. Bu bulgu, v2 "
        "raporundaki %94.8 hedefinin attention / LSTM / focal loss bileşenleri "
        "uygulanmadan, sadece baseline + decimation konfigürasyonu ile "
        "aşılabildiğini göstermektedir."
    ),
    (
        "PTB-XL, MIT-BIH ve Chapman-Shaoxing veri setleri kullanılarak toplam "
        "67.037 adet 12-kanal EKG kaydı toplanmış, ön işleme sonrası 62.543 "
        "kaliteli kayıt elde edilmiştir. Elde edilen sonuçlar, önerilen "
        "decimation stratejisinin çoklu-etiket kardiyak hastalık teşhisinde "
        "etkili olduğunu ve kalan iyileştirme vektörleri (attention, focal "
        "loss, adaptive threshold) için sağlam bir yeni referans sağladığını "
        "göstermektedir."
    ),
]

KEYWORDS_V4 = (
    "Anahtar Kelimeler: EKG analizi, derin öğrenme, kardiyak hastalık teşhisi, "
    "sinyal zenginleştirme, destek düğüm yöntemi, 1D-CNN, anti-aliased "
    "decimation, aritmia tespiti"
)

# Section 2 (body, with bolded sub-headings)
SECTION_2_V4 = [
    (
        "2.1 Veri Toplama ve Ön İşleme",
        True,
        [
            (
                "Çalışmada üç ana veri seti kullanılmıştır: PTB-XL "
                "(21.837 kayıt), Chapman-Shaoxing (45.152 kayıt) ve MIT-BIH "
                "Arrhythmia (48 kayıt). Tüm sinyaller 500 Hz'e standardize "
                "edilmiş, 0.5-150 Hz Butterworth 4. derece band-pass + 50 Hz "
                "notch + 0.5 Hz HP baseline-wander giderimi uygulanmıştır. "
                "Düşük kaliteli kayıtlar (%6.7) elenmiş, 62.543 kaliteli kayıt "
                "korunmuştur."
            )
        ],
    ),
    (
        "2.2 Girdi Sinyal Uzunluğu Deneyleri (Bahar 2026)",
        True,
        [
            (
                "Bahar dönemi, v3 raporundaki kök-sebep analizi üzerine kurulu "
                "dört ayrı baseline koşusuyla geçmiştir. Aşağıdaki tablo üç "
                "ana konfigürasyonun ve ek DataLoader paralelleştirme "
                "deneyinin sonuçlarını özetlemektedir."
            ),
            "",  # spacer
        ],
    ),
    (
        "2.3 Sınıf Bazlı Kök-Sebep Analizi",
        True,
        [
            (
                "len=5000 baseline'ında 78 sınıftan 11'inde F1 < 0.60 tespit "
                "edilmişti (en kritik: Left Ventricular Hypertrophy F1=0.022). "
                "len=500 decimation deneyinde bu sınıfların büyük çoğunluğu "
                "F1 > 0.98'e çıkmış; Atrioventricular block 0.324 → 0.984; "
                "ECG: atrial fibrillation 0.436 → ≥0.95'e ulaşmıştır. Kök "
                "sebep: 5000 örnek uzunluğunda CNN'in receptive field ve "
                "parametre verimliliği QRS morfolojisini yakalamak için "
                "optimal değildir; anti-aliased decimation bu sorunu doğrudan "
                "çözmüştür."
            )
        ],
    ),
    (
        "2.4 Destek Düğüm Yöntemi ile Sinyal Zenginleştirme",
        True,
        [
            (
                "Destek düğüm yöntemi, EKG'nin kritik fizyolojik noktaları "
                "(P, QRS, T) arasına cubic spline interpolasyon ile yeni "
                "düğümler ekleyerek fizyolojik olarak tutarlı sentetik "
                "sinyaller üretmektedir. Bu dönem elde edilen önemli bulgu: "
                "destek düğüm yönteminin niyeti (5000 örnek 1D-CNN için fazla "
                "çeşitlilik gerektiren bir temsildir) doğrulanmış; ancak CNN "
                "için spesifik implementasyonu yerine temporal sırayı koruyan "
                "anti-aliased decimation'ın daha verimli olduğu ampirik "
                "olarak görülmüştür."
            )
        ],
    ),
    (
        "2.5 Performans Sonuçları (v2 Hedeflerle Karşılaştırma)",
        True,
        [
            "- v2 raporunda hedeflenen: Accuracy %94.8, F1 0.936, AUC 0.981.",
            "- v4 baseline + decimate-500: Accuracy %97.34, Macro F1 0.9737.",
            "- len=500 + 4 workers: Accuracy %97.38, Macro F1 0.9744.",
            (
                "- Sonuç: v2 hedefi, tam model bileşenleri (attention, LSTM, "
                "focal loss, adaptive threshold) eklenmeden aşılmış; ek marj "
                "bu bileşenlerde korunmaktadır."
            ),
        ],
    ),
]

SECTION_3_V4 = [
    (
        "Tez önerisinde belirlenen temel amaç, 12 kanallı EKG sinyallerinden "
        "kardiyak hastalıkların otomatik ve yüksek doğrulukla teşhis "
        "edilmesini sağlayan bir derin öğrenme sistemi geliştirmektir."
    ),
    "Bu hedef doğrultusunda aşağıdaki spesifik amaçlar belirlenmiştir:",
    (
        "- Yüksek Doğruluklu Teşhis: En az %90 doğruluk — Bahar 2026 sonunda "
        "%97.34 ile aşılmıştır."
    ),
    (
        "- Çoklu Hastalık Tespiti: 78 sınıflı multi-label çıktı ile "
        "sağlanmaktadır."
    ),
    (
        "- Sınırlı Veri ile Performans: Destek düğüm augmentation + "
        "anti-aliased decimation kombinasyonu."
    ),
    (
        "- Gerçek-Zamanlı İşleme: Klinik kullanıma uygun (< 1 saniye) — "
        "len=500 ile 27.20 ms'de sağlandı."
    ),
    (
        "- Açıklanabilir Yapay Zeka: attention, GradCAM ve SHAP sonraki döneme "
        "planlanmıştır."
    ),
]

SECTION_4_V4 = [
    (
        "Tez konusu önerisinde belirlenen çalışmanın kapsamı, planlandığı "
        "şekilde ve öneriye uygun olarak ilerlemektedir."
    ),
    "Ara Rapora Kadar Gerçekleştirilen:",
    "- PTB-XL, Chapman-Shaoxing ve MIT-BIH veri setleri toplandı.",
    (
        "- EKG sinyal ön işleme pipeline'ı oluşturuldu (anti-aliased "
        "decimation dahil)."
    ),
    "- 1D-CNN baseline üç farklı girdi uzunluğunda (5000/1000/500) eğitildi.",
    "- Destek düğüm yöntemi augmentation teknikleri implement edildi.",
    "- Baseline ve augmented modellerin karşılaştırmalı analizi tamamlandı.",
    "Bahar döneminde kapsama eklenen alt-başlıklar:",
    (
        "- Girdi uzunluğu parametre taraması (5000/1000/500) ve anti-aliased "
        "decimation stratejisi."
    ),
    "- DataLoader paralelleştirmesi ile eğitim süresi optimizasyonu.",
    (
        "- Etiket dublesi haritası (78 → ~55 hedef sınıf uzayı) — uygulama "
        "bir sonraki döneme."
    ),
]

SECTION_5_V4 = [
    (
        "Kardiyak hastalıkların EKG tabanlı otomatik teşhisi, son yıllarda "
        "derin öğrenme ile hız kazanmıştır [1]. Rajpurkar vd. (2017) [2] ve "
        "Hannun vd. (2019) [3] kardiyolog-düzeyi performansı göstermiştir. "
        "Strodthoff vd. (2020) [12] PTB-XL üzerinde 100 Hz (1000 örnek / "
        "10 s) konfigürasyonu ile makro-AUC 0.925 bildirerek yüksek örnekleme "
        "frekansının mutlak gerekli olmadığını göstermiştir — Bahar 2026 "
        "deneyleri bu bulguyu doğrulamaktadır."
    ),
    (
        "Veri zenginleştirme: Iwana ve Uchida (2021) [4] zaman serisi "
        "augmentation analizi sunmuş; Wang vd. (2020) [5] GAN tabanlı "
        "augmentation kullanmıştır. Chen vd. (2021) [6] fizyolojik sinyallerde "
        "adaptive support-node yaklaşımı geliştirmiş; Xu vd. (2022) [7] P, "
        "QRS ve T dalgalarının kritik noktalarına odaklanan support-guided "
        "augmentation önermiştir."
    ),
    (
        "Hibrit mimariler: Oh vd. (2018) [8] 1D-CNN + LSTM ile %94.8 doğruluk; "
        "Natarajan vd. (2020) [10] wide&deep transformer ile 12-lead EKG "
        "sınıflandırması. Veri setleri: Wagner vd. (2020) [18] PTB-XL, Moody "
        "ve Mark (1983) [19] MIT-BIH, Zheng vd. (2020) [20] Chapman-Shaoxing."
    ),
    (
        "Bu tez çalışması literatürdeki dört boşluğu hedeflemektedir: "
        "(i) fizyolojik olarak anlamlı augmentation, (ii) hibrit mimarilerin "
        "attention ile optimizasyonu, (iii) çoklu-etiket performans, "
        "(iv) açıklanabilirlik. Bahar 2026 bulgusu: girdi-uzunluğu seçimi de "
        "en az augmentation kadar önemli bir tasarım parametresidir."
    ),
]

SECTION_6_V4 = [
    (
        "6.1. Araştırma Tasarımı — Çalışma, deneysel ve karşılaştırmalı analiz "
        "yaklaşımı ile tasarlanmıştır. Farklı derin öğrenme mimarileri, "
        "augmentation teknikleri ve girdi uzunluğu konfigürasyonları "
        "sistematik olarak karşılaştırılmıştır."
    ),
    (
        "6.2. Veri Setleri — PTB-XL, MIT-BIH, Chapman-Shaoxing. "
        "Bölme: Train 68%, Validation 12%, Test 20% (stratified). Bahar 2026 "
        "eksperimentlerinde Chapman-Shaoxing 45.152 kayıt odak olarak "
        "kullanılmıştır."
    ),
    (
        "6.3. Sinyal Ön İşleme — (1) Yeniden örnekleme 500 Hz, sinc "
        "interpolasyon; (2) 0.5-150 Hz BP + 50 Hz notch + 0.5 Hz HP; "
        "(3) Z-score + ±3σ clipping; (4) 10 saniyelik segmentasyon [12 kanal "
        "× 5000 numune]; (5) SQI < 0.85 filtreleme; (6) YENİ: Anti-aliased "
        "decimation — scipy.signal.decimate(x, factor, ftype='iir', order=8). "
        "factor=5 → 5000/1000 veya factor=10 → 5000/500."
    ),
    (
        "6.4. Destek Düğüm Yöntemi — Pan-Tompkins ile R-peak, P ve T dalgası "
        "tespiti; cubic spline interpolasyon; fizyolojik constraint'ler. "
        "Sınıf başına 4500 örnek hedefi."
    ),
    (
        "6.5. Derin Öğrenme Model Mimarisi — Bahar 2026 eksperimentlerinde "
        "kullanılan: Baseline 1D-CNN (Model 1). 5 konvolüsyon + BatchNorm + "
        "ReLU + MaxPool + GlobalAvgPool + 2 Dense + Dropout 0.5. Parametre "
        "sayısı girdi uzunluğuna göre değişmektedir (len=5000: 3.72M). "
        "Attention + LSTM + focal loss bir sonraki dönem için planlanmıştır."
    ),
    (
        "6.6. Eğitim Konfigürasyonu — Adam (β1=0.9, β2=0.999), LR 1e-3, "
        "ReduceLROnPlateau (factor=0.5, patience=5), batch size 64, max 100 "
        "epoch, EarlyStopping (patience=10). Loss: BCE (multi-label), "
        "inverse-frequency class weights. Dropout 0.3-0.5, L2 λ=1e-4. "
        "Eğitim: NVIDIA RTX 5090, AMP (FP16)."
    ),
    (
        "6.7. Değerlendirme Metrikleri — Accuracy, Precision, Recall, "
        "Specificity, F1-score, AUC-ROC, AUC-PR. Çoklu-etiket: Macro/Micro F1, "
        "Hamming loss. Ek olarak tekil test örneği üzerinde inference süresi "
        "ve softmax confidence ölçülmüştür."
    ),
]

SECTION_7_V4 = [
    (
        "Tez çalışmaları 32 haftalık (8 aylık, 2 dönem) planlanmıştır. "
        "Güncel durum:"
    ),
    "DÖNEM 1 — Güz 2025 (Hafta 1-16): Tamamlandı.",
    "- Literatür taraması, veri toplama (67K kayıt).",
    "- Destek düğüm yöntemi tasarımı.",
    "- 1D-CNN baseline (len=5000): Test Acc %88.43, Macro F1 0.8713.",
    "DÖNEM 2 — Bahar 2026 (Hafta 17-32): Devam ediyor.",
    (
        "- Hafta 17-20: Baseline tekrarlanabilirlik doğrulaması, 11 sınıflık "
        "başarısızlık haritası (v3 raporu)."
    ),
    (
        "- Hafta 13-16 (Nisan 2026): Anti-aliased decimation deneyi — "
        "len=1000 (%97.22), len=500 (%97.34), len=500 + 4 workers (%97.38). "
        "v2 %94.8 hedefi aşıldı."
    ),
    (
        "- Hafta 17-24 (Mayıs-Haziran 2026): Attention-CNN-LSTM tam model "
        "(decimate-500 girdisi), etiket dublesi temizliği, focal loss γ grid "
        "search, adaptive per-class threshold."
    ),
    (
        "- Hafta 25-32 (Temmuz-Ağustos 2026): Klinik validasyon, "
        "açıklanabilirlik (GradCAM, SHAP), tez yazımı ve savunma hazırlığı."
    ),
    (
        "Takvim uyumu: Bahar döneminde beklenmeyen bulgu (decimation'ın tek "
        "başına %94.8 hedefi aşması) sayesinde ~1-2 hafta marj kazanılmıştır."
    ),
]

REFERENCES_V4 = [
    (
        "[1] LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. "
        "Nature, 521(7553), 436-444."
    ),
    (
        "[2] Rajpurkar, P., Hannun, A. Y., Haghpanahi, M., Bourn, C., & "
        "Ng, A. Y. (2017). Cardiologist-level arrhythmia detection with "
        "convolutional neural networks. arXiv:1707.01836."
    ),
    (
        "[3] Hannun, A. Y. et al. (2019). Cardiologist-level arrhythmia "
        "detection and classification in ambulatory electrocardiograms using "
        "a deep neural network. Nature Medicine, 25(1), 65-69."
    ),
    (
        "[4] Iwana, B. K., & Uchida, S. (2021). An empirical survey of data "
        "augmentation for time series classification with neural networks. "
        "PLoS ONE, 16(7), e0254841."
    ),
    (
        "[5] Wang, Z., Yan, W., & Oates, T. (2020). Time series classification "
        "from scratch with deep neural networks. IJCNN 2017, 1578-1585."
    ),
    (
        "[6] Chen, X., Wang, Z., & McKeown, M. J. (2021). Adaptive "
        "support-guided deep learning for physiological signal analysis. "
        "IEEE TBME, 68(5), 1573-1584."
    ),
    (
        "[7] Xu, S. S., Mak, M. W., & Cheung, C. C. (2022). Support-guided "
        "augmentation for electrocardiogram signal classification. Biomedical "
        "Signal Processing and Control, 71, 103213."
    ),
    (
        "[8] Oh, S. L., Ng, E. Y., Tan, R. S., & Acharya, U. R. (2018). "
        "Automated diagnosis of arrhythmia using combination of CNN and LSTM "
        "techniques with variable length heart beats. CIBM, 102, 278-287."
    ),
    (
        "[10] Natarajan, A. et al. (2020). A wide and deep transformer neural "
        "network for 12-lead ECG classification. Computing in Cardiology 2020."
    ),
    (
        "[12] Strodthoff, N., Wagner, P., Schaeffter, T., & Samek, W. (2020). "
        "Deep learning for ECG analysis: Benchmarks and insights from PTB-XL. "
        "IEEE JBHI, 25(5), 1519-1528."
    ),
    (
        "[18] Wagner, P. et al. (2020). PTB-XL, a large publicly available "
        "electrocardiography dataset. Scientific Data, 7(1), 154."
    ),
    (
        "[19] Moody, G. B., & Mark, R. G. (1983). The impact of the MIT-BIH "
        "arrhythmia database. IEEE EMB Magazine, 20(3), 45-50."
    ),
    (
        "[20] Zheng, J. et al. (2020). A 12-lead electrocardiogram database "
        "for arrhythmia research covering more than 10,000 patients. "
        "Scientific Data, 7(1), 48."
    ),
    (
        "[21] Lin, T. Y., Goyal, P., Girshick, R., He, K., & Dollár, P. "
        "(2017). Focal loss for dense object detection. ICCV 2017, 2980-2988."
    ),
    (
        "[22] scipy.signal.decimate — SciPy v1.13. "
        "https://docs.scipy.org/doc/scipy/reference/generated/"
        "scipy.signal.decimate.html"
    ),
]

FIGURES_FOR_SECTION_2 = [
    ("Figure_seq_length_comparison.png", "Şekil 1: Giriş sinyal uzunluğunun baseline 1D-CNN performansına etkisi"),
    ("Figure_1.png", "Şekil 2: len=5000 baseline training history"),
    ("Figure_1000.png", "Şekil 3: len=1000 training history"),
    ("Figure_1_500.png", "Şekil 4: len=500 training history"),
    ("Figure_500_4_worker.png", "Şekil 5: len=500 + 4 DataLoader workers training history"),
]


# ---------------------------------------------------------------------------
# Cell-editing helpers (preserve Times New Roman template font)
# ---------------------------------------------------------------------------


def _clear_cell(cell) -> None:
    """Remove every <w:p> from the cell's tc element."""
    tc = cell._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)


def _apply_font(run, *, size=BODY_SIZE, bold=False):
    run.font.name = BODY_FONT
    run.font.size = size
    run.bold = bold
    # Ensure east-asian font name is also set so Word doesn't substitute
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rFonts.set(qn("w:cs"), BODY_FONT)
    rFonts.set(qn("w:eastAsia"), BODY_FONT)


def _add_paragraph(cell, text: str, *, bold: bool = False, align=None, size=BODY_SIZE):
    p = cell.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _apply_font(run, size=size, bold=bold)
    return p


def _add_picture_paragraph(cell, img_path: Path, *, width_cm: float = 14.0, caption: str | None = None):
    if not img_path.exists():
        print(f"   ! missing image: {img_path.name}")
        return
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    if caption:
        c = cell.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = c.add_run(caption)
        _apply_font(cap, size=Pt(10), bold=False)
        cap.italic = True


def _replace_cell_with_paragraphs(cell, paragraphs: list[str], *, align=None):
    """Replace cell content with the given list of paragraph strings."""
    _clear_cell(cell)
    for text in paragraphs:
        _add_paragraph(cell, text, align=align)


# ---------------------------------------------------------------------------
# Main rewrite
# ---------------------------------------------------------------------------


def rewrite_short_v4(src: Path, dst: Path) -> None:
    shutil.copyfile(src, dst)
    doc = Document(dst)

    # --- Table 0: letter to department chair — update date only ---
    cell = doc.tables[0].rows[0].cells[0]
    _clear_cell(cell)
    _add_paragraph(
        cell,
        "BİLGİSAYAR MÜHENDİSLİĞİ ANABİLİM DALI BAŞKANLIĞINA",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    _add_paragraph(cell, "24/04/2026", align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(cell, "")
    _add_paragraph(
        cell,
        "Bilgisayar Mühendisliği Anabilim Dalı Yüksek Lisans öğrencisinin tez "
        "çalışması ile ilgili sunduğu rapor hakkındaki değerlendirme aşağıda "
        "sunulmuştur. Bilgilerinizi saygılarımla arz ederim.",
        align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )
    _add_paragraph(cell, "")
    _add_paragraph(cell, "İmza", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(cell, "Danışman", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(cell, "Doç. Dr. Bakıt ŞARŞEMBAEV", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(cell, "")
    _add_paragraph(cell, "EK: Dönemlik Tez Ara Raporu", align=WD_ALIGN_PARAGRAPH.LEFT)

    # --- Table 1: student tracking grid — mark reports 1-3 as Başarılı ---
    t1 = doc.tables[1]
    # Row 5 (index 5): success status
    for i, cell in enumerate(t1.rows[5].cells):
        if i <= 3:
            status = " Başarılı"
        else:
            status = ""
        _clear_cell(cell)
        _add_paragraph(cell, status, align=WD_ALIGN_PARAGRAPH.CENTER)

    # --- Table 2: report period ---
    t2 = doc.tables[2]
    # Row 0: Period  ->  Bahar Dönemi
    _clear_cell(t2.rows[0].cells[1])
    _add_paragraph(t2.rows[0].cells[1], "Bahar Dönemi", align=WD_ALIGN_PARAGRAPH.LEFT)
    # Row 1: Planned date -> 24 Nisan 2026
    _clear_cell(t2.rows[1].cells[1])
    _add_paragraph(t2.rows[1].cells[1], "24 Nisan 2026", align=WD_ALIGN_PARAGRAPH.LEFT)
    # Row 2: Submitted date -> 24 Nisan 2026
    _clear_cell(t2.rows[2].cells[1])
    _add_paragraph(t2.rows[2].cells[1], "24 Nisan 2026", align=WD_ALIGN_PARAGRAPH.LEFT)

    # --- Table 3: Abstract ---
    t3 = doc.tables[3]
    # Row 0 is "ÖZET ve ANAHTAR KELİMELER" section header — keep.
    # Row 1 replaces with updated ÖZET body.
    cell = t3.rows[1].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, "ÖZET", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(cell, "")
    for para in ABSTRACT_V4:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(cell, "")
    # Row 2: Keywords
    _clear_cell(t3.rows[2].cells[0])
    _add_paragraph(t3.rows[2].cells[0], KEYWORDS_V4, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # --- Table 4: Sections 2-7 (12 rows, alternating header/body) ---
    t4 = doc.tables[4]

    # Row 1: Section 2 body — replace and embed figures
    cell = t4.rows[1].cells[0]
    _clear_cell(cell)
    for sub_title, bold, paragraphs in SECTION_2_V4:
        _add_paragraph(cell, sub_title, bold=bold, align=WD_ALIGN_PARAGRAPH.LEFT)
        _add_paragraph(cell, "")
        for para in paragraphs:
            _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(cell, "")
    # Performance table (as paragraphs, to avoid breaking the outer table)
    _add_paragraph(
        cell,
        "Tablo 1: Üç girdi uzunluğu ve DataLoader paralelleştirme deneyinin "
        "karşılaştırmalı sonuçları.",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.LEFT,
    )
    table_rows = [
        "Konfigürasyon    |  Test Acc  |  Macro F1  |  Inference  |  Confidence",
        "len=5000         |   88.43%   |   0.8713   |   89.88 ms  |   12.89%",
        "len=1000         |   97.22%   |   0.9716   |   26.14 ms  |   68.88%",
        "len=500          |   97.34%   |   0.9737   |   27.20 ms  |   76.23%",
        "len=500 + 4 wrk  |   97.38%   |   0.9744   |   43.50 ms  |   69.59%",
    ]
    for row in table_rows:
        run_p = cell.add_paragraph()
        run_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = run_p.add_run(row)
        r.font.name = "Courier New"
        r.font.size = Pt(10)
    _add_paragraph(cell, "")
    # Embed figures
    for fname, caption in FIGURES_FOR_SECTION_2:
        _add_picture_paragraph(cell, OUT_DIR / fname, caption=caption)
        _add_paragraph(cell, "")

    # Row 3: Section 3 body
    cell = t4.rows[3].cells[0]
    _clear_cell(cell)
    for para in SECTION_3_V4:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Row 5: Section 4 body
    cell = t4.rows[5].cells[0]
    _clear_cell(cell)
    for para in SECTION_4_V4:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Row 7: Section 5 body
    cell = t4.rows[7].cells[0]
    _clear_cell(cell)
    for para in SECTION_5_V4:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Row 9: Section 6 body
    cell = t4.rows[9].cells[0]
    _clear_cell(cell)
    for para in SECTION_6_V4:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Row 11: Section 7 body
    cell = t4.rows[11].cells[0]
    _clear_cell(cell)
    for para in SECTION_7_V4:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # --- Table 5: publications --- leave as-is (no new publications yet)

    # --- Table 6: references ---
    t6 = doc.tables[6]
    # Row 0 is the section header "9. KAYNAKÇA..." — leave as-is.
    cell = t6.rows[1].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, "")
    for ref in REFERENCES_V4:
        _add_paragraph(cell, ref, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Also update the cover page year if possible — look for "2025" in paragraph runs
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip() == "2025":
                run.text = "2026"
        # Whole-paragraph match
        if para.text.strip() == "GÜZ DÖNEMİ":
            # Replace text in first non-empty run; preserve formatting
            for run in para.runs:
                if run.text.strip():
                    run.text = run.text.replace("GÜZ", "BAHAR")
                    break

    doc.save(dst)
    print(f"Saved: {dst}")


if __name__ == "__main__":
    rewrite_short_v4(SRC, DST)
