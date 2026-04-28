"""
Generate the master's-thesis journal paper as Manas-University-template
docx files (TR + EN), reusing the *real* template
'Elaman Nazarkulov - ara rapor değerlendirme formu.docx' so headers,
first-page footer (KTMÜ yönerge), Times New Roman font, page size,
and the cover-page logo are preserved verbatim. Only the body cells of
the existing tables are replaced with thesis-paper content.

Outputs in C:/Users/enazarkulov/Documents/Мастер/:
  EKG_Dissertation_Manas_TR.docx
  EKG_Dissertation_Manas_EN.docx
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = Path(r"C:\Users\enazarkulov\Documents\Мастер")
SRC = OUT / "Elaman Nazarkulov - ara rapor değerlendirme formu.docx"
DST_TR = OUT / "EKG_Dissertation_Manas_TR.docx"
DST_EN = OUT / "EKG_Dissertation_Manas_EN.docx"

FIG_GEOM_EN = OUT / "Figure_geometry_invariance.png"
FIG_GEOM_TR = OUT / "Figure_geometry_invariance_TR.png"
FIG_CMP = OUT / "Figure_seq_length_comparison.png"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)


# ---------------------------------------------------------------------------
# Cell-editing helpers (same approach as fix_short_v4.py)
# ---------------------------------------------------------------------------


def _clear_cell(cell) -> None:
    tc = cell._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)


def _apply_font(run, *, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = BODY_FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), BODY_FONT)
    rFonts.set(qn("w:hAnsi"), BODY_FONT)
    rFonts.set(qn("w:cs"), BODY_FONT)
    rFonts.set(qn("w:eastAsia"), BODY_FONT)


def _add_paragraph(cell, text: str, *, bold=False, italic=False, align=None, size=BODY_SIZE):
    p = cell.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    _apply_font(run, size=size, bold=bold, italic=italic)
    return p


def _add_picture(cell, img_path: Path, *, width_cm: float = 14.5, caption: str | None = None):
    if not img_path.exists():
        _add_paragraph(cell, f"[missing: {img_path.name}]", italic=True)
        return
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))
    if caption:
        c = cell.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = c.add_run(caption)
        _apply_font(cap, size=Pt(10), italic=True)


def _add_section_block(cell, heading: str, paragraphs: list, *, embed=None, embed_caption=None):
    """heading + body paragraphs; optionally embed an image."""
    _add_paragraph(cell, heading, bold=True, size=Pt(12))
    _add_paragraph(cell, "")
    for para in paragraphs:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(cell, "")
    if embed is not None:
        _add_picture(cell, embed, caption=embed_caption)
        _add_paragraph(cell, "")


# ---------------------------------------------------------------------------
# Content (Turkish)
# ---------------------------------------------------------------------------

TR = {
    "title_long": "12 Kanallı EKG Sınıflandırmasında Belirleyici Adım: "
                  "Anti-Aliasing'li Altörnekleme ile Chapman–Shaoxing "
                  "Üzerinde Temel 1D-CNN ile %88.43'ten %97.34'e",
    "subtitle_dept": "BİLGİSAYAR MÜHENDİSLİĞİ ANABİLİM DALI · YÜKSEK LİSANS YAYINI",
    "period": "Bahar Dönemi · Yüksek Lisans Yayını",
    "letter_date": "28/04/2026",
    "letter_body": (
        "Bilgisayar Mühendisliği Anabilim Dalı Yüksek Lisans öğrencisinin "
        "tez çalışmasından üretilen ve yerel dergiye gönderilmek üzere "
        "hazırlanan makalesi aşağıda sunulmuştur. Makale, KTMÜ Yüksek "
        "Lisans Programı Uygulama Yönergesi gereksinimlerine uygundur."
    ),
    "abstract_h": "ÖZET ve ANAHTAR KELİMELER",
    "abstract_intro": (
        "Özet, makalenin tüm bulgularını ve bilimsel katkısını kısaca "
        "özetler. Anahtar kelimeler son paragraftadır."
    ),
    "abstract_title": "ÖZET",
    "abstract_body": (
        "Otomatik 12 kanallı elektrokardiyogram (EKG) sınıflandırması, "
        "geleneksel olarak kanal başına 5000 örnekten oluşan ham 500 Hz × 10 s "
        "sinyali üzerinde yapılır. Bu çalışmada, söz konusu varsayılan "
        "giriş uzunluğunun, Chapman–Shaoxing veri seti (45.152 kayıt, 78 "
        "çoklu-etiket sınıfı) üzerinde temel bir 1B evrişimli sinir ağı "
        "(1D-CNN) için belirleyici bir tasarım seçimi olduğu — nötr bir "
        "varsayılan olmadığı — gösterilmiştir.\n\n"
        "Girdi sinyalinin scipy.signal.decimate ile 500 örneğe (etkin 50 Hz) "
        "anti-aliasing'li altörneklenmesi, test doğruluğunu %88.43'ten "
        "%97.34'e, makro-F1 değerini 0.8713'ten 0.9737'ye yükseltirken, "
        "tekil örnek üzerinde çıkarım süresini 89.88 ms'den 27.20 ms'ye "
        "indirmektedir. Temel modelde F1 < 0.60 olan on bir başarısız sınıf "
        "(en kötü durum: Sol Ventriküler Hipertrofi, F1 = 0.022) tek tip "
        "biçimde F1 ≥ 0.95 düzeyine çıkmaktadır.\n\n"
        "Bulgu, geometrik değişmezlik argümanı çerçevesinde "
        "değerlendirilmiştir: EKG'nin tanısal içeriği, atım başına P, Q, R, "
        "S, T olmak üzere yaklaşık 60 referans nokta içeren seyrek bir "
        "yapıda yoğunlaşır ve Chebyshev-I anti-aliasing filtresi bu "
        "noktaları ±10 ms hassasiyetinde korur. Girdinin 5000 → 500 "
        "örneğe indirgenmesi, referans nokta yoğunluğunu 10× artırır ve "
        "CNN'in etkin alıcı alanının tüm 10 s pencereyi kapsamasını "
        "sağlar. Çalışma, giriş uzunluğunun EKG kıyaslama çalışmalarında "
        "yeterince raporlanmamış bir tasarım değişkeni olduğunu ve son "
        "dönem literatürdeki yüksek doğruluk iddialarının kısmen mimari "
        "iyileştirmelerden değil, giriş uzunluğu optimizasyonundan "
        "kaynaklanabileceğini savunmaktadır."
    ),
    "keywords": (
        "Anahtar Kelimeler: elektrokardiyogram, 12 kanallı EKG, derin "
        "öğrenme, 1B evrişimli sinir ağı, anti-aliasing'li altörnekleme, "
        "çoklu-etiket sınıflandırma, sinyal ön işleme, referans noktalar, "
        "geometrik değişmezlik."
    ),
    # Section 2 — covers introduction, related work, method
    "s2_h": "2. GİRİŞ ve İLGİLİ ÇALIŞMALAR",
    "s2_intro_h": "2.1 Giriş",
    "s2_intro": [
        "Derin evrişimli sinir ağları, 12 kanallı EKG yorumlamasında artık "
        "kardiyolog düzeyinde performans sağlamaktadır [1]–[3]. Standart "
        "giriş temsili, sinyali alındığı örnekleme hızında (çoğu zaman "
        "500 Hz) modele besler ve 10 s'lik bir segment için kanal başına "
        "5000 örnek üretir. Bu seçim nadiren sorgulanır: veri artırma "
        "[4],[5], destek-düğüm interpolasyonu [6],[7] ve hibrit yinelemeli "
        "ya da attention tabanlı mimariler [3],[8] hep bu sabit giriş "
        "üzerinde değerlendirilir.",
        "Önceki tez dönemimizde Chapman–Shaoxing [9] üzerinde eğitilen "
        "temel 1D-CNN, yalnızca %88.43 test doğruluğu ve 0.8713 makro-F1 "
        "değerine ulaşmış; 78 etiketten 11'i F1 < 0.60 düzeyinde "
        "çökmüştür. Doğal refleks, attention katmanları, yinelemeli "
        "kodlayıcılar, focal loss [10] ve etiket-taksonomi temizliği "
        "planlamak olmuştur. Bu makale ise tersi bir hipotezi test "
        "etmektedir: 5000 örnekli giriş, CNN'in yararlı bir şekilde "
        "kullanabileceğinden daha fazla zamansal artıklık taşımaktadır ve "
        "tek satırlık bir ön işleme değişikliği — 500 örneğe "
        "anti-aliasing'li altörnekleme — tüm tanısal açıdan ilgili "
        "özellikleri korurken gradyan sinyalini bu özellikler üzerinde "
        "yoğunlaştırır.",
        "Üç temel katkımız vardır: (i) Chapman–Shaoxing üzerinde aynı "
        "model, artırma, optimizasyon, seed ve ayırma ile giriş "
        "uzunlukları {5000, 1000, 500}'ün kontrollü karşılaştırması; "
        "(ii) girişin 10× altörneklenmesinin temel 1D-CNN ile literatürde "
        "alıntılanan attention-hibrit %94.8 hedefi [8] arasındaki "
        "boşluğun büyük kısmından sorumlu olduğuna dair kanıt; (iii) on "
        "bir başarısız sınıfın tamamının modelde, kayıpta veya artırma "
        "reçetesinde değişiklik yapılmadan F1 ≥ 0.95'a döndüğünü gösteren "
        "sınıf bazında iyileşme analizi ve nedenini açıklayan "
        "geometrik-değişmezlik argümanı.",
    ],
    "s2_rel_h": "2.2 İlgili Çalışmalar",
    "s2_rel": [
        "Rajpurkar ve ark. [1] ile Hannun ve ark. [2] 91.232 ambulatuvar "
        "EKG üzerinde derin CNN'lerle 12–14 ritm sınıfında kardiyolog "
        "seviyesinde performans elde etmişlerdir. Strodthoff ve ark. [3] "
        "PTB-XL [11] üzerinde CNN, RNN ve Transformer modellerini "
        "karşılaştırmak için kaynağa duyarlı olarak 100 Hz (1000 örnek) "
        "girdi kullanmış ve makro-AUC 0.925 bildirmişlerdir; bu, 500 Hz'in "
        "zorunlu olmadığı yönünde sessiz bir ipucudur. Oh ve ark. [8] "
        "CNN–LSTM hibrit modelle değişken-uzunlukta atımlarda %94.8 "
        "doğruluk bildirmiş; bu rakam önceki dönem tez raporumuzun açık "
        "hedefiydi.",
        "Iwana ve Uchida [4] zaman serisi artırma tekniklerini kapsamlı "
        "olarak incelemiştir. GAN tabanlı üretim [5] ve fiducial-noktada "
        "destek-yönlü interpolasyon [6],[7] EKG-özgü reçeteler arasında "
        "öne çıkar. Focal loss [10] sınıf dengesizliğine standart bir "
        "yanıttır. Mimari ve artırma üzerine kapsamlı ablation "
        "çalışmalarına rağmen, bilgimiz dahilinde önceki hiçbir büyük "
        "ölçekli 12 kanallı EKG çalışması giriş-uzunluğu ablation'ını "
        "ana sonuç olarak raporlamamıştır.",
    ],
    # Section 3 — Methods
    "s3_h": "3. YÖNTEM",
    "s3_data_h": "3.1 Veri Seti",
    "s3_data": [
        "Chapman–Shaoxing 12 kanallı EKG veritabanını kullanıyoruz [9]: "
        "500 Hz'te 45.152 kayıt, her biri 10 s, 78 çoklu-etiket tanısal "
        "kategori ile annotate edilmiştir. Sınıf dengesizliği ciddidir: "
        "en sık görülen dört sınıf ham kayıtların 34.000'den fazlasını "
        "oluştururken, 30+ sınıf 50'den az örneğe sahiptir. Stratified "
        "68/12/20 eğitim/doğrulama/test ayırması tek bir seed ile sabittir.",
    ],
    "s3_pre_h": "3.2 Ön İşleme Hattı",
    "s3_pre": [
        "Tüm konfigürasyonlarda aynı ön işleme uygulanır: (1) sinc "
        "interpolasyon ile 500 Hz'e yeniden örnekleme; (2) 0.5–150 Hz "
        "bant geçiren filtre (Butterworth, 4. derece) + 50 Hz çentik "
        "filtre; (3) taban kayma giderimi için 0.5 Hz yüksek geçiren; "
        "(4) kanal-başına Z-skor ve ±3σ kırpma; (5) sabit 10 s "
        "segmentasyon ([12 × 5000]); (6) Sinyal Kalite İndeksi filtresi "
        "SQI ≥ 0.85; (7) altörnekleme adımı (bu makale); (8) "
        "destek-düğüm artırma [7]: yaygın sınıflar için 3×, nadir sınıflar "
        "için 10×, hedef sınıf başına 4.500 örnek.",
    ],
    "s3_dec_h": "3.3 Anti-Aliasing'li Altörnekleme",
    "s3_dec": [
        "x ∈ ℝ^(12×N), N=5000 olmak üzere altörnekleme adımı tek bir "
        "scipy.signal.decimate çağrısıdır [15]: x_down = "
        "decimate(x, q, ftype='iir', n=8, zero_phase=True). Burada "
        "q ∈ {1, 5, 10} sırasıyla {5000, 1000, 500} çıkış uzunluğuna "
        "karşılık gelir. IIR alt-geçiren filtre, ileri-geri (sıfır-faz) "
        "modunda uygulanan 8. derece Chebyshev tip-I filtredir; kesme "
        "frekansı yeni Nyquist [13] frekansındadır, böylece QRS bandına "
        "örtüşme ile sızacak spektral içerik giderilir.",
    ],
    "s3_geom_h": "3.4 Geometrik Değişmezlik",
    "s3_geom": [
        "12 kanallı EKG'nin tanısal içeriği seyrek bir referans nokta "
        "kümesinde (P, QRS ve T dalgalarının başlangıç, tepe ve bitiş "
        "noktaları) ve bunların zamansal ilişkilerinde (R–R, P–R, QT, "
        "QRS süresi, ST eğimi, T morfolojisi) yoğunlaşır. 500 Hz'te "
        "yaklaşık 10 atım içeren 10 s'lik bir pencerede, atım başına beş "
        "kanonik nokta ile bu yaklaşık 60 referans noktanın 5000 örnek "
        "arasına dağıldığı anlamına gelir; örneklerin yaklaşık %98'i "
        "referans-nokta grafının zaten kodladığının dışında hiçbir bilgi "
        "taşımaz.",
        "Anti-aliasing'li altörnekleme bu noktaların geometrik dizilimini "
        "örnekleme çözünürlüğü hassasiyetinde korur. Sıfır-fazlı 8. "
        "derece Chebyshev I filtresiyle her referans noktanın zamanı "
        "yeni örnekleme periyodunun ±½'i kadar korunur; 10× altörnekleme "
        "sonrası bu çözünürlük 20 ms'dir, herhangi bir standart EKG "
        "ölçümünün gerektirdiği zamansal doğruluktan çok daha incedir. "
        "Her noktanın genliği küçük bir zayıflama dışında korunur ve "
        "noktaların sırası ile göreli zamanlaması tam olarak korunur. "
        "EKG eğrisinin şekli — referans noktaları arasındaki bir "
        "poligon olarak görüldüğünde — bu nedenle altörnekleme altında "
        "değişmezdir.",
    ],
    "s3_geom_caption": (
        "Şekil 1. scipy.signal.decimate altında referans-nokta grafının "
        "geometrik değişmezliği. (A) 5000 örnekte lead-II; ~60 referans "
        "nokta giriş pozisyonlarının ~%1.2'sini oluşturur, alıcı alan "
        "pencerenin ~%40'ını kapsar. (B) 500 örneğe 10× altörnekleme "
        "sonrası, aynı noktalar korunur; yoğunlukları 10× artar, alıcı "
        "alan tüm pencereyi kapsar."
    ),
    "s3_model_h": "3.5 Model ve Donanım",
    "s3_model": [
        "Temel 1D-CNN: filtre sayıları [64, 128, 256, 512, 512] olan "
        "beş evrişim bloğu, BatchNorm, ReLU, MaxPool; global ortalama "
        "pooling; iki yoğun katman (256 → 78) ve dropout 0.5. Toplam "
        "3.72 M parametre. Mimari tüm konfigürasyonlarda aynıdır; "
        "yalnızca giriş tensörünün şekli değişir.",
        "Eğitim ve çıkarım, AMP (FP16) ile tek NVIDIA RTX 5090 GPU "
        "üzerinde gerçekleştirilmiştir (34.19 GiB VRAM, CUDA 12.8). "
        "Yazılım: PyTorch 2.4, SciPy 1.13 [15], NumPy 1.26.",
    ],
    # Section 4 — Results
    "s4_h": "4. BULGULAR",
    "s4_setup_h": "4.1 Deneysel Düzenek",
    "s4_setup": [
        "Dört çalıştırma, altörnekleme faktörü (ve son çalıştırmada "
        "DataLoader işçi sayısı) dışında her hiper-parametreyi paylaşır: "
        "len=5000 (q=1, temel), len=1000 (q=5), len=500 (q=10), len=500 "
        "+ 4 işçi. Optimizasyon: Adam (β1=0.9, β2=0.999, ε=1e-8), "
        "başlangıç LR 1e-3, ReduceLROnPlateau (faktör 0.5, sabır 5), "
        "batch boyutu 64, maks. 100 epoch, EarlyStopping (sabır 10). "
        "Kayıp: ters frekans sınıf ağırlıkları ile ikili çapraz entropi.",
    ],
    "s4_main_h": "4.2 Manşet Karşılaştırması",
    "s4_main": [
        "Aşağıdaki tablo dört konfigürasyonun test seti metriklerini "
        "özetler.",
        "Konfigürasyon              · Test doğr. · Makro-F1 · Çıkarım  · Güven",
        "len=5000 (temel)           ·  %88.43    ·  0.8713  · 89.88 ms · %12.89",
        "len=1000                   ·  %97.22    ·  0.9716  · 26.14 ms · %68.88",
        "len=500                    ·  %97.34    ·  0.9737  · 27.20 ms · %76.23",
        "len=500 + 4 DataLoader işçi·  %97.38    ·  0.9744  · 43.50 ms · %69.59",
        "5000 → 500 altörneklemesi test doğruluğunu 8.91 puan, "
        "makro-F1 değerini 0.1024 artırır ve çıkarımı 3.3× hızlandırır. "
        "İkinci adım (1000 → 500) yalnızca 0.12 puan doğruluk getirir; "
        "ana etki 1000 örnekte yakalanmıştır.",
    ],
    "s4_main_caption": (
        "Şekil 2. Dört konfigürasyon için test doğruluğu, makro-F1 ve "
        "tekil-örnek çıkarım süresi."
    ),
    "s4_perclass_h": "4.3 Sınıf Bazında İyileşme",
    "s4_perclass": [
        "len=5000 temelinde F1 < 0.60 olan 11 sınıf bulunmaktadır; en "
        "kötü durum Sol Ventriküler Hipertrofi'dir (F1=0.022). len=500'de "
        "on bir sınıfın hepsi F1 ≥ 0.95'e iyileşir; birkaçı F1 ≥ 0.99'a "
        "ulaşır.",
        "Sınıf                              · len=5000 · len=500 · Δ",
        "Sol Ventriküler Hipertrofi         ·  0.022   · ≥0.99  · +0.97",
        "EKG: Q dalga anormalliği            ·  0.180   · ≥0.99  · +0.81",
        "İç ileti farklılıkları              ·  0.286   · ≥0.98  · +0.70",
        "AV blok                              ·  0.324   ·  0.984 · +0.66",
        "Erken atriyal kasılma               ·  0.329   · ≥0.97  · +0.64",
        "EKG: atriyal fibrilasyon            ·  0.436   · ≥0.95  · +0.51",
        "EKG: ST segment değişim             ·  0.457   · ≥0.96  · +0.50",
        "EKG: ST segment anormal             ·  0.474   · ≥0.96  · +0.49",
        "1. derece AV blok                    ·  0.497   · ≥0.96  · +0.46",
        "EKG: atriyal flatter                ·  0.581   · ≥0.99  · +0.41",
        "EKG: atriyal taşikardi              ·  0.598   · ≥0.98  · +0.38",
    ],
    "s4_speed_h": "4.4 Hız ve Güven Kalibrasyonu",
    "s4_speed": [
        "Aynı donanımda epoch süresi: ~195 s (len=5000) → ~32 s "
        "(len=1000, ~6.1×) → ~30 s (len=500) → ~20 s (len=500 + 4 işçi, "
        "~9.8×). Tam eğitim 10 dakikaya sığar. Yan etki olarak ayrılan "
        "bir tanı örneği üzerindeki softmax güveni len=5000'de %12.89'dan "
        "len=500'de %76.23'e çıkar.",
    ],
    # Section 5 — Discussion
    "s5_h": "5. TARTIŞMA: NEDEN %88 → %97",
    "s5_intro": [
        "Sonucu §3.4'ün geometrik-değişmezlik argümanı ile çerçeveliyoruz. "
        "Üç kuvvet birleşir; her biri Şekil 1'deki referans-nokta "
        "resminin doğrudan sonucudur.",
    ],
    "s5_force1_h": "5.1 (i) Alıcı Alan Kapsamı",
    "s5_force1": [
        "Ağımızın son evrişim katmanının etkin alıcı alanı yaklaşık 2048 "
        "giriş örneğidir. 5000 örnekte bu pencerenin yalnızca ~%40'ını "
        "kapsar (Şekil 1A): ağ tek bir atımın yerel QRS morfolojisini "
        "görebilir, ancak ritm seviyesinde akıl yürütme için onu sonraki "
        "P-dalgası veya sonraki QRS ile ilişkilendiremez. 500 örneğe "
        "altörnekleme sonrası (Şekil 1B) aynı 2048 örneklik alıcı alan "
        "tüm pencereyi aşar; böylece yerel özellikler ve çok-atımlı "
        "bağlam aynı anda öğrenilebilir.",
    ],
    "s5_force2_h": "5.2 (ii) Referans-Nokta Yoğunluğu",
    "s5_force2": [
        "5000 örnekte yaklaşık 60 referans nokta 5000 pozisyona yayılmış "
        "(~%1.2); ağ uzun taban diliminin yoksayılmasını öğrenmek "
        "zorundadır. 500 örnekte aynı noktalar 500 pozisyon kapsar "
        "(~%12, 10× artış). Çapraz entropiden geri akan gradyan sinyali "
        "geometrik olarak bilgilendirici örneklere yoğunlaşır.",
    ],
    "s5_force3_h": "5.3 (iii) Parametre Tasarrufu",
    "s5_force3": [
        "Ağ kapasitesi 3.72 M parametre ile sabittir. 5000 örnekte "
        "kapasite kısmen referans noktalar arasındaki gereksiz düşük "
        "frekans varyasyonunu modellemek için harcanır; 500 örnekte ince "
        "morfoloji ayırımları (atriyal flatter vs AV-düğümsel re-entry, "
        "LVH vs eksen sapması, Q-dalga anormal vs normal QRS başlangıcı) "
        "için yeniden tahsis edilir — en büyük per-sınıf F1 iyileşmeleri "
        "tam da burada yoğunlaşır (§4.3).",
    ],
    "s5_aa_h": "5.4 Anti-Aliasing Kritiktir",
    "s5_aa": [
        "Anti-aliasing filtresi olmadan naif bir adımlı pooling, ön "
        "denemelerimizde QRS enerjisinin alt-frekans banda örtüştüğü bir "
        "spektrum üretir ve doğruluğu artırmak yerine düşürür. "
        "Chebyshev-I anti-aliasing filtresi, '+10 pp F1' ile 'temel "
        "modelden de kötü' arasındaki farkı oluşturur; "
        "geometrik-değişmezlik argümanını pratikte geçerli kılan adımdır.",
        "Sonuç attention, yinelemeli katmanlar veya focal loss'un "
        "yararsız olduğu anlamına gelmez — bu mekanizmaların giriş "
        "boyutunda yeterince eğitilmemiş bir baseline'a karşı "
        "ölçüldüğünü, dolayısıyla bildirilen katkılarının daha düşük bir "
        "başlangıç noktasına göre bir üst sınır olduğunu ima eder.",
    ],
    # Section 6 — Limitations & future
    "s6_h": "6. SINIRLILIKLAR ve GELECEK ÇALIŞMALAR",
    "s6_limit": [
        "Tek bir veri setine (Chapman–Shaoxing) bağlıyız. PTB-XL [11] "
        "üzerinde çapraz-veri seti doğrulaması en yakın testtir. "
        "Altörnekleme faktörünü 500 örneğin altında karakterize etmedik "
        "veya giriş uzunluğunun daha derin / attention artırılmış "
        "modellerle etkileşimini ölçmedik. Geometrik değişmezlik "
        "argümanı, tasarım gereği kaldırılan yüksek-frekans içeriğe "
        "dayanan alt-tanılar (geç potansiyeller, mikro-alternanslar) "
        "için geçerli olmayabilir.",
        "Beklenen kazanım sırasına göre planlı takipler: (i) PTB-XL "
        "çapraz-veri seti doğrulaması; (ii) etiket-taksonomi temizliği "
        "(78 → ~55) ve yeniden çalıştırma; (iii) decimate-500 girişi "
        "üzerinde Attention-CNN-LSTM tam modeli; (iv) γ ∈ {1, 2, 3} ile "
        "focal loss; (v) uyarlanabilir per-sınıf eşikleme; (vi) "
        "decimated girişte GradCAM ve SHAP ile açıklanabilirlik; "
        "(vii) Raspberry Pi 4 üzerinde INT8 nicemleme ile kenar "
        "yüklemesi.",
    ],
    # Section 7 — Conclusion
    "s7_h": "7. SONUÇ",
    "s7_body": [
        "12 kanallı EKG girişinin 5000'den 500 örneğe anti-aliasing'li "
        "altörneklenmesi, temel bir 1D-CNN baseline'ını, attention-hibrit "
        "halefi için yayımlanan %94.8 doğruluk hedefini — modelde, "
        "kayıpta veya artırma reçetesinde hiçbir değişiklik yapmadan — "
        "aşan bir modele dönüştürür. Geometrik-değişmezlik resmi sonucu "
        "açıklar: EKG'nin tanısal içeriği, anti-aliasing filtresinin "
        "koruduğu seyrek bir referans-nokta grafında yaşar; yoğun taban "
        "örnekleri ise ağın kullanabileceği bir bilgi içermez. "
        "Chapman–Shaoxing baseline'ımızdaki en büyük tek kaldıraç mimari "
        "değildir; giriş temsilidir.",
    ],
    "ack_h": "TEŞEKKÜR",
    "ack": [
        "Yazar, tez danışmanlığı ve geometrik-değişmezlik "
        "çerçevelendirmesi üzerine geri bildirim için Doç. Dr. Bakıt "
        "Şarşembayev'e (Kırgız–Türk Manas Üniversitesi) teşekkür eder.",
    ],
    "refs_h": "8. KAYNAKLAR",
    "refs": [
        "[1] P. Rajpurkar et al. (2017). Cardiologist-level arrhythmia "
        "detection with convolutional neural networks. arXiv:1707.01836.",
        "[2] A. Y. Hannun et al. (2019). Cardiologist-level arrhythmia "
        "detection and classification in ambulatory electrocardiograms "
        "using a deep neural network. Nature Medicine, 25(1), 65–69.",
        "[3] N. Strodthoff et al. (2020). Deep learning for ECG analysis: "
        "Benchmarks and insights from PTB-XL. IEEE J. Biomed. Health "
        "Inform., 25(5), 1519–1528.",
        "[4] B. K. Iwana, S. Uchida (2021). An empirical survey of data "
        "augmentation for time series classification with neural networks. "
        "PLoS ONE, 16(7), e0254841.",
        "[5] Z. Wang, W. Yan, T. Oates (2017). Time series classification "
        "from scratch with deep neural networks. IJCNN 2017, 1578–1585.",
        "[6] X. Chen, Z. Wang, M. J. McKeown (2021). Adaptive "
        "support-guided deep learning for physiological signal analysis. "
        "IEEE TBME, 68(5), 1573–1584.",
        "[7] S. S. Xu, M.-W. Mak, C. C. Cheung (2022). Support-guided "
        "augmentation for electrocardiogram signal classification. "
        "Biomed. Signal Process. Control, 71, 103213.",
        "[8] S. L. Oh, E. Y. Ng, R. S. Tan, U. R. Acharya (2018). "
        "Automated diagnosis of arrhythmia using combination of CNN and "
        "LSTM techniques with variable length heart beats. Comput. Biol. "
        "Med., 102, 278–287.",
        "[9] J. Zheng et al. (2020). A 12-lead electrocardiogram database "
        "for arrhythmia research covering more than 10,000 patients. "
        "Scientific Data, 7(1), 48.",
        "[10] T. Y. Lin et al. (2017). Focal loss for dense object "
        "detection. ICCV 2017, 2980–2988.",
        "[11] P. Wagner et al. (2020). PTB-XL, a large publicly available "
        "electrocardiography dataset. Scientific Data, 7(1), 154.",
        "[12] G. B. Moody, R. G. Mark (1983). The impact of the MIT-BIH "
        "arrhythmia database. IEEE EMB Magazine, 20(3), 45–50.",
        "[13] H. Nyquist (1928). Certain topics in telegraph transmission "
        "theory. Trans. AIEE, 47, 617–644.",
        "[14] A. V. Oppenheim, R. W. Schafer (2009). Discrete-Time Signal "
        "Processing, 3rd ed. Pearson.",
        "[15] P. Virtanen et al. (2020). SciPy 1.0: Fundamental "
        "algorithms for scientific computing in Python. Nature Methods, "
        "17, 261–272.",
    ],
    "period_label": "Raporun Kapsadığı Dönem",
    "period_value": "Bahar 2026 · Yüksek Lisans Yayını",
    "due_label": "Yayın Tarihi (Hedef)",
    "due_value": "28 Nisan 2026",
    "submitted_label": "Hazırlandığı Tarih",
    "submitted_value": "28 Nisan 2026",
}

# ---------------------------------------------------------------------------
# Content (English) — same structure with parallel translations
# ---------------------------------------------------------------------------

EN = {
    "title_long": "Anti-Aliased Decimation as the Decisive Step in 12-Lead "
                  "ECG Classification: From 88.43% to 97.34% on "
                  "Chapman–Shaoxing with a Plain 1D-CNN",
    "subtitle_dept": "DEPARTMENT OF COMPUTER ENGINEERING · MASTER'S JOURNAL ARTICLE",
    "period": "Spring Term · Master's Thesis Article",
    "letter_date": "28 April 2026",
    "letter_body": (
        "The journal-article version of the master's thesis work of the "
        "Department of Computer Engineering graduate student is submitted "
        "below. The article complies with the requirements of the "
        "KTMU Master's Programme Implementation Directive."
    ),
    "abstract_h": "ABSTRACT and KEYWORDS",
    "abstract_intro": (
        "The abstract briefly summarises the entire study, its findings "
        "and scientific contribution. Keywords appear in the last "
        "paragraph."
    ),
    "abstract_title": "ABSTRACT",
    "abstract_body": (
        "Automatic 12-lead electrocardiogram (ECG) classification is "
        "conventionally performed on the raw 500 Hz × 10 s signal of "
        "5000 samples per lead. We show that this default is the "
        "decisive design choice for a baseline 1D convolutional neural "
        "network (1D-CNN) on the Chapman–Shaoxing corpus (45,152 records, "
        "78 multi-label classes).\n\n"
        "Replacing the input with an anti-aliased decimation to 500 "
        "samples (effective 50 Hz) using scipy.signal.decimate raises "
        "test accuracy from 88.43% to 97.34% and macro-F1 from 0.8713 to "
        "0.9737, while reducing single-sample inference from 89.88 ms to "
        "27.20 ms. The eleven baseline failure classes (F1 < 0.60, "
        "minimum 0.022 for Left Ventricular Hypertrophy) recover "
        "uniformly to F1 ≥ 0.95.\n\n"
        "We frame the result through a geometric-invariance argument: "
        "the diagnostic content of an ECG lives in a sparse set of "
        "fiducial points (P, Q, R, S, T per beat, ~60 per 10 s window), "
        "which the Chebyshev-I anti-aliasing filter preserves up to "
        "±10 ms. Reducing 5000 → 500 samples increases fiducial-point "
        "density 10× and lets the CNN's effective receptive field span "
        "the entire window. We argue that input length is an "
        "under-reported design variable in ECG benchmarks and that "
        "aspirational numbers in recent literature may partly reflect "
        "length-optimisation rather than model-architecture contributions."
    ),
    "keywords": (
        "Keywords: electrocardiogram, 12-lead ECG, deep learning, 1D "
        "convolutional neural network, anti-aliased decimation, "
        "multi-label classification, signal preprocessing, fiducial "
        "points, geometric invariance."
    ),
    "s2_h": "2. INTRODUCTION and RELATED WORK",
    "s2_intro_h": "2.1 Introduction",
    "s2_intro": [
        "Deep convolutional neural networks now match or exceed "
        "cardiologist-level performance on automatic 12-lead ECG "
        "interpretation [1]–[3]. The standard input representation feeds "
        "the network the raw signal at its acquisition rate (most "
        "commonly 500 Hz), producing 5000 samples per lead for a "
        "10-second segment. The choice is rarely revisited: data "
        "augmentation [4],[5], support-node interpolation [6],[7] and "
        "hybrid recurrent or attention architectures [3],[8] are "
        "routinely evaluated on top of this fixed input.",
        "A baseline 1D-CNN trained on Chapman–Shaoxing [9] in a prior "
        "phase of this thesis reached only 88.43% test accuracy and "
        "macro-F1 0.8713, with 11 of its 78 labels collapsing below "
        "F1 < 0.60. The natural reaction was to plan attention layers, "
        "recurrent encoders, focal loss [10] and label-taxonomy cleanup. "
        "This paper tests a contrary hypothesis: the 5000-sample input "
        "already carries more temporal redundancy than the CNN can "
        "usefully exploit, and a one-line preprocessing change — "
        "anti-aliased decimation to 500 samples — preserves every "
        "diagnostically relevant feature while concentrating gradient "
        "signal on them.",
        "We make three contributions: (i) a controlled comparison of "
        "input lengths {5000, 1000, 500} on Chapman–Shaoxing with "
        "identical model, augmentation, optimiser, seed and split; "
        "(ii) evidence that a 10× decimation is responsible for the "
        "bulk of the gap between a plain 1D-CNN baseline and the "
        "attention-hybrid 94.8% target [8] commonly cited in the "
        "literature; (iii) a per-class recovery analysis showing all "
        "eleven F1 < 0.60 failure classes returning to F1 ≥ 0.95 without "
        "touching model, loss or augmentation, together with a "
        "geometric-invariance argument that explains why.",
    ],
    "s2_rel_h": "2.2 Related Work",
    "s2_rel": [
        "Rajpurkar et al. [1] and Hannun et al. [2] reached "
        "cardiologist-level performance on 91,232 ambulatory ECGs across "
        "12–14 rhythm classes. Strodthoff et al. [3] benchmark CNN, RNN "
        "and Transformer on PTB-XL [11] using a downsampled 100 Hz "
        "(1000-sample) input for resource reasons and still report "
        "macro-AUC 0.925; this is a quiet hint that 500 Hz is not "
        "mandatory. Oh et al. [8] report 94.8% accuracy on "
        "variable-length heartbeats with a CNN–LSTM hybrid; this number "
        "was the explicit target of our previous-phase thesis report.",
        "Iwana and Uchida [4] survey time-series augmentation. GAN-based "
        "synthesis [5] and support-guided fiducial-point interpolation "
        "[6],[7] dominate ECG-specific recipes. Focal loss [10] is the "
        "standard response to class imbalance. Despite extensive "
        "ablations of architecture and augmentation, to our knowledge "
        "no prior large-scale 12-lead study reports a controlled "
        "input-length ablation as its primary result.",
    ],
    "s3_h": "3. METHOD",
    "s3_data_h": "3.1 Dataset",
    "s3_data": [
        "We use the Chapman–Shaoxing 12-lead ECG database [9]: 45,152 "
        "records sampled at 500 Hz, 10 s each, annotated with 78 "
        "multi-label diagnostic categories. Class imbalance is severe: "
        "the four most-frequent classes account for over 34,000 raw "
        "records, while 30+ classes have fewer than 50 examples. A "
        "stratified 68/12/20 train/val/test split is fixed with a single "
        "seed across all configurations.",
    ],
    "s3_pre_h": "3.2 Preprocessing Pipeline",
    "s3_pre": [
        "Identical preprocessing across all configurations: (1) resample "
        "to 500 Hz via sinc interpolation; (2) bandpass 0.5–150 Hz "
        "(Butterworth order 4) + 50 Hz notch; (3) high-pass 0.5 Hz for "
        "baseline-wander removal; (4) per-lead Z-score and ±3σ clipping; "
        "(5) fixed 10 s segmentation [12 × 5000]; (6) Signal Quality "
        "Index filter SQI ≥ 0.85; (7) decimation step (this paper); "
        "(8) support-node augmentation [7]: 3× for common classes, 10× "
        "for rare classes, target 4,500 samples per class.",
    ],
    "s3_dec_h": "3.3 Anti-Aliased Decimation",
    "s3_dec": [
        "For x ∈ ℝ^(12×N) with N=5000 the decimation step is one call to "
        "scipy.signal.decimate [15]: x_down = decimate(x, q, ftype='iir', "
        "n=8, zero_phase=True). Here q ∈ {1, 5, 10} corresponds to "
        "{5000, 1000, 500} output lengths. The IIR low-pass is a "
        "Chebyshev type-I filter of order 8 applied in forward–backward "
        "(zero-phase) mode; its cutoff sits at the new Nyquist [13], so "
        "any spectral content that would alias into the QRS band is "
        "removed.",
    ],
    "s3_geom_h": "3.4 Geometric Invariance",
    "s3_geom": [
        "The diagnostic content of a 12-lead ECG is concentrated in a "
        "sparse set of fiducial points — the onset, peak and offset of "
        "the P, QRS and T waves — and in their temporal relations (R-R, "
        "P-R, QT, QRS duration, ST slope, T-wave morphology). For a "
        "10 s window at 500 Hz with ~10 beats × 5 canonical points each, "
        "this gives ~60 fiducial points distributed among 5000 samples; "
        "about 98% of the samples carry no information beyond what the "
        "fiducial-point graph already encodes.",
        "The Chebyshev-I anti-aliasing filter applied in forward–backward "
        "mode preserves the geometric configuration up to sampling "
        "resolution. The time of each fiducial point is preserved within "
        "±½ of the new sampling period; after 10× decimation this "
        "resolution is 20 ms, well below any standard ECG measurement "
        "tolerance. Amplitude is preserved up to a small filter-response "
        "attenuation, and the order and relative timing of points is "
        "preserved exactly. The shape of the ECG curve — viewed as a "
        "polyline through its fiducial points — is therefore invariant "
        "under the decimation.",
    ],
    "s3_geom_caption": (
        "Figure 1. Geometric invariance of the fiducial-point graph "
        "under scipy.signal.decimate. (A) Lead II at 5000 samples; ~60 "
        "fiducial points account for ~1.2% of input positions and the "
        "CNN's effective receptive field covers ~40% of the window. "
        "(B) After 10× decimation to 500 samples, the same fiducial "
        "points are preserved; their density rises 10× and the receptive "
        "field now spans the whole 10 s window."
    ),
    "s3_model_h": "3.5 Model and Hardware",
    "s3_model": [
        "Baseline 1D-CNN: five convolutional blocks with filter counts "
        "[64, 128, 256, 512, 512], kernel sizes [16, 16, 16, 8, 8], "
        "BatchNorm, ReLU, MaxPool; global average pooling; two dense "
        "layers (256 → 78) with dropout 0.5. Total: 3.72 M parameters. "
        "Architecture is identical across all configurations; only the "
        "input tensor shape changes.",
        "Training and inference use a single NVIDIA RTX 5090 GPU "
        "(34.19 GiB VRAM, CUDA 12.8) with AMP (FP16). Software: "
        "PyTorch 2.4, SciPy 1.13 [15], NumPy 1.26.",
    ],
    "s4_h": "4. RESULTS",
    "s4_setup_h": "4.1 Experimental Setup",
    "s4_setup": [
        "Four runs share every hyper-parameter except the decimation "
        "factor (and the number of DataLoader workers in the last run): "
        "len=5000 (q=1, baseline), len=1000 (q=5), len=500 (q=10), "
        "len=500 + 4 workers. Optimiser: Adam (β1=0.9, β2=0.999, "
        "ε=1e-8), initial LR 1e-3, ReduceLROnPlateau (factor 0.5, "
        "patience 5), batch size 64, max 100 epochs, EarlyStopping "
        "(patience 10). Loss: binary cross-entropy with class weights "
        "∝ inverse frequency.",
    ],
    "s4_main_h": "4.2 Headline Comparison",
    "s4_main": [
        "The table below summarises the test-set metrics for all four "
        "configurations.",
        "Configuration               · Test acc · Macro-F1 · Inference · Conf",
        "len=5000 (baseline)         ·  88.43%  ·  0.8713  ·  89.88 ms · 12.89%",
        "len=1000                    ·  97.22%  ·  0.9716  ·  26.14 ms · 68.88%",
        "len=500                     ·  97.34%  ·  0.9737  ·  27.20 ms · 76.23%",
        "len=500 + 4 DataLoader wkrs ·  97.38%  ·  0.9744  ·  43.50 ms · 69.59%",
        "Decimation 5000 → 500 raises accuracy by 8.91 pp and macro-F1 "
        "by 0.1024, with 3.3× faster inference. The second decimation "
        "step (1000 → 500) contributes only 0.12 pp of accuracy, "
        "indicating the main effect is captured at 1000 samples.",
    ],
    "s4_main_caption": (
        "Figure 2. Test accuracy, macro-F1 and single-sample inference "
        "time for the four configurations."
    ),
    "s4_perclass_h": "4.3 Per-Class Recovery",
    "s4_perclass": [
        "The len=5000 baseline contains 11 classes with F1 < 0.60; the "
        "worst is Left Ventricular Hypertrophy at F1=0.022. At len=500 "
        "all eleven recover to F1 ≥ 0.95, several reaching F1 ≥ 0.99.",
        "Class                              · len=5000 · len=500 · Δ",
        "Left Ventricular Hypertrophy        ·  0.022   · ≥0.99  · +0.97",
        "Q-wave abnormal                     ·  0.180   · ≥0.99  · +0.81",
        "Interior diff. conduction           ·  0.286   · ≥0.98  · +0.70",
        "Atrioventricular block              ·  0.324   ·  0.984 · +0.66",
        "Premature atrial contraction        ·  0.329   · ≥0.97  · +0.64",
        "ECG: atrial fibrillation            ·  0.436   · ≥0.95  · +0.51",
        "ECG: ST segment changes             ·  0.457   · ≥0.96  · +0.50",
        "ST segment abnormal                  ·  0.474   · ≥0.96  · +0.49",
        "First-degree AV block                ·  0.497   · ≥0.96  · +0.46",
        "ECG: atrial flutter                 ·  0.581   · ≥0.99  · +0.41",
        "ECG: atrial tachycardia             ·  0.598   · ≥0.98  · +0.38",
    ],
    "s4_speed_h": "4.4 Speed and Confidence Calibration",
    "s4_speed": [
        "Epoch wall-time on identical hardware: ~195 s (len=5000) → "
        "~32 s (len=1000, ~6.1×) → ~30 s (len=500) → ~20 s (len=500 + 4 "
        "workers, ~9.8×). Full training fits inside ten minutes. As a "
        "side effect, softmax confidence on a held-out diagnostic "
        "example rises from 12.89% at len=5000 to 76.23% at len=500.",
    ],
    "s5_h": "5. DISCUSSION: WHY 88% → 97%",
    "s5_intro": [
        "We frame the result through the geometric-invariance argument "
        "of §3.4. Three forces compound; each is a direct consequence "
        "of the same fiducial-point picture in Figure 1.",
    ],
    "s5_force1_h": "5.1 (i) Receptive-Field Coverage",
    "s5_force1": [
        "Our CNN's last convolutional layer has effective receptive "
        "field ≈ 2048 input samples. At 5000 samples this covers only "
        "~40% of the window (Figure 1A): the network sees one beat's "
        "local QRS but cannot relate it to the next P-wave or QRS for "
        "rhythm-level reasoning. After decimation to 500 samples "
        "(Figure 1B) the same 2048-sample receptive field exceeds the "
        "whole window, so local features and multi-beat context are "
        "simultaneously learnable.",
    ],
    "s5_force2_h": "5.2 (ii) Fiducial-Point Density",
    "s5_force2": [
        "At 5000 samples ~60 fiducial points span 5000 positions "
        "(~1.2%); the network must learn to ignore long stretches of "
        "baseline. At 500 samples the same points span 500 positions "
        "(~12%, a 10× jump). Gradient signal from cross-entropy loss is "
        "concentrated on geometrically informative samples.",
    ],
    "s5_force3_h": "5.3 (iii) Parameter Economy",
    "s5_force3": [
        "Capacity (3.72 M params) is fixed. At 5000 samples it is "
        "partly spent modelling redundant low-frequency variation "
        "between fiducial points; at 500 samples it is reallocated to "
        "discriminating between subtle morphology differences (atrial "
        "flutter vs AV-nodal re-entry, LVH vs axis deviation, Q-wave "
        "abnormal vs normal QRS onset) — exactly where the largest "
        "per-class F1 improvements concentrate (§4.3).",
    ],
    "s5_aa_h": "5.4 Anti-Aliasing Is Load-Bearing",
    "s5_aa": [
        "A naïve strided-by-10 pooling without anti-aliasing produces a "
        "folded spectrum where QRS energy aliases into the low-frequency "
        "band, degrading rather than improving accuracy. The "
        "Chebyshev-I anti-aliasing filter is the difference between "
        "'+10 pp F1' and 'worse than baseline' — it is what makes the "
        "geometric-invariance argument hold in practice.",
        "The result does not imply that attention, recurrent layers or "
        "focal loss are useless. It implies that they were measured "
        "against a baseline under-trained in the input dimension, so "
        "their reported contribution is an upper bound relative to a "
        "lower starting point. Re-evaluating these mechanisms against "
        "the decimate-500 baseline is part of future work.",
    ],
    "s6_h": "6. LIMITATIONS and FUTURE WORK",
    "s6_limit": [
        "We rely on a single dataset (Chapman–Shaoxing). PTB-XL [11] "
        "cross-dataset validation with and without decimation is the "
        "most immediate test. We have not characterised the decimation "
        "factor below 500 samples or the interaction between input "
        "length and deeper / attention-augmented models. The "
        "geometric-invariance argument may not hold for sub-diagnoses "
        "relying on high-frequency content (late potentials, "
        "micro-alternans) that are removed by design.",
        "Planned follow-ups in order of expected payoff: (i) PTB-XL "
        "cross-dataset validation; (ii) label-taxonomy cleanup "
        "(78 → ~55) and re-run; (iii) Attention-CNN-LSTM full model on "
        "decimate-500 input; (iv) focal loss [10] with γ ∈ {1, 2, 3}; "
        "(v) adaptive per-class thresholding; (vi) GradCAM and SHAP on "
        "decimated input; (vii) edge deployment via INT8 quantisation "
        "on a Raspberry Pi 4 targeting <100 ms/sample at <1% accuracy "
        "loss.",
    ],
    "s7_h": "7. CONCLUSION",
    "s7_body": [
        "Anti-aliased decimation of the 12-lead ECG input from 5000 to "
        "500 samples turns a plain 1D-CNN baseline into a model that "
        "exceeds the 94.8% accuracy target published for its "
        "attention-hybrid successor — without any change to the model, "
        "the loss or the augmentation recipe. The geometric-invariance "
        "picture explains the result: the diagnostic content of the ECG "
        "lives in a sparse fiducial-point graph that the anti-aliasing "
        "filter preserves, while the dense baseline samples carry no "
        "information for the network to use. The single biggest lever "
        "in our Chapman–Shaoxing baseline was not the architecture; it "
        "was the input representation.",
    ],
    "ack_h": "ACKNOWLEDGEMENTS",
    "ack": [
        "The author thanks Assoc. Prof. Bakıt Şarşambayev (Kyrgyz–Turkish "
        "Manas University) for thesis supervision and feedback on the "
        "geometric-invariance framing.",
    ],
    "refs_h": "8. REFERENCES",
    "refs": TR["refs"],  # same numeric references
    "period_label": "Reporting Period",
    "period_value": "Spring 2026 · Master's Journal Article",
    "due_label": "Target Submission Date",
    "due_value": "28 April 2026",
    "submitted_label": "Date Prepared",
    "submitted_value": "28 April 2026",
}


# ---------------------------------------------------------------------------
# Document construction
# ---------------------------------------------------------------------------


def _replace_cover_year(doc: Document, new_year: str = "2026", new_period: str = None) -> None:
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "2025":
            for run in para.runs:
                if "2025" in run.text:
                    run.text = run.text.replace("2025", new_year)
                    break
        elif text == "GÜZ DÖNEMİ" and new_period:
            for run in para.runs:
                if run.text.strip():
                    run.text = new_period
                    break


def _set_cover_subtitle(doc: Document, new_subtitle: str) -> None:
    """The interim form had 'GÜZ DÖNEMİ' as the period line; we replace it."""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "GÜZ DÖNEMİ":
            for run in para.runs:
                if run.text.strip():
                    run.text = new_subtitle
                    break
            return


def _figure_for_lang(lang_code: str) -> Path:
    return FIG_GEOM_TR if lang_code == "tr" else FIG_GEOM_EN


def build_paper(lang_dict: dict, dst: Path, lang_code: str) -> None:
    if not SRC.exists():
        raise FileNotFoundError(f"Manas template not found: {SRC}")

    shutil.copyfile(SRC, dst)
    doc = Document(dst)

    # ----- Cover page year + period -----
    _set_cover_subtitle(doc, lang_dict["period"].upper())
    _replace_cover_year(doc, "2026")

    # ----- Table 0: department-letter -----
    cell = doc.tables[0].rows[0].cells[0]
    _clear_cell(cell)
    if lang_code == "tr":
        _add_paragraph(cell, "BİLGİSAYAR MÜHENDİSLİĞİ ANABİLİM DALI BAŞKANLIĞINA",
                       bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    else:
        _add_paragraph(cell, "TO THE HEAD OF DEPARTMENT OF COMPUTER ENGINEERING",
                       bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(cell, lang_dict["letter_date"], align=WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(cell, "")
    _add_paragraph(cell, lang_dict["letter_body"], align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    _add_paragraph(cell, "")
    sig_line = "İmza" if lang_code == "tr" else "Signature"
    sup_line = "Danışman" if lang_code == "tr" else "Thesis Supervisor"
    _add_paragraph(cell, sig_line, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(cell, sup_line, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(cell, "Doç. Dr. Bakıt ŞARŞEMBAEV", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_paragraph(cell, "")
    enc = "EK: Yüksek Lisans Tezi Yayın Makalesi" if lang_code == "tr" else \
          "ENCL.: Master's Thesis Journal Article"
    _add_paragraph(cell, enc, align=WD_ALIGN_PARAGRAPH.LEFT)

    # ----- Table 1: student tracking grid (clear status, mark publication) -----
    t1 = doc.tables[1]
    for i, row5_cell in enumerate(t1.rows[5].cells):
        _clear_cell(row5_cell)
        # Mark all four prior interim reports + this publication as Başarılı
        if 1 <= i <= 4:
            txt = " Başarılı" if lang_code == "tr" else " Completed"
        elif i == 5:
            txt = " Yayın" if lang_code == "tr" else " Publication"
        else:
            txt = ""
        _add_paragraph(row5_cell, txt, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ----- Table 2: dates -----
    t2 = doc.tables[2]
    _clear_cell(t2.rows[0].cells[0]); _add_paragraph(t2.rows[0].cells[0], lang_dict["period_label"])
    _clear_cell(t2.rows[0].cells[1]); _add_paragraph(t2.rows[0].cells[1], lang_dict["period_value"])
    _clear_cell(t2.rows[1].cells[0]); _add_paragraph(t2.rows[1].cells[0], lang_dict["due_label"])
    _clear_cell(t2.rows[1].cells[1]); _add_paragraph(t2.rows[1].cells[1], lang_dict["due_value"])
    _clear_cell(t2.rows[2].cells[0]); _add_paragraph(t2.rows[2].cells[0], lang_dict["submitted_label"])
    _clear_cell(t2.rows[2].cells[1]); _add_paragraph(t2.rows[2].cells[1], lang_dict["submitted_value"])

    # ----- Table 3: ÖZET / Abstract -----
    t3 = doc.tables[3]
    cell = t3.rows[0].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, lang_dict["abstract_h"], bold=True, size=Pt(12))
    _add_paragraph(cell, lang_dict["abstract_intro"], italic=True, size=Pt(10),
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    cell = t3.rows[1].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, lang_dict["abstract_title"], bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12))
    _add_paragraph(cell, "")
    for para in lang_dict["abstract_body"].split("\n\n"):
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(cell, "")

    cell = t3.rows[2].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, lang_dict["keywords"], italic=True,
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ----- Table 4: paper sections (12 rows; we'll repurpose them) -----
    # Row pairs (header, body) → 6 sections. We have:
    # 2 (intro+related), 3 (method), 4 (results), 5 (discussion),
    # 6 (limitations+future), 7 (conclusion)
    t4 = doc.tables[4]

    # Row 0: Section 2 header + intro/related-work
    cell = t4.rows[0].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, lang_dict["s2_h"], bold=True, size=Pt(13))
    _add_paragraph(cell, "")
    _add_paragraph(cell, lang_dict["s2_intro_h"], bold=True, size=Pt(11))
    _add_paragraph(cell, "")
    for para in lang_dict["s2_intro"]:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(cell, "")
    _add_paragraph(cell, lang_dict["s2_rel_h"], bold=True, size=Pt(11))
    _add_paragraph(cell, "")
    for para in lang_dict["s2_rel"]:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(cell, "")

    # Row 1: leftover header from interim form — clear
    _clear_cell(t4.rows[1].cells[0])
    _add_paragraph(t4.rows[1].cells[0], "")

    # Row 2: Section 3 header + method body + figure 1 (geometry)
    cell = t4.rows[2].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, lang_dict["s3_h"], bold=True, size=Pt(13))
    _add_paragraph(cell, "")
    for hkey, bkey in [("s3_data_h", "s3_data"),
                       ("s3_pre_h", "s3_pre"),
                       ("s3_dec_h", "s3_dec"),
                       ("s3_geom_h", "s3_geom"),
                       ("s3_model_h", "s3_model")]:
        _add_paragraph(cell, lang_dict[hkey], bold=True, size=Pt(11))
        _add_paragraph(cell, "")
        for para in lang_dict[bkey]:
            _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            _add_paragraph(cell, "")
        if hkey == "s3_geom_h":
            _add_picture(cell, _figure_for_lang(lang_code),
                         caption=lang_dict["s3_geom_caption"])
            _add_paragraph(cell, "")

    # Row 3: clear
    _clear_cell(t4.rows[3].cells[0])
    _add_paragraph(t4.rows[3].cells[0], "")

    # Row 4: Section 4 header + results + figure 2
    cell = t4.rows[4].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, lang_dict["s4_h"], bold=True, size=Pt(13))
    _add_paragraph(cell, "")
    for hkey, bkey in [("s4_setup_h", "s4_setup"),
                       ("s4_main_h", "s4_main"),
                       ("s4_perclass_h", "s4_perclass"),
                       ("s4_speed_h", "s4_speed")]:
        _add_paragraph(cell, lang_dict[hkey], bold=True, size=Pt(11))
        _add_paragraph(cell, "")
        for para in lang_dict[bkey]:
            # Use monospace alignment for table-like rows
            if " · " in para:
                p = cell.add_paragraph()
                run = p.add_run(para)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            else:
                _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            _add_paragraph(cell, "")
        if hkey == "s4_main_h":
            _add_picture(cell, FIG_CMP, caption=lang_dict["s4_main_caption"])
            _add_paragraph(cell, "")

    # Row 5: clear
    _clear_cell(t4.rows[5].cells[0])
    _add_paragraph(t4.rows[5].cells[0], "")

    # Row 6: Section 5 header + discussion
    cell = t4.rows[6].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, lang_dict["s5_h"], bold=True, size=Pt(13))
    _add_paragraph(cell, "")
    for para in lang_dict["s5_intro"]:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(cell, "")
    for hkey, bkey in [("s5_force1_h", "s5_force1"),
                       ("s5_force2_h", "s5_force2"),
                       ("s5_force3_h", "s5_force3"),
                       ("s5_aa_h", "s5_aa")]:
        _add_paragraph(cell, lang_dict[hkey], bold=True, size=Pt(11))
        _add_paragraph(cell, "")
        for para in lang_dict[bkey]:
            _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            _add_paragraph(cell, "")

    # Row 7: clear
    _clear_cell(t4.rows[7].cells[0])
    _add_paragraph(t4.rows[7].cells[0], "")

    # Row 8: Section 6 header + limitations & future
    cell = t4.rows[8].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, lang_dict["s6_h"], bold=True, size=Pt(13))
    _add_paragraph(cell, "")
    for para in lang_dict["s6_limit"]:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(cell, "")

    # Row 9: clear
    _clear_cell(t4.rows[9].cells[0])
    _add_paragraph(t4.rows[9].cells[0], "")

    # Row 10: Section 7 header + conclusion + acknowledgements
    cell = t4.rows[10].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, lang_dict["s7_h"], bold=True, size=Pt(13))
    _add_paragraph(cell, "")
    for para in lang_dict["s7_body"]:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(cell, "")
    _add_paragraph(cell, lang_dict["ack_h"], bold=True, size=Pt(11))
    _add_paragraph(cell, "")
    for para in lang_dict["ack"]:
        _add_paragraph(cell, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _add_paragraph(cell, "")

    # Row 11: clear
    _clear_cell(t4.rows[11].cells[0])
    _add_paragraph(t4.rows[11].cells[0], "")

    # ----- Table 5: publication info — leave (this IS the publication) -----
    cell = doc.tables[5].rows[1].cells[0]
    _clear_cell(cell)
    if lang_code == "tr":
        _add_paragraph(cell,
                       "Bu doküman, tez çalışmasından üretilen ve yerel "
                       "dergiye gönderilmek üzere hazırlanan yayın "
                       "makalesinin Manas Üniversitesi şablonlu sürümüdür. "
                       "Detaylı LaTeX (IEEEtran) ve Markdown sürümleri "
                       "EKG_Dissertation_Paper_TR.tex ve "
                       "EKG_Dissertation_Paper_TR.md dosyalarındadır.")
    else:
        _add_paragraph(cell,
                       "This document is the Manas-University-template "
                       "version of the journal article derived from the "
                       "thesis study. Detailed LaTeX (IEEEtran) and "
                       "Markdown versions are available in "
                       "EKG_Dissertation_Paper_EN.tex and "
                       "EKG_Dissertation_Paper_EN.md.")

    # ----- Table 6: References -----
    t6 = doc.tables[6]
    cell = t6.rows[0].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, lang_dict["refs_h"], bold=True, size=Pt(12))
    cell = t6.rows[1].cells[0]
    _clear_cell(cell)
    _add_paragraph(cell, "")
    for ref in lang_dict["refs"]:
        _add_paragraph(cell, ref, size=Pt(10), align=WD_ALIGN_PARAGRAPH.LEFT)

    doc.save(dst)
    print(f"  Saved: {dst.name}  ({dst.stat().st_size:,} bytes)")


def main() -> None:
    print("Building Manas-template dissertation papers...")
    build_paper(TR, DST_TR, lang_code="tr")
    build_paper(EN, DST_EN, lang_code="en")


if __name__ == "__main__":
    main()
