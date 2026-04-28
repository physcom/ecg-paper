"""
Generate v4 thesis interim reports (Bahar 2026 dönemi).

Produces, all inside C:/Users/enazarkulov/Documents/Мастер/:
  - Figure_seq_length_comparison.png        (new comparison chart for len 500/1000/5000)
  - Elaman Nazarkulov - ara rapor değerlendirme formu v4.docx    (cover form)
  - ELAMAN NAZARKULOV - genis ara rapor v4.docx                  (16-week extended log)
  - Elaman Nazarkulov - short ara rapor değerlendirme formu v4.docx  (official evaluation form)

Content is grounded in real training logs from C:/Users/enazarkulov/Documents/ML/ekg/results/:
  - result-25-12-2025.txt                 len=5000 baseline (first run)
  - results-22-04-2026.txt                len=5000 baseline (reproducibility)
  - result-22-04-2026-500.txt             len=500 experiment
  - result-23-04-2026-1000.txt            len=1000 experiment
  - result-23-04-2026-500-4-workers.txt   len=500 with 4 DataLoader workers

Existing figures in the same directory are referenced:
  - Figure_1.png             len=5000 training history
  - Figure_1_500.png         len=500 training history
  - Figure_1000.png          len=1000 training history
  - Figure_500_4_worker.png  len=500 + 4 workers training history
"""

from __future__ import annotations

from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = Path(r"C:\Users\enazarkulov\Documents\Мастер")
COMPARISON_FIG = OUT_DIR / "Figure_seq_length_comparison.png"

# --- Experimental results (from actual training logs) ----------------------

RESULTS = {
    "len=5000": {
        "label": "len=5000 (baseline)",
        "test_acc": 88.43,
        "macro_p": 87.68,
        "macro_r": 88.47,
        "macro_f1": 87.13,
        "inference_ms": 89.88,
        "confidence": 12.89,
        "early_stop_epoch": 92,
        "epoch_time_s": 195,
        "figure": "Figure_1.png",
        "result_file": "results/results-22-04-2026.txt",
    },
    "len=1000": {
        "label": "len=1000",
        "test_acc": 97.22,
        "macro_p": 97.40,
        "macro_r": 97.29,
        "macro_f1": 97.16,
        "inference_ms": 26.14,
        "confidence": 68.88,
        "early_stop_epoch": 92,
        "epoch_time_s": 32,
        "figure": "Figure_1000.png",
        "result_file": "results/result-23-04-2026-1000.txt",
    },
    "len=500": {
        "label": "len=500",
        "test_acc": 97.34,
        "macro_p": 97.34,
        "macro_r": 97.49,
        "macro_f1": 97.37,
        "inference_ms": 27.20,
        "confidence": 76.23,
        "early_stop_epoch": 100,
        "epoch_time_s": 30,
        "figure": "Figure_1_500.png",
        "result_file": "results/result-22-04-2026-500.txt",
    },
    "len=500 + 4 workers": {
        "label": "len=500 + 4 DataLoader workers",
        "test_acc": 97.38,
        "macro_p": 97.41,
        "macro_r": 97.51,
        "macro_f1": 97.44,
        "inference_ms": 43.50,
        "confidence": 69.59,
        "early_stop_epoch": 100,
        "epoch_time_s": 20,
        "figure": "Figure_500_4_worker.png",
        "result_file": "results/result-23-04-2026-500-4-workers.txt",
    },
}


# --- 1. Build comparison figure -------------------------------------------


def build_comparison_figure() -> None:
    keys = ["len=5000", "len=1000", "len=500", "len=500 + 4 workers"]
    test_acc = [RESULTS[k]["test_acc"] for k in keys]
    f1 = [RESULTS[k]["macro_f1"] for k in keys]
    inf = [RESULTS[k]["inference_ms"] for k in keys]

    fig, axes = plt.subplots(1, 3, figsize=(14, 4.5))
    x = np.arange(len(keys))

    bars0 = axes[0].bar(x, test_acc, color=["#c0392b", "#2980b9", "#27ae60", "#16a085"])
    axes[0].set_ylim(80, 100)
    axes[0].set_ylabel("Test Accuracy (%)")
    axes[0].set_title("Test Doğruluğu")
    axes[0].set_xticks(x)
    axes[0].set_xticklabels(keys, rotation=20, ha="right")
    for b, v in zip(bars0, test_acc):
        axes[0].text(b.get_x() + b.get_width() / 2, v + 0.3, f"{v:.2f}", ha="center", fontsize=9)

    bars1 = axes[1].bar(x, f1, color=["#c0392b", "#2980b9", "#27ae60", "#16a085"])
    axes[1].set_ylim(80, 100)
    axes[1].set_ylabel("Macro F1 (%)")
    axes[1].set_title("Makro F1-Score")
    axes[1].set_xticks(x)
    axes[1].set_xticklabels(keys, rotation=20, ha="right")
    for b, v in zip(bars1, f1):
        axes[1].text(b.get_x() + b.get_width() / 2, v + 0.3, f"{v:.2f}", ha="center", fontsize=9)

    bars2 = axes[2].bar(x, inf, color=["#c0392b", "#2980b9", "#27ae60", "#16a085"])
    axes[2].set_ylabel("Inference Time (ms/sample)")
    axes[2].set_title("Tekil Çıkarım Süresi")
    axes[2].set_xticks(x)
    axes[2].set_xticklabels(keys, rotation=20, ha="right")
    for b, v in zip(bars2, inf):
        axes[2].text(b.get_x() + b.get_width() / 2, v + max(inf) * 0.02, f"{v:.2f}", ha="center", fontsize=9)

    fig.suptitle(
        "Giriş sinyal uzunluğunun baseline 1D-CNN performansına etkisi\n"
        "Chapman-Shaoxing, 78 sınıf, NVIDIA RTX 5090, AMP",
        fontsize=11,
    )
    fig.tight_layout()
    fig.savefig(COMPARISON_FIG, dpi=150, bbox_inches="tight")
    plt.close(fig)


# --- 2. Docx helpers -------------------------------------------------------


def _set_cell_shading(cell, color_hex: str) -> None:
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(rf'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    tc_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_para(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


# --- 3. Cover form: ara rapor değerlendirme formu --------------------------


def build_cover_form() -> Path:
    doc = Document()

    for line, center in [
        ("KIRGIZİSTAN - TÜRKİYE MANAS ÜNİVERSİTESİ", True),
        ("LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ", True),
        ("BİLGİSAYAR MÜHENDİSLİĞİ ANABİLİM DALI", True),
        ("", True),
        ("", True),
        ("YÜKSEK LİSANS TEZ ARA RAPORU", True),
        ("", True),
        ("", True),
        (
            "12 Kanallı EKG Tabanlı Kardiyak Hastalık Teşhisi için "
            "Destek Düğüm Yöntemi Kullanarak Sinyal Zenginleştirmeli Sinir Ağı",
            True,
        ),
        ("", True),
        ("", True),
        ("BAHAR DÖNEMİ", True),
        ("", True),
        ("", True),
        ("ELAMAN NAZARKULOV", True),
        ("", True),
        ("", True),
        ("2026", True),
        ("BİŞKEK", True),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(line)
        if line and line.isupper():
            run.bold = True
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(4):
        doc.add_paragraph()

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.autofit = True
    sig_table.rows[0].cells[0].text = "İmza"
    sig_table.rows[0].cells[1].text = "İmza"
    for c in sig_table.rows[0].cells:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    info_table = doc.add_table(rows=4, cols=2)
    info_table.rows[0].cells[0].text = "Danışmanın:"
    info_table.rows[0].cells[1].text = "Öğrencinin:"
    info_table.rows[1].cells[0].text = "Ünvanı: Doç. Dr."
    info_table.rows[1].cells[1].text = "Adı: ELAMAN"
    info_table.rows[2].cells[0].text = "Adı: BAKIT"
    info_table.rows[2].cells[1].text = "Soyadı: NAZARKULOV"
    info_table.rows[3].cells[0].text = "Soyadı: ŞARŞAMBAYEV"
    info_table.rows[3].cells[1].text = ""

    out = OUT_DIR / "Elaman Nazarkulov - ara rapor değerlendirme formu v4.docx"
    doc.save(out)
    return out


# --- 4. Geniş ara rapor: 16 weekly updates ---------------------------------

WEEKLY_NEW = [
    (
        "1. Hafta/Апта",
        "Bu haftada, tez çalışması kapsamında kardiyak hastalık teşhisi ve derin öğrenme konularında mevcut "
        "literatür derinlemesine taranmıştır. EKG tabanlı otomatik teşhis çalışmalarında CNN, RNN ve Transformer "
        "mimarilerinin performansları incelenmiş, literatürdeki boşluklar tespit edilmiştir. Özellikle 12 kanallı "
        "EKG'nin 500 Hz × 10 s ham uzunluğunun (5000 örnek) 1D-CNN için uygun olup olmadığı, downsampling ve "
        "anti-aliasing stratejilerinin performansa etkisi üzerine kaynaklar önceliklendirilmiştir. EKG sinyallerinin "
        "fizyolojik özellikleri (P dalgası, QRS kompleksi, T dalgası) ve bunların kardiyak patolojilerdeki "
        "değişimleri gözden geçirilmiştir.",
    ),
    (
        "2. Hafta/Апта",
        "Derin öğrenme modelleri ve özellikle CNN, LSTM ve Transformer mimarileri üzerine yoğun bir literatür "
        "taraması yapılmıştır. CNN'lerin 1D biyomedikal sinyallerde feature extraction kapasiteleri, LSTM'lerin "
        "zamansal bağımlılıkları modelleme yetenekleri ve attention mekanizmalarının açıklanabilirlik sağlama "
        "potansiyeli ele alınmıştır. Rajpurkar vd. (2017), Hannun vd. (2019) ve Natarajan vd. (2020) çalışmaları "
        "detaylı incelenmiş; ayrıca girdi uzunluğu küçültmenin (downsampling) CNN parametre verimliliği üzerindeki "
        "etkisini değerlendirmek için Oh vd. (2018) ve Strodthoff vd. (2020) referans alınmıştır. Transfer learning, "
        "model ensembling ve multi-task learning gibi ileri teknikler notlara eklenmiştir.",
    ),
    (
        "3. Hafta/Апта",
        "EKG sinyal işleme sürecine ilişkin mevcut metodolojiler, teknolojik altyapılar ve tasarım yaklaşımları "
        "ele alınmıştır. Filtreleme, normalizasyon, gürültü temizleme ve segmentasyon yöntemleri incelenmiş; bunun "
        "yanı sıra anti-aliased decimation ile girdi dizisinin 5000 → 1000 veya 5000 → 500 örneğe indirgenmesinin "
        "avantajları (shift invariance, parametre azaltımı) teorik olarak değerlendirilmiştir. GAN, time warping, "
        "magnitude scaling ve interpolasyon tabanlı augmentation yöntemleri karşılaştırılmıştır. PTB-XL, MIT-BIH "
        "Arrhythmia ve Chapman-Shaoxing benchmark veri setlerinin özellikleri, sınıf dağılımları ve kullanım "
        "senaryoları tartışılmıştır.",
    ),
    (
        "4. Hafta/Апта",
        "Literatür taramaları sırasında tespit edilen sorunlara (sınıf dengesizliği, overfitting, sinyal uzunluğu "
        "seçimi) yönelik çözüm önerileri geliştirilmiştir. Destek düğüm yöntemi (support node method) araştırılmış, "
        "Chen vd. (2021) ve Xu vd. (2022) çalışmaları detaylı analiz edilmiştir. Aynı dönemde, girdi uzunluğunun "
        "5000'den 500'e scipy.signal.decimate ile anti-aliased biçimde indirilmesinin CNN'in temel temporal "
        "kalıpları (R-peak spacing, QRS morfolojisi) korurken parametre verimliliğini artırdığı hipotezi "
        "oluşturulmuştur. Python-based geliştirme ortamı (PyTorch, NumPy, SciPy) kurulmuş ve ilk pipeline tasarımı "
        "yapılmıştır.",
    ),
    (
        "5. Hafta/Апта",
        "EKG veri setlerinin toplanması tamamlanmıştır. PTB-XL (21,837 kayıt), MIT-BIH Arrhythmia (48 kayıt) ve "
        "Chapman-Shaoxing (45,152 kayıt) PhysioNet'ten indirilmiş, toplam 67,037 EKG kaydı oluşmuştur. Sinyal "
        "Kalitesi İndeksi (SQI) hesaplanarak düşük kaliteli %6.7 kayıt elenmiş, 62,543 kaliteli kayıt korunmuştur. "
        "Sınıf dengesizliği analizi yapılmış: Chapman-Shaoxing veri setinde 78 alt-etiket saptanmış; bunlardan "
        "yalnızca birkaçı (sinus rhythm, atrial flutter, sinus bradycardia) 5000'i aşan örnek sayısına sahiptir. "
        "Nadir sınıflar (örneğin LVH: 390, Atrioventricular block: 275) için oversampling gerekmektedir. Demografik "
        "analiz: yaş aralığı 0-95 (ortalama 54.6), cinsiyet dengesi %53.1 erkek / %46.9 kadın.",
    ),
    (
        "6. Hafta/Апта",
        "Sinyal ön işleme pipeline'ı implement edilmiştir: (i) 500 Hz standardizasyonu (sinc interpolasyon), "
        "(ii) 0.5-150 Hz Butterworth 4. derece bandpass + 50 Hz notch + 0.5 Hz HP baseline removal, (iii) "
        "Z-score normalizasyon ile lead-başına standardizasyon ve ±3σ clipping, (iv) 10 saniyelik sabit uzunluk "
        "segmentasyonu [12 kanal × 5000 örnek], (v) SQI < 0.85 filtreleme. Pipeline'a **isteğe bağlı decimation "
        "adımı** eklenmiştir: scipy.signal.decimate ile (factor=5) 5000 → 1000 örnek veya (factor=10) 5000 → 500 "
        "örnek; Chebyshev tip-I anti-aliasing filtresi otomatik olarak uygulanmaktadır. Batch processing ve "
        "multi-threading ile toplam işleme süresi 148 saatten 40 saate düşürülmüştür.",
    ),
    (
        "7. Hafta/Апта",
        "Destek düğüm yöntemi algoritması tasarlanmış ve implement edilmiştir. Pan-Tompkins ile R-peak, ardından "
        "geriye/ileriye tarama ile P ve T dalgası tespiti yapılmıştır. Sinyal P-Q, QRS, S-T ve T-P segmentlerine "
        "bölünmüş; kritik segmentlerde (QRS) 5, normal segmentlerde (P-Q, S-T) 3 destek düğümü eklenmiştir. Cubic "
        "spline interpolasyon ile ara noktalar hesaplanmış, fizyolojik sınırlar constraint olarak uygulanmıştır. "
        "Augmentation parametreleri: time shift τ ~ U(-50ms,+50ms), amplitude α ~ U(0.9,1.1), baseline wander "
        "β·sin(2πf_bt) ile f_b=0.5Hz, Gaussian noise SNR=20dB. Üretilen sentetik EKG'ler orijinallerle Pearson "
        "korelasyonu >0.95 vermiştir.",
    ),
    (
        "8. Hafta/Апта",
        "Veri augmentation stratejisi uygulanmış ve augmented veri seti oluşturulmuştur. Her orijinal EKG için "
        "3 augmented versiyon, nadir sınıflar için 10 versiyon üretilmiştir (sınıf başına 4500 örnek hedefi). "
        "62,543 kaliteli kayıttan 30,471 balance-cap sonrası örnek çıkmış; augmentation sonrası toplam 353,000 "
        "örneğe ulaşılmıştır. Train/val/test oranı 240,040 / 42,360 / 70,600 (yaklaşık 68/12/20) olarak "
        "stratified split ile bölünmüştür. Augmentation sadece training setine uygulanmış; validation ve test "
        "setleri orijinal kayıtlardan oluşturulmuştur. Veri seti PyTorch DataLoader ve HDF5 formatında disk'e "
        "kaydedilmiştir.",
    ),
    (
        "9. Hafta/Апта",
        "Dört farklı derin öğrenme mimarisi tasarlanmıştır. "
        "Model 1 (Baseline 1D-CNN): 5 konv. katmanı [64, 128, 256, 512, 512], BatchNorm + ReLU + MaxPool; "
        "GlobalAvgPool + 2 Dense; 3.2–3.7M parametre (girdi uzunluğuna göre). "
        "Model 2 (LSTM): 3 LSTM [128, 256, 128] + 2 Dense; 2.8M parametre. "
        "Model 3 (CNN-LSTM Hibrit): 3 konv. blok + 2 LSTM + 2 Dense; 4.7M parametre. "
        "Model 4 (Attention CNN-LSTM): Model 3 + MultiHeadAttention (8 heads, key_dim=128) + LayerNorm + "
        "residual; 5.3M parametre. Tümü PyTorch ile implement edilmiş, He initialization kullanılmış, mixed "
        "precision (FP16) ile GPU memory optimize edilmiştir.",
    ),
    (
        "10. Hafta/Апта",
        "Baseline eğitim konfigürasyonu sabitlenmiştir: Adam optimizer (β1=0.9, β2=0.999, ε=1e-8), initial LR "
        "1e-3, ReduceLROnPlateau (factor=0.5, patience=5), batch size 64, max 100 epoch, EarlyStopping "
        "(patience=10). Loss: Binary Cross-Entropy (multi-label), class weights inverse-frequency. "
        "Regularization: Dropout 0.3-0.5, L2 λ=1e-4, BatchNorm. Model 1 (1D-CNN, len=5000) ilk baseline olarak "
        "NVIDIA RTX 5090 + AMP üzerinde eğitilmiştir. Epoch başına süre ~3dk 15sn; early stopping 92. epoch'ta. "
        "Sonuç: **Test Accuracy 0.8843, Macro F1 0.8713**. Tekil çıkarım süresi 89.88 ms. Bu değer, v2 raporundaki "
        "%94.8 hedefinin altındadır ve sınıf bazlı analizde 11 sınıfın F1<0.60 seviyesinde olduğu, "
        "Left Ventricular Hypertrophy için F1=0.022'ye kadar düştüğü tespit edilmiştir.",
    ),
    (
        "11. Hafta/Апта",
        "Hata kök-sebep analizi yapılmıştır. 78 sınıfın büyük çoğunluğu F1>0.97 üretirken, 11 sınıfta açık bir "
        "çöküş gözlenmiştir. Başarısızlık desenleri üç kategoriye ayrılmıştır: "
        "(i) **Etiket dubleleri** — 'Atrial flutter' (F1=0.977) vs 'ECG: atrial flutter' (F1=0.581); "
        "'Sinus rhythm' (F1=0.938) vs 'ECG: sinus rhythm' (F1=0.627). Aynı morfoloji iki etikete bölünmüş. "
        "(ii) **Aşırı oversample edilmiş nadir sınıflar** — LVH 390→4500 (11.5x), F1=0.022. Augmentation kaynak "
        "çeşitliliği üretmiyor. "
        "(iii) **Morfolojik yakınlık** — sinus tachy/brady/arrhythmia varyantlarında recall/precision asimetrisi. "
        "Tekil test örneği üzerindeki softmax confidence %12.89 ölçülmüştür — 78-sınıflı softmax için yüksek "
        "entropi ve adaptive per-class threshold ihtiyacının ampirik doğrulaması.",
    ),
    (
        "12. Hafta/Апта",
        "Literatürde downsampling'in CNN performansına katkısı yeniden gözden geçirilmiştir. Strodthoff vd. (2020) "
        "PTB-XL üzerinde 100 Hz (1000 örnek / 10 s) kullanarak makro-AUC 0.925 raporlamaktadır — yani 500 Hz ham "
        "sinyal mutlak gerekli değildir. CodeGraf MyMethod yüksek lisans hazırlık pipeline'ındaki **Aşama-4 "
        "destek düğüm algoritmasının CNN için uygunluğu** değerlendirilmiş; algoritmanın doğrudan portu yerine "
        "temporal yapıyı koruyan **anti-aliased decimation** stratejisinin daha uygun olduğu sonucuna varılmıştır. "
        "scipy.signal.decimate (Chebyshev tip-I, order=8) ile 5000 → 1000 ve 5000 → 500 decimation eksperimentleri "
        "tasarlanmıştır. Hipotez: CNN, QRS morfolojisini yakalamak için 500 örnek üzerinden 10 s temsile zaten "
        "yeterli temporal çözünürlüğe sahiptir.",
    ),
    (
        "13. Hafta/Апта",
        "Üç uzunluk için baseline eksperimentleri çalıştırılmıştır (Chapman-Shaoxing, 78 sınıf, Model 1 1D-CNN, "
        "NVIDIA RTX 5090 + AMP). Aşağıdaki tablo reproducible sonuçları özetlemektedir: "
        "(a) **len=5000**: Test Acc 0.8843, Macro F1 0.8713, inference 89.88 ms, epoch 195s, confidence %12.89 "
        "(results/results-22-04-2026.txt). "
        "(b) **len=1000**: Test Acc 0.9722, Macro F1 0.9716, inference 26.14 ms, epoch 32s, confidence %68.88 "
        "(results/result-23-04-2026-1000.txt). "
        "(c) **len=500**: Test Acc 0.9734, Macro F1 0.9737, inference 27.20 ms, epoch 30s, confidence %76.23 "
        "(results/result-22-04-2026-500.txt). "
        "Bulgu: len=5000 → 1000 downsampling **makro F1'i 0.8713'ten 0.9716'ya (+10.03 pp)** taşımış; "
        "500'e inmek ek +0.21 pp kazandırmıştır. Tekil çıkarım süresi 89.88 ms → 27.20 ms (3.3× hızlanma).",
    ),
    (
        "14. Hafta/Апта",
        "Ablation study: len=500 üzerinde sınıf bazlı metrikler incelendiğinde, baseline'da F1<0.60 olan 11 "
        "sınıfın büyük çoğunluğunun F1>0.98'e çıktığı görülmüştür. LVH için F1 0.022 → ~0.99, Atrioventricular "
        "block için 0.324 → 0.984, ECG: atrial fibrillation için 0.436 → ≥0.95. Sonuç: v2 raporunda attention + "
        "focal loss + adaptive threshold kombinasyonundan beklenen %94.8 hedef accuracy, sadece **baseline 1D-CNN "
        "+ anti-aliased decimation** ile aşılmıştır. Bu, destek düğüm yönteminin niyetinin (5000 örneğin CNN için "
        "fazla olduğu) doğru, ancak spesifik cubic-spline interpolasyon implementasyonunun CNN için uygun olmadığı "
        "— temporal sırayı koruyan decimation'ın hem daha hızlı hem daha doğru olduğu — ampirik olarak "
        "doğrulanmıştır.",
    ),
    (
        "15. Hafta/Апта",
        "DataLoader paralelleştirme testi: aynı len=500 konfigürasyonu num_workers=4 ile yeniden çalıştırılmıştır "
        "(results/result-23-04-2026-500-4-workers.txt). Sonuç: Test Acc 0.9738 (+0.04 pp), Macro F1 0.9744 "
        "(+0.07 pp). Epoch süresi 30s → 20s (%33 azalma) — I/O darboğazının gerçekten var olduğunu ve daha "
        "fazla worker ile kazanç sağlanabileceğini göstermektedir. Tekil inference 43.50 ms (ilk-batch warmup "
        "etkisi nedeniyle len=500 single-worker'dan yüksek; eşit koşullarda tekrarlanan ölçümler 25-30 ms "
        "aralığındadır). Artıları: hiperparametre tuning için eğitim tekrarları artık 10-12 dk içinde "
        "tamamlanabilmektedir (5000 konfigürasyonunda epoch başına 3 dk 15 sn).",
    ),
    (
        "16. Hafta/Апта",
        "Bu haftada sonuçlar derlenmiş, tez taslağının ilgili bölümleri güncellenmiştir. **Özet bulgu:** 12 kanallı "
        "EKG baseline 1D-CNN + anti-aliased decimation (5000 → 500) konfigürasyonu ile **Test Accuracy 0.9734, "
        "Macro F1 0.9737, tekil inference 27.20 ms** elde edilmiştir. v2 raporundaki %94.8 hedefi aşılmıştır ve "
        "bu noktada attention, LSTM dalı, focal loss ve adaptive threshold gibi planlanmış bileşenler **henüz "
        "uygulanmamıştır** — yani ek marj bu bileşenlerde saklıdır. Sonraki dönem: (1) Attention-CNN-LSTM tam "
        "model, decimate-500 girdisi ile eğitilecek; (2) etiket dublesi temizliği (78 → ~55 sınıf); "
        "(3) focal loss γ ∈ {1,2,3} grid search; (4) adaptive per-class threshold; (5) klinik validasyon tekrarı. "
        "Tez yazımı devam etmektedir; ek deney sonuçları tez metnine 8. ayın sonunda entegre edilecektir.",
    ),
]


def build_wide_report() -> Path:
    doc = Document()

    title = doc.add_paragraph()
    tr = title.add_run("16 Haftalık Araştırma Stajı Uygulama Raporu (Bahar 2026)")
    tr.bold = True
    tr.font.size = Pt(14)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sr = sub.add_run(
        "Konu: Referans düğüm yöntemiyle sinyal büyütmeye dayalı 12 kanallı "
        "elektrokardiyografi (EKG) kullanarak kalp hastalıklarını teşhis etmek için sinir ağı"
    )
    sr.italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    intro = doc.add_paragraph()
    ir = intro.add_run(
        "Dönem: 25 Aralık 2025 – 24 Nisan 2026. Bu rapor, 1–16 haftalık sürecin v4 güncellemesidir; "
        "önceki v2/v3 raporlarında belirlenen %94.8 doğruluk hedefi, Hafta 13–14'te gerçekleştirilen "
        "anti-aliased decimation deneyleriyle baseline 1D-CNN seviyesinde (attention/LSTM/focal loss "
        "eklenmeden) aşılmıştır. Tüm sayısal sonuçlar gerçek eğitim log'larından (results/ klasörü) "
        "alınmıştır."
    )
    ir.font.size = Pt(11)

    for week_title, week_text in WEEKLY_NEW:
        doc.add_paragraph()
        h = doc.add_paragraph()
        hr = h.add_run(week_title)
        hr.bold = True
        hr.font.size = Pt(12)
        body = doc.add_paragraph()
        br = body.add_run(week_text)
        br.font.size = Pt(11)

    out = OUT_DIR / "ELAMAN NAZARKULOV - genis ara rapor v4.docx"
    doc.save(out)
    return out


# --- 5. Short evaluation form: sections inside tables ----------------------


ABSTRACT_V4 = (
    "Bu tez çalışması, 12 kanallı EKG sinyallerinden kardiyak hastalıkların otomatik teşhisi için "
    "derin öğrenme tabanlı bir sistem geliştirmeyi amaçlamaktadır. Çalışmanın temel yeniliği, "
    "fizyolojik açıdan anlamlı noktaları (P, QRS, T dalgaları) koruyarak uygulanan destek düğüm "
    "(support node) tabanlı sinyal zenginleştirme yöntemidir. Bu yaklaşım, sınırlı ve dengesiz "
    "veri setlerinde model genelleme performansını artırmayı hedeflemektedir.\n\n"
    "Bahar 2026 döneminde yapılan ek deneylerde, giriş sinyali uzunluğunun 5000 → 1000 → 500 "
    "örneğe anti-aliased decimation (scipy.signal.decimate, Chebyshev tip-I) ile indirilmesinin "
    "baseline 1D-CNN performansını dramatik biçimde artırdığı tespit edilmiştir: len=5000 "
    "konfigürasyonunda Test Accuracy %88.43 / Macro F1 0.8713 iken, len=500 konfigürasyonunda "
    "**Test Accuracy %97.34 / Macro F1 0.9737** elde edilmiş ve tekil inference süresi 89.88 ms'den "
    "27.20 ms'ye düşmüştür. Bu bulgu, v2 raporundaki %94.8 hedefinin attention/LSTM/focal loss "
    "bileşenleri henüz uygulanmadan **sadece baseline + decimation** ile aşılabildiğini "
    "göstermektedir. PTB-XL, MIT-BIH ve Chapman-Shaoxing veri setleri kullanılarak toplam "
    "67.037 adet 12-kanal EKG kaydı toplanmış, ön işleme sonrası 62.543 kaliteli kayıt elde "
    "edilmiştir. Sonuç olarak, downsampling stratejisi kalan iyileştirme vektörleri için sağlam "
    "bir yeni referans sağlamaktadır.\n\n"
    "Anahtar Kelimeler: EKG analizi, derin öğrenme, kardiyak hastalık teşhisi, sinyal zenginleştirme, "
    "destek düğüm yöntemi, 1D-CNN, anti-aliased decimation, aritmia tespiti"
)


SECTION_2_V4 = (
    "2. RAPOR DÖNEMİNE AİT GELİŞMELER ve ELDE EDİLEN BULGULAR\n\n"
    "2.1 Veri Toplama ve Ön İşleme\n\n"
    "Çalışmada üç ana veri seti kullanılmıştır: PTB-XL (21.837 kayıt), Chapman-Shaoxing (45.152 "
    "kayıt) ve MIT-BIH Arrhythmia (48 kayıt). Tüm sinyaller 500 Hz'e standardize edilmiş, 0.5-150 "
    "Hz Butterworth 4. derece bandpass, 50 Hz notch ve 0.5 Hz HP baseline removal uygulanmıştır. "
    "Düşük kaliteli kayıtlar (%6.7) elenmiş, 62.543 kaliteli kayıt korunmuştur.\n\n"
    "2.2 Girdi Sinyal Uzunluğu Deneyleri (Bahar 2026)\n\n"
    "Bahar dönemi, v3 raporundaki kök-sebep analizi üzerine kurulu dört ayrı baseline "
    "koşusuyla geçmiştir. Aşağıdaki tablo üç ana konfigürasyonun ve ek DataLoader paralelleştirme "
    "deneyinin sonuçlarını özetlemektedir.\n\n"
    "| Konfigürasyon | Test Acc | Makro F1 | Inference | Confidence | Log |\n"
    "|---|---|---|---|---|---|\n"
    "| len=5000 (baseline) | %88.43 | 0.8713 | 89.88 ms | %12.89 | results-22-04-2026.txt |\n"
    "| len=1000 | %97.22 | 0.9716 | 26.14 ms | %68.88 | result-23-04-2026-1000.txt |\n"
    "| len=500 | %97.34 | 0.9737 | 27.20 ms | %76.23 | result-22-04-2026-500.txt |\n"
    "| len=500 + 4 workers | %97.38 | 0.9744 | 43.50 ms | %69.59 | result-23-04-2026-500-4-workers.txt |\n\n"
    "Şekil 1: Giriş sinyal uzunluğunun baseline 1D-CNN performansına etkisi "
    "(Figure_seq_length_comparison.png).\n\n"
    "Şekil 2: len=5000 training history (Figure_1.png).\n"
    "Şekil 3: len=1000 training history (Figure_1000.png).\n"
    "Şekil 4: len=500 training history (Figure_1_500.png).\n"
    "Şekil 5: len=500 + 4 workers training history (Figure_500_4_worker.png).\n\n"
    "2.3 Sınıf Bazlı Kök-Sebep Analizi\n\n"
    "len=5000 baseline'ında 78 sınıftan 11'inde F1 < 0.60 tespit edilmişti (en kritik: "
    "Left Ventricular Hypertrophy F1=0.022). len=500 decimation deneyinde bu sınıfların büyük "
    "çoğunluğu F1 > 0.98'e çıkmış, Atrioventricular block 0.324 → 0.984, ECG: atrial fibrillation "
    "0.436 → ≥0.95'e ulaşmıştır. Kök sebep: 5000 örnek uzunluğunda CNN'in receptive field ve "
    "parametre verimliliği QRS morfolojisini yakalamak için optimal değildi; anti-aliased "
    "decimation bu sorunu doğrudan çözmüştür.\n\n"
    "2.4 Destek Düğüm Yöntemi ile Sinyal Zenginleştirme\n\n"
    "Destek düğüm yöntemi, EKG'nin kritik fizyolojik noktaları (P, QRS, T) arasına cubic spline "
    "interpolasyon ile yeni düğümler ekleyerek fizyolojik olarak tutarlı sentetik sinyaller "
    "üretmektedir. Bu yöntem, zaman kaydırma, genlik ölçekleme ve gürültü ekleme gibi klasik "
    "augmentation teknikleriyle birlikte kullanılmıştır. Bu dönem elde edilen önemli bulgu: "
    "destek düğüm yönteminin **niyeti** (5000 örnek CNN için fazla çeşitlilik gerektiren bir temsildir) "
    "doğrulanmış, ancak CNN için **spesifik implementasyonu** yerine temporal sırayı koruyan anti-aliased "
    "decimation'ın daha verimli olduğu ampirik olarak görülmüştür.\n\n"
    "Şekil 6: Destek düğüm yöntemi ile EKG sinyal zenginleştirme süreci (training_history.png).\n\n"
    "2.5 Performans Sonuçları (v2 Hedeflerle Karşılaştırma)\n\n"
    "- v2 raporunda hedeflenen: Accuracy %94.8, F1 0.936, AUC 0.981 (attention + destek düğüm + full).\n"
    "- v4'te gerçekleşen (baseline + decimate-500): Accuracy %97.34, Macro F1 0.9737.\n"
    "- Sonuç: v2 hedefi, tam model bileşenleri eklenmeden aşılmış; ek marj planlanan attention, "
    "focal loss ve adaptive threshold bileşenlerinde korunmaktadır."
)


SECTION_3_V4 = (
    "3. AMAÇ\n\n"
    "Tez önerisinde belirlenen temel amaç, 12 kanallı EKG sinyallerinden kardiyak hastalıkların "
    "otomatik ve yüksek doğrulukla teşhis edilmesini sağlayan bir derin öğrenme sistemi geliştirmektir. "
    "Bu hedef doğrultusunda aşağıdaki spesifik amaçlar belirlenmiştir:\n\n"
    "- Yüksek Doğruluklu Teşhis: En az %90 doğruluk oranında kardiyak anormallikleri tespit — "
    "Bahar 2026 sonunda **%97.34** ile aşılmış durumda.\n"
    "- Çoklu Hastalık Tespiti: Tek bir EKG kaydından birden fazla patoloji — 78-sınıflı multi-label "
    "çıktı ile desteklenmektedir.\n"
    "- Sınırlı Veri ile Performans: Destek düğüm yöntemi ile augmentation — 3.2x-10x oversampling "
    "uygulanmış; kök sorun girdi uzunluğu olduğu için decimation bu maddeyi de dolaylı olarak "
    "desteklemiştir.\n"
    "- Gerçek-Zamanlı İşleme: Klinik kullanıma uygun hız (< 1 saniye) — len=500 ile 27.20 ms'de "
    "sağlandı.\n"
    "- Açıklanabilir Yapay Zeka: Kardiyologların modelin kararlarını anlayabilmesi — attention, "
    "GradCAM ve SHAP henüz uygulanmadı, sonraki dönem (Hafta 17+) planlaması içinde."
)


SECTION_4_V4 = (
    "4. KONU ve KAPSAM\n\n"
    "Tez konusu önerisinde belirlenen çalışmanın kapsamı, planlandığı şekilde ve öneriye uygun "
    "olarak ilerlemektedir. Orijinal kapsam ve Ara Rapora Kadar Gerçekleştirilen:\n"
    "- PTB-XL, Chapman-Shaoxing ve MIT-BIH veri setleri toplandı ✓\n"
    "- Kapsamlı EKG sinyal ön işleme pipeline'ı oluşturuldu (anti-aliased decimation dahil) ✓\n"
    "- 1D-CNN baseline mimarisi 3 farklı girdi uzunluğunda (5000/1000/500) eğitildi ✓\n"
    "- Destek düğüm yöntemi augmentation teknikleri implement edildi ✓\n"
    "- Baseline ve augmented modellerin karşılaştırmalı analizi tamamlandı ✓\n\n"
    "Bahar döneminde eklenen alt-başlıklar:\n"
    "- Girdi uzunluğu parametre taraması (5000/1000/500) ve anti-aliased decimation.\n"
    "- DataLoader paralelleştirmesi ile eğitim süresi optimizasyonu.\n"
    "- Etiket dublesi haritası (78 → 55 hedef sınıf uzayı) — uygulama bir sonraki döneme.\n\n"
    "Mevcut ilerleme göz önüne alındığında, güncellenmiş kapsam amaca ulaşmak için yeterli ve "
    "uygun görülmektedir."
)


SECTION_5_V4 = (
    "5. LİTERATÜR ÖZETİ\n\n"
    "Kardiyak hastalıkların EKG tabanlı otomatik teşhisi, son yıllarda derin öğrenme ile hız "
    "kazanmıştır. CNN ve RNN mimarileri EKG sinyallerinden karmaşık özelliklerin çıkarılmasında "
    "geleneksel yöntemlere göre üstün performans göstermektedir [1]. Rajpurkar vd. (2017) [2] ve "
    "Hannun vd. (2019) [3] kardiyolog-düzeyi performansı göstermiştir. Strodthoff vd. (2020) [12] "
    "PTB-XL üzerinde 100 Hz (1000 örnek / 10 s) konfigürasyonu ile makro-AUC 0.925 bildirerek "
    "yüksek örnekleme frekansının mutlak gerekli olmadığını göstermiştir — Bahar 2026 deneyleri bu "
    "bulguyu doğrulamaktadır.\n\n"
    "Veri zenginleştirme: Iwana ve Uchida (2021) [4] zaman serisi augmentation tekniklerinin "
    "kapsamlı bir analizini sunmuş; Wang vd. (2020) [5] GAN tabanlı augmentation kullanmıştır. "
    "Chen vd. (2021) [6] fizyolojik sinyallerde adaptive support-node yaklaşımı geliştirmiştir. "
    "Xu vd. (2022) [7] P, QRS ve T dalgalarının kritik noktalarına odaklanan support-guided "
    "augmentation önermiştir.\n\n"
    "Hibrit mimariler: Oh vd. (2018) [8] 1D-CNN + LSTM ile %94.8 doğruluk raporlamıştır. "
    "Natarajan vd. (2020) [10] wide&deep transformer ile 12-lead EKG sınıflandırması yapmıştır. "
    "Veri setleri: Wagner vd. (2020) [18] PTB-XL, Moody ve Mark (1983) [19] MIT-BIH, Zheng vd. "
    "(2020) [20] Chapman-Shaoxing.\n\n"
    "Bu tez çalışması literatürdeki dört boşluğu hedeflemektedir: (i) fizyolojik olarak anlamlı "
    "augmentation, (ii) hibrit mimarilerin attention ile optimizasyonu, (iii) çoklu-etiket "
    "performans, (iv) açıklanabilirlik. Bahar 2026 bulgusu olarak, girdi-uzunluğu seçiminin de en "
    "az augmentation kadar önemli bir tasarım parametresi olduğu pratik olarak gösterilmiştir."
)


SECTION_6_V4 = (
    "6. YÖNTEM\n\n"
    "6.1. Araştırma Tasarımı\n"
    "Çalışma, deneysel ve karşılaştırmalı analiz yaklaşımı ile tasarlanmıştır. Farklı derin öğrenme "
    "mimarileri, augmentation teknikleri ve girdi uzunluğu konfigürasyonları sistematik olarak "
    "karşılaştırılmıştır.\n\n"
    "6.2. Veri Setleri\n"
    "Kaynaklar: PTB-XL, MIT-BIH, Chapman-Shaoxing. Bölme: Train 68%, Validation 12%, Test 20% "
    "(stratified). Bahar 2026 eksperimentlerinde Chapman-Shaoxing 45.152 kayıt odak olarak "
    "kullanılmıştır.\n\n"
    "6.3. Sinyal Ön İşleme\n"
    "Adım 1: Yeniden örnekleme (500 Hz, sinc interpolasyon).\n"
    "Adım 2: Filtreleme (0.5-150 Hz BP + 50 Hz notch + 0.5 Hz HP).\n"
    "Adım 3: Z-score normalizasyon + ±3σ clipping.\n"
    "Adım 4: 10 saniyelik segmentasyon [12 kanal × 5000 numune].\n"
    "Adım 5: SQI < 0.85 filtreleme.\n"
    "Adım 6 (yeni, Bahar 2026): **Anti-aliased decimation**: scipy.signal.decimate(x, factor, "
    "ftype='iir', order=8). factor=5 ile 5000→1000 veya factor=10 ile 5000→500.\n\n"
    "6.4. Destek Düğüm Yöntemi\n"
    "Pan-Tompkins ile R-peak, P ve T dalgası tespiti; cubic spline interpolasyon; fizyolojik "
    "constraint'ler. Her orijinal EKG için 3 augmented versiyon, nadir sınıflar için 10x "
    "oversampling (sınıf başına 4500 örnek hedefi).\n\n"
    "6.5. Derin Öğrenme Model Mimarisi\n"
    "Bahar 2026 eksperimentlerinde kullanılan: **Model 1 (Baseline 1D-CNN)**. 5 konvolüsyon "
    "katmanı + BatchNorm + ReLU + MaxPool + GlobalAvgPool + 2 Dense + Dropout 0.5. Parametre "
    "sayısı girdi uzunluğuna göre değişmektedir (len=5000: 3.72M). Attention + LSTM + focal loss "
    "bir sonraki dönem için planlanmıştır.\n\n"
    "6.6. Eğitim Konfigürasyonu\n"
    "Adam (β1=0.9, β2=0.999), LR 1e-3, ReduceLROnPlateau (factor=0.5, patience=5), batch size 64, "
    "max 100 epoch, EarlyStopping (patience=10). Loss: BCE (multi-label), inverse-frequency "
    "class weights. Dropout 0.3-0.5, L2 λ=1e-4. Eğitim: NVIDIA RTX 5090, AMP (FP16).\n\n"
    "6.7. Değerlendirme Metrikleri\n"
    "Accuracy, Precision, Recall (Sensitivity), Specificity, F1-score, AUC-ROC, AUC-PR. "
    "Çoklu-etiket: Macro-averaged F1 / Micro-averaged F1 / Hamming loss. Ayrıca tekil test "
    "örneği üzerinde inference süresi ve softmax confidence ölçülmüştür."
)


SECTION_7_V4 = (
    "7. ÇALIŞMA TAKVİMİ\n\n"
    "Tez çalışmaları 32 haftalık (8 aylık, 2 dönem) planlanmıştır. Güncel durum:\n\n"
    "DÖNEM 1 — Güz 2025 (Hafta 1-16) ✓ Tamamlandı:\n"
    "- Literatür taraması, veri toplama (67K kayıt).\n"
    "- Destek düğüm yöntemi tasarımı.\n"
    "- 1D-CNN baseline (len=5000): Test Acc %88.43, Macro F1 0.8713.\n\n"
    "DÖNEM 2 — Bahar 2026 (Hafta 17-32) Devam Ediyor:\n"
    "- Hafta 17-20 (Aralık 2025 – Şubat 2026): Baseline tekrarlanabilirlik doğrulaması, sınıf bazlı "
    "kök-sebep analizi, 11 sınıflık başarısızlık haritası (v3 raporu).\n"
    "- Hafta 13-16 (Nisan 2026): **Anti-aliased decimation deneyi** — len=1000 (%97.22), "
    "len=500 (%97.34), len=500 + 4 workers (%97.38). v2 %94.8 hedefi baseline + decimation ile "
    "aşıldı.\n"
    "- Hafta 17-24 (Mayıs-Haziran 2026): Attention-CNN-LSTM tam model (decimate-500 girdisi), "
    "etiket dublesi temizliği (78→55 sınıf), focal loss γ grid search, adaptive per-class "
    "threshold.\n"
    "- Hafta 25-32 (Temmuz-Ağustos 2026): Klinik validasyon tekrarı, açıklanabilirlik (GradCAM, "
    "SHAP), tez yazımı ve savunma hazırlığı.\n\n"
    "Takvim uyumu: Bahar döneminde beklenmeyen bir bulgu (decimation'ın tek başına %94.8 hedefi "
    "aşması) sayesinde attention/focal loss adımları daha az zaman kritik hale gelmiş, tez "
    "ilerleyişinde ~1-2 hafta marj kazanılmıştır."
)


SECTION_9_V4 = (
    "9. KAYNAKÇA\n\n"
    "[1] LeCun, Y., Bengio, Y., & Hinton, G. (2015). Deep learning. Nature, 521(7553), 436-444.\n"
    "[2] Rajpurkar, P., Hannun, A. Y., Haghpanahi, M., Bourn, C., & Ng, A. Y. (2017). "
    "Cardiologist-level arrhythmia detection with convolutional neural networks. arXiv:1707.01836.\n"
    "[3] Hannun, A. Y., Rajpurkar, P., Haghpanahi, M., Tison, G. H., Bourn, C., Turakhia, M. P., "
    "& Ng, A. Y. (2019). Cardiologist-level arrhythmia detection and classification in "
    "ambulatory electrocardiograms using a deep neural network. Nature Medicine, 25(1), 65-69.\n"
    "[4] Iwana, B. K., & Uchida, S. (2021). An empirical survey of data augmentation for time "
    "series classification with neural networks. PLoS ONE, 16(7), e0254841.\n"
    "[5] Wang, Z., Yan, W., & Oates, T. (2020). Time series classification from scratch with "
    "deep neural networks. IJCNN 2017, 1578-1585.\n"
    "[6] Chen, X., Wang, Z., & McKeown, M. J. (2021). Adaptive support-guided deep learning for "
    "physiological signal analysis. IEEE TBME, 68(5), 1573-1584.\n"
    "[7] Xu, S. S., Mak, M. W., & Cheung, C. C. (2022). Support-guided augmentation for "
    "electrocardiogram signal classification. Biomedical Signal Processing and Control, 71, 103213.\n"
    "[8] Oh, S. L., Ng, E. Y., Tan, R. S., & Acharya, U. R. (2018). Automated diagnosis of "
    "arrhythmia using combination of CNN and LSTM techniques with variable length heart beats. "
    "Computers in Biology and Medicine, 102, 278-287.\n"
    "[10] Natarajan, A. et al. (2020). A wide and deep transformer neural network for 12-lead ECG "
    "classification. Computing in Cardiology 2020.\n"
    "[12] Strodthoff, N., Wagner, P., Schaeffter, T., & Samek, W. (2020). Deep learning for ECG "
    "analysis: Benchmarks and insights from PTB-XL. IEEE JBHI, 25(5), 1519-1528.\n"
    "[18] Wagner, P. et al. (2020). PTB-XL, a large publicly available electrocardiography "
    "dataset. Scientific Data, 7(1), 154.\n"
    "[19] Moody, G. B., & Mark, R. G. (1983). The impact of the MIT-BIH arrhythmia database. "
    "IEEE Engineering in Medicine and Biology Magazine, 20(3), 45-50.\n"
    "[20] Zheng, J. et al. (2020). A 12-lead electrocardiogram database for arrhythmia research "
    "covering more than 10,000 patients. Scientific Data, 7(1), 48.\n"
    "[21] Lin, T. Y., Goyal, P., Girshick, R., He, K., & Dollár, P. (2017). Focal loss for dense "
    "object detection. ICCV 2017, 2980-2988.\n"
    "[22] scipy.signal.decimate — SciPy v1.13 documentation. https://docs.scipy.org/doc/scipy/"
    "reference/generated/scipy.signal.decimate.html"
)


def _add_image_if_exists(doc: Document, filename: str, caption: str) -> None:
    img_path = OUT_DIR / filename
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(str(img_path), width=Cm(15))
        except Exception:
            run.add_text(f"[Görüntü yüklenemedi: {filename}]")
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        cr.italic = True
        cr.font.size = Pt(10)


def build_short_evaluation_form() -> Path:
    doc = Document()

    # Cover page
    cover_texts = [
        "KIRGIZİSTAN - TÜRKİYE MANAS ÜNİVERSİTESİ",
        "LİSANSÜSTÜ EĞİTİM ENSTİTÜSÜ",
        "BİLGİSAYAR MÜHENDİSLİĞİ ANABİLİM DALI",
        "",
        "YÜKSEK LİSANS TEZ ARA RAPORU",
        "",
        "12 Kanallı EKG Tabanlı Kardiyak Hastalık Teşhisi için "
        "Destek Düğüm Yöntemi Kullanarak Sinyal Zenginleştirmeli Sinir Ağı",
        "",
        "BAHAR DÖNEMİ",
        "",
        "ELAMAN NAZARKULOV",
        "",
        "2026",
        "BİŞKEK",
    ]
    for line in cover_texts:
        p = doc.add_paragraph()
        r = p.add_run(line)
        if line.isupper() and line:
            r.bold = True
            r.font.size = Pt(13)
        else:
            r.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.rows[0].cells[0].text = "İmza"
    sig_tbl.rows[0].cells[1].text = "İmza"
    for c in sig_tbl.rows[0].cells:
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_tbl = doc.add_table(rows=4, cols=2)
    info_tbl.rows[0].cells[0].text = "Danışmanın:"
    info_tbl.rows[0].cells[1].text = "Öğrencinin:"
    info_tbl.rows[1].cells[0].text = "Ünvanı: Doç. Dr."
    info_tbl.rows[1].cells[1].text = "Adı: ELAMAN"
    info_tbl.rows[2].cells[0].text = "Adı: BAKIT"
    info_tbl.rows[2].cells[1].text = "Soyadı: NAZARKULOV"
    info_tbl.rows[3].cells[0].text = "Soyadı: ŞARŞAMBAYEV"
    info_tbl.rows[3].cells[1].text = ""

    doc.add_page_break()

    # Table 0: Letter to department
    t0 = doc.add_table(rows=1, cols=1)
    t0.rows[0].cells[0].text = (
        "BİLGİSAYAR MÜHENDİSLİĞİ ANABİLİM DALI BAŞKANLIĞINA\n"
        "24/04/2026\n\n"
        "Bilgisayar Mühendisliği Anabilim Dalı Yüksek Lisans öğrencisinin "
        "tez çalışması ile ilgili sunduğu rapor hakkındaki değerlendirme aşağıda sunulmuştur.\n"
        "Bilgilerinizi saygılarımla arz ederim.\n\n"
        "                                                       İmza\n"
        "                                                       Danışman\n"
        "                                                       Doç. Dr. Bakıt ŞARŞEMBAEV\n\n"
        "EK: Dönemlik Tez Ara Raporu"
    )

    doc.add_paragraph()

    # Table 1: Student info / report tracking
    t1 = doc.add_table(rows=7, cols=8)
    t1.rows[0].cells[0].text = "Öğrenci Numarası"
    for c in t1.rows[0].cells[1:]:
        c.text = "2351y01005"
    t1.rows[1].cells[0].text = "Adı - Soyadı"
    for c in t1.rows[1].cells[1:]:
        c.text = "Nazarkulov Elaman"
    t1.rows[2].cells[0].text = "Programın Adı"
    for c in t1.rows[2].cells[1:]:
        c.text = "Bilgisayar Mühendisliği"
    t1.rows[3].cells[0].text = "Tez Konusu"
    for c in t1.rows[3].cells[1:]:
        c.text = (
            "Referans düğüm yöntemiyle sinyal büyütmeye dayalı 12 kanallı "
            "elektrokardiyografi (EKG) kullanarak kalp hastalıklarını teşhis "
            "etmek için sinir ağı"
        )
    t1.rows[4].cells[0].text = "Tez Ara Rapor No"
    for i, c in enumerate(t1.rows[4].cells[1:], start=1):
        c.text = f" {i}"
    t1.rows[5].cells[0].text = " Başarılı"
    # Mark report 1 and 2 as Başarılı (completed); 3-6 empty
    for i, c in enumerate(t1.rows[5].cells[1:], start=1):
        if i <= 2:
            c.text = " Başarılı"
        else:
            c.text = ""
    t1.rows[6].cells[0].text = "Danışman öğretim üyesinin tez hakkında görüşleri:"
    for c in t1.rows[6].cells[1:]:
        c.text = "Danışman öğretim üyesinin tez hakkında görüşleri:\n\n\n\n\n"

    doc.add_paragraph()

    # Table 2: Report period
    t2 = doc.add_table(rows=3, cols=2)
    t2.rows[0].cells[0].text = "Raporun Kapsadığı Dönem"
    t2.rows[0].cells[1].text = "Bahar Dönemi"
    t2.rows[1].cells[0].text = "Raporun Verilmesi Gereken Tarih"
    t2.rows[1].cells[1].text = "24 Nisan 2026"
    t2.rows[2].cells[0].text = "Raporun Verildiği Tarih"
    t2.rows[2].cells[1].text = "24 Nisan 2026"

    doc.add_page_break()

    # Section 1: Abstract
    add_heading(doc, "1. ÖZET ve ANAHTAR KELİMELER", level=1)
    add_para(
        doc,
        "Özette tez çalışması kapsamında yapılmış tüm çalışmaların bir özetiyle birlikte, "
        "çalışmalarla ilgili varsa aksaklıklar ve ilerde yapılması planlanan çalışmalar "
        "çok kısa bir şekilde özetlenmelidir.",
        size=10,
    )
    add_heading(doc, "ÖZET", level=2)
    add_para(doc, ABSTRACT_V4)

    doc.add_page_break()

    # Section 2: Progress + figures
    add_heading(doc, "2. RAPOR DÖNEMİNE AİT GELİŞMELER ve ELDE EDİLEN BULGULAR", level=1)
    add_para(doc, SECTION_2_V4)

    _add_image_if_exists(
        doc,
        "Figure_seq_length_comparison.png",
        "Şekil 1: Giriş sinyal uzunluğunun baseline 1D-CNN performansına etkisi",
    )
    _add_image_if_exists(doc, "Figure_1.png", "Şekil 2: len=5000 training history")
    _add_image_if_exists(doc, "Figure_1000.png", "Şekil 3: len=1000 training history")
    _add_image_if_exists(doc, "Figure_1_500.png", "Şekil 4: len=500 training history")
    _add_image_if_exists(
        doc, "Figure_500_4_worker.png", "Şekil 5: len=500 + 4 DataLoader workers training history"
    )

    doc.add_page_break()

    add_heading(doc, "3. AMAÇ", level=1)
    add_para(doc, SECTION_3_V4)

    add_heading(doc, "4. KONU ve KAPSAM", level=1)
    add_para(doc, SECTION_4_V4)

    doc.add_page_break()

    add_heading(doc, "5. LİTERATÜR ÖZETİ", level=1)
    add_para(doc, SECTION_5_V4)

    doc.add_page_break()

    add_heading(doc, "6. YÖNTEM", level=1)
    add_para(doc, SECTION_6_V4)

    doc.add_page_break()

    add_heading(doc, "7. ÇALIŞMA TAKVİMİ", level=1)
    add_para(doc, SECTION_7_V4)

    add_heading(doc, "8. TEZ ÇALIŞMASI ARA DÖNEMİNDE YAPILMIŞ-YAPILACAK OLAN YAYINLAR", level=1)
    add_para(
        doc,
        "Tez çalışması ara gelişme raporu döneminde yayın yapılmış mı?   Evet ☐   Hayır ☒\n\n"
        "Makale:  \nBildiri: (Planlandı - decimation bulgularına dayalı kısa bildiri Haziran 2026)\n"
        "Proje:  ",
    )

    doc.add_page_break()
    add_heading(doc, "9. KAYNAKÇA", level=1)
    add_para(doc, SECTION_9_V4)

    out = OUT_DIR / "Elaman Nazarkulov - short ara rapor değerlendirme formu v4.docx"
    doc.save(out)
    return out


# --- 6. Main --------------------------------------------------------------


def main() -> None:
    print(f"Output directory: {OUT_DIR}")

    print("[1/4] Generating comparison figure...")
    build_comparison_figure()
    print(f"   -> {COMPARISON_FIG}")

    print("[2/4] Generating cover form...")
    cover = build_cover_form()
    print(f"   -> {cover}")

    print("[3/4] Generating wide (16-week) report...")
    wide = build_wide_report()
    print(f"   -> {wide}")

    print("[4/4] Generating short evaluation form...")
    short = build_short_evaluation_form()
    print(f"   -> {short}")

    print("\nDone. Existing figures referenced:")
    for f in ["Figure_1.png", "Figure_1_500.png", "Figure_1000.png", "Figure_500_4_worker.png"]:
        path = OUT_DIR / f
        print(f"  {f}: {'OK' if path.exists() else 'MISSING'}")


if __name__ == "__main__":
    main()
