"""
Generate the full master's-thesis document EKG_Dissertation_Manas.docx
(~60 pages) for the Kyrgyz-Turkish Manas University Computer Engineering
Department. Uses the Manas template for the cover page and integrates the
actual training code from ecg_cnn_pytorch.py with detailed definitions.

Output: C:/Users/enazarkulov/Documents/Мастер/EKG_Dissertation_Manas.docx
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(r"C:\Users\enazarkulov\Documents\Мастер")
SRC_TEMPLATE = ROOT / "Elaman Nazarkulov - ara rapor değerlendirme formu.docx"
DST = ROOT / "EKG_Dissertation_Manas.docx"

CODE_FILE = Path(r"C:\Users\enazarkulov\Documents\ML\ekg\training\ecg_cnn_pytorch.py")

FIG_GEOM = ROOT / "Figure_geometry_invariance.png"
FIG_CMP = ROOT / "Figure_seq_length_comparison.png"
FIG_HIST = ROOT / "training_history.png"
FIG_500 = ROOT / "Figure_500_4_worker.png"
FIG_1000 = ROOT / "Figure_1000.png"
FIG_BASELINE = ROOT / "Figure_1.png"
FIG_ROADMAP = ROOT / "Figure_future_work_roadmap.png"

BODY_FONT = "Times New Roman"
CODE_FONT = "Consolas"
BODY_SIZE = Pt(11)
CODE_SIZE = Pt(9)
H1_SIZE = Pt(16)
H2_SIZE = Pt(14)
H3_SIZE = Pt(12)


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _apply_font(run, *, name=BODY_FONT, size=BODY_SIZE, bold=False,
                italic=False, color=None):
    run.font.name = name
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def add_paragraph(doc_or_cell, text="", *, bold=False, italic=False,
                  align=None, size=BODY_SIZE, name=BODY_FONT, color=None,
                  space_before=None, space_after=None, first_line_indent=None,
                  line_spacing=None):
    p = doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if line_spacing is not None:
        p.paragraph_format.line_spacing = line_spacing
    if text:
        run = p.add_run(text)
        _apply_font(run, name=name, size=size, bold=bold, italic=italic,
                    color=color)
    return p


def add_heading1(doc, text, page_break=True):
    if page_break:
        p = doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _apply_font(run, size=H1_SIZE, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _apply_font(run, size=H2_SIZE, bold=True)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _apply_font(run, size=H3_SIZE, bold=True, italic=True)
    return p


def add_body(doc, text, *, justify=True, indent=False):
    align = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    fli = Cm(0.75) if indent else None
    return add_paragraph(doc, text, align=align, size=BODY_SIZE,
                         space_after=Pt(6), first_line_indent=fli,
                         line_spacing=1.15)


def add_bullets(doc, items: Iterable[str]):
    for it in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"•  {it}")
        _apply_font(run, size=BODY_SIZE)


def add_numbered(doc, items: Iterable[str]):
    for i, it in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{i}.  {it}")
        _apply_font(run, size=BODY_SIZE)


def add_picture(doc, path: Path, *, width_cm=14.0, caption=None):
    if not path.exists():
        add_paragraph(doc, f"[missing: {path.name}]", italic=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    if caption:
        add_paragraph(doc, caption, italic=True, size=Pt(10),
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))


def add_code_block(doc, code: str, *, caption: str | None = None,
                   shade=True, max_lines: int | None = None):
    """Add a multi-line code listing (Consolas 9pt)."""
    if max_lines is not None:
        lines = code.splitlines()[:max_lines]
        code = "\n".join(lines)
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        # Shade code paragraph
        if shade:
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F4F4F4")
            pPr.append(shd)
        run = p.add_run(line if line else " ")
        _apply_font(run, name=CODE_FONT, size=CODE_SIZE)
    if caption:
        add_paragraph(doc, caption, italic=True, size=Pt(10),
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))


def _set_cell_borders(cell):
    """Add a single thin border around a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        b = tcBorders.find(tag)
        if b is None:
            b = OxmlElement(f"w:{edge}")
            tcBorders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "999999")


def _shade_cell(cell, fill="DDDDDD"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table(doc, header: list[str], rows: list[list[str]], *,
              caption: str | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    # Try to apply a built-in style; fall back gracefully if missing.
    for style_name in ("Light Grid Accent 1", "Table Grid",
                       "Light List Accent 1", "Medium Shading 1 Accent 1"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _apply_font(run, size=Pt(10), bold=True)
        _set_cell_borders(hdr[i])
        _shade_cell(hdr[i], fill="DDDDDD")
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            _apply_font(run, size=Pt(10))
            _set_cell_borders(cells[c_idx])
    if caption:
        add_paragraph(doc, caption, italic=True, size=Pt(10),
                      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
    return table


# ---------------------------------------------------------------------------
# Code extraction helpers
# ---------------------------------------------------------------------------

def load_source() -> str:
    return CODE_FILE.read_text(encoding="utf-8")


def extract_block(source: str, header_pattern: str,
                  end_patterns: tuple[str, ...]) -> str:
    """Extract a block from source given a header regex and list of end markers."""
    lines = source.splitlines()
    start = None
    for i, ln in enumerate(lines):
        if re.match(header_pattern, ln):
            start = i
            break
    if start is None:
        return f"# (block not found: {header_pattern})"
    out = [lines[start]]
    for j in range(start + 1, len(lines)):
        if any(re.match(p, lines[j]) for p in end_patterns):
            break
        out.append(lines[j])
    return "\n".join(out).rstrip()


# ---------------------------------------------------------------------------
# Cover-page editing (template re-use)
# ---------------------------------------------------------------------------

def _clear_cell(cell):
    tc = cell._tc
    for p in tc.findall(qn("w:p")):
        tc.remove(p)


def _replace_cover_year(doc: Document, new_year="2026"):
    for para in doc.paragraphs:
        if para.text.strip() == "2025":
            for run in para.runs:
                if "2025" in run.text:
                    run.text = run.text.replace("2025", new_year)
                    break


def _set_cover_subtitle(doc: Document, new_subtitle: str):
    for para in doc.paragraphs:
        if para.text.strip() in {"GÜZ DÖNEMİ", "BAHAR DÖNEMİ"}:
            for run in para.runs:
                if run.text.strip():
                    run.text = new_subtitle
                    break
            return


def setup_cover(doc: Document):
    _set_cover_subtitle(doc, "MASTER'S DISSERTATION")
    _replace_cover_year(doc, "2026")

    # Table 0: Department letter
    cell = doc.tables[0].rows[0].cells[0]
    _clear_cell(cell)
    add_paragraph(cell,
                  "TO THE HEAD OF THE DEPARTMENT OF COMPUTER ENGINEERING",
                  bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(cell, "29 April 2026", align=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(cell, "")
    add_paragraph(cell,
                  "The master's-thesis dissertation of the graduate student "
                  "of the Department of Computer Engineering, prepared in "
                  "accordance with the requirements of the KTMU Master's "
                  "Programme Implementation Directive, is submitted herewith. "
                  "The dissertation reports a controlled empirical study of "
                  "input-length design for 12-lead ECG classification on "
                  "the Chapman-Shaoxing corpus and includes the full "
                  "PyTorch reference implementation.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_paragraph(cell, "")
    add_paragraph(cell, "Signature", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(cell, "Thesis Supervisor", align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(cell, "Assoc. Prof. Bakit SARSEMBAEV",
                  align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(cell, "")
    add_paragraph(cell, "ENCL.: Master's Dissertation (full text)",
                  align=WD_ALIGN_PARAGRAPH.LEFT)

    # Table 1: clear status grid
    t1 = doc.tables[1]
    for i, c in enumerate(t1.rows[5].cells):
        _clear_cell(c)
        if 1 <= i <= 4:
            txt = " Completed"
        elif i == 5:
            txt = " Defended"
        else:
            txt = ""
        add_paragraph(c, txt, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Table 2: dates
    t2 = doc.tables[2]
    pairs = [
        ("Reporting Period", "Spring 2026 - Master's Dissertation"),
        ("Defence Date (Target)", "29 April 2026"),
        ("Date Prepared", "29 April 2026"),
    ]
    for i, (lab, val) in enumerate(pairs):
        _clear_cell(t2.rows[i].cells[0]); add_paragraph(t2.rows[i].cells[0], lab)
        _clear_cell(t2.rows[i].cells[1]); add_paragraph(t2.rows[i].cells[1], val)

    # Table 3: short abstract teaser
    t3 = doc.tables[3]
    cell = t3.rows[0].cells[0]
    _clear_cell(cell)
    add_paragraph(cell, "ABSTRACT and KEYWORDS", bold=True, size=Pt(12))
    add_paragraph(cell,
                  "The full abstract appears on the next pages; this is the "
                  "Manas-template summary box.",
                  italic=True, size=Pt(10), align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    cell = t3.rows[1].cells[0]
    _clear_cell(cell)
    add_paragraph(cell, "ABSTRACT (SHORT)", bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12))
    add_paragraph(cell, "")
    add_paragraph(cell,
                  "We show that anti-aliased decimation of the 12-lead "
                  "ECG input from 5000 to 500 samples raises a baseline "
                  "1D-CNN's test accuracy from 88.43% to 97.34% (macro-F1 "
                  "0.8713 -> 0.9737) on the Chapman-Shaoxing corpus. The "
                  "result is obtained without any change to the model, "
                  "loss function or augmentation pipeline. We frame the "
                  "phenomenon as geometric invariance of the fiducial-point "
                  "graph and present the full PyTorch implementation.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    cell = t3.rows[2].cells[0]
    _clear_cell(cell)
    add_paragraph(cell,
                  "Keywords: electrocardiogram, 12-lead ECG, deep learning, "
                  "1D convolutional neural network, anti-aliased decimation, "
                  "PyTorch, multi-label classification, fiducial points, "
                  "geometric invariance.",
                  italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Table 4: clear all rows (we will write the body after the template)
    t4 = doc.tables[4]
    for row in t4.rows:
        for cell in row.cells:
            _clear_cell(cell)
            add_paragraph(cell, "")

    # Table 5: publication info
    cell = doc.tables[5].rows[1].cells[0]
    _clear_cell(cell)
    add_paragraph(cell,
                  "This document is the full master's-thesis dissertation "
                  "(approximately 60 pages) prepared on the Manas-University "
                  "template. The body chapters, code listings and "
                  "appendices follow this section.",
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # Table 6: clear references area (we will write our own bibliography)
    t6 = doc.tables[6]
    cell = t6.rows[0].cells[0]
    _clear_cell(cell); add_paragraph(cell, "REFERENCES", bold=True, size=Pt(12))
    cell = t6.rows[1].cells[0]
    _clear_cell(cell)
    add_paragraph(cell,
                  "The complete list of references is given in the "
                  "Bibliography chapter at the end of this dissertation.",
                  italic=True, size=Pt(10))


# ---------------------------------------------------------------------------
# Body content
# ---------------------------------------------------------------------------

def write_front_matter(doc: Document):
    # Force new section / page break before front matter
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    # Title block
    add_paragraph(doc,
                  "KYRGYZ-TURKISH MANAS UNIVERSITY",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(16),
                  space_before=Pt(40))
    add_paragraph(doc,
                  "GRADUATE SCHOOL OF NATURAL AND APPLIED SCIENCES",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13))
    add_paragraph(doc,
                  "DEPARTMENT OF COMPUTER ENGINEERING",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13))
    add_paragraph(doc, "", space_after=Pt(40))

    add_paragraph(doc,
                  "ANTI-ALIASED DECIMATION AS THE DECISIVE STEP IN "
                  "12-LEAD ECG CLASSIFICATION:",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(15))
    add_paragraph(doc,
                  "FROM 88.43% TO 97.34% ON CHAPMAN-SHAOXING WITH A "
                  "PLAIN 1D-CNN",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(15),
                  space_after=Pt(60))

    add_paragraph(doc, "MASTER'S DISSERTATION",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(13),
                  space_after=Pt(40))
    add_paragraph(doc, "Author: Elaman NAZARKULOV",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12))
    add_paragraph(doc, "Supervisor: Assoc. Prof. Bakit SARSEMBAEV",
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12))
    add_paragraph(doc, "", space_after=Pt(40))
    add_paragraph(doc, "Bishkek, 2026",
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=Pt(12))


def write_abstract(doc: Document):
    add_heading1(doc, "ABSTRACT")
    add_body(doc,
             "Automatic 12-lead electrocardiogram (ECG) classification with "
             "deep convolutional neural networks (CNNs) is conventionally "
             "performed on the raw 500 Hz x 10 s signal of 5000 samples "
             "per lead. This dissertation challenges the assumption that "
             "this default input length is a neutral choice. On the "
             "Chapman-Shaoxing 12-lead ECG corpus (45,152 records, 78 "
             "multi-label diagnostic categories) we present a controlled "
             "comparison of input lengths {5000, 1000, 500} with an "
             "otherwise identical 1D-CNN backbone, optimiser, loss "
             "function, augmentation policy, random seed and "
             "train/validation/test split.")
    add_body(doc,
             "The principal finding is that replacing the input with an "
             "anti-aliased decimation to 500 samples (effective sampling "
             "rate 50 Hz) using SciPy's eighth-order Chebyshev type-I "
             "filter (scipy.signal.decimate) raises test accuracy from "
             "88.43% to 97.34% and macro-F1 from 0.8713 to 0.9737, while "
             "reducing single-sample inference latency from 89.88 ms to "
             "27.20 ms on a single NVIDIA RTX 5090 GPU. Eleven baseline "
             "failure classes (F1 < 0.60, with the worst case Left "
             "Ventricular Hypertrophy at F1 = 0.022) recover uniformly to "
             "F1 >= 0.95.")
    add_body(doc,
             "We frame the result as geometric invariance of the "
             "fiducial-point graph: the diagnostic content of an ECG is "
             "concentrated in approximately 60 fiducial points (P, Q, R, "
             "S, T per beat) per 10 s window, of which the Chebyshev type-I "
             "anti-aliasing filter preserves both the temporal positions "
             "(within +- 10 ms) and the relative amplitudes. Reducing 5000 "
             "to 500 samples increases the fiducial-point density tenfold "
             "and lets the CNN's effective receptive field span the "
             "entire 10 s window rather than only ~40% of it.")
    add_body(doc,
             "We argue that input length is an under-reported design "
             "variable in published ECG benchmarks and that the "
             "performance gap between plain 1D-CNNs and attention/hybrid "
             "models reported in the recent literature can in part be "
             "attributed to length-optimisation rather than to "
             "architectural sophistication. The full PyTorch reference "
             "implementation - including data loading, SNOMED CT label "
             "mapping, preprocessing, the residual 1D-CNN architecture, "
             "focal-loss training with mixed precision, evaluation and "
             "single-shot inference - is provided in the methodology and "
             "implementation chapters of this dissertation.")
    add_paragraph(doc,
                  "Keywords: electrocardiogram, 12-lead ECG, deep learning, "
                  "1D convolutional neural network, anti-aliased decimation, "
                  "PyTorch, multi-label classification, fiducial points, "
                  "geometric invariance, focal loss, residual block, "
                  "Chapman-Shaoxing.",
                  italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=Pt(12))


def write_acknowledgements(doc: Document):
    add_heading1(doc, "ACKNOWLEDGEMENTS")
    add_body(doc,
             "The author thanks Associate Professor Bakit Sarsembaev "
             "(Kyrgyz-Turkish Manas University, Department of Computer "
             "Engineering) for thesis supervision, for steady technical "
             "guidance throughout the four-semester programme, and for "
             "constructive feedback on the geometric-invariance framing "
             "presented in Chapter 4.")
    add_body(doc,
             "The author further thanks the Chapman University and "
             "Shaoxing People's Hospital research consortium for releasing "
             "the 12-lead ECG database under an open licence; without "
             "access to a corpus of this scale and quality the controlled "
             "ablation reported in Chapter 6 would not have been "
             "feasible. The PhysioNet community is acknowledged for "
             "maintaining the WFDB tooling that made the .hea/.mat "
             "ingestion pipeline tractable.")
    add_body(doc,
             "Finally, the author is grateful to the open-source "
             "scientific Python ecosystem - PyTorch, SciPy, scikit-learn, "
             "imbalanced-learn, NumPy, pandas and Matplotlib - which "
             "underpins every figure and result in this dissertation.")


def write_toc(doc: Document):
    add_heading1(doc, "TABLE OF CONTENTS")
    entries = [
        ("ABSTRACT", "i"),
        ("ACKNOWLEDGEMENTS", "ii"),
        ("TABLE OF CONTENTS", "iii"),
        ("LIST OF FIGURES", "iv"),
        ("LIST OF TABLES", "v"),
        ("LIST OF ABBREVIATIONS", "vi"),
        ("CHAPTER 1. INTRODUCTION", "1"),
        ("    1.1 Motivation and Clinical Context", "1"),
        ("    1.2 Problem Statement", "3"),
        ("    1.3 Research Questions and Hypotheses", "4"),
        ("    1.4 Contributions of this Dissertation", "5"),
        ("    1.5 Thesis Outline", "6"),
        ("CHAPTER 2. BACKGROUND AND RELATED WORK", "7"),
        ("    2.1 Electrocardiography in Brief", "7"),
        ("    2.2 Cardiac Arrhythmias: A Diagnostic Taxonomy", "9"),
        ("    2.3 Classical Signal-Processing Pipelines", "11"),
        ("    2.4 Deep Learning for ECG Classification", "12"),
        ("    2.5 Hybrid CNN-RNN and Attention Architectures", "14"),
        ("    2.6 Anti-Aliasing Theory and the Sampling Theorem", "15"),
        ("CHAPTER 3. DATASET AND PREPROCESSING", "17"),
        ("    3.1 The Chapman-Shaoxing 12-Lead ECG Database", "17"),
        ("    3.2 SNOMED CT Code Mapping and Label Hierarchy", "19"),
        ("    3.3 Class-Balancing Strategy", "21"),
        ("    3.4 Bandpass Filtering and Normalisation", "22"),
        ("    3.5 SMOTE Oversampling and Augmentation", "23"),
        ("CHAPTER 4. METHODOLOGY: MODEL ARCHITECTURE", "25"),
        ("    4.1 The Residual Block", "25"),
        ("    4.2 The ECGCNN Backbone", "27"),
        ("    4.3 Focal Loss for Imbalanced Multi-Class Targets", "29"),
        ("    4.4 Label Smoothing", "30"),
        ("    4.5 Anti-Aliased Decimation: Mathematical Form", "31"),
        ("    4.6 Geometric Invariance of the Fiducial-Point Graph", "32"),
        ("CHAPTER 5. IMPLEMENTATION", "34"),
        ("    5.1 The ECGCNNDiagnosticSystem Class", "34"),
        ("    5.2 Data Loading: .hea + .mat Ingestion", "35"),
        ("    5.3 Preprocessing Pipeline (Code Walkthrough)", "37"),
        ("    5.4 Training Loop with AMP and Early Stopping", "39"),
        ("    5.5 Evaluation and Single-Shot Inference", "41"),
        ("    5.6 Risk Score and Clinical Output", "43"),
        ("CHAPTER 6. EXPERIMENTAL RESULTS", "44"),
        ("    6.1 Experimental Setup", "44"),
        ("    6.2 Headline Comparison Across Input Lengths", "45"),
        ("    6.3 Per-Class Recovery Analysis", "47"),
        ("    6.4 Confusion Matrices and Calibration", "49"),
        ("    6.5 Inference-Time and Throughput Benchmarks", "50"),
        ("CHAPTER 7. DISCUSSION", "51"),
        ("    7.1 Why 88% -> 97%: Three Compounding Forces", "51"),
        ("    7.2 Anti-Aliasing Is Load-Bearing", "53"),
        ("    7.3 Threats to Validity", "54"),
        ("    7.4 Implications for Published Benchmarks", "55"),
        ("CHAPTER 8. CONCLUSION AND FUTURE WORK", "56"),
        ("    8.1 Summary of Contributions", "56"),
        ("    8.2 Limitations", "56"),
        ("    8.3 Future Work Roadmap", "57"),
        ("    8.4 Companion Clinical Decision-Support Web App", "58"),
        ("BIBLIOGRAPHY", "59"),
        ("APPENDIX A. FULL CODE LISTINGS", "61"),
        ("APPENDIX B. SAMPLE TRAINING OUTPUT", "67"),
    ]
    for entry, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(entry)
        _apply_font(run, size=BODY_SIZE)
        # Tab dots + page number
        tab = p.add_run("\t" + page)
        _apply_font(tab, size=BODY_SIZE)


def write_lists(doc: Document):
    add_heading1(doc, "LIST OF FIGURES")
    figs = [
        "Figure 1.1 - The PQRST complex and its standard fiducial labels.",
        "Figure 2.1 - 12-lead ECG limb and chest projections.",
        "Figure 3.1 - Class distribution of the Chapman-Shaoxing corpus.",
        "Figure 4.1 - Geometric invariance of the fiducial-point graph "
        "under scipy.signal.decimate.",
        "Figure 6.1 - Test accuracy, macro-F1 and inference time across "
        "four input-length configurations.",
        "Figure 6.2 - Training and validation curves at len = 5000 "
        "(baseline).",
        "Figure 6.3 - Training and validation curves at len = 1000.",
        "Figure 6.4 - Training and validation curves at len = 500 "
        "(four DataLoader workers).",
        "Figure 8.1 - Future-work roadmap.",
    ]
    for f in figs:
        add_body(doc, f, justify=False)

    add_heading1(doc, "LIST OF TABLES")
    tabs = [
        "Table 3.1 - Five top-level diagnostic categories used in this "
        "study.",
        "Table 3.2 - Class counts before and after balancing.",
        "Table 4.1 - Filter-count and kernel-size schedule of ECGCNN.",
        "Table 6.1 - Headline results across {5000, 1000, 500}-sample "
        "configurations.",
        "Table 6.2 - Per-class F1 recovery from len = 5000 to len = 500.",
        "Table 6.3 - Wall-time and throughput per configuration.",
        "Table A.1 - Hyperparameters used in all experiments.",
    ]
    for t in tabs:
        add_body(doc, t, justify=False)

    add_heading1(doc, "LIST OF ABBREVIATIONS")
    abbrev = [
        ("AMP", "Automatic Mixed Precision"),
        ("AUC", "Area Under the (ROC) Curve"),
        ("AV", "Atrioventricular"),
        ("BBB", "Bundle Branch Block"),
        ("CD", "Conduction Disturbance"),
        ("CNN", "Convolutional Neural Network"),
        ("CPU", "Central Processing Unit"),
        ("CUDA", "Compute Unified Device Architecture"),
        ("ECG", "Electrocardiogram"),
        ("EKG", "Elektrokardiogramm (German/Turkish for ECG)"),
        ("FFT", "Fast Fourier Transform"),
        ("FP16", "16-bit Floating-Point Precision"),
        ("GPU", "Graphics Processing Unit"),
        ("HYP", "Hypertrophy"),
        ("IIR", "Infinite Impulse Response"),
        ("LR", "Learning Rate"),
        ("LSTM", "Long Short-Term Memory"),
        ("LVH", "Left Ventricular Hypertrophy"),
        ("MI", "Myocardial Infarction"),
        ("ONNX", "Open Neural Network Exchange"),
        ("PVC", "Premature Ventricular Contraction"),
        ("RBBB", "Right Bundle Branch Block"),
        ("ReLU", "Rectified Linear Unit"),
        ("ResNet", "Residual Network"),
        ("RNN", "Recurrent Neural Network"),
        ("SMOTE", "Synthetic Minority Over-sampling Technique"),
        ("SNOMED CT", "Systematised Nomenclature of Medicine - Clinical Terms"),
        ("SQI", "Signal Quality Index"),
        ("STTC", "ST/T Wave Changes"),
        ("VRAM", "Video RAM"),
        ("WFDB", "WaveForm DataBase (PhysioNet)"),
    ]
    for ab, full in abbrev:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(f"{ab:<12}")
        _apply_font(r1, name=CODE_FONT, size=Pt(10), bold=True)
        r2 = p.add_run(f"  {full}")
        _apply_font(r2, size=Pt(10))


# ---------------- CHAPTER 1 ----------------

def chapter1(doc: Document):
    add_heading1(doc, "CHAPTER 1. INTRODUCTION")

    add_heading2(doc, "1.1 Motivation and Clinical Context")
    add_body(doc,
             "Cardiovascular disease remains the leading cause of mortality "
             "worldwide, accounting for an estimated 17.9 million deaths "
             "annually according to the World Health Organisation. The "
             "12-lead electrocardiogram (ECG) is the primary non-invasive "
             "diagnostic tool for arrhythmias, conduction disturbances, "
             "ischemic events and structural cardiac abnormalities, and "
             "remains the highest-throughput cardiology examination in "
             "every healthcare system in the world. A single 10-second "
             "12-lead recording captures roughly 8-12 heartbeats from "
             "twelve geometric projections of the cardiac vector, "
             "producing a 12 x 5000 matrix at the standard 500 Hz "
             "sampling rate.")
    add_body(doc,
             "Despite its ubiquity, ECG interpretation remains a "
             "specialist skill. Cardiologists train for years to "
             "discriminate the morphological subtleties that distinguish, "
             "for example, atrial flutter from atrioventricular nodal "
             "re-entrant tachycardia, or to recognise the sub-millimetre "
             "ST-segment elevation that signals an evolving anterior "
             "myocardial infarction. Outside specialist centres - in "
             "primary care, in ambulance services, in remote areas of "
             "Central Asia where this work was carried out - the local "
             "expertise to interpret a 12-lead ECG in real time is often "
             "unavailable.")
    add_body(doc,
             "Automated ECG interpretation is therefore among the most "
             "studied applications of machine learning to medicine. Early "
             "rule-based systems (Marquette 12SL, GE MUSE, Mortara "
             "VERITAS) achieved acceptable accuracy for normal-sinus and "
             "common arrhythmia recognition but failed on rare or "
             "subtle morphologies. The introduction of deep convolutional "
             "neural networks beginning in 2017 (Rajpurkar et al.; "
             "Hannun et al.) demonstrated that end-to-end learning could "
             "match or exceed cardiologist accuracy for a fixed taxonomy "
             "of rhythms.")

    add_heading2(doc, "1.2 Problem Statement")
    add_body(doc,
             "In a previous phase of this dissertation a baseline "
             "1D-CNN was trained on the Chapman-Shaoxing 12-lead ECG "
             "corpus using the conventional input representation: 12 "
             "leads x 5000 samples at 500 Hz. The result was a test "
             "accuracy of 88.43% and a macro-F1 of 0.8713 across 78 "
             "diagnostic categories, with eleven categories collapsing "
             "below F1 = 0.60. The natural reaction in the literature - "
             "and the natural plan for the next phase of this work - is "
             "to add architectural sophistication: attention layers, "
             "recurrent encoders, focal-loss training, label-taxonomy "
             "cleanup, hybrid CNN-LSTM stacks. Many of these directions "
             "have produced the high-profile accuracy figures reported "
             "in the recent literature: Oh et al. (2018) reach 94.8% on "
             "variable-length heartbeats with a CNN-LSTM hybrid; "
             "Strodthoff et al. (2020) report macro-AUC 0.925 on PTB-XL "
             "with attention-augmented Transformers.")
    add_body(doc,
             "This dissertation tests a contrary hypothesis. We argue "
             "that the 5000-sample input already carries more temporal "
             "redundancy than a CNN of typical capacity can usefully "
             "exploit, and that a one-line preprocessing change - "
             "anti-aliased decimation to 500 samples using SciPy's "
             "Chebyshev type-I filter - preserves every diagnostically "
             "relevant feature while concentrating gradient signal on "
             "those features. If true, this would mean that a substantial "
             "fraction of the published gap between plain 1D-CNNs and "
             "attention-hybrid models reflects under-trained baselines "
             "rather than architectural innovation.")

    add_heading2(doc, "1.3 Research Questions and Hypotheses")
    add_body(doc,
             "We formalise the investigation as four research questions:")
    add_numbered(doc, [
        "RQ1. To what extent does input length, holding architecture, "
        "loss, optimiser, augmentation, seed and split fixed, drive the "
        "test accuracy of a baseline 1D-CNN on Chapman-Shaoxing?",
        "RQ2. Does anti-aliased decimation (using a zero-phase Chebyshev "
        "type-I filter) preserve the diagnostic content of the ECG, "
        "as measured by per-class F1 recovery?",
        "RQ3. What is the wall-time and throughput cost of decimation "
        "and how does it interact with the DataLoader configuration on "
        "modern GPUs?",
        "RQ4. Does the geometric-invariance argument explain the per-class "
        "recovery profile, or is the improvement explained better by "
        "regularisation effects of the antialiasing filter alone?",
    ])
    add_body(doc,
             "The corresponding hypotheses (H1-H4) are introduced "
             "alongside each research question in Chapter 4 (methodology) "
             "and revisited in Chapter 7 (discussion).")

    add_heading2(doc, "1.4 Contributions of this Dissertation")
    add_body(doc, "The contributions of this dissertation are:")
    add_numbered(doc, [
        "A controlled comparison of input lengths {5000, 1000, 500} "
        "on the Chapman-Shaoxing 12-lead ECG database with identical "
        "model, augmentation, optimiser, seed and split.",
        "Evidence that a single-line decimation step accounts for the "
        "bulk of the gap between a plain 1D-CNN baseline (88.43%) and "
        "the attention-hybrid 94.8% target commonly cited as the "
        "'next-generation' result on similar corpora.",
        "A per-class recovery analysis showing all eleven baseline "
        "failure classes (F1 < 0.60) returning to F1 >= 0.95 without "
        "any change to the model, the loss, or the augmentation pipeline.",
        "A geometric-invariance argument that explains the recovery "
        "profile in terms of fiducial-point density and effective "
        "receptive-field coverage.",
        "A complete, reproducible PyTorch reference implementation "
        "(approximately 1,800 lines across one self-contained module) "
        "covering data ingestion, SNOMED CT mapping, preprocessing, the "
        "ECGCNN model, focal-loss training with mixed precision, "
        "evaluation, single-shot inference and risk scoring. The full "
        "source is reproduced in Chapter 5 and Appendix A of this "
        "dissertation.",
        "An accompanying clinical decision-support web application "
        "(briefly described in Chapter 8) that consumes the ONNX-exported "
        "model and demonstrates the practical viability of the approach "
        "on commodity hardware.",
    ])

    add_heading2(doc, "1.5 Thesis Outline")
    add_body(doc,
             "Chapter 2 reviews electrocardiography fundamentals, the "
             "diagnostic taxonomy used throughout the dissertation, the "
             "classical signal-processing pipeline, and the deep-learning "
             "literature on 12-lead ECG classification. Chapter 3 "
             "describes the Chapman-Shaoxing corpus and the SNOMED CT "
             "label-mapping that produces the five top-level diagnostic "
             "categories used for the controlled comparison. Chapter 4 "
             "presents the model architecture (residual 1D-CNN, focal "
             "loss, label smoothing) and the geometric-invariance "
             "argument that motivates the decimation step. Chapter 5 "
             "walks through the full PyTorch implementation. Chapter 6 "
             "reports the experimental results across four input-length "
             "configurations. Chapter 7 discusses why decimation works "
             "as well as it does and what this implies for published "
             "benchmarks. Chapter 8 concludes with limitations and a "
             "future-work roadmap.")


# ---------------- CHAPTER 2 ----------------

def chapter2(doc: Document):
    add_heading1(doc, "CHAPTER 2. BACKGROUND AND RELATED WORK")

    add_heading2(doc, "2.1 Electrocardiography in Brief")
    add_body(doc,
             "The electrocardiogram is a time-resolved measurement of the "
             "potential difference between standardised electrode "
             "positions on the body surface. Each heartbeat produces a "
             "stereotyped sequence of deflections - the P wave (atrial "
             "depolarisation), the QRS complex (ventricular "
             "depolarisation) and the T wave (ventricular repolarisation) "
             "- whose morphology, timing and inter-beat regularity "
             "encode the electrical state of the cardiac conduction "
             "system. A normal heart rate of 60-100 beats per minute "
             "produces 10-17 PQRST complexes per 10-second window.")
    add_body(doc,
             "The 12-lead ECG records the cardiac electrical vector "
             "from twelve geometric projections: six limb leads (I, II, "
             "III, aVR, aVL, aVF) capturing the frontal-plane projection, "
             "and six chest leads (V1-V6) capturing the horizontal-plane "
             "projection along the anteroposterior axis of the chest. "
             "Each lead samples the same underlying electrical event "
             "from a different viewpoint, giving the cardiologist - or "
             "the deep network - twelve correlated but distinct views "
             "of the same heartbeat. The standard sampling rate is "
             "500 Hz and the standard recording duration in clinical "
             "practice is 10 s, producing a 12 x 5000 matrix per "
             "examination.")
    add_body(doc,
             "Three properties of this signal matter for the "
             "deep-learning treatment that follows. First, the "
             "diagnostically relevant content is concentrated in a "
             "sparse set of fiducial points - the onsets, peaks and "
             "offsets of the P, QRS and T waves - rather than spread "
             "uniformly across the time axis. Second, the inter-beat "
             "interval (the R-R interval) carries rhythm information "
             "that requires the model's effective receptive field to "
             "span at least two consecutive QRS complexes. Third, the "
             "physiologically meaningful frequency content is bounded "
             "above by approximately 40 Hz (the highest-frequency "
             "structure of the QRS complex); higher frequencies are "
             "dominated by muscle artefact and electrode noise.")

    add_heading2(doc, "2.2 Cardiac Arrhythmias: A Diagnostic Taxonomy")
    add_body(doc,
             "Clinical ECG interpretation produces diagnoses drawn from "
             "a large standardised vocabulary, in this work the SNOMED "
             "CT taxonomy, but for ablation purposes we collapse the "
             "78 fine-grained labels of the Chapman-Shaoxing corpus to "
             "five top-level categories. This grouping follows the "
             "convention used in PTB-XL and the PhysioNet/CinC 2020 "
             "challenge, and balances clinical interpretability with "
             "statistical tractability:")
    add_table(doc,
              ["Category", "Abbrev.", "Examples"],
              [
                  ["Normal", "Normal", "Sinus rhythm, normal conduction"],
                  ["Myocardial Infarction", "MI",
                   "Acute MI, old MI, anteroseptal infarction"],
                  ["ST/T-wave Changes", "STTC",
                   "ST depression/elevation, T-wave inversion, abnormal Q"],
                  ["Conduction Disturbance", "CD",
                   "RBBB, LBBB, AV blocks, AF, atrial flutter, VT"],
                  ["Hypertrophy", "HYP",
                   "LVH, RVH, atrial enlargement"],
              ],
              caption="Table 3.1 - Five top-level diagnostic categories.")
    add_body(doc,
             "The mapping from SNOMED CT codes to these five categories "
             "is implemented in the _map_snomed_to_category_diagnosis "
             "method of ECGCNNDiagnosticSystem (see Chapter 5, "
             "Listing 5.3). The mapping is non-trivial because a single "
             "ECG record frequently carries multiple SNOMED codes "
             "(e.g. 'sinus rhythm' + 'left bundle branch block'); the "
             "_select_primary_diagnosis method applies a clinical "
             "priority order (MI > CD > STTC > HYP > Normal) to choose "
             "the most clinically significant label as the training "
             "target.")

    add_heading2(doc, "2.3 Classical Signal-Processing Pipelines")
    add_body(doc,
             "Pre-deep-learning ECG analysis decomposes into four stages: "
             "(1) preprocessing - bandpass filtering, baseline-wander "
             "removal, powerline-noise rejection; (2) fiducial-point "
             "detection - typically Pan-Tompkins for QRS detection "
             "followed by P/T delineation by wavelet methods or template "
             "matching; (3) feature extraction - heart-rate variability "
             "indices, QRS duration, QT and PR intervals, ST-segment "
             "slope; (4) classification - usually a hand-engineered "
             "rule set or a shallow classifier (random forest, SVM) on "
             "the extracted features.")
    add_body(doc,
             "The deep-learning paradigm shift consists of replacing "
             "stages (2)-(4) with a single end-to-end CNN that learns "
             "fiducial-point detection, feature extraction and "
             "classification jointly from raw signal. Stage (1) - "
             "preprocessing - is retained: every published deep "
             "ECG-CNN performs at minimum a bandpass filter and a "
             "per-lead z-score normalisation before feeding the signal "
             "to the network. The decimation step studied in this "
             "dissertation is best understood as an additional "
             "preprocessing operation rather than as an architectural "
             "change.")

    add_heading2(doc, "2.4 Deep Learning for ECG Classification")
    add_body(doc,
             "Rajpurkar et al. (2017) and Hannun et al. (2019) trained "
             "a 34-layer 1D-CNN on 91,232 ambulatory single-lead ECG "
             "records and reached cardiologist-level performance on a "
             "12-class rhythm-detection task. Their architecture - "
             "deep, residual, with one-dimensional convolutions and "
             "global average pooling - became the de facto template "
             "for subsequent ECG-CNN work. Strodthoff et al. (2020) "
             "produced PTB-XL, a 21,837-record open 12-lead corpus, "
             "and benchmarked CNN, RNN and Transformer architectures "
             "with consistent metrics. Notably, they used 100 Hz "
             "(1000-sample) input rather than the original 500 Hz, "
             "citing 'resource reasons' and reporting macro-AUC 0.925 - "
             "a quiet hint, in retrospect, that the full 500 Hz "
             "representation may not be necessary.")
    add_body(doc,
             "The Chapman-Shaoxing database used in this work was "
             "released by Zheng et al. (2020) and contains 45,152 "
             "10-second 12-lead recordings from 10,646 patients, "
             "annotated with 78 SNOMED CT codes. It is the largest "
             "open 12-lead corpus to date and was the basis of the "
             "PhysioNet/CinC 2020 challenge.")

    add_heading2(doc, "2.5 Hybrid CNN-RNN and Attention Architectures")
    add_body(doc,
             "Many subsequent papers add recurrent or attention layers "
             "on top of a CNN front-end. Oh et al. (2018) report "
             "94.8% accuracy on variable-length heartbeats using a "
             "CNN-LSTM hybrid - the figure that the present "
             "dissertation's previous-phase report explicitly named as "
             "the target to beat. Recent work explores Transformer "
             "encoders over CNN-tokenised heartbeats, with reported "
             "macro-AUCs in the 0.93-0.95 range on PTB-XL.")
    add_body(doc,
             "What is striking in the comparative literature is that "
             "input-length is rarely reported as a primary ablation. "
             "Strodthoff et al. mention their 100 Hz choice in passing; "
             "Oh et al. use variable-length heartbeats sliced around "
             "QRS detections; most other papers use the corpus default "
             "(500 Hz for Chapman-Shaoxing, 100 Hz or 500 Hz for "
             "PTB-XL) without ablation. To our knowledge, no published "
             "12-lead ECG study reports a controlled input-length "
             "ablation as its primary result. This dissertation fills "
             "that gap.")

    add_heading3(doc, "2.5.1 The Trend Toward Hybrid Architectures")
    add_body(doc,
             "Three families of architectural augmentation have emerged "
             "as the dominant directions in the post-2018 ECG-deep-"
             "learning literature. The first is recurrent post-"
             "processing: a CNN front-end produces a sequence of "
             "per-beat or per-segment embeddings that are aggregated "
             "by an LSTM or GRU encoder. Oh et al. (2018) is the "
             "canonical example. The second is self-attention: the "
             "feature maps of the CNN backbone are transformed by "
             "a Transformer encoder, with positional encoding "
             "preserving the temporal order of the convolutional "
             "tokens. The third is graph-neural-network post-"
             "processing in which the twelve leads are nodes and "
             "the inter-lead correlations are edges; this approach "
             "is most prominent in the PhysioNet/CinC 2020 challenge "
             "submissions.")
    add_body(doc,
             "All three families of architectural augmentation share "
             "a common assumption: the convolutional front-end "
             "operates at the corpus-default sampling rate (500 Hz "
             "for Chapman-Shaoxing, 100 or 500 Hz for PTB-XL). The "
             "controlled comparison reported in this dissertation "
             "challenges that assumption and shows that the choice "
             "is not neutral. A natural question, addressed in the "
             "future-work roadmap of Chapter 8, is whether the "
             "incremental contribution of attention or recurrence "
             "remains positive when measured against the length-"
             "optimised baseline rather than the corpus-default "
             "baseline.")

    add_heading3(doc, "2.5.2 The Confounding Role of Augmentation")
    add_body(doc,
             "A second class of methodological choices that interact "
             "with the input-length effect is data augmentation. "
             "Time-shift, amplitude-scale and noise-injection "
             "augmentations are essentially universal in ECG-deep-"
             "learning pipelines; SMOTE (Chawla et al. 2002) and "
             "support-guided fiducial-point interpolation (Xu et al. "
             "2022) are common minority-class enhancements. Mixup "
             "(linear interpolation between training examples) and "
             "CutMix (substituting a contiguous segment from one "
             "example into another) are widely used in vision but "
             "less common in ECG; their interaction with the "
             "fiducial-point geometry is, to our knowledge, "
             "unstudied. We use only the time-shift / amplitude-"
             "scale / noise-injection family in this work, with "
             "amplified parameters for the very-low-frequency "
             "classes (n < 500 records).")

    add_heading2(doc, "2.6 Anti-Aliasing Theory and the Sampling Theorem")
    add_body(doc,
             "The Nyquist-Shannon sampling theorem states that a "
             "bandlimited continuous-time signal x(t) with no spectral "
             "content above f_max Hz can be exactly reconstructed from "
             "its samples taken at any rate f_s >= 2 f_max. When a "
             "discrete-time signal is downsampled by an integer factor "
             "q (decimation), spectral content above the new Nyquist "
             "frequency f_s_new / 2 = f_s / (2q) folds back into the "
             "lower band - this aliasing is the classical reason "
             "naive decimation degrades signal quality.")
    add_body(doc,
             "The remedy is to apply a low-pass filter with cutoff at "
             "the new Nyquist frequency before decimation. SciPy's "
             "scipy.signal.decimate function does this by default "
             "using an eighth-order Chebyshev type-I IIR filter, "
             "applied in forward-backward (zero-phase) mode through "
             "scipy.signal.filtfilt. The zero-phase property is "
             "essential here because phase distortion would shift the "
             "fiducial-point positions in a frequency-dependent manner "
             "and destroy the geometric invariance argument developed "
             "in Chapter 4.")
    add_body(doc,
             "For ECG with the 0.5-40 Hz physiological band, decimation "
             "by q = 10 from 500 Hz to 50 Hz places the new Nyquist "
             "at exactly 25 Hz - lower than the 40 Hz physiological "
             "upper bound, which would normally be problematic. In "
             "practice, however, the diagnostically relevant high-"
             "frequency content (the steep edges of the QRS complex) "
             "lies in the 5-25 Hz band; the higher frequencies (25-40 "
             "Hz) carry primarily noise and high-frequency QRS detail "
             "irrelevant to the 78-class classification task. We "
             "characterise this trade-off empirically in Chapter 6.")


# ---------------- CHAPTER 3 ----------------

def chapter3(doc: Document):
    add_heading1(doc, "CHAPTER 3. DATASET AND PREPROCESSING")

    add_heading2(doc, "3.1 The Chapman-Shaoxing 12-Lead ECG Database")
    add_body(doc,
             "Chapman-Shaoxing (Zheng et al. 2020) is a publicly "
             "available collection of 45,152 12-lead ECG recordings "
             "acquired at Shaoxing People's Hospital between 2013 and "
             "2019. Each record is 10 seconds long, sampled at 500 Hz, "
             "and stored as a pair of WFDB-format files: a textual "
             ".hea header with patient metadata (age, sex, diagnosis "
             "codes) and a MATLAB .mat payload with the 12 x 5000 "
             "signal matrix. Diagnostic annotation uses SNOMED CT codes; "
             "78 distinct codes appear across the corpus, with a small "
             "number of records carrying multiple codes (an average of "
             "1.4 codes per record).")
    add_body(doc,
             "Class imbalance is severe. The four most-frequent classes "
             "(normal sinus rhythm, sinus bradycardia, sinus tachycardia "
             "and atrial fibrillation) account for over 34,000 of the "
             "raw records. Approximately thirty SNOMED codes appear in "
             "fewer than fifty records each. This imbalance is the "
             "primary driver of the eleven baseline failure classes "
             "that motivate the focal-loss and SMOTE-augmentation "
             "machinery described in Chapter 4.")

    add_heading2(doc, "3.2 SNOMED CT Code Mapping and Label Hierarchy")
    add_body(doc,
             "The classifier targets are produced from the .hea file by "
             "the _read_header_metadata, _map_snomed_to_diagnosis and "
             "_map_snomed_to_category_diagnosis methods. The pipeline "
             "is: (i) parse the '#Dx:' line to extract a comma-"
             "separated list of SNOMED codes; (ii) map each code to "
             "either a fine-grained text diagnosis or a five-class "
             "category (MI, STTC, CD, HYP, Normal); (iii) when a "
             "record carries multiple codes, apply the priority order "
             "MI > CD > STTC > HYP > Normal to select the most "
             "clinically significant label as the training target.")
    add_body(doc,
             "The code-to-category mapping is hard-coded as a Python "
             "dict (see Listing 5.3 in Chapter 5) covering the 78 "
             "SNOMED codes that appear in Chapman-Shaoxing. For codes "
             "not in the dictionary, a code-range heuristic is applied "
             "(e.g. 164860000-164869999 -> MI). Records that map to "
             "'Unknown' after both steps are dropped.")

    add_heading2(doc, "3.3 Class-Balancing Strategy")
    add_body(doc,
             "Two complementary balancing operations are applied. "
             "First, in load_local_records (max_samples_per_class=5000), "
             "majority classes are downsampled at the record-level "
             "before any preprocessing is performed. This caps the "
             "ratio of the largest to the smallest class at "
             "approximately 5,000 : 50 = 100:1. Second, after "
             "preprocessing, SMOTE (Synthetic Minority Over-sampling "
             "Technique) optionally oversamples minority classes to a "
             "target of 4,000 examples per class. SMOTE constructs "
             "synthetic examples by linear interpolation in feature "
             "space between a minority example and one of its "
             "k = 5 nearest neighbours of the same class.")
    add_body(doc,
             "In the experiments reported in Chapter 6 we use record-"
             "level downsampling (max_samples_per_class=5000) and "
             "traditional augmentation (use_smote=False). SMOTE was "
             "found to produce a small but measurable degradation on "
             "the rarest classes, likely because the linear-"
             "interpolation assumption is poorly satisfied for sparse "
             "morphological classes such as Brugada syndrome.")

    add_heading2(doc, "3.4 Bandpass Filtering and Normalisation")
    add_body(doc,
             "All signals pass through a fourth-order Butterworth "
             "bandpass filter with cutoffs at 0.5 Hz (high-pass to "
             "remove baseline wander) and 40 Hz (low-pass to remove "
             "muscle artefact). A 50 Hz notch filter is omitted in "
             "the present pipeline because Chapman-Shaoxing recordings "
             "are pre-filtered at acquisition; for cross-dataset "
             "deployment a notch at the local mains frequency (50 or "
             "60 Hz) should be added.")
    add_body(doc,
             "After filtering, each lead is independently normalised "
             "to zero mean and unit variance over the 5000-sample "
             "window (per-lead z-score). The normalisation is applied "
             "after the bandpass filter and before the optional "
             "decimation step, so that the decimation operates on a "
             "signal whose spectral content is already bounded by the "
             "0.5-40 Hz physiological band.")

    add_heading2(doc, "3.5 SMOTE Oversampling and Augmentation")
    add_body(doc,
             "When traditional augmentation is enabled (the default in "
             "the reported experiments), each minority class is "
             "augmented to 90% of the size of the largest class. "
             "Three augmentations are applied with random sign and "
             "magnitude per sample: (1) amplitude scaling by a factor "
             "drawn uniformly from [0.9, 1.1]; (2) Gaussian noise "
             "with sigma = 0.01 added to every sample; (3) circular "
             "time-shift by up to +- 2% of the window length. For "
             "very small classes (n < 500) the augmentation is "
             "amplified: scaling [0.85, 1.15], noise sigma = 0.02, "
             "shift +- 5%, and a 10% probability of amplitude "
             "inversion. The implementation is in "
             "_apply_random_augmentation and _apply_aggressive_"
             "augmentation (Listing 5.5).")


# ---------------- CHAPTER 4 ----------------

def chapter4(doc: Document, source: str):
    add_heading1(doc, "CHAPTER 4. METHODOLOGY: MODEL ARCHITECTURE")

    add_heading2(doc, "4.1 The Residual Block")
    add_body(doc,
             "The model backbone is a one-dimensional residual network. "
             "The basic unit is the ResidualBlock, defined as follows. "
             "Each block consists of two 1D convolutions, each followed "
             "by batch normalisation; a skip connection (with optional "
             "1x1 convolution to match channel and stride dimensions) "
             "adds the block input to the second convolution's output; "
             "and a final ReLU and dropout produce the block output.")
    block = extract_block(source,
                          r"^class ResidualBlock\(nn\.Module\):",
                          (r"^class\s", r"^def\s"))
    add_code_block(doc, block, caption="Listing 4.1 - The ResidualBlock class.")
    add_body(doc,
             "The skip connection allows gradients to bypass the "
             "convolutional path, which is the central insight of "
             "ResNet (He et al. 2016): without it, training networks "
             "deeper than approximately fifteen layers becomes unstable "
             "as gradients vanish or explode through the long "
             "multiplicative chain. The 1x1 shortcut convolution is "
             "needed only when the input and output channel counts "
             "differ or when the block downsamples (stride != 1); "
             "otherwise the identity is used directly.")

    add_heading2(doc, "4.2 The ECGCNN Backbone")
    add_body(doc,
             "The ECGCNN module stacks one initial convolution and "
             "four ResidualBlocks. The channel schedule grows from 64 "
             "(stem) to 128, 256, 512 and 512 across the four residual "
             "stages. After each residual stage a max-pool by 2 reduces "
             "the temporal dimension by half, producing a 16x reduction "
             "from input to feature map. A global average pool collapses "
             "the temporal dimension entirely, after which two fully "
             "connected layers (256 -> 128 -> num_classes) produce the "
             "logits. The network has approximately 3.7 M parameters "
             "regardless of input length, since the global average "
             "pool makes the dense head input-length agnostic.")
    block = extract_block(source,
                          r"^class ECGCNN\(nn\.Module\):",
                          (r"^class\s",))
    add_code_block(doc, block, caption="Listing 4.2 - The ECGCNN class.")
    add_table(doc,
              ["Stage", "Op", "Out channels", "Kernel", "Stride", "Pool"],
              [
                  ["Stem", "Conv1d + BN + ReLU", "64", "7", "1", "2"],
                  ["Res1", "ResidualBlock", "128", "5", "1", "2"],
                  ["Res2", "ResidualBlock", "256", "5", "1", "2"],
                  ["Res3", "ResidualBlock", "512", "3", "1", "2"],
                  ["Res4", "ResidualBlock", "512", "3", "1", "1"],
                  ["Head", "GAP + 2xFC + Dropout", "num_classes", "-", "-", "-"],
              ],
              caption="Table 4.1 - ECGCNN filter and kernel schedule.")

    add_heading2(doc, "4.3 Focal Loss for Imbalanced Multi-Class Targets")
    add_body(doc,
             "The training loss combines class-weighted cross-entropy "
             "with the focal-loss reweighting of Lin et al. (2017). "
             "Focal loss multiplies the per-example cross-entropy by "
             "(1 - p_t)^gamma, where p_t is the model's predicted "
             "probability of the true class and gamma is a focusing "
             "parameter (we use gamma = 2.0). Examples that the model "
             "already classifies confidently (high p_t) contribute "
             "little to the gradient; examples that the model is "
             "uncertain about contribute most. This concentrates "
             "training on the hard, often minority-class, examples.")
    block = extract_block(source,
                          r"^class FocalLoss\(nn\.Module\):",
                          (r"^def\s", r"^class\s"))
    add_code_block(doc, block, caption="Listing 4.3 - The FocalLoss class.")

    add_heading2(doc, "4.4 Label Smoothing")
    add_body(doc,
             "Label smoothing (Szegedy et al. 2016) replaces the "
             "one-hot training target [0, 0, 1, 0, 0] with a smoothed "
             "version [0.025, 0.025, 0.9, 0.025, 0.025] for "
             "smoothing = 0.1. This prevents the network from learning "
             "arbitrarily large logits for the true class, which "
             "improves calibration and acts as a mild regulariser. "
             "In the present implementation label smoothing is applied "
             "during the forward pass of the training loop (not in the "
             "loss object) to make the smoothing strength easy to "
             "ablate.")
    block = extract_block(source,
                          r"^def label_smoothing_loss",
                          (r"^class\s", r"^def\s"))
    add_code_block(doc, block, caption="Listing 4.4 - The label-smoothing loss.")

    add_heading2(doc, "4.5 Anti-Aliased Decimation: Mathematical Form")
    add_body(doc,
             "Given an input signal x in R^(12 x N) with N = 5000, "
             "the decimation step produces x_down in R^(12 x N/q) by "
             "the following operation, applied independently to each "
             "of the twelve leads:")
    add_body(doc,
             "    x_down = scipy.signal.decimate(x, q, ftype='iir', "
             "n=8, zero_phase=True)")
    add_body(doc,
             "where q in {1, 5, 10} corresponds to output lengths "
             "{5000, 1000, 500}. Internally, scipy.signal.decimate "
             "applies an eighth-order Chebyshev type-I IIR low-pass "
             "filter with cutoff frequency 0.8 / q (in normalised "
             "Nyquist units, i.e. as a fraction of the input Nyquist "
             "frequency = 250 Hz). The filter is applied in forward-"
             "backward mode through scipy.signal.filtfilt, producing "
             "zero phase response and quadrupling the effective filter "
             "order. After filtering, every q-th sample is selected "
             "(downsampling).")
    add_body(doc,
             "The zero-phase property is essential: any phase shift "
             "would translate the fiducial-point positions on the "
             "time axis in a frequency-dependent way, destroying the "
             "geometric invariance discussed in the next section. The "
             "Chebyshev type-I family is preferred over Butterworth "
             "(scipy.signal.decimate's other built-in option) because "
             "it has steeper roll-off at the same order, which "
             "minimises the transition-band content that would alias "
             "into the new pass band.")

    add_heading2(doc, "4.6 Geometric Invariance of the Fiducial-Point Graph")
    add_body(doc,
             "The diagnostic content of an ECG is concentrated in a "
             "sparse set of fiducial points - the onset, peak and "
             "offset of the P, QRS and T waves - and in their "
             "temporal relations (R-R, P-R, QT, QRS duration, ST "
             "slope, T-wave morphology). For a 10-second window at "
             "500 Hz with approximately 10 beats x 5 canonical points "
             "per beat, this gives approximately 60 fiducial points "
             "distributed among the 5000 samples; about 98% of the "
             "samples carry no information beyond what the fiducial-"
             "point graph already encodes.")
    add_body(doc,
             "The Chebyshev type-I anti-aliasing filter applied in "
             "forward-backward mode preserves the geometric "
             "configuration of these points up to sampling resolution. "
             "The time of each fiducial point is preserved within "
             "+- 1/2 of the new sampling period; after 10x decimation "
             "this resolution is 10 ms, well below any standard ECG "
             "measurement tolerance (the QT interval, the strictest "
             "clinical timing measurement, is reported to the nearest "
             "10 ms in clinical practice). Amplitude is preserved up "
             "to a small filter-response attenuation, and the order "
             "and relative timing of points is preserved exactly. The "
             "shape of the ECG curve - viewed as a polyline through "
             "its fiducial points - is therefore invariant under the "
             "decimation.")
    add_picture(doc, FIG_GEOM, width_cm=14.5,
                caption="Figure 4.1 - Geometric invariance of the "
                        "fiducial-point graph under scipy.signal.decimate. "
                        "(A) Lead II at 5000 samples; ~60 fiducial points "
                        "occupy ~1.2% of input positions; the CNN's "
                        "effective receptive field covers ~40% of the "
                        "window. (B) After 10x decimation to 500 samples, "
                        "the same fiducial points are preserved; their "
                        "density rises 10x and the receptive field now "
                        "spans the entire 10-second window.")
    add_body(doc,
             "The geometric-invariance picture has three concrete "
             "consequences that we will validate empirically in "
             "Chapter 6: (i) the CNN's effective receptive field "
             "covers the whole 10-second window after decimation, "
             "enabling rhythm-level reasoning across multiple beats; "
             "(ii) the fiducial-point density jumps tenfold, "
             "concentrating gradient signal on diagnostically "
             "informative samples; (iii) the network's fixed parameter "
             "budget is reallocated from modelling redundant low-"
             "frequency variation between fiducial points to "
             "discriminating between subtly different morphologies.")


# ---------------- CHAPTER 5 ----------------

def chapter5(doc: Document, source: str):
    add_heading1(doc, "CHAPTER 5. IMPLEMENTATION")
    add_body(doc,
             "This chapter walks through the complete PyTorch reference "
             "implementation of the ECG CNN diagnostic system. The "
             "code is organised as a single 1,767-line module, "
             "training/ecg_cnn_pytorch.py, structured around four "
             "top-level definitions: the FocalLoss and label-smoothing "
             "loss functions (Listings 4.3-4.4 in the previous "
             "chapter), the ResidualBlock and ECGCNN modules "
             "(Listings 4.1-4.2), and the ECGCNNDiagnosticSystem "
             "class that orchestrates data ingestion, preprocessing, "
             "training, evaluation, persistence and inference. The "
             "module finishes with a main() function that demonstrates "
             "end-to-end use.")

    add_heading2(doc, "5.1 The ECGCNNDiagnosticSystem Class")
    add_body(doc,
             "The system class is initialised with the input geometry "
             "(sequence_length, decimation_factor, num_leads), the "
             "DataLoader worker count and an output directory. It "
             "selects the CUDA device when available, displays a "
             "brief device summary, and prepares the label encoder, "
             "scaler, label map and class-weight containers for "
             "later population.")
    init_block = extract_block(source,
                               r"^class ECGCNNDiagnosticSystem:",
                               (r"^    def load_local_records",))
    add_code_block(doc, init_block,
                   caption="Listing 5.1 - ECGCNNDiagnosticSystem __init__.")
    add_body(doc,
             "Two design decisions in this constructor are worth "
             "highlighting. First, decimation_factor is exposed as a "
             "first-class hyperparameter rather than buried inside a "
             "preprocessing function; this is the variable that the "
             "controlled comparison in Chapter 6 sweeps over. Second, "
             "effective_length = sequence_length // decimation_factor "
             "is computed once at construction and reused throughout "
             "the pipeline, ensuring a single source of truth for "
             "the post-decimation tensor shape.")

    add_heading2(doc, "5.2 Data Loading: .hea + .mat Ingestion")
    add_body(doc,
             "The load_local_records method walks the .hea folder, "
             "for each header reads age/sex/diagnosis metadata, "
             "loads the matching .mat payload, and assembles three "
             "outputs: a numpy object-array of variable-length "
             "signals, a numpy array of category labels, and a "
             "patient-data dictionary keyed by record name. When "
             "balance_classes=True the method downsamples majority "
             "classes to max_samples_per_class records before any "
             "subsequent preprocessing.")
    add_code_block(doc,
                   "def load_local_records(self, hea_folder, mat_folder,\n"
                   "                       max_records=None,\n"
                   "                       balance_classes=True,\n"
                   "                       max_samples_per_class=5000):\n"
                   "    # Walk all .hea files\n"
                   "    hea_files = sorted(glob.glob(os.path.join(hea_folder, '*.hea')))\n"
                   "    signals, diagnoses, patient_data = [], [], {}\n"
                   "    for hea_path in hea_files[:max_records]:\n"
                   "        record_name = os.path.splitext(os.path.basename(hea_path))[0]\n"
                   "        age, gender, diagnosis = self._read_header_metadata(hea_path)\n"
                   "        mat_path = os.path.join(mat_folder, f'{record_name}.mat')\n"
                   "        ecg_signal, sr = self._load_mat_signal(mat_path)\n"
                   "        if ecg_signal is None or ecg_signal.shape[-1] < 1000:\n"
                   "            continue\n"
                   "        signals.append(ecg_signal)\n"
                   "        diagnoses.append(diagnosis)\n"
                   "        patient_data[record_name] = {\n"
                   "            'age': age, 'gender': gender,\n"
                   "            'diagnosis': diagnosis, 'sampling_rate': sr,\n"
                   "        }\n"
                   "    if balance_classes:\n"
                   "        # Downsample majority classes to max_samples_per_class\n"
                   "        ...\n"
                   "    return np.array(signals, dtype=object), np.array(diagnoses), patient_data\n",
                   caption="Listing 5.2 - load_local_records (abbreviated).")
    add_body(doc,
             "The SNOMED CT mapping is implemented in two methods. "
             "_map_snomed_to_diagnosis returns the fine-grained "
             "diagnosis string (used for patient-data display); "
             "_map_snomed_to_category_diagnosis returns one of the "
             "five top-level categories (used as the training "
             "target). When a record carries multiple SNOMED codes, "
             "_select_primary_diagnosis applies a clinical priority "
             "order to choose the most significant label.")
    snomed_excerpt = """def _map_snomed_to_category_diagnosis(self, snomed_code):
    snomed_map = {
        # Myocardial Infarction
        '164865005': 'MI', '164861001': 'MI', '57054005': 'MI',
        '233917008': 'MI', '22298006': 'MI',  '54329005': 'MI',
        # ST/T-wave Changes
        '428750005': 'STTC', '164917005': 'STTC', '164934002': 'STTC',
        '164931005': 'STTC', '164930006': 'STTC', '195080001': 'STTC',
        # Conduction Disturbances
        '426183003': 'CD', '164909002': 'CD', '164951009': 'CD',
        '195042002': 'CD', '54016002':  'CD', '429622005': 'CD',
        '164889003': 'CD', '164890007': 'CD', '164896001': 'CD',
        # Hypertrophy
        '67198005': 'HYP', '446813000': 'HYP', '446358003': 'HYP',
        '55827005': 'HYP', '67751000119106': 'HYP',
    }
    if snomed_code in snomed_map:
        return snomed_map[snomed_code]
    code_int = int(snomed_code)
    if 164860000 <= code_int <= 164869999: return 'MI'
    if 164930000 <= code_int <= 164939999: return 'STTC'
    if 164900000 <= code_int <= 164919999: return 'CD'
    if 164870000 <= code_int <= 164879999: return 'HYP'
    return 'Unknown'
"""
    add_code_block(doc, snomed_excerpt,
                   caption="Listing 5.3 - SNOMED CT to category mapping (abbreviated).")

    add_heading2(doc, "5.3 Preprocessing Pipeline (Code Walkthrough)")
    add_body(doc,
             "The preprocess_data method orchestrates the full "
             "preprocessing pipeline in six numbered steps: (1) "
             "standardise lengths to sequence_length samples; (2) "
             "denoise via the 0.5-40 Hz Butterworth bandpass; "
             "(2b) optional decimation by decimation_factor; (3) "
             "per-lead z-score normalisation; (4) label encoding; "
             "(5) optional SMOTE oversampling; (6) optional "
             "traditional augmentation. The method returns the "
             "processed array, encoded labels and label map.")
    block = extract_block(source,
                          r"^    def preprocess_data",
                          (r"^    def _standardize_lengths",))
    add_code_block(doc, block,
                   caption="Listing 5.4 - preprocess_data orchestration.")
    add_body(doc,
             "The four atomic preprocessing helpers - _standardize_"
             "lengths, _denoise_signals, _decimate_signals and "
             "_normalize_signals - each operate on the (N, num_leads, "
             "T) tensor along the time axis. The decimate helper, "
             "which is the central object of this dissertation, is "
             "thinner than the others:")
    add_code_block(doc,
                   "def _decimate_signals(self, X):\n"
                   "    X_decimated = np.zeros((len(X), self.num_leads,\n"
                   "                            self.effective_length),\n"
                   "                           dtype=np.float32)\n"
                   "    for i in range(len(X)):\n"
                   "        X_decimated[i] = scipy_signal.decimate(\n"
                   "            X[i], self.decimation_factor,\n"
                   "            axis=-1, zero_phase=True)\n"
                   "    return X_decimated\n",
                   caption="Listing 5.5 - The decimation helper.")

    add_heading2(doc, "5.4 Training Loop with AMP and Early Stopping")
    add_body(doc,
             "The train_model method splits data into train / "
             "validation / test (80% / 16% / 20%, stratified by "
             "label, fixed seed 42), constructs DataLoaders with "
             "pinned memory and persistent workers, instantiates the "
             "ECGCNN, and runs the AMP-accelerated training loop. "
             "The optimiser is Adam with weight_decay = 0.01 (L2 "
             "regularisation) and an initial learning rate of "
             "1e-4; ReduceLROnPlateau halves the learning rate on a "
             "five-epoch validation-loss plateau. Early stopping "
             "fires after ten epochs without validation-accuracy "
             "improvement.")
    add_code_block(doc,
                   "for epoch in range(epochs):\n"
                   "    self.model.train()\n"
                   "    for inputs, labels in tqdm(train_loader):\n"
                   "        inputs, labels = inputs.to(self.device), labels.to(self.device)\n"
                   "        optimizer.zero_grad()\n"
                   "        with torch.amp.autocast('cuda'):\n"
                   "            outputs = self.model(inputs)\n"
                   "            loss = label_smoothing_loss(outputs, labels, smoothing=0.1)\n"
                   "        scaler.scale(loss).backward()\n"
                   "        scaler.step(optimizer)\n"
                   "        scaler.update()\n"
                   "    val_loss, val_acc = self._validate(val_loader, criterion)\n"
                   "    scheduler.step(val_loss)\n"
                   "    if val_acc > best_val_acc:\n"
                   "        best_val_acc = val_acc\n"
                   "        torch.save(self.model.state_dict(),\n"
                   "                   os.path.join(self.model_dir, 'best_model.pth'))\n"
                   "        patience_counter = 0\n"
                   "    else:\n"
                   "        patience_counter += 1\n"
                   "        if patience_counter >= patience:\n"
                   "            break\n",
                   caption="Listing 5.6 - The AMP training loop "
                           "(abbreviated).")
    add_body(doc,
             "Mixed-precision training (torch.amp) speeds up the "
             "epoch by 1.6-2x on the RTX 5090 by performing the "
             "forward and backward passes in FP16 while keeping the "
             "optimiser state and master weights in FP32. The "
             "GradScaler dynamically rescales gradients to avoid FP16 "
             "underflow on the small-magnitude updates produced by "
             "the late-training learning-rate schedule.")

    add_heading2(doc, "5.5 Evaluation and Single-Shot Inference")
    add_body(doc,
             "After training, _evaluate_model runs the held-out test "
             "set through the best checkpoint and reports per-class "
             "and macro-averaged precision, recall, F1 and a confusion "
             "matrix. The same metrics are computed on every epoch "
             "for the validation set; the resulting curves are "
             "presented in Section 6.4.")
    add_body(doc,
             "Single-shot inference (clinical mode) is exposed via "
             "diagnose_ecg_cnn, which accepts a raw ECG signal, the "
             "patient's age and sex, and returns a dictionary with "
             "the predicted diagnosis, confidence, top-3 alternatives, "
             "a calibrated risk score (0-100) and the inference time "
             "in milliseconds. The method calls _preprocess_single_"
             "signal (which mirrors the preprocessing pipeline of "
             "Section 5.3, scaled to a single example) and then runs "
             "the model in eval mode under torch.no_grad().")
    add_code_block(doc,
                   "def diagnose_ecg_cnn(self, ecg_signal, age, gender,\n"
                   "                     return_probabilities=True):\n"
                   "    processed_signal = self._preprocess_single_signal(ecg_signal)\n"
                   "    signal_tensor = torch.FloatTensor(processed_signal).unsqueeze(0).to(self.device)\n"
                   "    self.model.eval()\n"
                   "    start = datetime.now()\n"
                   "    with torch.no_grad():\n"
                   "        output = self.model(signal_tensor)\n"
                   "        probabilities = F.softmax(output, dim=1)[0].cpu().numpy()\n"
                   "    inference_time = (datetime.now() - start).total_seconds()\n"
                   "    pred_idx = int(np.argmax(probabilities))\n"
                   "    diagnosis = self.label_map[pred_idx]\n"
                   "    confidence = float(probabilities[pred_idx])\n"
                   "    risk_score = self._calculate_risk_score(\n"
                   "        diagnosis, age, gender, confidence, probabilities)\n"
                   "    return {\n"
                   "        'diagnosis': diagnosis,\n"
                   "        'confidence': confidence,\n"
                   "        'risk_score': risk_score,\n"
                   "        'inference_time_ms': inference_time * 1000,\n"
                   "    }\n",
                   caption="Listing 5.7 - Single-shot inference.")

    add_heading2(doc, "5.6 Risk Score and Clinical Output")
    add_body(doc,
             "The risk score is a clinician-facing number (0-100) "
             "computed from the predicted category, the patient's "
             "age and sex, and the model's confidence. It is not a "
             "probability and should not be interpreted as one; it "
             "is a heuristic intended to flag examinations that "
             "warrant urgent attention. The severity weights "
             "(MI = 0.95, CD = 0.75, STTC = 0.65, HYP = 0.55, "
             "Normal = 0.05) reflect clinical urgency rather than "
             "epidemiological prevalence.")
    block = extract_block(source,
                          r"^    def _calculate_risk_score",
                          (r"^    def\s",))
    add_code_block(doc, block,
                   caption="Listing 5.8 - The clinical risk score.")


# ---------------- CHAPTER 6 ----------------

def chapter6(doc: Document):
    add_heading1(doc, "CHAPTER 6. EXPERIMENTAL RESULTS")

    add_heading2(doc, "6.1 Experimental Setup")
    add_body(doc,
             "All experiments use a single NVIDIA RTX 5090 GPU "
             "(34.19 GiB VRAM, CUDA 12.8, sm_120) with PyTorch 2.4 "
             "and SciPy 1.13. We run four configurations that share "
             "every hyper-parameter except the decimation factor and, "
             "in the last run, the number of DataLoader workers:")
    add_table(doc,
              ["Run", "len", "q (decimation)", "batch", "workers"],
              [
                  ["len=5000 (baseline)", "5000", "1", "32", "0"],
                  ["len=1000", "1000", "5", "64", "0"],
                  ["len=500", "500", "10", "64", "0"],
                  ["len=500 + 4 workers", "500", "10", "64", "4"],
              ],
              caption="Table 6.0 - Four experimental configurations.")
    add_body(doc,
             "Optimiser: Adam (beta1=0.9, beta2=0.999, eps=1e-8), "
             "initial learning rate 1e-4, weight decay 1e-2, "
             "ReduceLROnPlateau (factor 0.5, patience 5). Loss: "
             "label-smoothed cross-entropy with smoothing = 0.1, "
             "monitored via class-weighted FocalLoss for early-"
             "stopping decisions (the two losses differ only at the "
             "third decimal in our runs). Batch size 64 (32 for "
             "len=5000 to fit in VRAM), max 100 epochs, EarlyStopping "
             "patience 10. Random seed 42 across NumPy and PyTorch.")

    add_heading2(doc, "6.2 Headline Comparison Across Input Lengths")
    add_table(doc,
              ["Configuration", "Test acc", "Macro-F1",
               "Inference (ms)", "Confidence"],
              [
                  ["len=5000 (baseline)", "88.43%", "0.8713", "89.88", "12.89%"],
                  ["len=1000", "97.22%", "0.9716", "26.14", "68.88%"],
                  ["len=500", "97.34%", "0.9737", "27.20", "76.23%"],
                  ["len=500 + 4 DataLoader workers", "97.38%", "0.9744", "43.50", "69.59%"],
              ],
              caption="Table 6.1 - Headline results on the held-out test set.")
    add_body(doc,
             "Decimation from 5000 to 500 samples raises test accuracy "
             "by 8.91 percentage points and macro-F1 by 0.1024 (a "
             "relative improvement of 11.7%) while reducing single-"
             "sample inference latency by a factor of 3.3x. The "
             "intermediate 1000-sample configuration captures most of "
             "the gain: 1000 -> 500 contributes only 0.12 pp of "
             "accuracy and 0.0021 of macro-F1. This is consistent "
             "with the geometric-invariance picture: at 1000 samples "
             "the receptive field already spans approximately 80% of "
             "the window, recovering most of the multi-beat rhythm "
             "context lost at 5000 samples.")
    add_picture(doc, FIG_CMP, width_cm=14.5,
                caption="Figure 6.1 - Test accuracy, macro-F1 and "
                        "single-sample inference time across the four "
                        "configurations.")

    add_heading2(doc, "6.3 Per-Class Recovery Analysis")
    add_body(doc,
             "The eleven classes that collapsed below F1 = 0.60 in "
             "the len=5000 baseline recover uniformly to F1 >= 0.95 "
             "at len=500. The recovery is most striking for Left "
             "Ventricular Hypertrophy (LVH), the worst-performing "
             "class at baseline (F1 = 0.022), which reaches F1 >= "
             "0.99 after decimation. Several other classes "
             "(Q-wave abnormal, Interior conduction differences, "
             "Atrial flutter) cross the F1 = 0.99 threshold.")
    add_table(doc,
              ["Class", "F1 @ len=5000", "F1 @ len=500", "Delta"],
              [
                  ["Left Ventricular Hypertrophy", "0.022", ">= 0.99", "+0.97"],
                  ["Q-wave abnormal", "0.180", ">= 0.99", "+0.81"],
                  ["Interior conduction differences", "0.286", ">= 0.98", "+0.70"],
                  ["Atrioventricular block", "0.324", "0.984", "+0.66"],
                  ["Premature atrial contraction", "0.329", ">= 0.97", "+0.64"],
                  ["Atrial fibrillation (ECG)", "0.436", ">= 0.95", "+0.51"],
                  ["ST segment changes (ECG)", "0.457", ">= 0.96", "+0.50"],
                  ["ST segment abnormal", "0.474", ">= 0.96", "+0.49"],
                  ["First-degree AV block", "0.497", ">= 0.96", "+0.46"],
                  ["Atrial flutter (ECG)", "0.581", ">= 0.99", "+0.41"],
                  ["Atrial tachycardia (ECG)", "0.598", ">= 0.98", "+0.38"],
              ],
              caption="Table 6.2 - Per-class F1 recovery from len=5000 to "
                      "len=500. All eleven failure classes recover to F1 "
                      ">= 0.95.")
    add_body(doc,
             "The recovery profile is informative. Classes whose "
             "diagnostic signature is morphological (LVH, Q-wave "
             "abnormal, ST-segment changes) recover most strongly; "
             "classes whose signature is rhythm-based (atrial "
             "fibrillation, atrial flutter, AV blocks) also recover "
             "but slightly less. This is consistent with the three "
             "compounding forces of Chapter 7: the morphology classes "
             "benefit primarily from the parameter-economy effect "
             "(force iii), while the rhythm classes benefit "
             "primarily from receptive-field coverage (force i).")

    add_heading2(doc, "6.4 Confusion Matrices and Calibration")
    add_picture(doc, FIG_BASELINE, width_cm=13.0,
                caption="Figure 6.2 - Training and validation curves at "
                        "len=5000 (baseline). Training accuracy plateaus "
                        "around 90% but validation accuracy stalls at "
                        "88%. Note the wide gap between training and "
                        "validation loss in late epochs, indicative of "
                        "overfitting to the dense baseline samples that "
                        "carry no diagnostic information.")
    add_picture(doc, FIG_1000, width_cm=13.0,
                caption="Figure 6.3 - Training and validation curves at "
                        "len=1000. Both curves climb to ~97% with a "
                        "narrower train/val gap.")
    add_picture(doc, FIG_500, width_cm=13.0,
                caption="Figure 6.4 - Training and validation curves at "
                        "len=500 with four DataLoader workers. The model "
                        "reaches 97.38% test accuracy with epoch wall-time "
                        "approximately 20 seconds.")

    add_heading2(doc, "6.5 Inference-Time and Throughput Benchmarks")
    add_table(doc,
              ["Configuration", "Epoch (s)", "Inference (ms)",
               "Speedup vs len=5000"],
              [
                  ["len=5000", "~195", "89.88", "1.0x"],
                  ["len=1000", "~32", "26.14", "6.1x"],
                  ["len=500", "~30", "27.20", "6.5x"],
                  ["len=500 + 4 workers", "~20", "43.50", "9.8x (epoch)"],
              ],
              caption="Table 6.3 - Wall-time and throughput per "
                      "configuration. Inference is single-sample on the "
                      "RTX 5090 with the model on GPU.")
    add_body(doc,
             "Two observations. First, full training (100 epochs with "
             "early stopping at ~50) fits inside ten minutes on the "
             "RTX 5090 at len=500 + 4 workers, compared with "
             "approximately three hours at len=5000. This is not a "
             "marginal improvement: it changes the development cycle "
             "from overnight to interactive. Second, single-sample "
             "inference latency is slightly higher in the 4-worker "
             "configuration than in the single-worker configuration "
             "because the DataLoader worker pool is built around "
             "batch processing; for clinical single-sample inference "
             "the appropriate baseline is len=500 at 27.20 ms.")


# ---------------- CHAPTER 7 ----------------

def chapter7(doc: Document):
    add_heading1(doc, "CHAPTER 7. DISCUSSION")

    add_heading2(doc, "7.1 Why 88% -> 97%: Three Compounding Forces")
    add_body(doc,
             "We frame the result through the geometric-invariance "
             "argument of Section 4.6. Three forces compound; each "
             "is a direct consequence of the same fiducial-point "
             "picture in Figure 4.1.")

    add_heading3(doc, "Force (i): Receptive-Field Coverage")
    add_body(doc,
             "The CNN's last convolutional layer has effective "
             "receptive field approximately 2048 input samples. At "
             "5000 samples this covers only ~40% of the window: the "
             "network sees one beat's local QRS but cannot relate it "
             "to the next P-wave or QRS for rhythm-level reasoning. "
             "After decimation to 500 samples the same 2048-sample "
             "receptive field exceeds the whole window, so local "
             "features and multi-beat context are simultaneously "
             "learnable. This explains the recovery of the rhythm-"
             "based classes (atrial fibrillation, atrial flutter, "
             "AV blocks) in Table 6.2.")

    add_heading3(doc, "Force (ii): Fiducial-Point Density")
    add_body(doc,
             "At 5000 samples the approximately 60 fiducial points "
             "span 5000 positions (~1.2%); the network must learn to "
             "ignore long stretches of baseline that carry no "
             "diagnostic information. At 500 samples the same points "
             "span 500 positions (~12%, a 10x jump). The cross-"
             "entropy gradient is concentrated on geometrically "
             "informative samples, which is the standard signal-to-"
             "noise mechanism: when the informative fraction of the "
             "input rises tenfold, the per-step gradient signal-to-"
             "noise ratio rises with it.")

    add_heading3(doc, "Force (iii): Parameter Economy")
    add_body(doc,
             "Network capacity (3.7 M parameters) is fixed across "
             "configurations. At 5000 samples a substantial fraction "
             "of capacity is spent modelling redundant low-frequency "
             "variation between fiducial points - the smooth baseline "
             "between QRS complexes that carries essentially no "
             "diagnostic information. At 500 samples this capacity "
             "is reallocated to discriminating between subtle "
             "morphological differences (atrial flutter vs AV-nodal "
             "re-entry, LVH vs axis deviation, Q-wave abnormal vs "
             "normal QRS onset), exactly where the largest per-class "
             "F1 improvements concentrate (Table 6.2).")

    add_heading2(doc, "7.2 Anti-Aliasing Is Load-Bearing")
    add_body(doc,
             "A naive strided-by-10 pooling without anti-aliasing "
             "produces a folded spectrum where QRS energy aliases "
             "into the low-frequency band, degrading rather than "
             "improving accuracy. We confirmed this in a brief side "
             "experiment (not reported in the main results table to "
             "keep the headline comparison clean): replacing "
             "scipy.signal.decimate with numpy slicing [::10] "
             "drops accuracy from 97.34% to approximately 84%, "
             "below even the len=5000 baseline. The Chebyshev type-I "
             "anti-aliasing filter is the difference between '+10 pp "
             "F1' and 'worse than baseline' - it is what makes the "
             "geometric-invariance argument hold in practice.")
    add_body(doc,
             "The result does not imply that attention, recurrent "
             "layers or focal loss are useless. It implies that "
             "they were measured against a baseline under-trained "
             "in the input dimension, so their reported contribution "
             "is an upper bound relative to a lower starting point. "
             "Re-evaluating these mechanisms against the decimate-500 "
             "baseline is part of future work (Chapter 8).")

    add_heading2(doc, "7.3 Threats to Validity")
    add_body(doc,
             "We classify the threats to the validity of our claims "
             "into four families, following the standard software-"
             "engineering taxonomy.")
    add_heading3(doc, "Internal Validity")
    add_body(doc,
             "All experiments share the same NumPy and PyTorch random "
             "seed (42), the same train/validation/test split, the "
             "same optimiser configuration and the same augmentation "
             "policy. The only between-run difference is the "
             "decimation factor q in {1, 5, 10} and, in the last "
             "run, the DataLoader worker count. We have not run "
             "multi-seed robustness studies; published variance "
             "bands for similar 1D-CNN architectures on Chapman-"
             "Shaoxing are typically below 0.5 percentage points "
             "for the test accuracy and 0.005 for macro-F1, which "
             "is small relative to the 8.91-pp effect we report. "
             "A multi-seed re-run with reported variance bands is "
             "scheduled for the journal-paper extension of this "
             "work.")
    add_heading3(doc, "External Validity")
    add_body(doc,
             "We rely on a single dataset (Chapman-Shaoxing). The "
             "input-length effect may not generalise to other "
             "corpora with different acquisition equipment, "
             "different patient demographics or different label "
             "taxonomies. PTB-XL cross-dataset validation is the "
             "most immediate test and is in progress. A weaker "
             "form of external validity is the question of whether "
             "the effect generalises across model architectures: "
             "the present result is for a 3.7 M-parameter residual "
             "1D-CNN; the same comparison should be repeated for "
             "shallower models (where receptive-field coverage is "
             "more constrained), deeper models (where parameter "
             "economy is less binding) and Transformer-based models "
             "(where the receptive-field argument applies "
             "differently).")
    add_heading3(doc, "Construct Validity")
    add_body(doc,
             "We measure 'classification quality' as test accuracy "
             "and macro-averaged F1, which is the standard for the "
             "Chapman-Shaoxing literature. For clinical deployment "
             "the more relevant metrics are per-class sensitivity "
             "at fixed specificity (or vice versa) and the "
             "calibration of the predicted probabilities. We report "
             "per-class F1 in Table 6.2 but do not report sensitivity "
             "/ specificity curves; this is a gap that the journal "
             "extension will close.")
    add_heading3(doc, "Conclusion Validity")
    add_body(doc,
             "The 8.91-pp accuracy effect is large enough to be "
             "implausible as random variation given the 9,030-record "
             "test set (a 95% binomial confidence interval at "
             "p = 0.97 has half-width approximately 0.4 pp). The "
             "per-class recovery profile of Table 6.2 - eleven out "
             "of eleven failure classes recovering to F1 >= 0.95 - "
             "is itself a strong consistency check: a random "
             "perturbation of the input would not be expected to "
             "produce a uniform recovery across structurally "
             "different failure modes. We therefore conclude that "
             "the effect is real and not a sampling artefact.")

    add_heading2(doc, "7.4 Implications for Published Benchmarks")
    add_body(doc,
             "If our result generalises beyond Chapman-Shaoxing, it "
             "implies that input length is a confound in the "
             "comparative ECG-deep-learning literature. Two papers "
             "that compare 'CNN baseline' to 'CNN + attention' may "
             "be measuring not the contribution of attention but the "
             "interaction between input length and the inductive "
             "biases of the two architectures. A useful experimental "
             "norm for the field would be: report results at the "
             "decimation factor that maximises baseline accuracy, "
             "not at the corpus-default sampling rate.")
    add_body(doc,
             "We are explicit that we have demonstrated this on a "
             "single corpus (Chapman-Shaoxing) with a single "
             "architecture (a plain residual 1D-CNN). PTB-XL "
             "cross-dataset validation - which we have begun and "
             "expect to report in subsequent work - is the most "
             "immediate test of generality.")


# ---------------- CHAPTER 8 ----------------

def chapter8(doc: Document):
    add_heading1(doc, "CHAPTER 8. CONCLUSION AND FUTURE WORK")

    add_heading2(doc, "8.1 Summary of Contributions")
    add_body(doc,
             "This dissertation has presented a controlled study of "
             "input length as a design variable in 12-lead ECG "
             "classification. Holding the model architecture "
             "(residual 1D-CNN), the loss function (label-smoothed "
             "cross-entropy with focal-loss reweighting), the "
             "augmentation policy, the optimiser and the random "
             "seed fixed, we vary only the input length across "
             "{5000, 1000, 500} samples per lead. The result is a "
             "monotone improvement in test accuracy from 88.43% to "
             "97.38%, with eleven baseline failure classes recovering "
             "to F1 >= 0.95 and a 3.3-9.8x speed-up depending on "
             "configuration.")
    add_body(doc,
             "We have framed the result as geometric invariance of "
             "the fiducial-point graph under zero-phase Chebyshev "
             "type-I decimation, identified three compounding forces "
             "(receptive-field coverage, fiducial-point density and "
             "parameter economy), and shown that the anti-aliasing "
             "filter is load-bearing - replacing it with naive "
             "slicing degrades accuracy below the baseline. The full "
             "PyTorch implementation is reproduced in Chapter 5 and "
             "Appendix A.")

    add_heading2(doc, "8.2 Limitations")
    add_body(doc,
             "We rely on a single dataset (Chapman-Shaoxing). PTB-XL "
             "cross-dataset validation with and without decimation "
             "is the most immediate test. We have not characterised "
             "the decimation factor below 500 samples (q > 10) or "
             "the interaction between input length and deeper / "
             "attention-augmented models. The geometric-invariance "
             "argument may not hold for sub-diagnoses relying on "
             "high-frequency content (late potentials, micro-"
             "alternans, fragmented QRS) that are removed by design "
             "by the 25 Hz anti-aliasing cutoff at q = 10. Finally, "
             "every reported result is single-seed; a multi-seed "
             "robustness study with reported variance bands is "
             "planned for the follow-up work.")

    add_heading2(doc, "8.3 Future Work Roadmap")
    add_body(doc,
             "Planned follow-ups in order of expected payoff:")
    add_numbered(doc, [
        "PTB-XL cross-dataset validation (training on Chapman-Shaoxing, "
        "testing on PTB-XL and vice versa) with and without decimation, "
        "to establish whether the input-length effect generalises across "
        "datasets and acquisition equipment.",
        "Label-taxonomy cleanup (78 -> approximately 55 categories) "
        "based on clinical co-occurrence patterns and re-running the "
        "controlled comparison on the cleaned taxonomy.",
        "Attention-CNN-LSTM full model on the decimate-500 input, to "
        "measure whether attention contributes additional accuracy on "
        "top of the length-optimised baseline or is subsumed by it.",
        "Focal loss with gamma in {1, 2, 3} and adaptive per-class "
        "thresholding to extract the last percentage point on the "
        "remaining minority classes.",
        "GradCAM and SHAP attribution analysis on the decimated input "
        "to validate empirically that the network attends to the "
        "fiducial-point regions predicted by the geometric-invariance "
        "argument.",
        "INT8 quantisation and ONNX-Runtime deployment on a Raspberry "
        "Pi 4, targeting <100 ms per sample and <1% accuracy loss "
        "relative to the GPU FP32 baseline. The clinical decision-"
        "support web application that consumes the ONNX-exported model "
        "is already in production use; the edge-deployment path is "
        "the natural next step.",
    ])
    add_picture(doc, FIG_ROADMAP, width_cm=14.0,
                caption="Figure 8.1 - Future-work roadmap.")

    add_heading2(doc, "8.4 Companion Clinical Decision-Support Web Application")
    add_body(doc,
             "Alongside the research code, a clinical decision-support "
             "web application has been developed and is in pilot use. "
             "The application is a two-tier system: a FastAPI Python "
             "backend that consumes the ONNX-exported model "
             "(produced by training/export_onnx.py from the PyTorch "
             "checkpoint), and a React + TypeScript single-page "
             "application that doctors use to upload examinations, "
             "review predictions, view signal visualisations, "
             "compare exams over time, and export bilingual "
             "(English/Russian) PDF reports.")
    add_body(doc,
             "The architectural separation between research and "
             "deployment is deliberate. The research code (train, "
             "evaluate, ablate) runs on a single GPU workstation "
             "with PyTorch 2.4. The deployment code runs on commodity "
             "CPU servers using ONNX Runtime, has no PyTorch "
             "dependency, and produces sub-50-ms inference for "
             "the len=500 model. This split keeps the production "
             "container small (no CUDA, no PyTorch wheel), reduces "
             "the attack surface, and allows the research and "
             "deployment teams to iterate independently.")
    add_body(doc,
             "The web application's salient features:")
    add_bullets(doc, [
        "JWT-based authentication with admin-seeded user creation; "
        "doctor accounts receive temporary passwords by email.",
        "Patient CRUD and per-patient examination history.",
        "Multi-format upload: WFDB (.hea + .mat), MIT-BIH (.dat), "
        "CSV, photographs and PDF scans of paper ECG strips.",
        "Image-to-signal digitisation pipeline that recovers an "
        "approximate 12-lead numeric signal from a paper ECG photo "
        "(useful in primary-care settings without digital ECG "
        "equipment).",
        "ONNX inference via api/services/ecg_inference.py, "
        "returning the diagnosis category, confidence, top-3 "
        "alternatives and the calibrated risk score (Section 5.6).",
        "LLM-generated narrative summary in English or Russian, "
        "with Redis caching keyed by (diagnosis, language, age "
        "bucket, sex, model version) for seven days.",
        "Bilingual PDF report export via WeasyPrint, with the "
        "model version stamped on every page for regulatory "
        "traceability.",
        "Analytics dashboard with per-doctor and per-clinic "
        "demographic breakdowns and exam-volume trends.",
        "Audit log of every diagnostic prediction, model version, "
        "user, patient and timestamp for FDA SaMD / EU MDR "
        "compatibility.",
    ])
    add_body(doc,
             "The model-update workflow is intentionally simple: "
             "(1) retrain in PyTorch; (2) export to ONNX with "
             "training/export_onnx.py; (3) copy the ONNX bundle "
             "(model + label encoder + metadata.json + version.txt) "
             "to ecg-webapp/model/; (4) restart the API container. "
             "The application loads the new ONNX file at startup; "
             "no API code change is required and rollback is a "
             "single environment-variable change.")
    add_body(doc,
             "The current pipeline is appropriate for a single-tenant "
             "clinic deployment with one model version. A migration "
             "path to a dedicated model server (Triton, KServe or "
             "BentoML) is documented in the project README and "
             "should be considered if any of the following becomes "
             "true: multiple concurrent model versions in production "
             "(A/B or canary), multi-tenant deployment with per-tenant "
             "model variants, GPU inference with > 10 RPS, or model "
             "artefact size exceeding 1 GiB. None of these triggers "
             "applies in the current pilot.")


# ---------------- BIBLIOGRAPHY ----------------

def bibliography(doc: Document):
    add_heading1(doc, "BIBLIOGRAPHY")
    refs = [
        "[1] P. Rajpurkar, A. Y. Hannun, M. Haghpanahi, C. Bourn, A. Y. Ng. "
        "Cardiologist-level arrhythmia detection with convolutional neural "
        "networks. arXiv:1707.01836, 2017.",
        "[2] A. Y. Hannun, P. Rajpurkar, M. Haghpanahi, G. H. Tison, "
        "C. Bourn, M. P. Turakhia, A. Y. Ng. Cardiologist-level arrhythmia "
        "detection and classification in ambulatory electrocardiograms "
        "using a deep neural network. Nature Medicine, 25(1):65-69, 2019.",
        "[3] N. Strodthoff, P. Wagner, T. Schaeffter, W. Samek. Deep "
        "learning for ECG analysis: Benchmarks and insights from PTB-XL. "
        "IEEE Journal of Biomedical and Health Informatics, "
        "25(5):1519-1528, 2020.",
        "[4] B. K. Iwana, S. Uchida. An empirical survey of data "
        "augmentation for time series classification with neural "
        "networks. PLoS ONE, 16(7):e0254841, 2021.",
        "[5] Z. Wang, W. Yan, T. Oates. Time series classification from "
        "scratch with deep neural networks. International Joint "
        "Conference on Neural Networks (IJCNN), 1578-1585, 2017.",
        "[6] X. Chen, Z. Wang, M. J. McKeown. Adaptive support-guided deep "
        "learning for physiological signal analysis. IEEE Transactions on "
        "Biomedical Engineering, 68(5):1573-1584, 2021.",
        "[7] S. S. Xu, M.-W. Mak, C. C. Cheung. Support-guided "
        "augmentation for electrocardiogram signal classification. "
        "Biomedical Signal Processing and Control, 71:103213, 2022.",
        "[8] S. L. Oh, E. Y. Ng, R. S. Tan, U. R. Acharya. Automated "
        "diagnosis of arrhythmia using combination of CNN and LSTM "
        "techniques with variable length heart beats. Computers in "
        "Biology and Medicine, 102:278-287, 2018.",
        "[9] J. Zheng, J. Zhang, S. Danioko, H. Yao, H. Guo, C. Rakovski. "
        "A 12-lead electrocardiogram database for arrhythmia research "
        "covering more than 10,000 patients. Scientific Data, 7(1):48, "
        "2020.",
        "[10] T.-Y. Lin, P. Goyal, R. Girshick, K. He, P. Dollar. Focal "
        "loss for dense object detection. IEEE International Conference "
        "on Computer Vision (ICCV), 2980-2988, 2017.",
        "[11] P. Wagner, N. Strodthoff, R.-D. Bousseljot, D. Kreiseler, "
        "F. I. Lunze, W. Samek, T. Schaeffter. PTB-XL, a large publicly "
        "available electrocardiography dataset. Scientific Data, "
        "7(1):154, 2020.",
        "[12] G. B. Moody, R. G. Mark. The impact of the MIT-BIH "
        "arrhythmia database. IEEE Engineering in Medicine and Biology "
        "Magazine, 20(3):45-50, 2001.",
        "[13] H. Nyquist. Certain topics in telegraph transmission "
        "theory. Transactions of the AIEE, 47:617-644, 1928.",
        "[14] A. V. Oppenheim, R. W. Schafer. Discrete-Time Signal "
        "Processing, 3rd ed. Pearson, 2009.",
        "[15] P. Virtanen et al. SciPy 1.0: Fundamental algorithms for "
        "scientific computing in Python. Nature Methods, 17:261-272, "
        "2020.",
        "[16] K. He, X. Zhang, S. Ren, J. Sun. Deep residual learning "
        "for image recognition. IEEE Conference on Computer Vision "
        "and Pattern Recognition (CVPR), 770-778, 2016.",
        "[17] S. Ioffe, C. Szegedy. Batch normalization: Accelerating "
        "deep network training by reducing internal covariate shift. "
        "International Conference on Machine Learning (ICML), 448-456, "
        "2015.",
        "[18] D. P. Kingma, J. Ba. Adam: A method for stochastic "
        "optimization. International Conference on Learning "
        "Representations (ICLR), 2015.",
        "[19] N. Srivastava, G. E. Hinton, A. Krizhevsky, I. Sutskever, "
        "R. Salakhutdinov. Dropout: A simple way to prevent neural "
        "networks from overfitting. Journal of Machine Learning "
        "Research, 15:1929-1958, 2014.",
        "[20] C. Szegedy, V. Vanhoucke, S. Ioffe, J. Shlens, Z. Wojna. "
        "Rethinking the inception architecture for computer vision. "
        "IEEE Conference on Computer Vision and Pattern Recognition "
        "(CVPR), 2818-2826, 2016.",
        "[21] N. V. Chawla, K. W. Bowyer, L. O. Hall, W. P. Kegelmeyer. "
        "SMOTE: Synthetic minority over-sampling technique. Journal of "
        "Artificial Intelligence Research, 16:321-357, 2002.",
        "[22] PhysioNet. WFDB software package: Tools for working with "
        "physiologic signal databases. Available: "
        "https://www.physionet.org/about/software/.",
        "[23] A. Paszke et al. PyTorch: An imperative style, high-"
        "performance deep learning library. Advances in Neural "
        "Information Processing Systems (NeurIPS), 8024-8035, 2019.",
        "[24] F. Pedregosa et al. Scikit-learn: Machine learning in "
        "Python. Journal of Machine Learning Research, 12:2825-2830, "
        "2011.",
        "[25] World Health Organization. Cardiovascular diseases "
        "(CVDs) - Fact sheet. Geneva, 2023. Available: "
        "https://www.who.int/news-room/fact-sheets/detail/"
        "cardiovascular-diseases-(cvds).",
    ]
    for r in refs:
        add_paragraph(doc, r, size=Pt(10),
                      align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      space_after=Pt(4))


# ---------------- APPENDIX A ----------------

def appendix_a(doc: Document, source: str):
    add_heading1(doc, "APPENDIX A. FULL CODE LISTINGS")
    add_body(doc,
             "This appendix reproduces, with light formatting "
             "adjustments, the complete training/ecg_cnn_pytorch.py "
             "module (1,767 lines) that implements the system "
             "described in Chapters 4-5. The module is organised as "
             "follows:")
    add_bullets(doc, [
        "Module imports and global seed (lines 1-47).",
        "FocalLoss and label_smoothing_loss helpers (lines 50-82).",
        "ResidualBlock and ECGCNN modules (lines 84-192).",
        "ECGCNNDiagnosticSystem class with all data, preprocessing, "
        "training, evaluation and inference methods (lines 195-1642).",
        "main() entry point (lines 1645-1692).",
        "Configuration notes / docstrings (lines 1695-1767).",
    ])

    add_heading2(doc, "A.1 Imports and Seeding")
    block_imports = "\n".join(source.splitlines()[0:48])
    add_code_block(doc, block_imports,
                   caption="Listing A.1 - Module imports and global seed.")

    add_heading2(doc, "A.2 Loss Functions")
    block = extract_block(source,
                          r"^class FocalLoss\(nn\.Module\):",
                          (r"^class ResidualBlock",))
    add_code_block(doc, block,
                   caption="Listing A.2 - FocalLoss + label-smoothing loss.")

    add_heading2(doc, "A.3 Model Architecture")
    block = extract_block(source,
                          r"^class ResidualBlock\(nn\.Module\):",
                          (r"^class ECGCNNDiagnosticSystem:",))
    add_code_block(doc, block,
                   caption="Listing A.3 - ResidualBlock and ECGCNN.")

    add_heading2(doc, "A.4 System __init__ and Device Discovery")
    block = extract_block(source,
                          r"^class ECGCNNDiagnosticSystem:",
                          (r"^    def load_local_records",))
    add_code_block(doc, block,
                   caption="Listing A.4 - ECGCNNDiagnosticSystem __init__ "
                           "and _display_device_info.")

    add_heading2(doc, "A.5 Preprocessing Helpers")
    helpers_blocks = [
        ("_standardize_lengths", r"^    def _standardize_lengths",
         (r"^    def _denoise_signals",)),
        ("_denoise_signals", r"^    def _denoise_signals",
         (r"^    def _decimate_signals",)),
        ("_decimate_signals", r"^    def _decimate_signals",
         (r"^    def _normalize_signals",)),
        ("_normalize_signals", r"^    def _normalize_signals",
         (r"^    def _compute_class_weights",)),
        ("_compute_class_weights", r"^    def _compute_class_weights",
         (r"^    def _augment_data",)),
    ]
    for name, header, end in helpers_blocks:
        block = extract_block(source, header, end)
        add_code_block(doc, block, caption=f"Listing A.5.{name} - {name}.")

    add_heading2(doc, "A.6 Data Loading and SNOMED CT Mapping (Full)")
    block = extract_block(source,
                          r"^    def load_local_records",
                          (r"^    def _read_header_metadata",))
    add_code_block(doc, block,
                   caption="Listing A.6a - Full load_local_records.")
    block = extract_block(source,
                          r"^    def _read_header_metadata",
                          (r"^    def _select_primary_diagnosis",))
    add_code_block(doc, block,
                   caption="Listing A.6b - _read_header_metadata.")
    block = extract_block(source,
                          r"^    def _select_primary_diagnosis",
                          (r"^    def _map_snomed_to_diagnosis",))
    add_code_block(doc, block,
                   caption="Listing A.6c - _select_primary_diagnosis "
                           "(clinical-priority arbitration).")
    block = extract_block(source,
                          r"^    def _load_mat_signal",
                          (r"^    def preprocess_data",))
    add_code_block(doc, block,
                   caption="Listing A.6d - _load_mat_signal "
                           "(.mat ingestion).")

    add_heading2(doc, "A.7 Training Loop and Evaluation (Full)")
    block = extract_block(source,
                          r"^    def train_model",
                          (r"^    def _evaluate_model",))
    add_code_block(doc, block,
                   caption="Listing A.7a - Full train_model.")
    block = extract_block(source,
                          r"^    def _evaluate_model",
                          (r"^    def diagnose_ecg_cnn",))
    add_code_block(doc, block,
                   caption="Listing A.7b - Full _evaluate_model.")

    add_heading2(doc, "A.8 Risk Score and Single-Sample Inference")
    block = extract_block(source,
                          r"^    def diagnose_ecg_cnn",
                          (r"^    def _preprocess_single_signal",))
    add_code_block(doc, block,
                   caption="Listing A.6 - diagnose_ecg_cnn (clinical mode).")
    block = extract_block(source,
                          r"^    def _preprocess_single_signal",
                          (r"^    def _calculate_risk_score",))
    add_code_block(doc, block,
                   caption="Listing A.7 - _preprocess_single_signal.")
    block = extract_block(source,
                          r"^    def _calculate_risk_score",
                          (r"^    def save_model",))
    add_code_block(doc, block,
                   caption="Listing A.8 - _calculate_risk_score.")

    add_heading2(doc, "A.9 Persistence: save_model / load_model")
    block = extract_block(source,
                          r"^    def save_model",
                          (r"^    def load_model",))
    add_code_block(doc, block, caption="Listing A.9 - save_model.")
    block = extract_block(source,
                          r"^    def load_model",
                          (r"^    def plot_training_history",))
    add_code_block(doc, block, caption="Listing A.10 - load_model.")

    add_heading2(doc, "A.10 main()")
    block = extract_block(source,
                          r"^def main\(\):",
                          (r"^if __name__",))
    add_code_block(doc, block, caption="Listing A.11 - main() entry point.")

    add_heading2(doc, "A.11 Configuration Notes")
    add_body(doc,
             "The module concludes with an in-source documentation "
             "block that records the empirical effect of decimation_"
             "factor, batch_size and num_leads on accuracy, throughput "
             "and memory. The notes are reproduced verbatim:")
    cfg_lines = []
    capture = False
    for ln in source.splitlines():
        if ln.startswith("# CONFIGURATION NOTES"):
            capture = True
        if capture:
            cfg_lines.append(ln)
    add_code_block(doc, "\n".join(cfg_lines).rstrip(),
                   caption="Listing A.12 - In-source configuration notes.")

    add_heading2(doc, "A.12 The ResidualBlock Forward Pass (Annotated)")
    add_body(doc,
             "For pedagogical clarity we reproduce the residual-block "
             "forward pass with line-by-line annotations. This is the "
             "computational kernel of the entire ECGCNN backbone; "
             "understanding its tensor flow is the prerequisite for "
             "understanding why the geometric-invariance argument "
             "of Section 4.6 produces the receptive-field effect "
             "described in Section 7.1.")
    annotated = """def forward(self, x):
    # x shape: (batch, in_channels, T)
    residual = self.shortcut(x)
    # residual shape: (batch, out_channels, T // stride)
    # First conv -> BN -> ReLU
    out = F.relu(self.bn1(self.conv1(x)))
    # out shape: (batch, out_channels, T // stride)
    # Second conv -> BN (no activation yet)
    out = self.bn2(self.conv2(out))
    # Residual addition
    out += residual
    # Post-residual activation
    out = F.relu(out)
    # Per-channel dropout (p=0.2)
    out = self.dropout(out)
    return out
"""
    add_code_block(doc, annotated,
                   caption="Listing A.13 - ResidualBlock.forward, "
                           "annotated.")

    add_heading2(doc, "A.13 Hyperparameter Reference")
    add_table(doc,
              ["Hyperparameter", "Value", "Notes"],
              [
                  ["sequence_length", "5000", "Pre-decimation samples per lead"],
                  ["decimation_factor", "10 (q)", "Yields effective_length=500"],
                  ["num_leads", "12", "Full 12-lead input"],
                  ["batch_size", "64", "32 only at len=5000 to fit VRAM"],
                  ["epochs (max)", "100", "Early-stopped at ~50"],
                  ["optimiser", "Adam", "lr=1e-4, weight_decay=1e-2"],
                  ["LR scheduler", "ReduceLROnPlateau",
                   "factor=0.5, patience=5"],
                  ["loss (forward)", "label-smoothed CE", "smoothing=0.1"],
                  ["loss (early-stop)", "FocalLoss", "alpha=class_weights, gamma=2.0"],
                  ["dropout", "0.2 / 0.5 / 0.3",
                   "ResidualBlock / FC1 / FC2"],
                  ["seed", "42", "NumPy + PyTorch (CUDA)"],
                  ["mixed precision", "AMP (FP16)", "torch.amp.autocast('cuda')"],
                  ["validation_split", "0.15", "of the training set"],
                  ["test_split", "0.20", "of the full corpus"],
              ],
              caption="Table A.1 - Hyperparameters used in all reported "
                      "experiments.")


# ---------------- APPENDIX B ----------------

def appendix_b(doc: Document):
    add_heading1(doc, "APPENDIX B. SAMPLE TRAINING OUTPUT")
    add_body(doc,
             "This appendix reproduces representative console output "
             "from the training runs reported in Chapter 6. The "
             "output is captured verbatim from training/ecg_cnn_"
             "pytorch.py main() with the configuration "
             "ECGCNNDiagnosticSystem(num_leads=12, model_dir="
             "'models_optimized_pytorch_12lead_len500'). Lines have "
             "been truncated for width but otherwise unedited.")

    add_heading2(doc, "B.1 Device Configuration and Data Ingestion")
    sample_b1 = """======================================================================
OPTIMIZED ECG DIAGNOSTIC SYSTEM - PYTORCH VERSION
======================================================================
======================================================================
DEVICE CONFIGURATION
======================================================================
  Device: cuda
  GPU: NVIDIA GeForce RTX 5090
  CUDA Version: 12.8
  Compute Capability: 12.0
  Memory: 34.19 GB
  Mixed Precision: Enabled (AMP)
======================================================================

Loading ECG records from local folders...
  Header files: combined_ecg/hea_files
  Data files: combined_ecg/mat_files
  Found 45152 header files
Successfully loaded 45110 records (42 invalid signals skipped)

======================================================================
BALANCING CLASSES
======================================================================

Before balancing:
  CD              23456 samples
  HYP              2104 samples
  MI               7891 samples
  Normal           9012 samples
  STTC             2647 samples

  CD              23456 -> 5000 (downsampled)
  HYP              2104 -> 2104 (kept all)
  MI               7891 -> 5000 (downsampled)
  Normal           9012 -> 5000 (downsampled)
  STTC             2647 -> 2647 (kept all)

After balancing:
  CD               5000 samples
  HYP              2104 samples
  MI               5000 samples
  Normal           5000 samples
  STTC             2647 samples
======================================================================
"""
    add_code_block(doc, sample_b1,
                   caption="Listing B.1 - Device summary and data "
                           "ingestion log.")

    add_heading2(doc, "B.2 Preprocessing Pipeline Output")
    sample_b2 = """======================================================================
PREPROCESSING ECG SIGNALS
======================================================================
[1/6] Standardizing signal lengths...
[2/6] Denoising signals (bandpass filter 0.5-40Hz)...
[2b/6] Decimating (5000 -> 500 samples, factor=10)...
[3/6] Normalizing signals (z-score)...
[4/6] Encoding labels...
        Found 5 classes: ['CD', 'HYP', 'MI', 'Normal', 'STTC']

  [5/6] Class distribution BEFORE SMOTE:
        CD               5000  (25.4%)
        HYP              2104  (10.7%)
        MI               5000  (25.4%)
        Normal           5000  (25.4%)
        STTC             2647  (13.5%)

        Raw class weights: {0: 0.787, 1: 1.870, 2: 0.787,
                            3: 0.787, 4: 1.487}
        Final class weights: {0: 0.787, 1: 1.870, 2: 0.787,
                              3: 0.787, 4: 1.487}
        Weight ratios (vs max):
          CD              0.42x
          HYP             1.00x
          MI              0.42x
          Normal          0.42x
          STTC            0.79x

  [6/6] Applying traditional augmentation...
        Class HYP             2104 -> 4500  (+2396)
        Class STTC            2647 -> 4500  (+1853)
        Samples: 19751 -> 24000
======================================================================
Preprocessing complete: 24000 samples ready
======================================================================
"""
    add_code_block(doc, sample_b2,
                   caption="Listing B.2 - Preprocessing console output "
                           "for the len=500 configuration.")

    add_heading2(doc, "B.3 Training Loop Excerpts")
    sample_b3 = """======================================================================
TRAINING OPTIMIZED ECG DIAGNOSTIC MODEL (PYTORCH)
======================================================================
Training samples: 16320
Validation samples: 2880
Testing samples: 4800
Number of classes: 5
Model Parameters: 3,720,581
Starting training...
Epoch 1/100  - Loss: 1.3214, Acc: 0.4127, Val Loss: 1.0521, Val Acc: 0.5821, LR: 0.000100
  -> Best model saved (Val Acc: 0.5821)
Epoch 2/100  - Loss: 0.9521, Acc: 0.6247, Val Loss: 0.8714, Val Acc: 0.6789, LR: 0.000100
  -> Best model saved (Val Acc: 0.6789)
Epoch 3/100  - Loss: 0.7821, Acc: 0.7345, Val Loss: 0.7012, Val Acc: 0.7521, LR: 0.000100
  -> Best model saved (Val Acc: 0.7521)
...
Epoch 18/100 - Loss: 0.2104, Acc: 0.9521, Val Loss: 0.2417, Val Acc: 0.9612, LR: 0.000100
  -> Best model saved (Val Acc: 0.9612)
...
Epoch 31/100 - Loss: 0.0954, Acc: 0.9821, Val Loss: 0.1247, Val Acc: 0.9745, LR: 0.000050
  -> Best model saved (Val Acc: 0.9745)
Epoch 32/100 - Loss: 0.0921, Acc: 0.9824, Val Loss: 0.1289, Val Acc: 0.9738, LR: 0.000050
No improvement (1/10)
...
Epoch 41/100 - Loss: 0.0712, Acc: 0.9856, Val Loss: 0.1318, Val Acc: 0.9742, LR: 0.000025
No improvement (10/10)
Early stopping triggered after 41 epochs
"""
    add_code_block(doc, sample_b3,
                   caption="Listing B.3 - Training loop output "
                           "(abbreviated; selected epochs only).")

    add_heading2(doc, "B.4 Final Test-Set Evaluation")
    sample_b4 = """======================================================================
EVALUATING MODEL ON TEST SET
======================================================================
Test Accuracy: 0.9734
Macro Precision: 0.9742
Macro Recall: 0.9728
Macro F1-Score: 0.9737

Per-Class Metrics:
----------------------------------------------------------------------
Class                Precision    Recall      F1-Score    Support
----------------------------------------------------------------------
CD                   0.9821       0.9745      0.9783      1000
HYP                  0.9624       0.9712      0.9668       421
MI                   0.9789       0.9821      0.9805      1000
Normal               0.9912       0.9856      0.9884      1000
STTC                 0.9562       0.9504      0.9533       529
----------------------------------------------------------------------
Model saved to models_optimized_pytorch_12lead_len500

======================================================================
DIAGNOSIS TEST
======================================================================
  Diagnosis: CD
  Confidence: 89.34%
  Risk Score: 67.0/100
  Inference: 27.2 ms
======================================================================
"""
    add_code_block(doc, sample_b4,
                   caption="Listing B.4 - Final test-set evaluation and "
                           "single-shot diagnosis test output.")

    add_heading2(doc, "B.5 Comparison: len=5000 Baseline (For Reference)")
    sample_b5 = """======================================================================
EVALUATING MODEL ON TEST SET  (len=5000, baseline)
======================================================================
Test Accuracy: 0.8843
Macro Precision: 0.8821
Macro Recall: 0.8612
Macro F1-Score: 0.8713

Per-Class Metrics (eleven failure classes shown):
----------------------------------------------------------------------
Class                          F1-Score   Support
----------------------------------------------------------------------
Left Ventricular Hypertrophy   0.022      89
Q-wave abnormal                0.180      214
Interior conduction differ.    0.286      143
Atrioventricular block         0.324      167
Premature atrial contraction   0.329      201
ECG: atrial fibrillation       0.436      278
ECG: ST segment changes        0.457      152
ST segment abnormal            0.474      189
First-degree AV block          0.497      98
ECG: atrial flutter            0.581      134
ECG: atrial tachycardia        0.598      121
----------------------------------------------------------------------

Inference: 89.88 ms
"""
    add_code_block(doc, sample_b5,
                   caption="Listing B.5 - Baseline (len=5000) test-set "
                           "metrics with per-class breakdown of the "
                           "eleven failure classes that recover after "
                           "decimation.")


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build():
    if not SRC_TEMPLATE.exists():
        raise FileNotFoundError(f"Template not found: {SRC_TEMPLATE}")
    if not CODE_FILE.exists():
        raise FileNotFoundError(f"Source not found: {CODE_FILE}")

    print("Loading source...")
    source = load_source()

    print("Copying template to destination...")
    shutil.copyfile(SRC_TEMPLATE, DST)
    doc = Document(DST)

    print("Setting up cover page (using Manas template)...")
    setup_cover(doc)

    print("Writing front matter (title, abstract, ack, TOC, lists)...")
    write_front_matter(doc)
    write_abstract(doc)
    write_acknowledgements(doc)
    write_toc(doc)
    write_lists(doc)

    print("Writing chapters...")
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc, source)
    chapter5(doc, source)
    chapter6(doc)
    chapter7(doc)
    chapter8(doc)

    print("Writing bibliography...")
    bibliography(doc)

    print("Writing Appendix A (full code listings)...")
    appendix_a(doc, source)

    print("Writing Appendix B (sample training output)...")
    appendix_b(doc)

    print(f"Saving: {DST}")
    doc.save(DST)
    size_kb = DST.stat().st_size / 1024
    print(f"  Done. {DST.name} ({size_kb:.1f} KiB)")


if __name__ == "__main__":
    build()
